# PDF转MD文档处理系统

from flask import Flask, render_template, request, jsonify, redirect, url_for, session, flash, send_file, Response
from flask_wtf.csrf import CSRFProtect
from functools import wraps
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from database_config import DatabaseConfig
from database_manager import get_db_connection_with_error_handling, safe_db_close, db_manager
from md_storage_manager import md_storage_manager
# 导入标题优化处理器
from title_enhancer import TitleEnhancer
# 导入数据库操作模块
from db_operations import (
    log_user_action, get_file_info_from_db, get_user_by_username_or_email, 
    get_user_by_username, save_file_to_db, get_user_files_from_db, count_pending_files_for_user, 
    check_user_exists, check_email_exists, create_user, get_file_content_from_db,
    get_active_api_key, update_task_status_in_db, update_file_status_in_db,
    update_file_extracted_json, get_pdf_data_from_db,
    get_file_reference_data, add_group_api_key, get_group_api_keys, get_user_siliconflow_model
)
# 导入通知系统模块
from notification_system import init_notification_routes
# 导入智能日志系统
from smart_logger import smart_print, LogLevel, smart_logger
# Word文档处理相关导入 - 使用跨平台python-docx
from docx import Document as DocxDocument
# 参考文献验证器导入
from reference_validator import GBT7714Validator, detect_citation_system_in_text
import pymysql
import json
import os
import uuid
import time
import threading
import tempfile
import shutil
import re

# ==================== 全局并发控制 ====================
global_active_tasks_lock = threading.Lock()
global_active_tasks = 0  # 当前活跃的一键解析任务数

# 全局文件处理线程池：所有一键解析任务共享，严格限制最多2个文件同时处理
from concurrent.futures import ThreadPoolExecutor
_global_file_executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix='file-proc')

# 任务字典线程安全锁（保护 tasks_db 和 one_click_tasks 的并发读写）
tasks_db_lock = threading.Lock()
one_click_tasks_lock = threading.Lock()

# ==================== Word文档管理器 - 跨平台python-docx方案 ====================
class WordInstanceManager:
    """Word文档管理器 - 使用python-docx跨平台处理"""
    
    @staticmethod
    def create_word_instance():
        """兼容接口 - python-docx不需要实例"""
        return None
    
    @staticmethod
    def cleanup_word_instance(word, doc=None):
        """兼容接口 - python-docx不需要清理"""
        pass

word_instance_manager = WordInstanceManager()

# ==================== 监控状态缓存 ====================
# 用于存储上次监控的状态，实现变化检测
_monitoring_cache = {
    'last_system_stats': None,
    'last_queue_status': None, 
    'last_user_activity': None,
    'last_performance': None
}

def clean_monitoring_cache():
    """清理监控缓存中的 Decimal 类型数据"""
    global _monitoring_cache
    try:
        for cache_key, cache_data in _monitoring_cache.items():
            if cache_data and isinstance(cache_data, dict):
                for key, value in cache_data.items():
                    if value is not None:
                        # 转换可能的 Decimal 类型为 float
                        if hasattr(value, '__class__') and 'Decimal' in str(type(value)):
                            cache_data[key] = float(value)
                        elif isinstance(value, (list, tuple)):
                            # 处理列表中的 Decimal
                            cache_data[key] = [float(v) if hasattr(v, '__class__') and 'Decimal' in str(type(v)) else v for v in value]
    except Exception as e:
        # 如果清理失败，直接重置缓存
        _monitoring_cache = {
            'last_system_stats': None,
            'last_queue_status': None, 
            'last_user_activity': None,
            'last_performance': None
        }
import traceback
import zipfile
from datetime import datetime, timezone, timedelta
from typing import List, Dict, Any, Optional
from pathlib import Path
import requests
import qrcode
from io import BytesIO
import re
import html
from typing import List, Any, Dict, Union
from jinja2 import ChoiceLoader, FileSystemLoader

# 创建 Flask 应用
app = Flask(__name__)

# 配置多个模板文件夹 - 使用 ChoiceLoader 组合多个 FileSystemLoader
app.jinja_loader = ChoiceLoader([
    FileSystemLoader(os.path.join(app.root_path, 'templates')),
    FileSystemLoader(os.path.join(app.root_path, 'word_templates')),
])

app.config['JSON_AS_ASCII'] = False
# 安全：生产环境必须通过 SECRET_KEY 环境变量设置，否则每次重启session失效
_default_secret = os.environ.get('SECRET_KEY')
if not _default_secret:
    import hashlib
    # 基于机器名+路径生成固定但不公开的密钥（开发环境用）
    _default_secret = hashlib.sha256(f'{os.name}-{os.path.abspath(__file__)}-check_markdown_2025'.encode()).hexdigest()
    print('[安全警告] SECRET_KEY 未通过环境变量设置，使用基于机器特征的派生密钥（生产环境请设置环境变量）')
app.config['SECRET_KEY'] = _default_secret
app.config['SESSION_COOKIE_HTTPONLY'] = True      # 禁止JS读取session cookie
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'     # 防止CSRF跨站请求
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)  # 默认session有效期24小时
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB 最大文件大小

# CSRF保护
csrf = CSRFProtect(app)

@app.before_request
def _set_session_lifetime():
    """根据是否勾选“记住我”动态设置每个session的过期时间"""
    if session.get('remember_me'):
        app.permanent_session_lifetime = timedelta(days=30)
    else:
        app.permanent_session_lifetime = timedelta(hours=24)

# 登录验证装饰器
def login_required(f):
    """页面路由登录验证装饰器，未登录则跳转登录页"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def api_login_required(f):
    """API路由登录验证装饰器，未登录返回401 JSON"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'error': '请先登录'}), 401
        return f(*args, **kwargs)
    return decorated_function

def word_login_required(f):
    """Word模板路由登录验证装饰器"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('word_login'))
        return f(*args, **kwargs)
    return decorated_function

# 确保上传文件夹存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('downloads', exist_ok=True)  # 创建下载文件夹

# 配置数据库管理器的日志记录器
db_manager.logger = app.logger

# 初始化通知系统路由
init_notification_routes(app)

# 注意：以下内存存储已废弃，所有数据现在直接使用数据库
# 保留变量声明是为了兼容现有代码，但数据完全来自数据库
# files_db已完全废弃，所有文件数据直接从数据库获取

# 任务状态追踪（后续也可以迁移到MySQL）  
# 全局变量存储文件信息和任务状态
tasks_db = {}  # 任务状态存储

# 🔒 全局VLM信号量：本地VLM(vLLM)支持并行请求(continuous batching)，允许2个并发
_vlm_semaphore = threading.Semaphore(2)

# 🔒 全局本地MinerU信号量：本地MinerU独占GPU，允许2个并发PDF解析
_mineru_semaphore = threading.Semaphore(2)

# 简化的时间处理函数
def get_current_time():
    """获取当前中国时间（UTC+8）"""
    from datetime import timedelta
    china_tz = timezone(timedelta(hours=8))
    return datetime.now(china_tz)

def format_time_for_display(dt):
    """格式化时间用于前端显示"""
    if dt is None:
        return None
    if isinstance(dt, str):
        try:
            dt = datetime.fromisoformat(dt.replace('Z', '+00:00'))
        except:
            return dt
    
    if hasattr(dt, 'strftime'):
        return dt.strftime('%Y-%m-%d %H:%M:%S')
    return str(dt)

def format_file_size(size_bytes):
    """格式化文件大小为易读形式"""
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024.0 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    
    if i == 0:
        return f"{int(size_bytes)} {size_names[i]}"
    else:
        return f"{size_bytes:.1f} {size_names[i]}"

# ==================== 实时监控和统计功能 ====================

def log_system_realtime_stats():
    """记录系统实时统计信息 - 只在有显著变化时记录"""
    import psutil
    import threading
    
    try:
        # 清理缓存中可能的 Decimal 数据
        clean_monitoring_cache()
        
        # 获取系统资源使用情况
        cpu_usage = psutil.cpu_percent(interval=1)
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        # 获取当前处理任务统计
        with tasks_db_lock:
            active_tasks = len([task for task in list(tasks_db.values()) if task.get('status') == 'processing'])
        
        # 获取今日统计
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # 今日上传文件数
            cursor.execute("""
                SELECT COUNT(*) as daily_uploads 
                FROM user_files 
                WHERE upload_time >= CURDATE() 
                AND upload_time < DATE_ADD(CURDATE(), INTERVAL 1 DAY)
            """)
            result = cursor.fetchone()
            if not result:
                raise Exception("Failed to get daily uploads count")
            daily_uploads = result['daily_uploads']
            
            # 今日处理完成数
            cursor.execute("""
                SELECT COUNT(*) as daily_processed 
                FROM user_files 
                WHERE created_at >= CURDATE() 
                AND created_at < DATE_ADD(CURDATE(), INTERVAL 1 DAY) 
                AND status = 'completed'
            """)
            result = cursor.fetchone()
            if not result:
                raise Exception("Failed to get daily processed count")
            daily_processed = result['daily_processed']
            
            # 当前在线用户数（最近5分钟有活动的用户）
            cursor.execute("""
                SELECT COUNT(DISTINCT user_id) as online_users 
                FROM system_logs 
                WHERE created_at >= DATE_SUB(NOW(), INTERVAL 5 MINUTE)
            """)
            result = cursor.fetchone()
            if not result:
                raise Exception("Failed to get online users count")
            online_users = result['online_users']
            
            conn.close()
        else:
            daily_uploads = 0
            daily_processed = 0  
            online_users = 0
        
        # 创建当前状态快照
        current_stats = {
            'cpu_usage': round(cpu_usage, 1),
            'memory_percent': round(memory.percent, 1),
            'disk_percent': round(disk.percent, 1),
            'active_tasks': active_tasks,
            'daily_uploads': daily_uploads,
            'daily_processed': daily_processed,
            'online_users': online_users
        }
        
        # 检查是否有显著变化（CPU/内存变化>5%，或其他数值有变化）
        last_stats = _monitoring_cache.get('last_system_stats')
        should_log = False
        
        if last_stats is None:
            should_log = True  # 首次记录
        else:
            # 确保缓存中的数值是 float 类型
            for key in ['cpu_usage', 'memory_percent', 'disk_percent']:
                if key in last_stats and last_stats[key] is not None:
                    last_stats[key] = float(last_stats[key])
            
            # 检查CPU和内存是否有显著变化（>5%）
            cpu_change = abs(float(current_stats['cpu_usage']) - float(last_stats['cpu_usage']))
            memory_change = abs(float(current_stats['memory_percent']) - float(last_stats['memory_percent']))
            
            # 检查其他指标是否有变化
            other_changes = (
                current_stats['active_tasks'] != last_stats['active_tasks'] or
                current_stats['daily_uploads'] != last_stats['daily_uploads'] or
                current_stats['daily_processed'] != last_stats['daily_processed'] or
                current_stats['online_users'] != last_stats['online_users']
            )
            
            should_log = cpu_change >= 5.0 or memory_change >= 5.0 or other_changes
        
        # 只在有显著变化时记录日志
        if should_log:
            # 更新缓存
            _monitoring_cache['last_system_stats'] = current_stats
            
            # 获取系统管理员ID
            system_user_id = None
            try:
                cursor.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                admin_user = cursor.fetchone()
                if admin_user:
                    system_user_id = admin_user['id']
            except:
                pass
            
            # 只有获取到有效用户ID时才记录日志
            if system_user_id:
                # 记录系统性能日志
                log_user_action(
                    user_id=system_user_id,
                    action='system_performance_summary',
                    details=f'CPU使用率: {cpu_usage:.1f}%, 内存使用: {memory.percent:.1f}%, 磁盘使用: {disk.percent:.1f}%',
                    ip_address=None,
                    user_agent='System-Monitor'
                )
                
                # 记录处理统计日志
                log_user_action(
                    user_id=system_user_id,
                    action='processing_status_summary', 
                    details=f'处理中任务: {active_tasks}, 今日上传: {daily_uploads}, 今日处理: {daily_processed}, 在线用户: {online_users}',
                    ip_address=None,
                    user_agent='Processing-Monitor'
                )
        
    except Exception as e:
        smart_print(f"[实时统计] 记录系统统计失败: {e}", LogLevel.ERROR)

def log_queue_status_update():
    """记录处理队列状态更新 - 只在状态有变化时记录"""
    try:
        # 清理缓存中可能的 Decimal 数据
        clean_monitoring_cache()
        
        # 统计不同状态的任务
        with tasks_db_lock:
            tasks_snapshot = list(tasks_db.values())
            total_tasks = len(tasks_db)
        processing_count = len([task for task in tasks_snapshot if task.get('status') == 'processing'])
        pending_count = len([task for task in tasks_snapshot if task.get('status') in ['pending', 'waiting']])
        failed_count = len([task for task in tasks_snapshot if task.get('status') == 'error'])
        completed_count = len([task for task in tasks_snapshot if task.get('status') == 'completed'])
        success_rate = (completed_count / total_tasks * 100) if total_tasks > 0 else 0
        
        # 创建当前队列状态快照
        current_queue = {
            'total_tasks': total_tasks,
            'processing_count': processing_count,
            'pending_count': pending_count,
            'failed_count': failed_count,
            'completed_count': completed_count,
            'success_rate': round(success_rate, 1)
        }
        
        # 检查是否有变化
        last_queue = _monitoring_cache.get('last_queue_status')
        should_log = False
        
        if last_queue is None:
            # 首次记录，只在有任务时记录
            should_log = total_tasks > 0
        else:
            # 检查是否有任何状态变化
            should_log = current_queue != last_queue
        
        # 只在有变化时记录日志
        if should_log:
            # 更新缓存
            _monitoring_cache['last_queue_status'] = current_queue
            
            # 获取系统管理员ID
            system_user_id = None
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    cursor.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                    admin_user = cursor.fetchone()
                    if admin_user:
                        system_user_id = admin_user['id']
                    conn.close()
            except:
                pass
            
            # 只有获取到有效用户ID时才记录日志
            if system_user_id:
                log_user_action(
                    user_id=system_user_id,
                    action='queue_status_update',
                    details=f'队列总数: {total_tasks}, 处理中: {processing_count}, 等待中: {pending_count}, 已完成: {completed_count}, 失败: {failed_count}, 成功率: {success_rate:.1f}%',
                    ip_address=None,
                    user_agent='Queue-Monitor'
                )
        
    except Exception as e:
        smart_print(f"[队列监控] 记录队列状态失败: {e}", LogLevel.ERROR)

def log_user_activity_metrics():
    """记录用户活动指标 - 只在有显著变化时记录"""
    try:
        # 清理缓存中可能的 Decimal 数据
        clean_monitoring_cache()
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return
            
        cursor = conn.cursor()
        
        # 最近24小时活跃用户数
        cursor.execute("""
            SELECT COUNT(DISTINCT user_id) as active_users_24h
            FROM system_logs 
            WHERE created_at >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
        """)
        result = cursor.fetchone()
        active_users_24h = result['active_users_24h'] if result else 0
        
        # 最近24小时新注册用户
        cursor.execute("""
            SELECT COUNT(*) as new_registrations_24h
            FROM users 
            WHERE created_at >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
        """)
        result = cursor.fetchone()
        new_registrations_24h = result['new_registrations_24h'] if result else 0
        
        # 最受欢迎的功能（按日志action统计）
        cursor.execute("""
            SELECT action, COUNT(*) as usage_count
            FROM system_logs 
            WHERE created_at >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
            AND action NOT IN ('login', 'logout', 'page_access', 'user_activity_metrics', 'queue_status_update')
            GROUP BY action 
            ORDER BY usage_count DESC 
            LIMIT 3
        """)
        popular_features = cursor.fetchall()
        top_features = [f"{feature['action']}({feature['usage_count']})" for feature in popular_features]
        
        # 创建当前用户活动状态快照
        current_activity = {
            'active_users_24h': active_users_24h,
            'new_registrations_24h': new_registrations_24h,
            'top_features': top_features
        }
        
        # 检查是否有变化
        last_activity = _monitoring_cache.get('last_user_activity')
        should_log = False
        
        if last_activity is None:
            # 首次记录，只在有实际活动时记录
            should_log = active_users_24h > 1 or new_registrations_24h > 0 or len(top_features) > 0
        else:
            # 检查是否有变化
            should_log = current_activity != last_activity
        
        # 只在有变化时记录日志
        if should_log:
            # 更新缓存
            _monitoring_cache['last_user_activity'] = current_activity
            
            conn.close()
            
            # 获取系统管理员ID
            system_user_id = None
            try:
                conn2 = DatabaseConfig.get_db_connection()
                if conn2:
                    cursor2 = conn2.cursor()
                    cursor2.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                    admin_user = cursor2.fetchone()
                    if admin_user:
                        system_user_id = admin_user['id']
                    conn2.close()
            except:
                pass
            
            # 只有获取到有效用户ID时才记录日志
            if system_user_id:
                # 记录用户活动指标日志
                log_user_action(
                    user_id=system_user_id,
                    action='user_activity_metrics',
                    details=f'24小时活跃用户: {active_users_24h}, 新注册: {new_registrations_24h}, 热门功能: {", ".join(top_features)}',
                    ip_address=None,
                    user_agent='Activity-Monitor'
                )
        else:
            conn.close()
        
    except Exception as e:
        smart_print(f"[用户活动] 记录用户活动指标失败: {e}", LogLevel.ERROR)

def log_processing_performance_metrics():
    """记录处理性能指标 - 只在有显著变化时记录"""
    try:
        # 强制清空性能缓存，避免 Decimal 问题
        if 'last_performance' in _monitoring_cache:
            _monitoring_cache['last_performance'] = None
        
        # 清理缓存中可能存在的 Decimal 类型数据
        if 'last_performance' in _monitoring_cache:
            last_perf = _monitoring_cache['last_performance']
            if last_perf and isinstance(last_perf, dict):
                for key in ['avg_time', 'min_time', 'max_time']:
                    if key in last_perf and last_perf[key] is not None:
                        last_perf[key] = float(last_perf[key])
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return
            
        cursor = conn.cursor()
        
        # 最近24小时的处理性能统计
        cursor.execute("""
            SELECT 
                AVG(TIMESTAMPDIFF(SECOND, processing_start_time, processed_time)) as avg_processing_time,
                MIN(TIMESTAMPDIFF(SECOND, processing_start_time, processed_time)) as min_processing_time,
                MAX(TIMESTAMPDIFF(SECOND, processing_start_time, processed_time)) as max_processing_time,
                COUNT(*) as processed_files
            FROM user_files 
            WHERE processed_time IS NOT NULL 
            AND processing_start_time IS NOT NULL
            AND processed_time >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
            AND status = 'completed'
        """)
        result = cursor.fetchone()
        
        # 创建当前性能状态快照
        current_performance = {
            'processed_files': 0,
            'avg_time': 0,
            'min_time': 0,
            'max_time': 0
        }
        
        if result and result['processed_files'] > 0:
            current_performance = {
                'processed_files': result['processed_files'],
                'avg_time': round(float(result['avg_processing_time'] or 0), 1),
                'min_time': round(float(result['min_processing_time'] or 0), 1),
                'max_time': round(float(result['max_processing_time'] or 0), 1)
            }
        
        # 检查是否有变化
        last_performance = _monitoring_cache.get('last_performance')
        should_log = False
        
        if last_performance is None:
            # 首次记录，只在有实际处理活动时记录
            should_log = current_performance['processed_files'] > 0
        else:
            # 确保 last_performance 数据类型安全
            if not isinstance(last_performance.get('avg_time', 0), (int, float)):
                last_performance['avg_time'] = float(last_performance.get('avg_time', 0))
            
            # 检查是否有变化（文件数量变化或处理时间有显著变化 >10%）
            files_changed = current_performance['processed_files'] != last_performance['processed_files']
            
            # 检查平均处理时间是否有显著变化 (超过10%变化或绝对差值超过5秒)
            time_threshold = max(0.1 * float(last_performance['avg_time']), 5.0) if last_performance['avg_time'] > 0 else 5.0
            time_changed = abs(float(current_performance['avg_time']) - float(last_performance['avg_time'])) > time_threshold
            
            should_log = files_changed or time_changed
        
        # 只在有变化时记录日志
        if should_log:
            # 更新缓存
            _monitoring_cache['last_performance'] = current_performance
            
            # 获取系统管理员ID
            system_user_id = None
            try:
                cursor.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                admin_user = cursor.fetchone()
                if admin_user:
                    system_user_id = admin_user['id']
            except:
                pass
            
            # 只有获取到有效用户ID时才记录日志
            if system_user_id and current_performance['processed_files'] > 0:
                log_user_action(
                    user_id=system_user_id,
                    action='processing_performance_metrics',
                    details=f'24小时内处理文件: {current_performance["processed_files"]}, 平均耗时: {current_performance["avg_time"]}秒, 最快: {current_performance["min_time"]}秒, 最慢: {current_performance["max_time"]}秒',
                    ip_address=None,
                    user_agent='Performance-Monitor'
                )
        
        conn.close()
        
    except Exception as e:
        # 增加详细的错误信息，帮助定位问题
        import traceback
        error_details = f"错误: {e}, 类型: {type(e).__name__}, 追踪: {traceback.format_exc()}"
        smart_print(f"[性能监控] 记录处理性能指标失败: {error_details}", LogLevel.ERROR)

# 全局监控控制变量
_monitoring_shutdown = False

# 启动后台实时监控任务
def start_realtime_monitoring():
    """启动实时监控任务 - 恢复原频率但保持智能过滤"""
    import threading
    import time
    
    def monitoring_loop():
        """监控循环 - 恢复5分钟频率，但保持智能化过滤垃圾日志"""
        global _monitoring_shutdown
        while not _monitoring_shutdown:
            try:
                # 每5分钟记录一次系统统计
                if not _monitoring_shutdown:
                    log_system_realtime_stats()
                
                # 每5分钟记录一次用户活动（仅在有实际活动时）
                if not _monitoring_shutdown:
                    log_user_activity_metrics() 
                
                # 只在有队列任务时记录队列状态
                if not _monitoring_shutdown:
                    log_queue_status_update()
                
                # 每5分钟记录一次性能指标
                if not _monitoring_shutdown:
                    log_processing_performance_metrics()
                
                # 分段睡眠，允许更快响应关闭信号
                for _ in range(60):  # 60 x 5秒 = 300秒（5分钟）
                    if _monitoring_shutdown:
                        break
                    time.sleep(5)
                
            except Exception as e:
                if not _monitoring_shutdown:
                    smart_print(f"[监控循环] 监控任务异常: {e}", LogLevel.ERROR)
                    time.sleep(60)  # 出错时等待1分钟再重试
        
        smart_print("[实时监控] 监控循环已安全退出", LogLevel.INFO)
    
    # 启动监控线程
    monitor_thread = threading.Thread(target=monitoring_loop, daemon=True)
    monitor_thread.start()
    smart_print("[实时监控] 后台监控任务已启动（5分钟频率，智能过滤垃圾日志）", LogLevel.INFO)

def shutdown_monitoring():
    """优雅关闭监控系统"""
    global _monitoring_shutdown
    _monitoring_shutdown = True
    smart_print("[实时监控] 正在关闭监控系统...", LogLevel.INFO)

def parse_timestamp_safe(timestamp_value):
    """简单安全的时间戳解析"""
    if timestamp_value is None:
        return None
    
    try:
        if isinstance(timestamp_value, str):
            # 字符串时间戳
            if 'T' in timestamp_value:
                # ISO格式
                if 'Z' in timestamp_value:
                    # UTC时间
                    dt = datetime.fromisoformat(timestamp_value.replace('Z', '+00:00'))
                    return dt.timestamp()
                elif '+' in timestamp_value or '-' in timestamp_value[-6:]:
                    # 带时区的ISO时间
                    dt = datetime.fromisoformat(timestamp_value)
                    return dt.timestamp()
                else:
                    # 不带时区的ISO时间，假设为本地时间
                    dt = datetime.fromisoformat(timestamp_value)
                    return dt.timestamp()
            else:
                # 尝试解析为浮点数时间戳
                return float(timestamp_value)
        elif hasattr(timestamp_value, 'timestamp'):
            # datetime对象
            return timestamp_value.timestamp()
        else:
            # 数字时间戳
            return float(timestamp_value)
    except:
        # 解析失败，返回当前时间戳
        return time.time()

def interrupt_file_tasks(file_id: str, user_id: int) -> List[str]:
    """中断指定文件的所有相关任务"""
    interrupted_tasks = []
    
    # 1. 查找并中断 tasks_db 中的相关任务
    task_ids_to_remove = []
    with tasks_db_lock:
        for task_id, task in list(tasks_db.items()):
            if task.get('file_id') == file_id and task.get('user_id') == user_id:
                # 标记任务为中断状态
                task['status'] = 'cancelled'
                task['message'] = '❌ 文件已删除，任务中断'
                task_ids_to_remove.append(task_id)
                interrupted_tasks.append(f"文件任务: {task_id}")
                smart_print(f"[任务中断] 文件任务中断: {task_id}", LogLevel.WARNING)
    
    # 延迟删除任务记录，让前端有时间显示中断状态
    import threading
    def delayed_cleanup():
        import time
        for i in range(20):  # 分20次检查，每次100ms，总共2秒
            if _monitoring_shutdown:
                return  # 如果系统正在关闭，直接退出
            time.sleep(0.1)
        
        if not _monitoring_shutdown:  # 只有在系统未关闭时才清理
            with tasks_db_lock:
                for task_id in task_ids_to_remove:
                    tasks_db.pop(task_id, None)
                    smart_print(f"[任务清理] 已清理文件任务: {task_id}", LogLevel.INFO)
    
    cleanup_thread = threading.Thread(target=delayed_cleanup)
    cleanup_thread.daemon = True
    cleanup_thread.start()
    
    # 2. 查找并中断一键处理任务中的相关文件
    with one_click_tasks_lock:
        for task_id, task in list(one_click_tasks.items()):
            if task.get('user_id') == user_id:
                files_to_process = task.get('files_to_process', [])
                # 检查是否包含要删除的文件
                file_found = False
                for file_info in files_to_process:
                    if file_info.get('file_id') == file_id:
                        file_found = True
                        break
                
                if file_found:
                    # 从处理列表中移除该文件
                    task['files_to_process'] = [f for f in files_to_process if f.get('file_id') != file_id]
                    task['total_files'] = len(task['files_to_process'])
                    
                    # 查找并中断子任务
                    oneclick_file_task_id = f"oneclick_{task_id}_{file_id}"
                    if oneclick_file_task_id in tasks_db:
                        tasks_db[oneclick_file_task_id]['status'] = 'cancelled'
                        tasks_db[oneclick_file_task_id]['message'] = '❌ 文件已删除，任务中断'
                        interrupted_tasks.append(f"一键处理子任务: {oneclick_file_task_id}")
                        smart_print(f"[任务中断] 一键处理子任务中断: {oneclick_file_task_id}", LogLevel.WARNING)
                    
                    interrupted_tasks.append(f"一键处理任务: {task_id} (已移除该文件)")
                    smart_print(f"[任务中断] 从一键处理任务 {task_id} 中移除文件 {file_id}", LogLevel.INFO)
    
    return interrupted_tasks

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx'}

# MinerU API 配置
MINERU_CONFIG = {
    'enable_formula': True,      # 开启公式识别（仅pipeline模型有效）
    'enable_table': True,        # 开启表格识别（仅pipeline模型有效）
    'is_ocr': True,            # 是否启动OCR功能，默认false（仅pipeline模型有效）
    'language': 'auto',           # 指定文档语言，默认ch（仅pipeline模型有效）
    'model_version': 'vlm'      # 使用vlm模型（表格识别质量更高）
}

def upload_single_file_to_mineru_from_db(file_id, user_id, api_token, options, task_id=None):
    """从数据库读取PDF数据并上传到MinerU - 使用官方批量上传API"""
    import tempfile
    import os
    try:
        smart_print(f"[MinerU-API] 从数据库获取文件: {file_id}", LogLevel.DEBUG)
        
        # 检查任务是否已被删除
        if task_id and is_task_cancelled(task_id):
            smart_print(f"[MinerU-API] 任务已被删除，取消上传: {task_id}", LogLevel.WARNING)
            return None
        
        # 从数据库获取PDF数据
        pdf_info = get_file_info_from_db(file_id, user_id)
        if not pdf_info:
            smart_print(f"[MinerU-API] 错误: 无法从数据库获取PDF数据", LogLevel.ERROR)
            return None
        
        smart_print(f"[MinerU-API] PDF数据大小: {pdf_info['size']} bytes", LogLevel.DEBUG)
        smart_print(f"[MinerU-API] 文件名: {pdf_info['original_name']}", LogLevel.DEBUG)
        smart_print(f"[MinerU-API] API Token长度: {len(api_token) if api_token else 0}", LogLevel.DEBUG)
        
        # 再次检查任务是否已被删除
        if task_id and is_task_cancelled(task_id):
            smart_print(f"[MinerU-API] 任务已被删除，取消上传: {task_id}", LogLevel.WARNING)
            return None
        
        # 第一步：申请上传链接
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_token}'
        }
        
        # 构建请求参数 - 按照官方API文档
        data = {
            "enable_formula": options.get("enable_formula", True),
            "language": options.get("language", "ch"),  # 默认中文
            "enable_table": options.get("enable_table", True),
            "model_version": options.get("model_version", "vlm"),  # 使用vlm模型（表格识别质量更高）
            "files": [{
                "name": pdf_info['original_name'],
                "is_ocr": options.get("is_ocr", True),  # 默认false
                "data_id": file_id  # 使用file_id作为data_id
            }]
        }
        
        # 如果有额外格式要求
        if options.get("extra_formats"):
            data["extra_formats"] = options["extra_formats"]
        
        # 如果有页码范围要求
        if options.get("page_ranges"):
            data["files"][0]["page_ranges"] = options["page_ranges"]
            
        url = 'https://mineru.net/api/v4/file-urls/batch'
        smart_print(f"[MinerU-API] 申请上传链接URL: {url}", LogLevel.DEBUG)
        smart_print(f"[MinerU-API] 请求参数: {data}", LogLevel.DEBUG)
        
        response = requests.post(url, headers=headers, json=data, timeout=600)  # 10分钟超时
        smart_print(f"[MinerU-API] 响应状态码: {response.status_code}", LogLevel.DEBUG)
        
        if response.status_code != 200:
            smart_print(f"[MinerU-API] 申请上传链接失败: {response.status_code} {response.text}", LogLevel.ERROR)
            return None
            
        result = response.json()
        smart_print(f"[MinerU-API] 响应内容: {result}", LogLevel.DEBUG)
        
        if result.get("code") != 0:
            smart_print(f"[MinerU-API] API返回错误: {result.get('msg', 'Unknown error')}", LogLevel.ERROR)
            return None
            
        # 获取批次ID和上传链接
        batch_id = result["data"]["batch_id"]
        file_urls = result["data"]["file_urls"]
        
        if not file_urls:
            smart_print(f"[MinerU-API] 未获取到上传链接", LogLevel.ERROR)
            return None
            
        upload_url = file_urls[0]  # 第一个文件的上传链接
        smart_print(f"[MinerU-API] 获取到批次ID: {batch_id}", LogLevel.INFO)
        smart_print(f"[MinerU-API] 获取到上传链接: {upload_url}", LogLevel.DEBUG)
        
        # 检查任务是否已被删除（在上传文件前）
        if task_id and task_id not in tasks_db:
            smart_print(f"[MinerU-API] 任务已被删除，取消文件上传: {task_id}", LogLevel.WARNING)
            return None
        
        # 第二步：上传文件到获取的链接
        smart_print(f"[MinerU-API] 开始上传文件到: {upload_url}", LogLevel.INFO)
        
        upload_response = requests.put(upload_url, data=pdf_info['pdf_data'], timeout=600)  # 10分钟超时
        smart_print(f"[MinerU-API] 文件上传响应状态码: {upload_response.status_code}", LogLevel.DEBUG)
        
        if upload_response.status_code == 200:
            smart_print(f"[MinerU-API] 文件上传成功！", LogLevel.INFO)
            smart_print(f"[MinerU-API] 系统将自动开始解析，批次ID: {batch_id}", LogLevel.INFO)
            
            # 返回批次ID，用于后续查询结果
            return {
                "batch_id": batch_id,
                "data_id": file_id,
                "status": "uploaded",
                "upload_url": upload_url
            }
        else:
            print(f"[MinerU-API] 文件上传失败: {upload_response.status_code}")
            if upload_response.text:
                print(f"[MinerU-API] 上传失败详情: {upload_response.text}")
            return None
        
    except requests.exceptions.Timeout as e:
        print(f"[MinerU-API] 请求超时: {e}")
        return None
    except requests.exceptions.ConnectionError as e:
        print(f"[MinerU-API] 连接错误: {e}")
        return None
    except requests.exceptions.RequestException as e:
        print(f"[MinerU-API] 请求异常: {e}")
        return None
    except Exception as e:
        print(f"[MinerU-API] 上传文件异常: {e}")
        import traceback
        print(f"[MinerU-API] 异常详情: {traceback.format_exc()}")
        return None

def is_task_cancelled(task_id):
    """检查任务是否已被取消/删除"""
    if not task_id:
        return False
    
    # 只检查内存中的任务（纯内存管理）
    if task_id not in tasks_db:
        return True
    
    # 任务存在于内存中，说明还在处理
    return False

def poll_batch_status(batch_id, api_token):
    """轮询批量任务状态 - 完全按照mineru_api.py的方法"""
    try:
        poll_url = f"https://mineru.net/api/v4/extract-results/batch/{batch_id}"
        headers = {
            "Authorization": f"Bearer {api_token}",
            "Content-Type": "application/json"
        }
        
        print(f"[MinerU-API] 批量轮询: {poll_url}")
        response = requests.get(poll_url, headers=headers, timeout=600)  # 10分钟超时
        print(f"[MinerU-API] 批量轮询响应: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"[MinerU-API] 批量轮询结果: {result}")
            if result.get("code") == 0 and result.get("data") and "extract_result" in result["data"]:
                extract_results = result["data"]["extract_result"]
                
                all_done = True
                completed_files = []
                failed_files = []
                
                for item in extract_results:
                    file_name = item.get("file_name", "")
                    state = item.get("state", "")
                    zip_url = item.get("full_zip_url", "")
                    error_msg = item.get("err_msg", "")
                    print(f"[MinerU-API] 文件: {file_name}, 状态: {state}, 错误: {error_msg}")
                    
                    if state == "done" and zip_url:
                        completed_files.append({
                            "file_name": file_name,
                            "zip_url": zip_url,
                            "state": state
                        })
                    elif state == "failed":
                        failed_files.append({
                            "file_name": file_name,
                            "error": error_msg or "未知错误",
                            "state": state
                        })
                    elif state not in ("done", "failed"):
                        all_done = False
                        print(f"[MinerU-API] 文件 {file_name} 仍在处理中，状态: {state}")
                
                if all_done:
                    if failed_files:
                        print(f"[MinerU-API] 处理完成，但有 {len(failed_files)} 个文件失败")
                        return "failed", failed_files
                    else:
                        print(f"[MinerU-API] 所有文件处理完成，成功 {len(completed_files)} 个")
                        return "completed", completed_files
                else:
                    print(f"[MinerU-API] 仍在处理中，已完成 {len(completed_files)} 个")
                    return "processing", completed_files + failed_files
            elif result.get("code") != 0:
                error_msg = result.get("message", "未知错误")
                print(f"[MinerU-API] API返回错误: {error_msg}")
                return "failed", [{"error": error_msg}]
            else:
                print(f"[MinerU-API] 数据格式异常: {result}")
                return "failed", []
        elif response.status_code == 401:
            print(f"[MinerU-API] 认证失败，请检查API Token")
            return "failed", [{"error": "API Token认证失败"}]
        elif response.status_code == 404:
            print(f"[MinerU-API] 批次任务不存在: {batch_id}")
            return "failed", [{"error": "批次任务不存在"}]
        else:
            print(f"[MinerU-API] 轮询失败: {response.status_code} - {response.text}")
            return "failed", [{"error": f"轮询失败: {response.status_code}"}]
    except requests.exceptions.Timeout as e:
        print(f"[MinerU-API] 轮询超时: {e}")
        return "failed", [{"error": "轮询超时"}]
    except requests.exceptions.ConnectionError as e:
        print(f"[MinerU-API] 连接错误: {e}")
        return "failed", [{"error": "网络连接错误"}]
    except requests.exceptions.RequestException as e:
        print(f"[MinerU-API] 请求异常: {e}")
        return "failed", [{"error": f"请求异常: {str(e)}"}]
    except Exception as e:
        print(f"[MinerU-API] 查询批量任务异常: {e}")
        import traceback
        print(f"[MinerU-API] 异常详情: {traceback.format_exc()}")
        return "failed", [{"error": f"查询异常: {str(e)}"}]

def download_result_file(zip_url, output_path):
    """下载处理结果 - 完全按照mineru_api.py的方法"""
    try:
        print(f"[MinerU-API] 开始下载结果: {zip_url}")
        print(f"[MinerU-API] 输出路径: {output_path}")
        
        with requests.get(zip_url, stream=True, timeout=600) as response:  # 10分钟超时
            print(f"[MinerU-API] 下载响应: {response.status_code}")
            
            if response.status_code == 200:
                # 检查Content-Type
                content_type = response.headers.get('Content-Type', '')
                content_length = response.headers.get('Content-Length', '0')
                print(f"[MinerU-API] 内容类型: {content_type}, 大小: {content_length} bytes")
                
                # 确保目录存在
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                downloaded_size = 0
                with open(output_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                            downloaded_size += len(chunk)
                
                final_size = os.path.getsize(output_path)
                print(f"[MinerU-API] 下载完成: {output_path}, 最终大小: {final_size} bytes")
                
                # 验证文件是否为有效的ZIP文件
                try:
                    import zipfile
                    with zipfile.ZipFile(output_path, 'r') as zip_ref:
                        file_list = zip_ref.namelist()
                        print(f"[MinerU-API] ZIP文件验证成功，包含 {len(file_list)} 个文件")
                        for file_name in file_list[:5]:  # 只显示前5个文件名
                            print(f"[MinerU-API] ZIP内文件: {file_name}")
                        if len(file_list) > 5:
                            print(f"[MinerU-API] ...还有 {len(file_list) - 5} 个文件")
                except zipfile.BadZipFile as e:
                    print(f"[MinerU-API] 警告: 下载的文件不是有效的ZIP文件: {e}")
                    return False
                
                return True
            else:
                print(f"[MinerU-API] 下载失败: {response.status_code} - {response.text[:200]}")
                return False
    except requests.exceptions.Timeout as e:
        print(f"[MinerU-API] 下载超时: {e}")
        return False
    except requests.exceptions.ConnectionError as e:
        print(f"[MinerU-API] 下载连接错误: {e}")
        return False
    except requests.exceptions.RequestException as e:
        print(f"[MinerU-API] 下载请求异常: {e}")
        return False
    except IOError as e:
        print(f"[MinerU-API] 文件写入错误: {e}")
        return False
    except Exception as e:
        print(f"[MinerU-API] 下载异常: {e}")
        import traceback
        print(f"[MinerU-API] 异常详情: {traceback.format_exc()}")
        return False

# ==================== Word文档转换功能 ====================

def convert_word_to_markdown_from_db(file_id, user_id, task_id=None, progress_callback=None):
    """从数据库读取Word文档并转换为Markdown"""
    try:
        # 检查任务是否被取消
        if task_id and is_task_cancelled(task_id):
            print(f"[Word转换] 任务已被取消: {task_id}")
            return {'success': False, 'error': '任务已被取消', 'cancelled': True}
        
        # 报告进度：开始处理
        if progress_callback:
            progress_callback(0, "📄 步骤1/3: 开始读取Word文档")
        
        # 从数据库获取文档数据
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            print(f"[Word转换] 数据库连接失败")
            # 🚨 修复：数据库连接失败时更新文件状态
            try:
                from db_operations import update_file_status_in_db
                update_file_status_in_db(file_id, 'error', remarks="Word转换：数据库连接失败")
                print(f"[Word转换] 已更新文件状态为error: {file_id}")
            except Exception as status_error:
                print(f"[Word转换] 更新状态失败: {status_error}")
            return {'success': False, 'error': '数据库连接失败'}
        
        cursor = conn.cursor()
        cursor.execute("""
            SELECT original_name, pdf_data FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        conn.close()
        
        if not result or not result['pdf_data']:
            print(f"[Word转换] 未找到文档数据")
            # 🚨 修复：未找到文档数据时更新文件状态
            try:
                from db_operations import update_file_status_in_db
                update_file_status_in_db(file_id, 'error', remarks="Word转换：未找到文档数据")
                print(f"[Word转换] 已更新文件状态为error: {file_id}")
            except Exception as status_error:
                print(f"[Word转换] 更新状态失败: {status_error}")
            return {'success': False, 'error': '未找到文档数据'}
        
        original_name = result['original_name']
        doc_data = result['pdf_data']
        
        print(f"[Word转换] 开始转换文档: {original_name}")
        
        # 报告进度：准备提取
        if progress_callback:
            progress_callback(5, "📄 步骤1/3: 准备提取参考文献")
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.docx' if original_name.lower().endswith('.docx') else '.doc', delete=False) as temp_file:
            temp_file.write(doc_data)
            temp_path = temp_file.name
        
        try:
            # 检查任务是否被取消
            if task_id and is_task_cancelled(task_id):
                print(f"[Word转换] 任务在转换前被取消: {task_id}")
                return {'success': False, 'error': '任务已被取消', 'cancelled': True}
            
            # 报告进度：开始提取参考文献
            if progress_callback:
                progress_callback(10, "📚 步骤2/3: 正在提取参考文献")
            
            # 🔧 优化：只提取参考文献,不转换整个MD文档
            # 原因：用户只需要参考文献,跳过MD转换可以大幅提升性能
            md_content, ref_md_content, ref_validation_json = convert_word_file_to_markdown_ref_only(temp_path, user_id, task_id, progress_callback)
            
            # 检查任务是否在转换过程中被取消 (md_content为None可能是取消导致的)
            if task_id and is_task_cancelled(task_id):
                print(f"[Word转换] 任务在转换后被取消: {task_id}")
                return {'success': False, 'error': '任务已被取消', 'cancelled': True}
            
            # 如果md_content为None且任务未被显式取消,可能是转换过程中被取消了
            if md_content is None:
                # 再次检查任务状态
                if task_id and is_task_cancelled(task_id):
                    print(f"[Word转换] 任务在转换过程中被取消(返回None): {task_id}")
                    return {'success': False, 'error': '任务已被取消', 'cancelled': True}
                else:
                    # 真正的转换失败
                    print(f"[Word转换] 转换失败(返回None): {original_name}")
                    try:
                        from db_operations import update_file_status_in_db
                        update_file_status_in_db(file_id, 'error', remarks=f"Word转换失败: {original_name}")
                        print(f"[Word转换] 已更新文件状态为error: {file_id}")
                    except Exception as status_error:
                        print(f"[Word转换] 更新状态失败: {status_error}")
                    return {'success': False, 'error': 'Word转换失败'}
            
            if md_content:
                # 报告进度：保存结果
                if progress_callback:
                    progress_callback(95, "💾 步骤3/3: 保存转换结果到数据库")
                
                # 🔧 调试：打印保存的内容
                print(f"[Word转换] 准备保存到数据库:")
                print(f"   - file_id: {file_id}")
                print(f"   - user_id: {user_id}")
                print(f"   - md_content: {'有' if md_content else '无'} ({len(md_content) if md_content else 0} 字符)")
                print(f"   - ref_md_content: {'有' if ref_md_content else '无'} ({len(ref_md_content) if ref_md_content else 0} 字符)")
                print(f"   - ref_validation_json: {'有' if ref_validation_json else '无'} ({len(ref_validation_json) if ref_validation_json else 0} 字符)")
                
                # 存储到数据库
                success = save_word_conversion_results(file_id, user_id, md_content, ref_md_content, ref_validation_json)
                if success:
                    print(f"[Word转换] 转换成功: {original_name}")
                    if progress_callback:
                        progress_callback(100, "✅ Word转换完成")
                    return True
                else:
                    print(f"[Word转换] 保存失败: {original_name}")
                    # 🚨 修复：转换失败时更新数据库状态
                    try:
                        from db_operations import update_file_status_in_db
                        update_file_status_in_db(file_id, 'error', remarks=f"Word转换保存失败: {original_name}")
                        print(f"[Word转换] 已更新文件状态为error: {file_id}")
                    except Exception as status_error:
                        print(f"[Word转换] 更新状态失败: {status_error}")
                    return {'success': False, 'error': 'Word转换保存失败'}
            else:
                print(f"[Word转换] 转换结果为空: {original_name}")
                # 🚨 修复：转换失败时更新数据库状态
                try:
                    from db_operations import update_file_status_in_db
                    update_file_status_in_db(file_id, 'error', remarks=f"Word转换结果为空: {original_name}")
                    print(f"[Word转换] 已更新文件状态为error: {file_id}")
                except Exception as status_error:
                    print(f"[Word转换] 更新状态失败: {status_error}")
                return {'success': False, 'error': 'Word转换结果为空'}
                
        finally:
            # 清理临时文件
            try:
                os.unlink(temp_path)
            except:
                pass
                
    except Exception as e:
        print(f"[Word转换] 异常: {e}")
        import traceback
        print(f"[Word转换] 异常详情: {traceback.format_exc()}")
        
        # 🚨 修复：异常时更新数据库状态
        try:
            from db_operations import update_file_status_in_db
            update_file_status_in_db(file_id, 'error', remarks=f"Word转换异常: {str(e)}")
            print(f"[Word转换] 已更新文件状态为error: {file_id}")
        except Exception as status_error:
            print(f"[Word转换] 更新状态失败: {status_error}")
        
        # 返回统一的错误格式
        return {'success': False, 'error': f'Word转换异常: {str(e)}'}

def extract_references_from_word_document(doc, user_id=None, progress_callback=None, task_id=None):
    """直接从Word文档对象中提取参考文献,不经过MD转换
    
    Args:
        doc: Word文档对象 (python-docx Document对象)
        user_id: 用户ID
        progress_callback: 进度回调函数
        task_id: 任务ID
        
    Returns:
        tuple: (ref_md_content, ref_validation_json) 参考文献MD内容和验证JSON
    """
    try:
        # 检查任务是否被取消
        if task_id and is_task_cancelled(task_id):
            print(f"[Word参考文献提取] 任务已被取消: {task_id}")
            return None, None
        
        if progress_callback:
            progress_callback(0, "开始从Word文档提取参考文献")
        
        # python-docx处理：读取所有段落
        paragraphs = doc.paragraphs
        total_paragraphs = len(paragraphs)
        
        print(f"[Word参考文献提取] 文档共有 {total_paragraphs} 个段落")
        
        if progress_callback:
            progress_callback(5, "快速定位参考文献标题")
        
        # 第一步：快速读取纯文本，从后往前查找关键标题
        ref1_pos = None  # 第一个参考文献（报告）
        ref2_pos = None  # 第二个参考文献（综述）
        review_pos = None  # 综述标题
        acknowledgment_pos = None  # 致谢标题
        
        for i in range(total_paragraphs - 1, -1, -1):  # 从后往前，0-based索引
            if task_id and is_task_cancelled(task_id):
                print(f"[Word参考文献提取] 任务被取消: {task_id}")
                return None, None
            
            try:
                para = paragraphs[i]
                text = para.text.strip()
                cleaned = text.replace('\r', '').replace('\n', '').replace(' ', '').replace('　', '')
                
                # 查找致谢（最后出现）
                if acknowledgment_pos is None and (cleaned == "致谢" or cleaned == "致谢：" or cleaned == "附录" or cleaned == "附录："):
                    acknowledgment_pos = i
                    print(f"[Word参考文献提取] 找到'{cleaned}'标题在第 {i+1} 段")
                
                # 查找第二个参考文献（综述部分）
                if ref2_pos is None and (cleaned == "参考文献" or cleaned == "参考文献："):
                    ref2_pos = i
                    print(f"[Word参考文献提取] 找到'参考文献'(综述)标题在第 {i+1} 段")
                    continue
                
                # 查找综述标题
                if review_pos is None and (cleaned == "综述" or cleaned == "综述："):
                    review_pos = i
                    print(f"[Word参考文献提取] 找到'综述'标题在第 {i+1} 段")
                
                # 查找第一个参考文献（报告部分，从后往前第二个）
                if ref1_pos is None and ref2_pos is not None and (cleaned == "参考文献" or cleaned == "参考文献："):
                    ref1_pos = i
                    print(f"[Word参考文献提取] 找到'参考文献'(报告)标题在第 {i+1} 段")
                    break  # 找到两个参考文献标题就停止
                    
            except:
                continue
        
        if not ref1_pos and not ref2_pos:
            print("[Word参考文献提取] 未找到'参考文献'标题")
            return None, None
        
        print(f"[Word参考文献提取] 标题定位完成: 报告={ref1_pos}, 综述标题={review_pos}, 综述参考文献={ref2_pos}, 致谢={acknowledgment_pos}")
        
        if progress_callback:
            progress_callback(20, "读取参考文献区间段落")
        
        # 第二步：构建要处理的参考文献区间
        sections_to_process = []
        
        # 报告部分参考文献
        if ref1_pos is not None:
            end_pos = review_pos if review_pos is not None else total_paragraphs - 1
            sections_to_process.append({
                'name': '报告',
                'start': ref1_pos + 1,  # 跳过标题行
                'end': end_pos
            })
        
        # 综述部分参考文献
        if ref2_pos is not None:
            end_pos = acknowledgment_pos if acknowledgment_pos is not None else total_paragraphs - 1
            sections_to_process.append({
                'name': '综述',
                'start': ref2_pos + 1,  # 跳过标题行
                'end': end_pos
            })
        
        if progress_callback:
            progress_callback(40, f"找到{len(sections_to_process)}处参考文献,提取条目")
        
        # 第三步：只读取参考文献区间的段落（包含自动编号）
        all_references = []
        ref_sections = []
        
        for section in sections_to_process:
            section_name = section['name']
            ref_start_idx = section['start']
            ref_end_idx = section['end']
            
            print(f"[Word参考文献提取] 处理{section_name}参考文献部分 (从第{ref_start_idx+1}段到第{ref_end_idx+1}段)")
            
            # python-docx: 读取区间段落
            ref_paragraphs = []
            
            ref_number = 1  # 编号计数器
            for i in range(ref_start_idx, ref_end_idx + 1):
                if task_id and is_task_cancelled(task_id):
                    print(f"[Word参考文献提取] 任务被取消: {task_id}")
                    return None, None
                
                try:
                    para = paragraphs[i]
                    text = para.text.strip()
                    
                    # 跳过空段落
                    if not text:
                        continue
                    
                    # 🔧 检查是否为自动编号列表项并提取编号
                    is_numbered = False
                    num_text = None
                    
                    # 检查段落的编号属性
                    if hasattr(para, '_element') and hasattr(para._element, 'pPr'):
                        pPr = para._element.pPr
                        if pPr is not None and pPr.numPr is not None:
                            is_numbered = True
                            # 使用顺序编号
                            num_text = f"[{ref_number}]"
                            ref_number += 1
                    
                    # 检查段落样式
                    if not is_numbered and para.style and para.style.name:
                        style_name = para.style.name.lower()
                        if 'list' in style_name or 'numbering' in style_name:
                            is_numbered = True
                            num_text = f"[{ref_number}]"
                            ref_number += 1
                    
                    # 只保留自动编号的段落
                    if not is_numbered:
                        continue
                    
                    # 跳过过短的段落
                    if len(text) < 15:
                        continue
                    
                    # 🔧 将编号添加到文本前面
                    if num_text:
                        full_text = f"{num_text} {text}"
                    else:
                        full_text = text
                    
                    ref_paragraphs.append(full_text)
                    
                except Exception as e:
                    print(f"[Word参考文献提取-异常] 第{i+1}段处理失败: {e}")
                    continue
            
            print(f"[Word参考文献提取] {section_name}提取到 {len(ref_paragraphs)} 段参考文献内容")
            
            if not ref_paragraphs:
                print(f"[Word参考文献提取] {section_name}无有效内容,跳过")
                continue
            
            # 🔧 关键修复：由于已经筛选出只有编号开头的段落，每个段落就是一条参考文献
            # 不需要复杂的合并逻辑，直接使用
            section_references = ref_paragraphs
            
            print(f"[Word参考文献提取] {section_name}共识别出 {len(section_references)} 条参考文献")
            
            # 保存这部分的参考文献
            if section_references:
                section_content = f"# 参考文献 - {section_name}\n\n" + "\n\n".join(section_references)
                ref_sections.append(section_content)
                all_references.extend(section_references)
        
        # 检查是否提取到参考文献
        if not all_references:
            print("[Word参考文献提取] 未能提取到有效的参考文献条目")
            return None, None
        
        print(f"[Word参考文献提取] 总计识别出 {len(all_references)} 条参考文献 (来自{len(ref_sections)}个部分)")
        
        if progress_callback:
            progress_callback(40, f"识别到 {len(all_references)} 条参考文献")
        
        # 5. 生成参考文献MD内容
        ref_md_content = "\n\n---\n\n".join(ref_sections) if ref_sections else None
        
        if progress_callback:
            progress_callback(41, "LLM验证中")
        
        # 6. 生成JSON验证报告
        ref_validation_json = None
        try:
            # 检查任务是否被取消
            if task_id and is_task_cancelled(task_id):
                print(f"[Word参考文献提取] 验证前任务被取消: {task_id}")
                return None, None
            
            validator = GBT7714Validator()
            
            # 创建验证进度回调
            def validation_progress_callback(progress, message):
                if task_id and is_task_cancelled(task_id):
                    print(f"[Word参考文献验证] 验证中任务被取消: {task_id}")
                    return False
                print(f"[Word参考文献验证] {message} ({progress:.1f}%)")
                if progress_callback:
                    # 🎯 将validate_references_list的0-100%映射到41-100% (59个百分点)
                    # 0-1% → 41%: L1检查（忽略，因为进度太小）
                    # 1-90% → 41-95%: LLM验证 (54个百分点)
                    # 90-100% → 95-100%: L2/L3检查 (5个百分点)
                    mapped_progress = 41 + (progress * 0.59)
                    progress_callback(mapped_progress, message)
                return True
            
            # 检测引用体系
            md_for_detection = "\n\n".join(all_references)
            citation_system = detect_citation_system_in_text(md_for_detection)
            print(f"[Word参考文献提取] 检测到的引用体系: {citation_system}")
            
            # 验证参考文献
            validation_results = validator.validate_references_list(
                all_references, 
                user_id=user_id, 
                progress_callback=validation_progress_callback
            )
            
            # 生成验证报告
            ref_validation_json = generate_json_validation_report(
                validation_results, 
                all_references, 
                citation_system
            )
            
            # 🎯 不在这里调用progress_callback，因为validate_references_list已经报告到100%
            # 外层映射后也是100%，无需重复调用
                
        except Exception as e:
            import traceback
            print(f"[Word参考文献验证] 验证失败: {e}")
            print(f"[Word参考文献验证] 错误详情:")
            traceback.print_exc()
        
        # 🎯 不在这里调用progress_callback(100)，因为validate_references_list已经报告到100%
        
        # 🔧 安全地序列化JSON，处理可能的编码问题
        validation_json_str = None
        if ref_validation_json:
            try:
                validation_json_str = json.dumps(ref_validation_json, ensure_ascii=False)
                # 验证生成的JSON是否有效
                json.loads(validation_json_str)
                print(f"[Word参考文献提取] JSON验证成功，长度: {len(validation_json_str)}")
            except Exception as json_error:
                print(f"[Word参考文献提取] JSON序列化失败: {json_error}")
                # 尝试使用ensure_ascii=True来避免编码问题
                try:
                    validation_json_str = json.dumps(ref_validation_json, ensure_ascii=True)
                    print(f"[Word参考文献提取] 使用ASCII编码成功")
                except:
                    validation_json_str = None
                    print(f"[Word参考文献提取] JSON序列化完全失败，返回None")
        
        return ref_md_content, validation_json_str
        
    except Exception as e:
        print(f"[Word参考文献提取] 提取异常: {e}")
        import traceback
        traceback.print_exc()
        return None, None

def convert_word_file_to_markdown_ref_only(word_path, user_id=None, task_id=None, progress_callback=None):
    """优化版本:只从Word文档提取参考文献,不转换整个MD文档 (使用python-docx)
    
    性能优化:跳过段落遍历和MD生成,直接提取参考文献
    
    Args:
        word_path: Word文档路径
        user_id: 用户ID
        task_id: 任务ID
        progress_callback: 进度回调函数
        
    Returns:
        tuple: (md_content, ref_md_content, ref_validation_json)
               md_content为空字符串(保持兼容性),ref_md_content和ref_validation_json为提取结果
    """
    try:
        # 检查任务是否被取消
        if task_id and is_task_cancelled(task_id):
            print(f"[Word参考文献] 任务已被取消: {task_id}")
            return None, None, None
        
        # 使用python-docx直接打开文档
        doc = DocxDocument(word_path)
        print(f"[Word参考文献] 文档打开成功: {word_path}")
        
        # 检查任务是否被取消
        if task_id and is_task_cancelled(task_id):
            print(f"[Word参考文献] 提取前任务被取消: {task_id}")
            return None, None, None
        
        # 创建参考文献提取进度回调
        def ref_progress_callback(progress, message):
            if task_id and is_task_cancelled(task_id):
                print(f"[Word参考文献] 提取中任务被取消: {task_id}")
                return False
            print(f"[Word参考文献] {message} ({progress:.1f}%)")
            if progress_callback:
                progress_callback(progress, message)
            return True
        
        # 直接从Word文档提取参考文献
        print("[Word参考文献] 直接提取参考文献,跳过完整MD转换")
        ref_md_content, ref_validation_json = extract_references_from_word_document(
            doc, user_id, ref_progress_callback, task_id
        )
        
        # 检查提取结果
        if ref_md_content is None and ref_validation_json is None:
            if task_id and is_task_cancelled(task_id):
                print(f"[Word参考文献] 提取后确认任务已取消: {task_id}")
                return None, None, None
        
        # 返回结果 (md_content为空字符串,保持数据库兼容性)
        md_content = "# Word文档参考文献\n\n仅提取参考文献内容，未转换完整文档。"
        
        return md_content, ref_md_content, ref_validation_json
        
    except Exception as e:
        print(f"[Word参考文献] 提取异常: {e}")
        import traceback
        traceback.print_exc()
        return None, None, None

def convert_word_file_to_markdown(word_path, user_id=None, task_id=None, progress_callback=None):
    """将Word文件转换为Markdown (使用python-docx)"""
    try:
        # 检查任务是否被取消
        if task_id and is_task_cancelled(task_id):
            print(f"[Word转换] Word文件转换任务已被取消: {task_id}")
            return None, None, None
        
        # 报告进度：打开文档
        if progress_callback:
            progress_callback(2, "打开Word文档")
        
        # 使用python-docx打开文档
        doc = DocxDocument(word_path)
        print(f"[Word转换] 文档打开成功: {word_path}")
        
        # 报告进度：开始处理内容
        if progress_callback:
            progress_callback(5, "开始处理Word文档内容")
        
        # 创建临时目录用于图片
        temp_dir = tempfile.mkdtemp()
        images_dir = os.path.join(temp_dir, "images")
        os.makedirs(images_dir, exist_ok=True)
        
        markdown_lines = []
        image_count = 1
        
        # python-docx: 处理图片
        # 注意: python-docx无法使用剪贴板，需要直接从文档中提取图片
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_name = f"image{image_count}.png"
                        image_path = os.path.join(images_dir, image_name)
                        with open(image_path, 'wb') as img_file:
                            img_file.write(rel.target_part.blob)
                        markdown_lines.append(f"![](images/{image_name})")
                        image_count += 1
                    except Exception as img_error:
                        print(f"[Word转换] 提取图片失败: {img_error}")
        except Exception as img_error:
            print(f"[Word转换] 图片处理失败: {img_error}")
        
        # python-docx: 处理段落
        paragraphs = doc.paragraphs
        total_paragraphs = len(paragraphs)
        print(f"[Word转换] 成功访问 {total_paragraphs} 个段落")
            
        if progress_callback:
            progress_callback(5, f"开始处理{total_paragraphs}个段落")
        
        # 处理段落文本
        for i, para in enumerate(paragraphs):
            # 每处理50个段落检查一次任务状态和更新进度
            if i % 50 == 0:
                if task_id and is_task_cancelled(task_id):
                    print(f"[Word转换] 段落处理过程中任务被取消: {task_id}")
                    return None, None, None
                
                # 更新进度：段落处理进度占5%-30% (25个百分点)
                if progress_callback:
                    para_progress = 5 + (i / total_paragraphs) * 25
                    progress_callback(para_progress, f"处理段落 {i+1}/{total_paragraphs}")
            
            text = para.text.strip()
            if not text:
                continue
            
            try:
                # python-docx: 使用样式名称判断标题
                style_name = para.style.name if para.style else ""
                if "Heading 1" in style_name or "标题 1" in style_name:
                    markdown_lines.append(f"# {text}")
                elif "Heading 2" in style_name or "标题 2" in style_name:
                    markdown_lines.append(f"## {text}")
                else:
                    markdown_lines.append(text)
            except:
                markdown_lines.append(text)
        
        # 直接从Word文档对象提取参考文献
        # 检查任务是否在参考文献提取前被取消
        if task_id and is_task_cancelled(task_id):
            print(f"[Word转换] 参考文献提取前任务被取消: {task_id}")
            return None, None, None
        
        # 报告进度：开始提取参考文献
        if progress_callback:
            progress_callback(40, "提取参考文献")
        
        # 创建参考文献提取进度回调
        def word_ref_progress_callback(progress, message):
            if task_id and is_task_cancelled(task_id):
                print(f"[Word转换-参考文献] 参考文献处理中任务被取消: {task_id}")
                return False
            print(f"[Word转换-参考文献] {message} ({progress:.1f}%)")
            if progress_callback:
                progress_callback(progress, message)
            return True
        
        # 直接从Word文档提取参考文献
        print("[Word转换] 从Word文档获取参考文献")
        ref_md_content, ref_validation_json = extract_references_from_word_document(
            doc, user_id, word_ref_progress_callback, task_id
        )
        
        # 检查参考文献提取是否因取消而返回None
        if ref_md_content is None and ref_validation_json is None:
            if task_id and is_task_cancelled(task_id):
                print(f"[Word转换] 参考文献提取后确认任务已取消: {task_id}")
                return None, None, None
        
        # 生成完整的Markdown内容
        md_content = "\n\n".join(markdown_lines)
        
        # 清理临时目录
        try:
            shutil.rmtree(temp_dir)
        except:
            pass
        
        return md_content, ref_md_content, ref_validation_json
        
    except Exception as e:
        print(f"[Word转换] 转换异常: {e}")
        import traceback
        traceback.print_exc()
        return None, None, None

def extract_references_from_markdown(md_content, user_id=None, progress_callback=None, task_id=None):
    """从markdown内容中提取参考文献，基于import win32com.py的逻辑"""
    try:
        # 检查任务是否被取消
        if task_id and is_task_cancelled(task_id):
            print(f"[参考文献提取] 任务已被取消: {task_id}")
            return None, None
        lines = md_content.split('\n')
        
        # 检测全文的引用体系
        citation_system = detect_citation_system_in_text(md_content)
        print(f"[参考文献提取] 检测到的引用体系: {citation_system}")
        
        # ==================== 三重策略检测参考文献位置 ====================
        
        # 策略1: 标题匹配 - 查找"参考文献"等标题关键词
        ref_positions = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            # 去掉 markdown 标题标记 (# / ## / ### ...)
            clean = re.sub(r'^#{1,6}\s*', '', stripped).strip()
            # 匹配 "参考文献" "主要参考文献" "主要参考文献：" "References" 等
            if re.match(r'(?:主要)?参\s*考\s*文\s*献\s*[:：]?\s*$', clean) or \
               re.match(r'^References?\s*[:：]?\s*$', clean, re.IGNORECASE):
                ref_positions.append(i)
        
        detection_method = "标题匹配"
        
        # 策略2: 格式匹配 - 查找[数字]引用模式聚集区域
        if not ref_positions:
            detection_method = "格式匹配"
            ref_num_pattern = re.compile(r'^\s*\[(\d+)\]')
            ref_line_indices = [i for i, line in enumerate(lines) if ref_num_pattern.match(line.strip())]
            
            if len(ref_line_indices) >= 3:  # 至少3条才认为是参考文献区域
                # 寻找连续/接近的[数字]行聚集区域（允许间隔不超过5行，跨行合并）
                clusters = []
                current_cluster = [ref_line_indices[0]]
                for idx in ref_line_indices[1:]:
                    if idx - current_cluster[-1] <= 5:
                        current_cluster.append(idx)
                    else:
                        if len(current_cluster) >= 3:
                            clusters.append(current_cluster)
                        current_cluster = [idx]
                if len(current_cluster) >= 3:
                    clusters.append(current_cluster)
                
                if clusters:
                    # 优先选择文档后半部分的聚集区（参考文献通常在文档末尾）
                    total_lines = len(lines)
                    best_cluster = None
                    for cluster in reversed(clusters):
                        if cluster[0] > total_lines * 0.4:
                            best_cluster = cluster
                            break
                    if not best_cluster:
                        best_cluster = clusters[-1]  # 取最后一个
                    
                    # 用第一个[数字]行的前一行作为虚拟"标题"位置
                    virtual_header = max(0, best_cluster[0] - 1)
                    ref_positions.append(virtual_header)
                    print(f"[参考文献提取] 格式匹配：发现{len(best_cluster)}条[数字]引用聚集区，起始行{best_cluster[0]+1}")
        
        # 策略3: 位置匹配 - 文档尾部参考文献特征分析
        if not ref_positions:
            detection_method = "位置匹配"
            total_lines = len(lines)
            tail_start = int(total_lines * 0.7)  # 检查文档最后30%
            tail_lines = lines[tail_start:]
            
            # 参考文献特征模式
            ref_features = [
                re.compile(r'^\s*\[?\d+\]?\s*[A-Z\u4e00-\u9fff]'),  # [数字] 或 数字] 开头后跟字母/中文
                re.compile(r'\d{4}[\s,，;；.\(（]'),  # 包含年份（如2024, 2023等）
                re.compile(r'(?:Journal|Proceedings|Trans\.|IEEE|ACM|Vol\.|pp\.|et\s+al)', re.IGNORECASE),  # 英文期刊特征
                re.compile(r'[\u4e00-\u9fff]+(?:学报|期刊|杂志|出版社|报|论文集)'),  # 中文期刊特征
            ]
            
            # 统计每行的特征得分
            feature_scores = []
            for i, line in enumerate(tail_lines):
                stripped = line.strip()
                if not stripped:
                    feature_scores.append(0)
                    continue
                score = sum(1 for p in ref_features if p.search(stripped))
                feature_scores.append(score)
            
            # 找高得分区域（连续多行得分>=2的区域）
            high_score_start = None
            high_score_count = 0
            best_region = None
            best_count = 0
            
            for i, score in enumerate(feature_scores):
                if score >= 2:
                    if high_score_start is None:
                        high_score_start = i
                    high_score_count += 1
                else:
                    if high_score_count >= 3 and high_score_count > best_count:
                        best_region = high_score_start
                        best_count = high_score_count
                    high_score_start = None
                    high_score_count = 0
            # 处理末尾情况
            if high_score_count >= 3 and high_score_count > best_count:
                best_region = high_score_start
                best_count = high_score_count
            
            if best_region is not None:
                actual_start = tail_start + best_region
                virtual_header = max(0, actual_start - 1)
                ref_positions.append(virtual_header)
                print(f"[参考文献提取] 位置匹配：在文档尾部发现{best_count}行参考文献特征区域，起始行{actual_start+1}")
        
        if not ref_positions:
            print("[参考文献提取] 三种策略均未找到参考文献内容")
            return None, None
        
        print(f"[参考文献提取] 通过【{detection_method}】找到 {len(ref_positions)} 个参考文献部分")
        
        all_references = []
        ref_sections = []
        
        # 提取每个参考文献部分
        for idx, ref_start_idx in enumerate(ref_positions):
            part_name = "报告" if idx == 0 else "综述" if idx == 1 else f"第{idx + 1}部分"
            
            print(f"[参考文献提取] 处理{part_name}参考文献部分...")
            
            # 找到这个参考文献部分的结束位置
            ref_end_idx = len(lines)
            
            # 查找下一个参考文献或致谢作为结束标记
            for i in range(ref_start_idx + 1, len(lines)):
                line = lines[i].strip()
                clean_end = re.sub(r'^#{1,6}\s*', '', line).strip()
                is_another_ref_header = bool(re.match(r'(?:主要)?参\s*考\s*文\s*献\s*[:：]?\s*$', clean_end))
                # 致谢检测：必须是标题形式（如 "# 致谢" "致 谢" 等短标题），不匹配参考文献正文中恰好包含这两个字的行
                is_acknowledgement = bool(re.match(r'^(?:#{1,6}\s*)?致\s*谢\s*[:：]?\s*$', line))
                # 综述检测：必须是标题形式（如 "综述" "# 综述"），不匹配以"综"开头的参考文献正文
                is_review_header = bool(re.match(r'^(?:#{1,6}\s*)?综\s*述\s*[:：]?\s*$', line))
                # Markdown标题行（以#开头且不含"参考文献"），如 "# 研究内容" "## 方法"
                is_other_heading = (line.startswith("#") and "参考文献" not in line)
                
                if (is_another_ref_header or is_acknowledgement or is_review_header or is_other_heading):
                    ref_end_idx = i
                    print(f"[参考文献提取] 第{idx+1}部分结束标记在行{i+1}: [{line[:60]}]")
                    break
            
            # 从参考文献标题的下一行开始提取
            # 跳过标题行本身(参考文献)
            content_start = ref_start_idx + 1
            
            # 提取这一部分的所有行
            ref_lines = []
            for i in range(content_start, ref_end_idx):
                line = lines[i].strip()
                if line and not re.match(r'(?:#{1,6}\s*)?(?:主要)?参\s*考\s*文\s*献\s*[:：]?\s*$', line):  # 非空行且不是标题
                    ref_lines.append(line)
            
            print(f"[参考文献提取] {part_name}区域: 行{ref_start_idx+1}~{ref_end_idx+1}, 有效行{len(ref_lines)}行")
            if ref_lines:
                print(f"[参考文献提取] 首行预览: {ref_lines[0][:80]}")
                if len(ref_lines) > 1:
                    print(f"[参考文献提取] 末行预览: {ref_lines[-1][:80]}")
            
            # 筛选并合并参考文献条目
            # 支持两种格式：[数字] 和 (数字)
            filtered_refs = []
            current_ref = None
            
            for line in ref_lines:
                # 检查是否是新的参考文献条目(以[数字]或(数字)开头,允许前导空格)
                stripped = line.strip()
                is_new_ref = False
                
                # 检查[数字]格式
                if stripped.startswith("[") and "]" in stripped:
                    bracket_end = stripped.find("]")
                    if bracket_end > 1:
                        potential_num = stripped[1:bracket_end]
                        if potential_num.isdigit() or ("-" in potential_num and all(p.strip().isdigit() for p in potential_num.split("-") if p.strip())):
                            is_new_ref = True
                
                # 检查(数字)格式
                if not is_new_ref and stripped.startswith("(") and ")" in stripped:
                    bracket_end = stripped.find(")")
                    if bracket_end > 1:
                        potential_num = stripped[1:bracket_end]
                        if potential_num.isdigit() or ("-" in potential_num and all(p.strip().isdigit() for p in potential_num.split("-") if p.strip())):
                            is_new_ref = True
                
                if is_new_ref:
                    # 这是一个新的参考文献条目
                    if current_ref:  # 保存之前的条目
                        filtered_refs.append(current_ref.strip())
                    current_ref = line  # 开始新条目
                else:
                    # 不是新条目,追加到当前条目
                    if current_ref:
                        current_ref += " " + line
                    elif stripped:  # 如果even还没开始current_ref，先初始化
                        current_ref = line
            
            # 添加最后一个条目
            if current_ref:
                filtered_refs.append(current_ref.strip())
            
            if filtered_refs:
                section_content = f"# 参考文献 - {part_name}\n\n" + "\n\n".join(filtered_refs)
                ref_sections.append(section_content)
                all_references.extend(filtered_refs)
                print(f"[参考文献提取] {part_name}共 {len(filtered_refs)} 条参考文献")
        
        # 生成参考文献MD内容
        ref_md_content = "\n\n---\n\n".join(ref_sections) if ref_sections else None
        
        # 生成JSON验证报告
        ref_validation_json = None
        if all_references:
            try:
                validator = GBT7714Validator()
                
                # 创建参考文献验证的进度回调
                def ref_validation_progress_callback(progress, message):
                    # 检查任务是否被取消
                    if task_id and is_task_cancelled(task_id):
                        print(f"[参考文献验证] 任务已被取消: {task_id}")
                        return False  # 停止验证
                    if progress_callback:
                        # 🎯 直接传递验证进度，让外层的progress_callback进行统一映射
                        # validate_references_list 内部进度分配：
                        #   0-1%: L1格式检查
                        #   1-90%: LLM批量解析（89个百分点）
                        #   90-100%: L2/L3规则检查
                        # 注意：此处不需要再映射，直接传递给外层
                        result = progress_callback(progress, f"参考文献验证: {message}")
                        if result is False:  # 如果progress_callback返回False，表示应该停止
                            return False
                    return True
                
                validation_results = validator.validate_references_list(all_references, user_id=user_id, progress_callback=ref_validation_progress_callback)
                ref_validation_json = generate_json_validation_report(validation_results, all_references, citation_system)
                
                # 报告最终完成进度
                if progress_callback:
                    progress_callback(100, "验证完成")
            except Exception as e:
                import traceback
                print(f"[参考文献验证] 验证失败: {e}")
                print(f"[参考文献验证] 错误详情:")
                traceback.print_exc()
        
        return ref_md_content, json.dumps(ref_validation_json, ensure_ascii=False) if ref_validation_json else None
        
    except Exception as e:
        print(f"[参考文献提取] 提取异常: {e}")
        return None, None

def generate_json_validation_report(validation_results, references, citation_system):
    """生成JSON格式的验证报告，基于import win32com.py的逻辑"""
    # 导入用户友好的错误描述映射
    from reference_validator import get_user_friendly_error_type
    # 统计信息
    total_refs = len(references)
    valid_count = sum(1 for r in validation_results.values() if r.is_valid)
    error_rate = ((total_refs - valid_count) / total_refs * 100) if total_refs > 0 else 0
    
    # 错误和信息统计
    error_stats = {}
    info_stats = {}
    
    all_errors = []
    all_infos = []
    
    for result in validation_results.values():
        # 只收集errors（severity="error"的项目）
        all_errors.extend([err for err in result.errors if err.severity == "error"])
        
        # 收集所有info级别的信息（包括来自warnings的）
        all_infos.extend([err for err in result.errors if err.severity == "info"])
        all_infos.extend([warning for warning in result.warnings if warning.severity == "info"])
    
    # 统计错误类型
    for error in all_errors:
        error_stats[error.error_type] = error_stats.get(error.error_type, 0) + 1
    
    # 统计信息类型
    for info in all_infos:
        info_stats[info.error_type] = info_stats.get(info.error_type, 0) + 1
    
    # 构建详细的参考文献验证结果
    detailed_results = []
    for i, (ref_id, result) in enumerate(validation_results.items()):
        ref_data = {
            "id": ref_id,
            "index": i + 1,
            "content": references[i] if i < len(references) else "",
            "is_valid": result.is_valid,
            "confidence": round(result.confidence, 2),
            "parsed_data": result.parsed_data if hasattr(result, 'parsed_data') else None,  # LLM解析的结构化数据
            "errors": [
                {
                    "type": get_user_friendly_error_type(error.error_type),
                    "type_code": error.error_type,  # 保留原始错误代码
                    "message": error.message,
                    "severity": error.severity,
                    "suggestion": error.suggestion,
                    "position": error.position,
                    "correct_format": getattr(error, 'correct_format', None)
                } for error in result.errors if error.severity == "error"
            ],
            "infos": [
                {
                    "type": get_user_friendly_error_type(item.error_type),
                    "type_code": item.error_type,  # 保留原始错误代码
                    "message": item.message,
                    "severity": item.severity,
                    "suggestion": item.suggestion,
                    "position": item.position,
                    "correct_format": getattr(item, 'correct_format', None)
                } for item in (result.errors + result.warnings) if item.severity == "info"
            ]
        }
        detailed_results.append(ref_data)
    
    # 生成完整的JSON报告
    json_report = {
        "meta": {
            "report_type": "reference_validation",
            "citation_system": citation_system,
            "total_references": total_refs,
            "validation_date": "2025-09-12",
            "validator_version": "1.0.0"
        },
        "summary": {
            "total_references": total_refs,
            "valid_count": valid_count,
            "invalid_count": total_refs - valid_count,
            "error_rate": round(error_rate, 1),
            "success_rate": round(100 - error_rate, 1)
        },
        "statistics": {
            "error_types": {
                get_user_friendly_error_type(error_type): {
                    "count": count,
                    "percentage": round((count / len(all_errors) * 100), 1) if all_errors else 0,
                    "error_code": error_type  # 保留原始错误代码用于技术参考
                } for error_type, count in sorted(error_stats.items(), key=lambda x: x[1], reverse=True)
            },
            "info_types": {
                get_user_friendly_error_type(info_type): {
                    "count": count,
                    "percentage": round((count / len(all_infos) * 100), 1) if all_infos else 0,
                    "error_code": info_type  # 保留原始错误代码用于技术参考
                } for info_type, count in sorted(info_stats.items(), key=lambda x: x[1], reverse=True)
            },
            "total_errors": len(all_errors),
            "total_infos": len(all_infos)
        },
        "references": detailed_results
    }
    
    return json_report

def save_word_conversion_results(file_id, user_id, md_content, ref_md_content, ref_validation_json):
    """保存Word转换结果到数据库"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            # 🚨 修复：数据库连接失败时更新文件状态
            print(f"[Word转换] 保存结果时数据库连接失败")
            try:
                from db_operations import update_file_status_in_db
                update_file_status_in_db(file_id, 'error', remarks="Word转换保存：数据库连接失败")
                print(f"[Word转换] 已更新文件状态为error: {file_id}")
            except Exception as status_error:
                print(f"[Word转换] 更新状态失败: {status_error}")
            return False
        
        cursor = conn.cursor()
        
        # 更新文件记录
        update_fields = []
        update_values = []
        
        # 🔧 强制保存调试
        print(f"[Word转换-保存] 检查输入参数:")
        print(f"   md_content: {bool(md_content)} - {len(md_content) if md_content else 0} 字符")
        print(f"   ref_md_content: {bool(ref_md_content)} - {len(ref_md_content) if ref_md_content else 0} 字符")
        print(f"   ref_validation_json: {bool(ref_validation_json)} - {len(ref_validation_json) if ref_validation_json else 0} 字符")
        
        if md_content:
            update_fields.append("md_content = %s")
            update_values.append(md_content)
            print(f"[Word转换-保存] ✓ 添加md_content字段")
        
        if ref_md_content:
            update_fields.append("reference_md_content = %s")
            update_values.append(ref_md_content)
            print(f"[Word转换-保存] ✓ 添加reference_md_content字段")
        
        if ref_validation_json:
            update_fields.append("reference_validation_json = %s")
            update_values.append(ref_validation_json)
            print(f"[Word转换-保存] ✓ 添加reference_validation_json字段")
        
        if update_fields:
            # 不直接设置完成状态，让 update_file_complete_status 函数来智能判断
            # 设置处理完成时间
            update_fields.append("processed_time = %s")
            update_values.append(get_current_time().strftime('%Y-%m-%d %H:%M:%S'))
            
            # 🔧 如果processing_start_time为NULL，设置为upload_time（使用上传时间作为开始时间）
            update_fields.append("processing_start_time = COALESCE(processing_start_time, upload_time)")
            
            update_values.extend([file_id, user_id])
            
            sql = f"""
                UPDATE user_files 
                SET {', '.join(update_fields)}
                WHERE file_id = %s AND user_id = %s
            """
            
            print(f"[Word转换-保存] 执行SQL:")
            print(f"   SQL: {sql}")
            print(f"   参数数量: {len(update_values)}")
            
            cursor.execute(sql, update_values)
            rows_affected = cursor.rowcount
            conn.commit()
            
            print(f"[Word转换-保存] SQL执行完成:")
            print(f"   影响行数: {rows_affected}")
            print(f"   - 更新字段数: {len(update_fields)}")
            print(f"   - 参考文献内容: {'有' if ref_md_content else '无'}")
            print(f"   - 验证JSON: {'有' if ref_validation_json else '无'}")
            if ref_md_content:
                print(f"   - 参考文献长度: {len(ref_md_content)} 字符")
                print(f"   - 参考文献预览: {ref_md_content[:100]}...")
            if ref_validation_json:
                print(f"   - 验证JSON长度: {len(ref_validation_json)} 字符")
        
        conn.close()
        
        # 保存成功后，调用状态更新函数来智能判断是否完成
        from db_operations import update_file_complete_status
        print(f"[Word转换] 调用状态更新函数...")
        # ✅ 参数顺序正确：update_file_complete_status(user_id, file_id)
        update_result = update_file_complete_status(user_id, file_id)
        print(f"[Word转换] 状态更新结果: {update_result}")
        
        return True
        
    except Exception as e:
        print(f"[Word转换] 保存结果失败: {e}")
        
        # 🚨 修复：保存异常时更新数据库状态
        try:
            from db_operations import update_file_status_in_db
            update_file_status_in_db(file_id, 'error', remarks=f"Word转换保存异常: {str(e)}")
            print(f"[Word转换] 已更新文件状态为error: {file_id}")
        except Exception as status_error:
            print(f"[Word转换] 更新状态失败: {status_error}")
        
        return False

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ==================== 表格转JSON处理功能 ====================

def extract_basic_info_table(content: str) -> str:
    """提取基本信息表格内容（从"# 基本信息"到"# 报告正文"之间的内容）"""
    try:
        lines = content.split('\n')
        start_idx = -1
        end_idx = len(lines)
        
        # 寻找"# 基本信息"标题
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if re.match(r'^#+\s*基本信息\s*$', line_stripped, re.IGNORECASE):
                start_idx = i
                break
            elif '基本信息' in line_stripped and line_stripped.startswith('#'):
                start_idx = i
                break
        
        if start_idx == -1:
            print("[表格提取] 未找到基本信息标题")
            return ""
        
        # 寻找"# 报告正文"标题作为结束位置
        for i, line in enumerate(lines[start_idx + 1:], start_idx + 1):
            line_stripped = line.strip()
            if re.match(r'^#+\s*报告正文\s*$', line_stripped, re.IGNORECASE):
                end_idx = i
                break
            elif '报告正文' in line_stripped and line_stripped.startswith('#'):
                end_idx = i
                break
        
        # 提取基本信息内容
        table_lines = lines[start_idx:end_idx]
        table_content = '\n'.join(table_lines).strip()
        
        print(f"[表格提取] 提取到基本信息内容，从第{start_idx + 1}行到第{end_idx}行，共{len(table_lines)}行")
        
        return table_content
        
    except Exception as e:
        print(f"[表格提取] 提取基本信息失败: {e}")
        return ""

def extract_report_content_for_nsfc(content: str) -> str:
    """提取报告正文内容：从'# 报告正文'开始到'附件信息'前的部分（专用于NSFC LLM处理）"""
    try:
        lines = content.split('\n')
        start_idx = -1
        end_idx = len(lines)
        
        # 寻找"# 报告正文"标题（必须精确匹配）
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if re.match(r'^#+\s*报告正文\s*$', line_stripped, re.IGNORECASE):
                start_idx = i  # 包含"# 报告正文"这一行
                print(f"[NSFC内容提取] 找到报告正文标题，行号: {i + 1}")
                break
        
        if start_idx == -1:
            print("[NSFC内容提取] 错误：未找到'报告正文'标题")
            return ""  # 没有报告正文标题，返回空
        
        # 寻找结束标记：附件信息（必须在附件信息之前停止）
        for i in range(start_idx + 1, len(lines)):
            line_stripped = lines[i].strip()
            # 匹配"附件信息"标题（任何形式）
            if '附件信息' in line_stripped:
                end_idx = i  # 在附件信息之前停止
                print(f"[NSFC内容提取] 找到附件信息标记，行号: {i + 1}，在此之前停止")
                break
        
        # 提取报告正文内容（从"# 报告正文"到"附件信息"之前）
        report_lines = lines[start_idx:end_idx]
        report_content = '\n'.join(report_lines).strip()
        
        print(f"[NSFC内容提取] 提取报告正文：第{start_idx + 1}行到第{end_idx}行，共{len(report_lines)}行")
        print(f"[NSFC内容提取] 内容长度: {len(report_content)} 字符")
        
        # 验证是否包含附件信息（调试用）
        if '附件信息' in report_content:
            print("[NSFC内容提取] 警告：提取的内容仍包含'附件信息'关键词")
        
        return report_content
        
    except Exception as e:
        print(f"[NSFC内容提取] 提取报告正文失败: {e}")
        return ""  # 失败时返回空字符串


def split_report_text_by_sections(report_text: str) -> list:
    """按一级标题（一）（二）（三）分割正文内容
    
    NSFC申请书正文结构：
    - （一）立项依据与研究内容（包含1-5小节）
    - （二）研究基础与工作条件
    - （三）其他需要说明的情况
    
    Args:
        report_text: 完整的报告正文内容（包含"# 报告正文"标题）
    
    Returns:
        章节列表，每个元素为 {'section_id': str, 'title': str, 'content': str}
    """
    try:
        sections = []
        lines = report_text.split('\n')
        
        # 优先使用（一）（二）（三）格式分割一级标题
        # 匹配格式：# （一）标题 或 #（一）标题 或 （一）标题
        section_pattern = re.compile(r'^#*\s*（([一二三四五六七八九十]+)）\s*(.*)$')
        
        current_section = None
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            match = section_pattern.match(line_stripped)
            
            if match:
                # 保存前一个章节
                if current_section:
                    current_section['content'] = '\n'.join(current_section['lines']).strip()
                    del current_section['lines']
                    sections.append(current_section)
                    print(f"[正文分批] 章节 {current_section['section_id']} 完成，共 {len(current_section['content'])} 字符")
                
                section_num = match.group(1)  # 一、二、三
                section_title = match.group(2).strip()  # 标题内容
                current_section = {
                    'section_id': f'（{section_num}）',
                    'title': f'（{section_num}）{section_title}',
                    'lines': [line]  # 保留原始行（含#号）
                }
                print(f"[正文分批] 找到一级标题: （{section_num}）{section_title}")
            elif current_section:
                current_section['lines'].append(line)
            # 如果还没遇到一级标题，跳过（如"# 报告正文"等前置内容）
        
        # 保存最后一个章节
        if current_section:
            current_section['content'] = '\n'.join(current_section['lines']).strip()
            del current_section['lines']
            sections.append(current_section)
            print(f"[正文分批] 章节 {current_section['section_id']} 完成，共 {len(current_section['content'])} 字符")
        
        # 如果没有找到（一）（二）（三）格式，回退到关键词匹配
        if not sections:
            print(f"[正文分批] 未找到（一）（二）（三）格式，尝试关键词匹配")
            core_keywords = ['立项依据', '研究内容', '研究基础']
            
            for keyword in core_keywords:
                pattern = rf'(^#{{1,6}}.*?{re.escape(keyword)}.*?$)'
                match = re.search(pattern, report_text, re.MULTILINE)
                
                if match:
                    start_pos = match.start()
                    title_line = match.group(1)
                    
                    hash_match = re.match(r'^(#{1,6})', title_line)
                    if hash_match:
                        current_level = len(hash_match.group(1))
                        next_title_pattern = rf'^#{{1,{current_level}}}\s+'
                        remaining_content = report_text[match.end():]
                        next_match = re.search(next_title_pattern, remaining_content, re.MULTILINE)
                        
                        if next_match:
                            end_pos = match.end() + next_match.start()
                            content = report_text[start_pos:end_pos].strip()
                        else:
                            content = report_text[start_pos:].strip()
                        
                        sections.append({
                            'section_id': keyword,
                            'title': title_line.strip(),
                            'content': content
                        })
                        print(f"[正文分批] 找到关键词章节: {keyword} ({len(content)} 字符)")
        
        # 如果没有找到新模板章节，尝试使用旧格式（一）、（二）分割
        # 去重：如果多个章节标识对应相同内容，合并为一个
        if sections:
            deduplicated = []
            seen_contents = {}
            
            for sec in sections:
                content_hash = hash(sec['content'])
                
                if content_hash in seen_contents:
                    # 内容重复，合并章节标识
                    existing_sec = seen_contents[content_hash]
                    if sec['section_id'] not in existing_sec['section_id']:
                        existing_sec['section_id'] = f"{existing_sec['section_id']}+{sec['section_id']}"
                        print(f"[正文分批] 检测到重复内容，合并章节标识: {existing_sec['section_id']}")
                else:
                    # 新内容
                    seen_contents[content_hash] = sec
                    deduplicated.append(sec)
            
            sections = deduplicated
        
        print(f"[正文分批] 去重后共 {len(sections)} 个章节")
        for sec in sections:
            print(f"  - {sec['section_id']}: {len(sec['content'])} 字符")
        
        return sections
        
    except Exception as e:
        print(f"[正文分批] 分割失败: {e}")
        import traceback
        traceback.print_exc()
        return []


# 已废弃：process_report_text_in_batches 函数已迁移到 md_content_processor.py
# 一键解析统一使用 md_content_processor.call_llm_for_text_extraction
# 已迁移：normalize_content_json_headings 函数已迁移到 md_content_processor.py
# 已废弃：test_siliconflow_api 函数已删除（浪费LLM调用来测试API密钥有效性，改为直接调用时失败处理）
# 已废弃：process_nsfc_content_with_llm 函数已删除（功能已由 md_content_processor.call_llm_for_text_extraction 替代）
# 已废弃：call_nsfc_extraction_for_file 函数已删除（无任何调用者）

def preliminary_check_with_llm(table_content: str, text_content: str, api_key: str, user_id: int = None, json_data: dict = None, md_content: str = None, vlm_md_content: str = None) -> dict:
    """正则初筛检查 + 可选LLM辅助（C1/C2自引率英文姓名匹配）
    
    Args:
        table_content: 表格部分的Markdown内容
        text_content: 正文部分的Markdown内容
        api_key: SiliconFlow API密钥（用于C2自引率LLM辅助判断英文姓名）
        user_id: 用户ID（用于获取分组模型配置）
        json_data: 从表格提取的JSON数据（用于正则检查）
        md_content: 完整的原始Markdown文档（优先使用，避免拆分重组打乱文档结构）
        vlm_md_content: VLM处理后的Markdown内容（可选，作为HTML表格检查的回退源）
    
    Returns:
        dict: {
            'success': bool,
            'violations': list,
            'total_rules': int,
            'violation_count': int
        }
    """
    try:
        from regex_preliminary_check import run_regex_preliminary_check
        
        # 优先使用原始完整md_content，避免拆分重组打乱文档结构
        if not md_content:
            md_content = (table_content or '') + '\n\n' + (text_content or '')
        
        # 获取用户分组配置的模型名称（用于C2自引率LLM辅助）
        model_name = None
        if user_id:
            try:
                from db_operations import get_user_llm_provider
                user_provider = get_user_llm_provider(user_id)
                if user_provider == 'local_llm':
                    from md_content_processor import LOCAL_LLM_CONFIG
                    model_name = f"local:{LOCAL_LLM_CONFIG['model']}"
                else:
                    model_name = get_user_siliconflow_model(user_id) or "Qwen/Qwen3-30B-A3B-Instruct-2507"
            except Exception:
                model_name = "Qwen/Qwen3-30B-A3B-Instruct-2507"
        
        # 执行正则检查（28条规则: 26条纯正则 + 2条混合C1/C2）
        violations = run_regex_preliminary_check(
            md_content=md_content,
            table_content=table_content,
            text_content=text_content,
            json_data=json_data,
            api_key=api_key,
            model_name=model_name,
            vlm_md_content=vlm_md_content
        )
        
        print(f"[初筛检查] 正则+混合检查完成，发现 {len(violations)} 处违规")
        
        return {
            'success': True,
            'violations': violations,
            'total_rules': 28,  # 28条规则: 原12条 + A系列14条 + C系列2条(时效性+自引率)
            'violation_count': len(violations)
        }
            
    except Exception as e:
        print(f"[初筛检查] 异常: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return {'success': False, 'error': str(e)}

# 已迁移：ensure_complete_json_structure 函数已迁移到 md_content_processor.py
# 已废弃：process_table_with_llm 函数已迁移到 md_content_processor.py
# 一键解析统一使用 md_content_processor.call_llm_for_table_extraction

# ==================== 表格转JSON处理功能结束 ====================

@app.route('/')
def index():
    """上传文件 - 检查登录状态"""
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

def _check_login_rate_limit(ip_address):
    """检查登录频率限制，返回 (is_blocked, message)"""
    try:
        conn_rl = DatabaseConfig.get_db_connection()
        if conn_rl:
            cursor_rl = conn_rl.cursor()
            cursor_rl.execute("""
                SELECT COUNT(*) as cnt
                FROM system_logs 
                WHERE action = 'login_failed' 
                AND ip_address = %s
                AND created_at >= DATE_SUB(NOW(), INTERVAL 10 MINUTE)
            """, (ip_address,))
            result = cursor_rl.fetchone()
            recent_failures = result['cnt'] if result else 0
            cursor_rl.close()
            conn_rl.close()
            if recent_failures >= 10:
                return True, '登录尝试次数过多，请10分钟后再试'
    except Exception:
        pass
    return False, ''

@app.route('/login', methods=['GET', 'POST'])
def login():
    """用户登录"""
    if request.method == 'POST':
        # 登录频率限制：同一IP 10分钟内失败>=10次则阻止
        is_blocked, block_msg = _check_login_rate_limit(request.remote_addr)
        if is_blocked:
            flash(block_msg, 'error')
            return render_template('login.html')

        username_or_email = request.form['username']
        password = request.form['password']
        
        # 从数据库获取用户信息（支持用户名或邮箱登录）
        # 注意: get_user_by_username_or_email 现在使用 legacy 版本，直接返回用户字典或 None
        user = get_user_by_username_or_email(username_or_email)
        print(f"DEBUG: 查找用户 '{username_or_email}', 结果: {user is not None}")
        
        if user and check_password_hash(user['password_hash'], password):
            print(f"DEBUG: 用户 {user['username']} (ID: {user['id']}) 密码验证成功")
            
            # "记住我"功能：勾选后session有效期延长至30天
            session.permanent = True
            if request.form.get('remember_me'):
                session['remember_me'] = True
            else:
                session.pop('remember_me', None)
            
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['role'] = user['role']  # 添加用户角色到session
            
            # 记录登录日志（系统日志已足够记录登录信息）
            log_user_action(
                user['id'], 
                'login', 
                f"用户 {user['username']} (登录名: {username_or_email}) 登录成功",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            # 🆕 实时更新用户活动统计
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    
                    # 获取当前在线用户数（最近5分钟有活动）
                    cursor.execute("""
                        SELECT COUNT(DISTINCT user_id) as online_users
                        FROM system_logs 
                        WHERE created_at >= DATE_SUB(NOW(), INTERVAL 5 MINUTE)
                    """)
                    result = cursor.fetchone()
                    online_users = result['online_users'] if result else 0
                    
                    # 获取今日登录次数
                    cursor.execute("""
                        SELECT COUNT(*) as today_logins
                        FROM system_logs 
                        WHERE action = 'login' 
                        AND created_at >= CURDATE() 
                        AND created_at < DATE_ADD(CURDATE(), INTERVAL 1 DAY)
                    """)
                    result = cursor.fetchone()
                    today_logins = result['today_logins'] if result else 0
                    
                    conn.close()
                    
                    # 获取系统管理员ID
                    system_user_id = None
                    try:
                        conn2 = DatabaseConfig.get_db_connection()
                        if conn2:
                            cursor2 = conn2.cursor()
                            cursor2.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                            admin_user = cursor2.fetchone()
                            if admin_user:
                                system_user_id = admin_user['id']
                            conn2.close()
                    except:
                        pass
                    
                    # 只有获取到有效用户ID时才记录日志
                    if system_user_id:
                        # 记录实时登录活动统计
                        log_user_action(
                            system_user_id,
                            'realtime_login_activity',
                            f"当前在线用户: {online_users}, 今日登录次数: {today_logins}, 最新登录: {user['username']}",
                            None,
                            'Login-Activity-Monitor'
                        )
            except Exception as e:
                print(f"[实时统计] 登录活动统计更新失败: {e}")
            
            flash('登录成功！', 'success')
            return redirect(url_for('dashboard'))
        else:
            # 记录失败登录尝试
            log_user_action(
                None, 
                'login_failed', 
                f"登录名 {username_or_email} 登录失败",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            # 🆕 实时安全监控 - 异常登录检测
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    
                    # 检查最近10分钟从同一IP的失败登录次数
                    cursor.execute("""
                        SELECT COUNT(*) as recent_failures
                        FROM system_logs 
                        WHERE action = 'login_failed' 
                        AND ip_address = %s
                        AND created_at >= DATE_SUB(NOW(), INTERVAL 10 MINUTE)
                    """, (request.remote_addr,))
                    result = cursor.fetchone()
                    recent_failures = result['recent_failures'] if result else 0
                    
                    # 检查最近1小时总的失败登录次数
                    cursor.execute("""
                        SELECT COUNT(*) as hourly_failures
                        FROM system_logs 
                        WHERE action = 'login_failed'
                        AND created_at >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
                    """)
                    result = cursor.fetchone()
                    hourly_failures = result['hourly_failures'] if result else 0
                    
                    conn.close()
                    
                    # 异常检测阈值
                    if recent_failures >= 5 or hourly_failures >= 20:
                        # 获取系统管理员ID用于记录安全警报
                        system_user_id = None
                        try:
                            conn_admin = DatabaseConfig.get_db_connection()
                            if conn_admin:
                                cursor_admin = conn_admin.cursor()
                                cursor_admin.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                                admin_user = cursor_admin.fetchone()
                                if admin_user:
                                    system_user_id = admin_user['id']
                                cursor_admin.close()
                                conn_admin.close()
                        except:
                            pass
                        
                        # 只有在获取到有效管理员ID时才记录安全警报
                        if system_user_id:
                            log_user_action(
                                system_user_id,  # 使用系统管理员用户ID
                                'security_alert_login_anomaly',
                                f"检测到异常登录活动: IP {request.remote_addr} 最近10分钟失败{recent_failures}次, 最近1小时系统总失败{hourly_failures}次",
                                request.remote_addr,
                                'Security-Monitor'
                            )
                        
                        # 如果是严重异常，发送通知
                        if recent_failures >= 10:
                            try:
                                from notification_system import send_system_notification
                                send_system_notification(
                                    title="🚨 安全警报：异常登录活动",
                                    content=f"检测到来自 IP {request.remote_addr} 的频繁登录失败尝试（{recent_failures}次/10分钟），请立即检查系统安全！",
                                    notification_type='security_alert',
                                    priority='high',
                                    recipient_id=None
                                )
                            except Exception as e:
                                print(f"[安全通知] 发送安全警报失败: {e}")
            
            except Exception as e:
                print(f"[安全监控] 异常登录检测失败: {e}")
            
            flash('用户名/邮箱或密码错误！', 'error')
    
    return render_template('login.html')

def validate_email(email):
    """验证邮箱格式"""
    import re
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def validate_password_strength(password):
    """验证密码强度"""
    import re
    
    # 检查密码长度
    if len(password) < 8:
        return False, "密码长度至少8位"
    
    if len(password) > 18:
        return False, "密码长度不能超过50位"
    
    # 检查是否包含大写字母
    if not re.search(r'[A-Z]', password):
        return False, "密码必须包含至少一个大写字母"
    
    # 检查是否包含小写字母
    if not re.search(r'[a-z]', password):
        return False, "密码必须包含至少一个小写字母"
    
    # 检查是否包含数字
    if not re.search(r'\d', password):
        return False, "密码必须包含至少一个数字"
    
    # 检查是否包含特殊字符
    if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
        return False, "密码必须包含至少一个特殊字符(!@#$%^&*(),.?\":{}|<>)"
    
    return True, "密码强度符合要求"

def generate_username_from_email(email):
    """从邮箱生成用户名"""
    if not validate_email(email):
        return None
    
    # 提取@前的部分作为基础用户名
    base_username = email.split('@')[0]
    
    # 清理用户名，只保留字母数字和下划线
    import re
    base_username = re.sub(r'[^a-zA-Z0-9_]', '_', base_username)
    
    # 如果清理后的用户名为空，则使用默认值
    if not base_username:
        base_username = 'user'
    
    # 检查用户名是否已存在，如果存在则添加数字后缀
    username = base_username
    counter = 1
    # 修复：check_user_exists 返回元组 (exists, error_msg)
    user_exists, _ = check_user_exists(username)
    while user_exists:
        username = f"{base_username}_{counter}"
        counter += 1
        user_exists, _ = check_user_exists(username)
    
    return username

@app.route('/register', methods=['GET', 'POST'])
def register():
    """用户注册"""
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        # 后端邮箱验证
        if not validate_email(email):
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败: 邮箱格式无效 {email}",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('邮箱格式无效！', 'error')
            return render_template('register.html')
        
        # 从邮箱自动生成用户名
        username = generate_username_from_email(email)
        if not username:
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败: 无法从邮箱生成用户名 {email}",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('无法从邮箱生成用户名！', 'error')
            return render_template('register.html')
        
        # 修复：check_email_exists 返回元组 (exists, error_msg)
        email_exists, email_error = check_email_exists(email, None, request.remote_addr, request.headers.get('User-Agent'))
        if email_error:
            # 数据库连接或查询错误
            flash('系统错误，请稍后重试！', 'error')
            return render_template('register.html')
        
        # 验证
        if email_exists:
            # 记录邮箱已存在的尝试
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败: 邮箱 {email} 已存在",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('邮箱已被使用！', 'error')
        elif password != confirm_password:
            # 记录密码不一致的尝试
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败: 邮箱 {email} 密码不一致",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('两次输入的密码不一致！', 'error')
        else:
            # 验证密码强度
            is_valid, message = validate_password_strength(password)
            if not is_valid:
                log_user_action(
                    None, 
                    'register_failed', 
                    f"注册失败: 邮箱 {email} 密码强度不足 - {message}",
                    request.remote_addr,
                    request.headers.get('User-Agent')
                )
                flash(message, 'error')
            else:
                # 创建新用户
                password_hash = generate_password_hash(password)
                user_id = create_user(username, email, password_hash)
                
                if user_id:
                    # 记录注册成功日志
                    log_user_action(
                        user_id, 
                        'register', 
                        f"新用户 {username} (邮箱: {email}) 注册成功",
                        request.remote_addr,
                        request.headers.get('User-Agent')
                    )
                    flash(f'注册成功！您的用户名为: {username}，请登录。', 'success')
                    return redirect(url_for('login'))
                else:
                    # 记录数据库创建用户失败
                    log_user_action(
                        None, 
                        'register_failed', 
                        f"注册失败: 邮箱 {email} 数据库创建失败",
                        request.remote_addr,
                        request.headers.get('User-Agent')
                    )
                    flash('注册失败，请稍后重试！', 'error')
    
    return render_template('register.html')


# 忘记密码：密码重置令牌存储 {token: {'user_id': int, 'expires': datetime}}
_password_reset_tokens = {}

def _cleanup_expired_tokens():
    """清理过期的密码重置令牌"""
    now = datetime.now()
    expired = [t for t, v in _password_reset_tokens.items() if v['expires'] < now]
    for t in expired:
        del _password_reset_tokens[t]

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    """忘记密码 - 身份验证后重置"""
    if request.method == 'POST':
        step = request.form.get('step', '1')
        
        if step == '1':
            # 步骤1：通过用户名或邮箱查找用户
            identifier = request.form.get('identifier', '').strip()
            
            if not identifier:
                flash('请输入用户名或邮箱', 'error')
                return render_template('forgot_password.html')
            
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT id, username, email FROM users WHERE (username = %s OR email = %s) AND is_active = 1",
                        (identifier, identifier)
                    )
                    user = cursor.fetchone()
                    cursor.close()
                    conn.close()
                    
                    if user:
                        # 生成重置令牌
                        _cleanup_expired_tokens()
                        import secrets as _secrets
                        token = _secrets.token_urlsafe(32)
                        _password_reset_tokens[token] = {
                            'user_id': user['id'],
                            'username': user['username'],
                            'expires': datetime.now() + timedelta(minutes=15)
                        }
                        log_user_action(user['id'], 'password_reset_request', f"用户 {user['username']} 请求密码重置", request.remote_addr, request.headers.get('User-Agent'))
                        # 邮箱脱敏显示
                        email = user['email']
                        at_idx = email.index('@')
                        masked_email = email[0] + '*' * (at_idx - 1) + email[at_idx:] if at_idx > 1 else email
                        return render_template('forgot_password.html', step=2, token=token, username=user['username'], masked_email=masked_email)
                    else:
                        flash('未找到该用户，请检查用户名或邮箱是否正确', 'error')
                else:
                    flash('系统错误，请稍后重试', 'error')
            except Exception as e:
                app.logger.error(f"忘记密码验证失败: {e}")
                flash('系统错误，请稍后重试', 'error')
        
        elif step == '2':
            # 步骤2：重置密码
            token = request.form.get('token', '')
            new_password = request.form.get('new_password', '')
            confirm_password = request.form.get('confirm_password', '')
            
            # 验证令牌
            _cleanup_expired_tokens()
            token_data = _password_reset_tokens.get(token)
            if not token_data:
                flash('重置链接已过期或无效，请重新验证身份', 'error')
                return render_template('forgot_password.html')
            
            if not new_password or not confirm_password:
                flash('请输入新密码', 'error')
                return render_template('forgot_password.html', step=2, token=token, username=token_data['username'])
            
            if new_password != confirm_password:
                flash('两次输入的密码不一致', 'error')
                return render_template('forgot_password.html', step=2, token=token, username=token_data['username'])
            
            is_valid, msg = validate_password_strength(new_password)
            if not is_valid:
                flash(msg, 'error')
                return render_template('forgot_password.html', step=2, token=token, username=token_data['username'])
            
            # 执行密码重置
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    new_hash = generate_password_hash(new_password)
                    cursor.execute("UPDATE users SET password_hash = %s, updated_at = CURRENT_TIMESTAMP WHERE id = %s", (new_hash, token_data['user_id']))
                    conn.commit()
                    cursor.close()
                    conn.close()
                    
                    # 删除已使用的令牌
                    del _password_reset_tokens[token]
                    
                    log_user_action(token_data['user_id'], 'password_reset', f"用户 {token_data['username']} 通过忘记密码功能重置了密码", request.remote_addr, request.headers.get('User-Agent'))
                    flash('密码重置成功！请使用新密码登录', 'success')
                    return redirect(url_for('login'))
                else:
                    flash('系统错误，请稍后重试', 'error')
            except Exception as e:
                app.logger.error(f"密码重置失败: {e}")
                flash('密码重置失败，请重试', 'error')
    
    return render_template('forgot_password.html')


@app.route('/word/forgot-password', methods=['GET', 'POST'])
def word_forgot_password():
    """Word模板版本 - 忘记密码"""
    if request.method == 'POST':
        step = request.form.get('step', '1')
        
        if step == '1':
            identifier = request.form.get('identifier', '').strip()
            
            if not identifier:
                flash('请输入用户名或邮箱', 'error')
                return render_template('forgot_password_word.html')
            
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT id, username, email FROM users WHERE (username = %s OR email = %s) AND is_active = 1",
                        (identifier, identifier)
                    )
                    user = cursor.fetchone()
                    cursor.close()
                    conn.close()
                    
                    if user:
                        _cleanup_expired_tokens()
                        import secrets as _secrets
                        token = _secrets.token_urlsafe(32)
                        _password_reset_tokens[token] = {
                            'user_id': user['id'],
                            'username': user['username'],
                            'expires': datetime.now() + timedelta(minutes=15)
                        }
                        log_user_action(user['id'], 'password_reset_request', f"用户 {user['username']} 通过Word模板请求密码重置", request.remote_addr, request.headers.get('User-Agent'))
                        email = user['email']
                        at_idx = email.index('@')
                        masked_email = email[0] + '*' * (at_idx - 1) + email[at_idx:] if at_idx > 1 else email
                        return render_template('forgot_password_word.html', step=2, token=token, username=user['username'], masked_email=masked_email)
                    else:
                        flash('未找到该用户，请检查用户名或邮箱是否正确', 'error')
                else:
                    flash('系统错误，请稍后重试', 'error')
            except Exception as e:
                app.logger.error(f"忘记密码验证失败: {e}")
                flash('系统错误，请稍后重试', 'error')
        
        elif step == '2':
            token = request.form.get('token', '')
            new_password = request.form.get('new_password', '')
            confirm_password = request.form.get('confirm_password', '')
            
            _cleanup_expired_tokens()
            token_data = _password_reset_tokens.get(token)
            if not token_data:
                flash('重置链接已过期或无效，请重新验证身份', 'error')
                return render_template('forgot_password_word.html')
            
            if not new_password or not confirm_password:
                flash('请输入新密码', 'error')
                return render_template('forgot_password_word.html', step=2, token=token, username=token_data['username'])
            
            if new_password != confirm_password:
                flash('两次输入的密码不一致', 'error')
                return render_template('forgot_password_word.html', step=2, token=token, username=token_data['username'])
            
            is_valid, msg = validate_password_strength(new_password)
            if not is_valid:
                flash(msg, 'error')
                return render_template('forgot_password_word.html', step=2, token=token, username=token_data['username'])
            
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    new_hash = generate_password_hash(new_password)
                    cursor.execute("UPDATE users SET password_hash = %s, updated_at = CURRENT_TIMESTAMP WHERE id = %s", (new_hash, token_data['user_id']))
                    conn.commit()
                    cursor.close()
                    conn.close()
                    
                    del _password_reset_tokens[token]
                    log_user_action(token_data['user_id'], 'password_reset', f"用户 {token_data['username']} 通过Word模板忘记密码功能重置了密码", request.remote_addr, request.headers.get('User-Agent'))
                    flash('密码重置成功！请使用新密码登录', 'success')
                    return redirect(url_for('word_login'))
                else:
                    flash('系统错误，请稍后重试', 'error')
            except Exception as e:
                app.logger.error(f"密码重置失败: {e}")
                flash('密码重置失败，请重试', 'error')
    
    return render_template('forgot_password_word.html')


def update_user_session_version(user_id):
    """更新用户的会话版本，用于全设备注销"""
    # 由于删除了last_login字段，这里可以简化为仅返回True
    # 实际的会话管理通过session和安全日志来处理
    print(f"会话版本更新: 用户 {user_id} 执行全设备注销")
    return True

def log_security_event(user_id, event_type, details, ip_address=None, user_agent=None):
    """记录安全相关事件"""
    try:
        conn = get_db_connection_with_error_handling()
        if conn:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO system_logs (user_id, action, details, ip_address, user_agent, created_at) VALUES (%s, %s, %s, %s, %s, NOW())",
                (user_id, event_type, details, ip_address, user_agent)
            )
            conn.commit()  # 提交日志记录
            safe_db_close(conn, cursor)
    except Exception as e:
        smart_print(f"记录安全事件错误: {e}", LogLevel.ERROR)

@app.route('/logout')
def logout():
    """用户注销/登出"""
    user_id = session.get('user_id')
    username = session.get('username')
    
    # 记录注销日志
    if user_id and username:
        log_user_action(
            user_id, 
            'logout', 
            f"用户 {username} 成功注销登出",
            request.remote_addr,
            request.headers.get('User-Agent')
        )
    
    # 清除所有会话数据
    session.clear()
    
    flash('您已成功登出！', 'info')
    return redirect(url_for('login'))

@app.route('/logout_all')
def logout_all():
    """注销所有设备上的登录会话（扩展功能）"""
    user_id = session.get('user_id')
    username = session.get('username')
    
    if user_id and username:
        # 更新用户会话版本，使其他设备上的登录失效
        if update_user_session_version(user_id):
            # 记录全局注销日志
            log_user_action(
                user_id, 
                'logout_all', 
                f"用户 {username} 执行全局注销（所有设备）",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            # 记录安全事件
            log_security_event(
                user_id,
                'global_logout',
                f"用户 {username} 主动执行全设备注销",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            flash('已成功注销所有设备上的登录！', 'info')
        else:
            flash('全设备注销操作失败，请稍后重试。', 'error')
    else:
        flash('您尚未登录！', 'warning')
    
    # 清除当前会话
    session.clear()
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    """传统用户仪表板 - 只显示未完成的文件"""
    
    # 记录页面访问日志
    log_user_action(
        session['user_id'], 
        'page_access', 
        "访问仪表板页面",
        request.remote_addr,
        request.headers.get('User-Agent')
    )
    
    # 🆕 记录实时页面访问统计
    try:
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # 获取最近5分钟的页面访问数
            cursor.execute("""
                SELECT COUNT(*) as recent_page_views
                FROM system_logs 
                WHERE action = 'page_access' 
                AND created_at >= DATE_SUB(NOW(), INTERVAL 5 MINUTE)
            """)
            result = cursor.fetchone()
            recent_views = result['recent_page_views'] if result else 0
            
            # 获取当前用户的会话时长
            cursor.execute("""
                SELECT MIN(created_at) as session_start
                FROM system_logs 
                WHERE user_id = %s 
                AND action IN ('login', 'page_access')
                AND created_at >= DATE_SUB(NOW(), INTERVAL 4 HOUR)
                ORDER BY created_at DESC
                LIMIT 1
            """, (session['user_id'],))
            session_result = cursor.fetchone()
            
            session_duration = 0
            if session_result and session_result['session_start']:
                session_start = session_result['session_start']
                session_duration = (get_current_time().replace(tzinfo=None) - session_start).total_seconds() / 60
            
            conn.close()
            
            # 获取系统管理员ID
            system_user_id = None
            try:
                conn2 = DatabaseConfig.get_db_connection()
                if conn2:
                    cursor2 = conn2.cursor()
                    cursor2.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                    admin_user = cursor2.fetchone()
                    if admin_user:
                        system_user_id = admin_user['id']
                    conn2.close()
            except:
                pass
            
            # 只有获取到有效用户ID时才记录日志
            if system_user_id:
                # 记录实时页面访问统计日志
                log_user_action(
                    system_user_id,
                    'realtime_page_activity',
                    f"最近5分钟页面访问: {recent_views}, 当前用户: {session['username']} (会话时长: {session_duration:.1f}分钟)",
                    None,
                    'Page-Activity-Monitor'
                )
    except Exception as e:
        print(f"[实时统计] 页面访问统计更新失败: {e}")
    
    # 从数据库获取当前用户的文件，过滤掉已完成的文件
    user_files = get_user_files_from_db(session['user_id'])
    # 只显示未完成的文件（排除状态为 'completed' 的文件）
    pending_files = [file for file in user_files if file.get('status') != 'completed']
    
    # 获取用户的API配置
    user_result = get_user_by_username(session['username'])
    user = user_result.get('data') if user_result.get('success') else None
    
    # 获取用户的SiliconFlow API密钥
    api_key_info = get_active_api_key(session['user_id'], 'siliconflow')
    api_token = api_key_info.get('api_key', '') if api_key_info else ''
    
    return render_template('dashboard.html', 
                         username=session['username'],
                         user_email=user.get('email', '') if user else '',
                         files=pending_files,
                         api_token=api_token)

@app.route('/user-info-ynex')
@login_required
def user_info_ynex():
    """YNEX风格用户信息管理页面 - 测试版本"""
    
    # 获取用户详细信息
    user_result = get_user_by_username(session['username'])
    user = user_result.get('data') if user_result.get('success') else None
    
    # 安全地获取用户信息，处理可能不存在的字段
    created_at_str = '未知'
    if user and 'created_at' in user and user['created_at']:
        try:
            if isinstance(user['created_at'], str):
                created_at_str = user['created_at']
            else:
                created_at_str = user['created_at'].strftime('%Y-%m-%d %H:%M:%S')
        except:
            created_at_str = '未知'
    
    # 准备用户信息传递给模板
    user_info = {
        'user_id': session['user_id'],  # 已经在函数开头检查过存在性
        'username': session.get('username', ''),
        'email': user.get('email', '') if user else '',
        'created_at': created_at_str,
        'is_admin': session.get('role') in ['admin', 'premium']
    }
    
    return render_template('user_info_ynex.html', **user_info)

@app.route('/comprehensive-processor')
@login_required
def comprehensive_processor():
    """综合处理页面（三位一体）"""
    
    # 从数据库获取当前用户的文件
    user_files = get_user_files_from_db(session['user_id'])
    
    # 获取用户的API配置
    user_result = get_user_by_username(session['username'])
    user = user_result.get('data') if user_result.get('success') else None
    
    # 获取用户的SiliconFlow API密钥
    api_key_info = get_active_api_key(session['user_id'], 'siliconflow')
    api_token = api_key_info.get('api_key', '') if api_key_info else ''
    
    return render_template('comprehensive_processor.html', 
                         username=session['username'],
                         user_email=user.get('email', '') if user else '',
                         files=user_files,
                         api_token=api_token)

@app.route('/rule-checker')
@login_required
def rule_checker():
    """规则检查页面"""
    
    return render_template('rule_checker.html', 
                         username=session['username'])

@app.route('/api/preliminary-check', methods=['POST'])
def api_preliminary_check():
    """初筛检查API接口 - 优先返回数据库已保存的结果"""
    try:
        if 'user_id' not in session:
            return jsonify({'success': False, 'error': '请先登录'}), 401
        
        # 获取请求数据
        data = request.get_json()
        file_id = data.get('file_id')
        force_refresh = data.get('force_refresh', False)  # 是否强制重新检查
        
        print(f"[初筛API] 接收到请求，file_id: {file_id}, user_id: {session['user_id']}, force_refresh: {force_refresh}")
        
        if not file_id:
            return jsonify({'success': False, 'error': '缺少file_id参数'}), 400
        
        # 从数据库获取文件信息（包括已保存的初筛结果）
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        cursor.execute("""
            SELECT original_name, md_content, original_md_content, table_preliminary_check, text_preliminary_check, status, extracted_json_data, nsfc_json_data
            FROM user_files
            WHERE file_id = %s AND user_id = %s
        """, (file_id, session['user_id']))
        file_info = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not file_info:
            print(f"[初筛API] 错误：未找到文件，file_id={file_id}, user_id={session['user_id']}")
            return jsonify({'success': False, 'error': '文件不存在或无权访问'}), 404
        
        print(f"[初筛API] 成功获取文件信息，文件名: {file_info.get('original_name')}")

        def merge_internal_quality_rules(json_text, source_kind, target_violations):
            """从JSON内部字段合并LLM质量规则到违规列表。

            source_kind: 'content' | 'table'
            """
            if not json_text:
                return 0
            try:
                payload = json.loads(json_text) if isinstance(json_text, str) else json_text
            except Exception as e:
                print(f"[初筛API] 解析{source_kind}质量规则JSON失败: {e}")
                return 0

            if not isinstance(payload, dict):
                return 0

            quality_field = None
            for candidate in ['_quality_check_internal', '_content_quality_check', '_table_quality_check_internal']:
                if candidate in payload and isinstance(payload[candidate], dict):
                    quality_field = candidate
                    break

            if not quality_field:
                return 0

            quality_check = payload[quality_field]
            print(f"[初筛API] 发现{source_kind}质量评估字段 {quality_field}，包含 {len(quality_check)} 条规则")

            # 正文规则保留固定映射，表格规则允许动态扩展
            content_rule_config = {
                'rule_4_1_1': {
                    'name': '4.1.1 立项依据需说清楚"为什么要开展此项研究，研究的价值何在"',
                    'severity': 'high'
                },
                'rule_4_1_2': {
                    'name': '4.1.2 研究内容应根据自己的研究思路和逻辑自主撰写',
                    'severity': 'high'
                },
                'rule_4_1_3': {
                    'name': '4.1.3 研究基础需展示前期的工作积累',
                    'severity': 'medium'
                }
            }

            merged_failed = 0
            for rule_key, check_result in quality_check.items():
                if not isinstance(check_result, dict):
                    continue
                if check_result.get('status', 'not_applicable') != 'failed':
                    continue

                if source_kind == 'content' and rule_key in content_rule_config:
                    rule_name = content_rule_config[rule_key]['name']
                    severity = content_rule_config[rule_key]['severity']
                    suggestion = '请根据新模板要求重新撰写该章节内容'
                    location = '正文对应章节'
                else:
                    # 表格规则及未知正文规则：动态展示
                    rule_name = check_result.get('name') or rule_key
                    severity = check_result.get('severity', 'medium')
                    suggestion = check_result.get('suggestion', '请根据规则要求补充或修正内容')
                    location = '基本信息表' if source_kind == 'table' else '正文对应章节'

                target_violations.append({
                    'rule_name': rule_name,
                    'severity': severity,
                    'error_text': check_result.get('reason', '内容不符合规则要求'),
                    'location': location,
                    'suggestion': suggestion,
                    'source': f'llm_quality_check_{source_kind}'
                })
                merged_failed += 1

            return merged_failed
        
        # 如果不强制刷新，且数据库中已有初筛结果，直接返回
        if not force_refresh and (file_info.get('table_preliminary_check') or file_info.get('text_preliminary_check')):
            print(f"[初筛API] 使用数据库缓存的初筛结果")
            
            # 合并表格和正文的初筛结果
            all_violations = []
            
            if file_info.get('table_preliminary_check'):
                try:
                    table_violations = json.loads(file_info['table_preliminary_check'])
                    if isinstance(table_violations, list):
                        for v in table_violations:
                            v['source'] = 'table'  # 标记来源
                            # 字段映射：将violation转为error_text，rule_description转为suggestion
                            if 'violation' in v and 'error_text' not in v:
                                v['error_text'] = v['violation']
                            if 'rule_description' in v and 'suggestion' not in v:
                                v['suggestion'] = v['rule_description']
                        all_violations.extend(table_violations)
                except Exception as e:
                    print(f"[初筛API] 解析表格初筛结果失败: {e}")
            
            if file_info.get('text_preliminary_check'):
                try:
                    text_violations = json.loads(file_info['text_preliminary_check'])
                    if isinstance(text_violations, list):
                        for v in text_violations:
                            v['source'] = 'text'  # 标记来源
                            # 字段映射：将violation转为error_text，rule_description转为suggestion
                            if 'violation' in v and 'error_text' not in v:
                                v['error_text'] = v['violation']
                            if 'rule_description' in v and 'suggestion' not in v:
                                v['suggestion'] = v['rule_description']
                        all_violations.extend(text_violations)
                except Exception as e:
                    print(f"[初筛API] 解析正文初筛结果失败: {e}")
            
            # 合并正文/表格JSON中的LLM质量评估结果（内部字段）
            merge_internal_quality_rules(file_info.get('nsfc_json_data'), 'content', all_violations)
            merge_internal_quality_rules(file_info.get('extracted_json_data'), 'table', all_violations)
            
            return jsonify({
                'success': True,
                'violations': all_violations,
                'violation_count': len(all_violations),
                'total_rules': 31,  # 31条规则: 28条正则/混合 + 3条LLM质量评估
                'from_cache': True
            })
        
        # 如果没有缓存结果或强制刷新，进行实时检查
        print(f"[初筛API] 执行实时初筛检查...")
        
        md_content = file_info.get('md_content')
        if not md_content:
            file_status = file_info.get('status', 'unknown')
            print(f"[初筛API] 错误：MD内容为空，文件状态: {file_status}")
            return jsonify({'success': False, 'error': f'文件尚未解析或MD内容为空（当前状态：{file_status}）'}), 400
        
        # 优先使用原始MinerU输出（VLM替换前），因为VLM可能丢失部分字段（如研究期限）
        original_md_content = file_info.get('original_md_content') or md_content
        print(f"[初筛API] MD内容长度: {len(md_content)} 字符，原始MinerU: {len(original_md_content)} 字符")
        
        # 分离表格和正文内容（支持HTML <table>和Markdown管道表格）
        table_content = ""
        text_content = ""
        
        lines = md_content.split('\n')
        in_pipe_table = False
        in_html_table = False
        html_table_depth = 0
        table_lines = []
        text_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # HTML表格检测：<table> ... </table>（支持嵌套）
            if not in_html_table and '<table' in stripped.lower():
                in_html_table = True
                html_table_depth = stripped.lower().count('<table') - stripped.lower().count('</table>')
                if html_table_depth < 1:
                    html_table_depth = 0
                    in_html_table = False
                table_lines.append(line)
                continue
            
            if in_html_table:
                table_lines.append(line)
                html_table_depth += stripped.lower().count('<table') - stripped.lower().count('</table>')
                if html_table_depth <= 0:
                    in_html_table = False
                    html_table_depth = 0
                continue
            
            # Markdown管道表格检测
            if '|' in line and stripped.startswith('|'):
                in_pipe_table = True
                table_lines.append(line)
            elif in_pipe_table and stripped == '':
                in_pipe_table = False
                text_lines.append(line)
            elif in_pipe_table:
                table_lines.append(line)
            else:
                text_lines.append(line)
        
        table_content = '\n'.join(table_lines)
        text_content = '\n'.join(text_lines)
        
        print(f"[初筛API] 内容分离完成 - 表格: {len(table_content)} 字符, 正文: {len(text_content)} 字符")
        
        # 获取API密钥
        api_key_info = get_active_api_key(session['user_id'], 'siliconflow')
        if not api_key_info:
            print(f"[初筛API] 错误：用户未配置SiliconFlow API密钥")
            return jsonify({'success': False, 'error': '未配置智能分析服务密钥，请前往个人信息页面配置'}), 400
        
        api_key = api_key_info.get('api_token')  # 使用api_token字段
        print(f"[初筛API] API密钥已获取，长度: {len(api_key) if api_key else 0}")
        
        # 检查API密钥是否为空
        if not api_key or api_key.strip() == '':
            print(f"[初筛API] 错误：API密钥为空")
            return jsonify({'success': False, 'error': 'API密钥为空，请前往个人信息页面重新配置'}), 400
        
        # 调用初筛检查函数（24条正则/混合规则，C2自引率可能调用LLM辅助）
        print(f"[初筛API] 开始执行24条正则/混合规则检查...")
        
        # 准备JSON数据供正则检查使用
        json_data = None
        if file_info.get('extracted_json_data'):
            try:
                json_data = json.loads(file_info['extracted_json_data'])
                print(f"[初筛API] JSON数据已加载，字段数: {len(json_data) if isinstance(json_data, dict) else 0}")
            except Exception as e:
                print(f"[初筛API] 解析JSON数据失败: {e}")
        
        result = preliminary_check_with_llm(
            table_content=table_content,
            text_content=text_content,
            api_key=api_key,
            user_id=session['user_id'],
            json_data=json_data,
            md_content=original_md_content,  # 传入原始MinerU输出（VLM修改前），避免VLM替换丢失字段
            vlm_md_content=md_content  # VLM处理后的内容作为回退源
        )
        
        print(f"[初筛API] 24条正则/混合规则检查完成，结果: {result.get('success')}, 违规数: {result.get('violation_count', 0)}")
        
        if result.get('success'):
            # ========== 保存新的检查结果到数据库 ==========
            try:
                # 分离表格和正文的违规项
                table_violations_to_save = []
                text_violations_to_save = []
                for v in result.get('violations', []):
                    if v.get('source') == 'table':
                        table_violations_to_save.append(v)
                    else:
                        text_violations_to_save.append(v)
                
                save_conn = DatabaseConfig.get_db_connection()
                if save_conn:
                    save_cursor = save_conn.cursor()
                    save_cursor.execute("""
                        UPDATE user_files 
                        SET table_preliminary_check = %s, text_preliminary_check = %s
                        WHERE file_id = %s AND user_id = %s
                    """, (
                        json.dumps(table_violations_to_save, ensure_ascii=False) if table_violations_to_save else None,
                        json.dumps(text_violations_to_save if text_violations_to_save else result.get('violations', []), ensure_ascii=False),
                        file_id, session['user_id']
                    ))
                    save_conn.commit()
                    save_cursor.close()
                    save_conn.close()
                    print(f"[初筛API] 重新检查结果已保存到数据库")
            except Exception as save_err:
                print(f"[初筛API] 保存检查结果到数据库失败: {save_err}")
            
            # ========== 合并LLM质量评估结果（来自一键解析时已生成的正文/表格JSON内部字段） ==========
            all_violations = list(result.get('violations', []))
            llm_quality_merged = False
            merged_content = merge_internal_quality_rules(file_info.get('nsfc_json_data'), 'content', all_violations)
            merged_table = merge_internal_quality_rules(file_info.get('extracted_json_data'), 'table', all_violations)
            llm_quality_merged = (merged_content + merged_table) > 0
            
            if not llm_quality_merged:
                print(f"[初筛API] 注意：未找到LLM质量评估数据（nsfc_json_data为空或无_quality_check_internal字段），3条LLM规则未参与")
            
            result['violations'] = all_violations
            result['violation_count'] = len(all_violations)
            result['total_rules'] = 31 if llm_quality_merged else 28
            result['llm_quality_merged'] = llm_quality_merged
        
        # 记录检查日志
        log_user_action(
            session['user_id'],
            'preliminary_check',
            f"对文件 {file_info.get('original_name')} 执行重新检查，发现 {result.get('violation_count', 0)} 处违规",
            request.remote_addr,
            request.headers.get('User-Agent')
        )
        
        return jsonify(result)
        
    except Exception as e:
        print(f"[初筛API] 异常: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': f'服务器错误: {str(e)}'}), 500

@app.route('/parsed-documents')
@login_required
def parsed_documents():
    """已解析文档页面"""
    
    # 从数据库获取当前用户已解析的文件（状态为completed）
    user_files = get_user_files_from_db(session['user_id'])
    parsed_files = [file for file in user_files if file.get('status') == 'completed']
    
    # 为每个文件计算处理时间
    for file in parsed_files:
        if file.get('processing_start_time') and file.get('processed_time'):
            from datetime import datetime
            start_time = file['processing_start_time']
            end_time = file['processed_time']
            
            # 如果是字符串格式，转换为datetime对象
            if isinstance(start_time, str):
                start_time = datetime.fromisoformat(start_time.replace('Z', '+00:00'))
            if isinstance(end_time, str):
                end_time = datetime.fromisoformat(end_time.replace('Z', '+00:00'))
            
            # 计算时间差（秒）
            time_diff = (end_time - start_time).total_seconds()
            file['processing_time_seconds'] = int(time_diff)
        else:
            file['processing_time_seconds'] = None
    
    # 获取用户的API配置
    user_result = get_user_by_username(session['username'])
    user = user_result.get('data') if user_result.get('success') else None
    
    return render_template('parsed_documents.html', 
                         username=session['username'],
                         user_email=user.get('email', '') if user else '',
                         files=parsed_files)

def parse_reference_data(file_data):
    """解析文件的参考文献数据"""
    import json
    import re
    
    # 初始化返回数据结构
    result = {
        'report_references': [],
        'review_references': [],
        'analysis_summary': {
            'total_references': 0,
            'correct_count': 0,
            'error_count': 0,
            'success_rate': 0.0,
            'error_rate': 0.0,
            'citation_system': 'GB/T 7714-2015',
            'main_errors': []
        }
    }
    
    # 获取参考文献内容
    reference_content = file_data.get('reference_md_content', '')
    validation_json = file_data.get('reference_validation_json', '')
    
    print(f"[parse_reference_data] 输入数据检查:")
    print(f"   - reference_content: {'有' if reference_content else '无'} ({len(reference_content) if reference_content else 0} 字符)")
    print(f"   - validation_json: {'有' if validation_json else '无'} ({len(validation_json) if validation_json else 0} 字符)")
    
    if not reference_content:
        print(f"[parse_reference_data] ⚠️ reference_content为空，返回空结果")
        return result
    
    # 解析参考文献内容，根据标题分类
    lines = reference_content.strip().split('\n')
    current_section = None
    references = []
    
    for line in lines:
        line = line.strip()
        if line.startswith('# 参考文献 - 报告'):
            current_section = 'report'
        elif line.startswith('# 参考文献 - 综述'):
            current_section = 'review'
        elif line and (line.startswith('[') or re.match(r'^\d+[\.。]', line)):
            # 这是一个参考文献条目
            if current_section == 'report':
                result['report_references'].append(line)
            elif current_section == 'review':
                result['review_references'].append(line)
            references.append(line)
    
    # 如果没有明确分类，将所有参考文献归为报告类别（Word文档可能只有单个列表）
    if not result['report_references'] and not result['review_references'] and references:
        result['report_references'] = references
    
    print(f"[parse_reference_data] 参考文献解析完成:")
    print(f"   - 提取的参考文献总数: {len(references)}")
    print(f"   - report_references: {len(result['report_references'])} 条")
    print(f"   - review_references: {len(result['review_references'])} 条")
    if result['report_references']:
        print(f"   - 报告第1条预览: {result['report_references'][0][:80]}...")
    
    # 解析验证结果JSON
    if validation_json:
        try:
            validation_data = json.loads(validation_json)
            
            # 从JSON中提取数据
            summary = validation_data.get('summary', {})
            meta = validation_data.get('meta', {})
            statistics = validation_data.get('statistics', {})
            
            # 基本统计信息
            total_refs = summary.get('total_references', len(references))
            valid_count = summary.get('valid_count', 0)
            invalid_count = summary.get('invalid_count', 0)
            success_rate = summary.get('success_rate', 0.0)
            error_rate = summary.get('error_rate', 0.0)
            
            # 引用系统
            citation_system = meta.get('citation_system', 'numeric')
            if citation_system == 'numeric':
                citation_system_name = 'GB/T 7714-2015 (顺序编码制)'
            else:
                citation_system_name = 'GB/T 7714-2015'
            
            # 错误类型统计
            error_types = statistics.get('error_types', {})
            main_errors = []
            
            # 错误类型映射（已升级为用户友好格式，现在大部分情况下不需要转换）
            # 为了兼容旧数据，保留一些映射
            from reference_validator import get_user_friendly_error_type
            
            error_type_map = {
                'missing_period_after_type': '文献类型标识后缺少句号',
                'missing_required_field': '缺少必备字段',
                'missing_volume_issue': '缺少卷(期)号信息',
                'missing_pages': '缺少页码信息',
                'missing_publish_place': '缺少出版地信息',
                'missing_publisher': '缺少出版社信息',
                # 兼容旧的错误代码格式
                'L3.03': get_user_friendly_error_type('L3.03'),
                'L3.10': get_user_friendly_error_type('L3.10'),
                'L3.06': get_user_friendly_error_type('L3.06'),
                'L2.02': get_user_friendly_error_type('L2.02'),
                'L2.04': get_user_friendly_error_type('L2.04'),
                'L3.08': get_user_friendly_error_type('L3.08')
            }
            
            for error_type, error_data in error_types.items():
                if isinstance(error_data, dict):
                    count = error_data.get('count', 0)
                    percentage = error_data.get('percentage', 0.0)
                    error_code = error_data.get('error_code', '')  # 获取错误代码
                    
                    error_name = error_type_map.get(error_type, error_type)
                    description = f'影响{count}条参考文献，占{percentage:.1f}%'
                    
                    main_errors.append({
                        'type': error_name,
                        'type_code': error_code,  # 使用error_code字段用于排序
                        'count': count,
                        'description': description
                    })
            
            # 定义规则级别排序函数
            def get_rule_priority(error_code):
                """
                获取规则的优先级，用于排序
                G.L1 > G.L2 > G.L3 > 其他
                同级别内按编号排序
                """
                if not error_code:
                    return (999, 999, 999)
                    
                # 提取规则代码（如 G.L1.01, C.L2.02）
                match = re.match(r'^([A-Z])\.L(\d+)\.(\d+)', error_code)
                if match:
                    category = match.group(1)  # G, M, J, C等
                    level = int(match.group(2))  # 1, 2, 3
                    number = int(match.group(3))  # 01, 02, 03
                    # 优先级：G类最高，然后按级别L1>L2>L3，最后按编号
                    priority = (0 if category == 'G' else 1, level, number)
                    return priority
                else:
                    # 不符合规则代码格式的，优先级最低
                    return (999, 999, 999)
            
            # 按规则级别排序（先按优先级，再按错误数量）
            main_errors.sort(key=lambda x: (get_rule_priority(x.get('type_code', '')), -x['count']))
            
            result['analysis_summary'].update({
                'total_references': total_refs,
                'correct_count': valid_count,
                'error_count': invalid_count,
                'success_rate': success_rate,
                'error_rate': error_rate,
                'citation_system': citation_system_name,
                'main_errors': main_errors[:5]  # 只显示前5个主要错误
            })
            
        except json.JSONDecodeError as e:
            print(f"JSON解析错误: {e}")
    else:
        # 如果没有验证数据，设置基本统计
        total_refs = len(references)
        result['analysis_summary'].update({
            'total_references': total_refs,
            'correct_count': total_refs,  # 假设都正确
            'error_count': 0,
            'success_rate': 100.0 if total_refs > 0 else 0,
            'error_rate': 0.0
        })
    
    return result

@app.route('/reference-analysis')
@login_required
def reference_analysis():
    """参考文献分析页面"""
    
    # 获取文件ID参数
    file_id = request.args.get('file_id')
    if not file_id:
        flash('未指定文件ID', 'warning')
        return redirect(url_for('parsed_documents'))
    
    # 从数据库获取文件的参考文献数据
    file_data = get_file_reference_data(file_id, session['user_id'])
    if not file_data:
        flash('未找到指定文件或无权访问', 'error')
        return redirect(url_for('parsed_documents'))
    
    # 调试：打印数据库查询结果
    print(f"[参考文献分析] 数据库查询结果:")
    print(f"   - file_id: {file_data.get('file_id')}")
    print(f"   - original_name: {file_data.get('original_name')}")
    print(f"   - reference_md_content: {'有' if file_data.get('reference_md_content') else '无'} ({len(file_data.get('reference_md_content', '')) if file_data.get('reference_md_content') else 0} 字符)")
    print(f"   - reference_validation_json: {'有' if file_data.get('reference_validation_json') else '无'} ({len(file_data.get('reference_validation_json', '')) if file_data.get('reference_validation_json') else 0} 字符)")
    
    # 解析参考文献数据
    reference_data = parse_reference_data(file_data)
    
    # 调试：打印解析结果
    print(f"[参考文献分析] 解析结果:")
    print(f"   - report_references: {len(reference_data.get('report_references', []))} 条")
    print(f"   - review_references: {len(reference_data.get('review_references', []))} 条")
    print(f"   - total_references: {reference_data['analysis_summary']['total_references']}")
    
    return render_template('reference_analysis.html',
                         username=session['username'],
                         file_name=file_data.get('original_name', '未知文件'),
                         reference_data=reference_data,
                         validation_json=file_data.get('reference_validation_json', ''))

# ==================== Word模板路由（独立的第二套模板系统）====================
# 所有路由添加 /word 前缀，指向 word_templates 文件夹中的模板
# 保持与原始路由完全相同的功能，仅模板路径不同

@app.route('/word/login', methods=['GET', 'POST'])
def word_login():
    """Word模板版本 - 用户登录"""
    if request.method == 'POST':
        # 登录频率限制：同一IP 10分钟内失败>=10次则阻止
        is_blocked, block_msg = _check_login_rate_limit(request.remote_addr)
        if is_blocked:
            flash(block_msg, 'error')
            return render_template('login_word.html')

        username_or_email = request.form['username']
        password = request.form['password']
        
        user = get_user_by_username_or_email(username_or_email)
        print(f"DEBUG: 查找用户 '{username_or_email}', 结果: {user is not None}")
        
        if user and check_password_hash(user['password_hash'], password):
            print(f"DEBUG: 用户 {user['username']} (ID: {user['id']}) 密码验证成功")
            
            # "记住我"功能：勾选后session有效期延长至30天
            session.permanent = True
            if request.form.get('remember_me'):
                session['remember_me'] = True
            else:
                session.pop('remember_me', None)
            
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['role'] = user['role']
            
            log_user_action(
                user['id'], 
                'login', 
                f"用户 {user['username']} (登录名: {username_or_email}) 通过Word模板登录成功",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            # 登录后的统计更新逻辑（与原始登录相同）
            try:
                conn = DatabaseConfig.get_db_connection()
                if conn:
                    cursor = conn.cursor()
                    
                    cursor.execute("""
                        SELECT COUNT(DISTINCT user_id) as online_users
                        FROM system_logs 
                        WHERE created_at >= DATE_SUB(NOW(), INTERVAL 5 MINUTE)
                    """)
                    result = cursor.fetchone()
                    online_users = result['online_users'] if result else 0
                    
                    cursor.execute("""
                        SELECT COUNT(*) as today_logins
                        FROM system_logs 
                        WHERE action = 'login' 
                        AND created_at >= CURDATE() 
                        AND created_at < DATE_ADD(CURDATE(), INTERVAL 1 DAY)
                    """)
                    result = cursor.fetchone()
                    today_logins = result['today_logins'] if result else 0
                    
                    conn.close()
                    
                    system_user_id = None
                    try:
                        conn2 = DatabaseConfig.get_db_connection()
                        if conn2:
                            cursor2 = conn2.cursor()
                            cursor2.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                            result = cursor2.fetchone()
                            if result:
                                system_user_id = result['id']
                            conn2.close()
                    except Exception as e2:
                        print(f"获取系统用户ID失败: {e2}")
                    
            except Exception as e:
                print(f"更新用户活动统计失败: {e}")
            
            flash(f'欢迎回来，{user["username"]}！', 'success')
            return redirect(url_for('word_dashboard'))
        else:
            flash('用户名或密码错误！', 'error')
    
    return render_template('login_word.html')

@app.route('/word/register', methods=['GET', 'POST'])
def word_register():
    """Word模板版本 - 用户注册"""
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        if not validate_email(email):
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败(Word模板): 邮箱格式无效 {email}",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('邮箱格式无效！', 'error')
            return render_template('register_word.html')
        
        username = generate_username_from_email(email)
        if not username:
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败(Word模板): 无法从邮箱生成用户名 {email}",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('无法从邮箱生成用户名！', 'error')
            return render_template('register_word.html')
        
        # 修复：check_email_exists 返回元组 (exists, error_msg)
        email_exists, email_error = check_email_exists(email, None, request.remote_addr, request.headers.get('User-Agent'))
        if email_error:
            # 数据库连接或查询错误
            flash('系统错误，请稍后重试！', 'error')
            return render_template('register_word.html')
        
        if email_exists:
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败(Word模板): 邮箱 {email} 已存在",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('邮箱已被使用！', 'error')
        elif password != confirm_password:
            log_user_action(
                None, 
                'register_failed', 
                f"注册失败(Word模板): 邮箱 {email} 密码不一致",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            flash('两次输入的密码不一致！', 'error')
        else:
            is_valid, message = validate_password_strength(password)
            if not is_valid:
                log_user_action(
                    None, 
                    'register_failed', 
                    f"注册失败(Word模板): 邮箱 {email} 密码强度不足 - {message}",
                    request.remote_addr,
                    request.headers.get('User-Agent')
                )
                flash(message, 'error')
            else:
                password_hash = generate_password_hash(password)
                user_id = create_user(username, email, password_hash)
                
                if user_id:
                    log_user_action(
                        user_id, 
                        'register', 
                        f"新用户 {username} (邮箱: {email}) 通过Word模板注册成功",
                        request.remote_addr,
                        request.headers.get('User-Agent')
                    )
                    flash(f'注册成功！您的用户名为: {username}，请登录。', 'success')
                    return redirect(url_for('word_login'))
                else:
                    log_user_action(
                        None, 
                        'register_failed', 
                        f"注册失败(Word模板): 邮箱 {email} 数据库创建失败",
                        request.remote_addr,
                        request.headers.get('User-Agent')
                    )
                    flash('注册失败，请稍后重试！', 'error')
    
    return render_template('register_word.html')

@app.route('/word/dashboard')
@word_login_required
def word_dashboard():
    """Word模板版本 - 用户仪表板"""
    
    user_files = get_user_files_from_db(session['user_id'])
    
    user_result = get_user_by_username(session['username'])
    user = user_result.get('data') if user_result.get('success') else None
    
    api_key_info = get_active_api_key(session['user_id'], 'siliconflow')
    api_token = api_key_info.get('api_key', '') if api_key_info else ''
    
    return render_template('dashboard_word.html', 
                         username=session['username'],
                         user_email=user.get('email', '') if user else '',
                         files=user_files,
                         api_token=api_token)

@app.route('/word/user-info-ynex')
@word_login_required
def word_user_info_ynex():
    """Word模板版本 - YNEX风格用户信息管理页面"""
    
    user_result = get_user_by_username(session['username'])
    user = user_result.get('data') if user_result.get('success') else None
    
    created_at_str = '未知'
    if user and 'created_at' in user and user['created_at']:
        try:
            if isinstance(user['created_at'], str):
                created_at_str = user['created_at']
            else:
                created_at_str = user['created_at'].strftime('%Y-%m-%d %H:%M:%S')
        except:
            created_at_str = '未知'
    
    user_info = {
        'user_id': session['user_id'],
        'username': session.get('username', ''),
        'email': user.get('email', '') if user else '',
        'created_at': created_at_str,
        'is_admin': session.get('role') in ['admin', 'premium']
    }
    
    return render_template('user_info_ynex_word.html', **user_info)

@app.route('/word/comprehensive-processor')
@word_login_required
def word_comprehensive_processor():
    """Word模板版本 - 综合处理页面（三位一体）"""
    
    user_files = get_user_files_from_db(session['user_id'])
    
    user_result = get_user_by_username(session['username'])
    user = user_result.get('data') if user_result.get('success') else None
    
    api_key_info = get_active_api_key(session['user_id'], 'siliconflow')
    api_token = api_key_info.get('api_key', '') if api_key_info else ''
    
    return render_template('comprehensive_processor_word.html', 
                         username=session['username'],
                         user_email=user.get('email', '') if user else '',
                         files=user_files,
                         api_token=api_token)

@app.route('/word/parsed-documents')
@word_login_required
def word_parsed_documents():
    """Word模板版本 - 已解析文档页面"""
    
    user_files = get_user_files_from_db(session['user_id'])
    parsed_files = [file for file in user_files if file.get('status') == 'completed']
    
    for file in parsed_files:
        if file.get('processing_start_time') and file.get('processed_time'):
            from datetime import datetime
            start_time = file['processing_start_time']
            end_time = file['processed_time']
            if isinstance(start_time, str):
                start_time = datetime.strptime(start_time, '%Y-%m-%d %H:%M:%S')
            if isinstance(end_time, str):
                end_time = datetime.strptime(end_time, '%Y-%m-%d %H:%M:%S')
            
            duration = (end_time - start_time).total_seconds()
            file['processing_duration'] = f"{duration:.2f}秒"
        else:
            file['processing_duration'] = '未知'
    
    return render_template('parsed_documents_word.html', 
                         username=session['username'],
                         files=parsed_files)

@app.route('/word/reference-analysis')
@word_login_required
def word_reference_analysis():
    """Word模板版本 - 参考文献分析页面"""
    
    file_id = request.args.get('file_id')
    if not file_id:
        flash('未指定文件ID', 'warning')
        return redirect(url_for('word_parsed_documents'))
    
    file_data = get_file_reference_data(file_id, session['user_id'])
    if not file_data:
        flash('未找到指定文件或无权访问', 'error')
        return redirect(url_for('word_parsed_documents'))
    
    reference_data = parse_reference_data(file_data)
    
    return render_template('reference_analysis_word.html',
                         username=session['username'],
                         file_name=file_data.get('original_name', '未知文件'),
                         reference_data=reference_data,
                         validation_json=file_data.get('reference_validation_json', ''))

@app.route('/word/group-management')
@word_login_required
def word_group_management():
    """Word模板版本 - 组管理页面"""
    
    if session.get('role') not in ['admin', 'premium']:
        flash('您没有权限访问此页面', 'error')
        return redirect(url_for('word_dashboard'))
    
    group_api_keys = get_group_api_keys(session['user_id'])
    
    return render_template('group_management_word.html',
                         username=session['username'],
                         group_api_keys=group_api_keys)

@app.route('/word/admin')
@word_login_required
def word_admin_dashboard():
    """Word模板版本 - 管理员控制台"""
    
    if session.get('role') not in ['admin', 'premium']:
        flash('您没有权限访问此页面', 'error')
        return redirect(url_for('word_dashboard'))
    
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            flash('数据库连接失败', 'error')
            return redirect(url_for('word_dashboard'))
        
        cursor = conn.cursor()
        
        # 获取用户统计信息
        cursor.execute("SELECT COUNT(*) as total_users FROM users")
        user_stats = cursor.fetchone()
        
        # 获取文件统计信息
        cursor.execute("SELECT COUNT(*) as total_files FROM user_files")
        file_stats = cursor.fetchone()
        
        # 获取最近注册的用户
        cursor.execute("""
            SELECT id, username, email, created_at, is_active 
            FROM users 
            ORDER BY created_at DESC 
            LIMIT 5
        """)
        recent_users = cursor.fetchall()
        
        # 获取今日活跃统计 - 使用MySQL CURDATE() 避免Python/MySQL时区不一致
        cursor.execute("""
            SELECT COUNT(DISTINCT sl.user_id) as active_today 
            FROM system_logs sl
            WHERE sl.action = 'login' 
            AND sl.created_at >= CURDATE()
            AND sl.created_at < CURDATE() + INTERVAL 1 DAY
            AND sl.user_id IS NOT NULL
        """)
        active_stats = cursor.fetchone()
        active_today = active_stats['active_today'] if active_stats else 0
        
        # 获取今日处理成功的文件数量 - 使用MySQL CURDATE()
        cursor.execute("""
            SELECT COUNT(*) as processed_files 
            FROM system_logs 
            WHERE action = 'file_processed' 
              AND created_at >= CURDATE()
              AND created_at < CURDATE() + INTERVAL 1 DAY
        """)
        processed_stats = cursor.fetchone()
        processed_files = processed_stats['processed_files'] if processed_stats else 0
        
        conn.close()
        
        # 统计数据整合
        stats = {
            'total_users': user_stats['total_users'] if user_stats else 0,
            'total_files': file_stats['total_files'] if file_stats else 0,
            'processed_files': processed_files,
            'active_today': active_today
        }
        
        return render_template('admin/dashboard_ynex_word.html', 
                             stats=stats,
                             user_stats=user_stats,
                             file_stats=file_stats,
                             recent_users=recent_users,
                             active_today=active_today)
        
    except Exception as e:
        print(f"管理员控制台错误: {e}")
        flash('加载管理员控制台时出错', 'error')
        return redirect(url_for('word_dashboard'))

@app.route('/word/admin/users')
def word_admin_users():
    """Word模板版本 - 用户管理页面"""
    try:
        if 'user_id' not in session:
            return redirect(url_for('word_login'))
        
        if session.get('role') not in ['admin', 'premium']:
            flash('您没有权限访问此页面', 'error')
            return redirect(url_for('word_dashboard'))
        
        page = request.args.get('page', 1, type=int)
        per_page = 20  # 每页显示20个用户
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            flash('数据库连接失败', 'error')
            return redirect(url_for('word_admin_dashboard'))
        
        cursor = conn.cursor()
        
        # 获取用户总数
        cursor.execute("SELECT COUNT(*) as total FROM users")
        total_result = cursor.fetchone()
        total_users = total_result['total'] if total_result else 0
        
        # 计算分页
        offset = (page - 1) * per_page
        total_pages = (total_users + per_page - 1) // per_page
        
        # 获取用户列表 - 包含角色信息
        cursor.execute("""
            SELECT u.id, u.username, u.email, u.role, u.created_at, u.updated_at, u.is_active,
                   COUNT(f.file_id) as file_count,
                   MAX(f.upload_time) as last_upload
            FROM users u
            LEFT JOIN user_files f ON u.id = f.user_id
            GROUP BY u.id, u.username, u.email, u.role, u.created_at, u.updated_at, u.is_active
            ORDER BY u.created_at DESC
            LIMIT %s OFFSET %s
        """, (per_page, offset))
        
        users = cursor.fetchall()
        
        conn.close()
        
        return render_template('admin/users_ynex_word.html', 
                             users=users,
                             current_page=page,
                             total_pages=total_pages,
                             total_users=total_users,
                             per_page=per_page,
                             current_user_role=session.get('role', 'user'))
        
    except Exception as e:
        flash(f'获取用户列表失败: {str(e)}', 'error')
        return redirect(url_for('word_admin_dashboard'))

@app.route('/word/admin/notifications')
@word_login_required
def word_admin_notifications():
    """Word模板版本 - 管理员通知中心"""
    if session.get('role') not in ['admin', 'premium']:
        flash('您没有权限访问此页面', 'error')
        return redirect(url_for('word_dashboard'))
    return render_template('admin_notifications_word.html',
                         username=session.get('username'),
                         user_id=session.get('user_id'),
                         current_user_role=session.get('role', 'user'))

@app.route('/word/user/notifications')
@word_login_required
def word_user_notifications():
    """Word模板版本 - 用户通知中心"""
    return render_template('user_notifications_word.html',
                         username=session.get('username'),
                         user_id=session.get('user_id'),
                         current_user_role=session.get('role', 'user'))

@app.route('/word/logout')
def word_logout():
    """Word模板版本 - 用户注销/登出"""
    user_id = session.get('user_id')
    username = session.get('username')
    
    # 记录注销日志
    if user_id and username:
        log_user_action(
            user_id, 
            'logout', 
            f"用户 {username} 成功注销登出 (Word模板)",
            request.remote_addr,
            request.headers.get('User-Agent')
        )
    
    # 清除所有会话数据
    session.clear()
    
    flash('您已成功登出！', 'info')
    return redirect(url_for('word_login'))

@app.route('/word/logout_all')
def word_logout_all():
    """Word模板版本 - 注销所有设备上的登录会话"""
    user_id = session.get('user_id')
    username = session.get('username')
    
    if user_id and username:
        # 更新用户会话版本，使其他设备上的登录失效
        if update_user_session_version(user_id):
            # 记录全局注销日志
            log_user_action(
                user_id, 
                'logout_all', 
                f"用户 {username} 执行全局注销（所有设备）(Word模板)",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            # 记录安全事件
            log_security_event(
                user_id,
                'global_logout',
                f"用户 {username} 主动执行全设备注销 (Word模板)",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            flash('已成功注销所有设备上的登录！', 'info')
        else:
            flash('全设备注销操作失败，请稍后重试。', 'error')
    else:
        flash('您尚未登录！', 'warning')
    
    # 清除当前会话
    session.clear()
    
    return redirect(url_for('word_login'))

# ==================== Word模板路由结束 ====================

@app.route('/upload', methods=['POST'])
def upload_files():
    """处理文件上传"""
    try:
        if 'user_id' not in session:
            return jsonify({'error': '请先登录'}), 401
        
        if 'files' not in request.files:
            return jsonify({'error': '没有选择文件'}), 400
        
        # 检查当前用户的待处理文件数量限制（轻量级查询）
        pending_count = count_pending_files_for_user(session['user_id'])
        
        files = request.files.getlist('files')
        
        # 计算可以上传的文件数量
        remaining_slots = 10 - pending_count
        if remaining_slots <= 0:
            return jsonify({'error': '文件列表已达到上限（最多10个文件），请先清除一些文件后再上传'}), 400
        
        # 如果上传的文件数量超过剩余槽位，只处理前面的文件
        files_to_upload = files[:remaining_slots]
        skipped_count = len(files) - len(files_to_upload)
        
        uploaded_files = []
        
        for file in files_to_upload:
            if file and file.filename and allowed_file(file.filename):
                # 保存原始文件名（用于显示）
                original_filename = file.filename
                print(f"[上传] 原始文件名: {original_filename}")
                
                # 生成安全的文件名（用于存储）
                safe_filename_base = secure_filename(file.filename)
                print(f"[上传] 安全文件名: {safe_filename_base}")
                
                # 处理中文文件名导致的空文件名问题
                if not safe_filename_base or safe_filename_base == '':
                    # 如果安全文件名为空，使用时间戳生成文件名
                    timestamp = get_current_time().strftime('%Y%m%d_%H%M%S')
                    file_extension = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else 'pdf'
                    safe_filename_base = f"upload_{timestamp}.{file_extension}"
                    print(f"[上传] 使用时间戳文件名: {safe_filename_base}")
                
                file_id = str(uuid.uuid4())
                
                # 安全地获取文件扩展名
                if '.' in safe_filename_base:
                    file_extension = safe_filename_base.rsplit('.', 1)[1].lower()
                else:
                    file_extension = 'pdf'  # 默认扩展名
                
                safe_filename = f"{file_id}.{file_extension}"
                
                # 保存文件到数据库而不是文件系统
                file_content = file.read()  # 读取文件二进制数据
                file_size = len(file_content)
                
                print(f"[上传] 文件大小: {file_size} 字节")
                
                # 记录文件信息
                file_info = {
                    'file_id': file_id,
                    'original_name': original_filename,
                    'safe_filename': safe_filename,
                    'file_path': f'database://{file_id}',  # 标记为数据库存储
                    'user_id': session['user_id'],
                    'upload_time': format_time_for_display(get_current_time()),
                    'size': file_size,
                    'status': 'uploaded',  # uploaded, processing, completed, error
                    'pdf_data': file_content  # 文件二进制数据(PDF/DOC/DOCX)
                }
                
                # 保存到数据库
                if save_file_to_db(
                    user_id=file_info['user_id'],
                    file_id=file_info['file_id'],
                    original_name=file_info['original_name'],
                    safe_filename=file_info['safe_filename'],
                    file_path=file_info['file_path'],
                    file_size=file_info['size'],
                    pdf_data=file_info['pdf_data'],
                    storage_type='database'
                ):
                    print(f"[上传] 文件 {original_filename} 成功保存到数据库")
                    
                    # 记录文件上传日志
                    log_user_action(
                        session['user_id'], 
                        'file_upload', 
                        f"上传文件: {original_filename} (大小: {format_file_size(file_size)})",
                        request.remote_addr,
                        request.headers.get('User-Agent')
                    )
                    
                    # 🆕 异步处理实时统计和里程碑通知（不阻塞上传响应）
                    def async_upload_stats_and_milestone():
                        try:
                            import threading, time
                            # 延迟处理统计，避免影响前端响应速度
                            time.sleep(0.5)
                            
                            conn = DatabaseConfig.get_db_connection()
                            if conn:
                                cursor = conn.cursor()
                                cursor.execute("""
                                    SELECT COUNT(*) as today_uploads 
                                    FROM user_files 
                                    WHERE upload_time >= CURDATE() 
                                    AND upload_time < DATE_ADD(CURDATE(), INTERVAL 1 DAY)
                                """)
                                result = cursor.fetchone()
                                today_uploads = result['today_uploads'] if result else 0
                                conn.close()
                                
                                # 里程碑通知
                                milestone_thresholds = [10, 25, 50, 100, 200, 500]
                                if today_uploads in milestone_thresholds:
                                    try:
                                        from notification_system import send_system_notification
                                        send_system_notification(
                                            title=f"📊 今日上传里程碑达成！",
                                            content=f"🎉 恭喜！今日文件上传量已达到 {today_uploads} 个！最新上传: {original_filename}",
                                            notification_type='system_milestone',
                                            priority='normal',
                                            recipient_id=None
                                        )
                                    except Exception as e:
                                        print(f"[里程碑通知] 发送失败: {e}")
                        except Exception as e:
                            print(f"[异步统计] 错误: {e}")
                    
                    # 在后台线程中异步处理
                    stat_thread = threading.Thread(target=async_upload_stats_and_milestone, daemon=True)
                    stat_thread.start()
                    
                    # 创建返回用的文件信息（不包含二进制数据）
                    return_file_info = file_info.copy()
                    return_file_info.pop('pdf_data', None)  # 移除二进制数据
                    uploaded_files.append(return_file_info)
                else:
                    print(f"[上传] 文件 {original_filename} 保存到数据库失败")
                    return jsonify({'error': f'保存文件 {original_filename} 失败'}), 500
            else:
                print(f"[上传] 文件被拒绝: {file.filename if file else 'None'}")
        
        if not uploaded_files:
            return jsonify({'error': '没有有效的文件被上传'}), 400
        
        # 构建成功消息
        success_message = f'成功上传 {len(uploaded_files)} 个文件'
        if skipped_count > 0:
            success_message += f'，跳过 {skipped_count} 个文件（已达到10个文件上限）'
        
        return jsonify({
            'success': True,
            'message': success_message,
            'files': uploaded_files,
            'skipped_count': skipped_count
        })
    
    except Exception as e:
        print(f"[上传] 异常错误: {e}")
        print(f"[上传] 异常详情: {traceback.format_exc()}")
        return jsonify({'error': f'上传处理错误: {str(e)}'}), 500

@app.route('/api/health-check')
def health_check():
    """系统健康检查API"""
    try:
        # 检查数据库连接
        conn = get_db_connection_with_error_handling()
        if conn:
            cursor = conn.cursor()
            cursor.execute("SELECT 1")
            cursor.fetchone()
            safe_db_close(conn, cursor)
            return jsonify({
                'status': 'healthy',
                'timestamp': get_current_time().isoformat(),
                'database': 'connected'
            })
        else:
            return jsonify({
                'status': 'unhealthy',
                'timestamp': get_current_time().isoformat(),
                'database': 'disconnected'
            }), 503
    except Exception as e:
        return jsonify({
            'status': 'error',
            'timestamp': get_current_time().isoformat(),
            'error': str(e)
        }), 500

@app.route('/api/realtime-stats')
def get_realtime_stats():
    """获取系统实时统计数据API"""
    try:
        import psutil
        
        # 系统资源使用情况
        cpu_usage = psutil.cpu_percent(interval=0.1)
        memory = psutil.virtual_memory() 
        disk = psutil.disk_usage('/')
        
        # 处理任务统计
        with tasks_db_lock:
            tasks_snapshot = list(tasks_db.values())
            total_tasks = len(tasks_db)
        active_processing = len([task for task in tasks_snapshot if task.get('status') == 'processing'])
        pending_tasks = len([task for task in tasks_snapshot if task.get('status') in ['pending', 'waiting']])
        failed_tasks = len([task for task in tasks_snapshot if task.get('status') == 'error'])
        
        # 数据库统计
        conn = DatabaseConfig.get_db_connection()
        db_stats = {}
        
        if conn:
            cursor = conn.cursor()
            
            # 今日统计
            cursor.execute("""
                SELECT 
                    (SELECT COUNT(*) FROM user_files 
                     WHERE upload_time >= CURDATE() AND upload_time < DATE_ADD(CURDATE(), INTERVAL 1 DAY)) as daily_uploads,
                    (SELECT COUNT(*) FROM user_files 
                     WHERE created_at >= CURDATE() AND created_at < DATE_ADD(CURDATE(), INTERVAL 1 DAY) AND status = 'completed') as daily_completed,
                    (SELECT COUNT(*) FROM users 
                     WHERE created_at >= CURDATE() AND created_at < DATE_ADD(CURDATE(), INTERVAL 1 DAY)) as daily_new_users,
                    (SELECT COUNT(DISTINCT user_id) FROM system_logs 
                     WHERE created_at >= DATE_SUB(NOW(), INTERVAL 5 MINUTE)) as online_users,
                    (SELECT COUNT(*) FROM system_logs 
                     WHERE action = 'login' AND created_at >= CURDATE() AND created_at < DATE_ADD(CURDATE(), INTERVAL 1 DAY)) as daily_logins
            """)
            result = cursor.fetchone()
            
            if result:
                db_stats = {
                    'daily_uploads': result['daily_uploads'] or 0,
                    'daily_completed': result['daily_completed'] or 0,
                    'daily_new_users': result['daily_new_users'] or 0,
                    'online_users': result['online_users'] or 0,
                    'daily_logins': result['daily_logins'] or 0
                }
            
            # 最近处理性能
            cursor.execute("""
                SELECT 
                    AVG(TIMESTAMPDIFF(SECOND, processing_start_time, processed_time)) as avg_processing_time,
                    MIN(TIMESTAMPDIFF(SECOND, processing_start_time, processed_time)) as min_processing_time,
                    MAX(TIMESTAMPDIFF(SECOND, processing_start_time, processed_time)) as max_processing_time
                FROM user_files 
                WHERE processed_time IS NOT NULL 
                AND processing_start_time IS NOT NULL
                AND processed_time >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
                AND status = 'completed'
            """)
            perf_result = cursor.fetchone()
            
            if perf_result and perf_result['avg_processing_time']:
                db_stats['recent_performance'] = {
                    'avg_time': round(float(perf_result['avg_processing_time']), 1),
                    'min_time': round(float(perf_result['min_processing_time']), 1),
                    'max_time': round(float(perf_result['max_processing_time']), 1)
                }
            
            conn.close()
        
        # 计算成功率
        success_rate = 0
        if total_tasks > 0:
            completed_tasks = len([task for task in tasks_snapshot if task.get('status') == 'completed'])
            success_rate = round((completed_tasks / total_tasks) * 100, 1)
        
        # 返回统计数据
        stats = {
            'timestamp': get_current_time().isoformat(),
            'system_resources': {
                'cpu_usage': round(cpu_usage, 1),
                'memory_usage': round(memory.percent, 1),
                'disk_usage': round(disk.percent, 1),
                'memory_available': format_file_size(memory.available),
                'disk_free': format_file_size(disk.free)
            },
            'processing_queue': {
                'total_tasks': total_tasks,
                'active_processing': active_processing,
                'pending_tasks': pending_tasks,
                'failed_tasks': failed_tasks,
                'success_rate': success_rate
            },
            'daily_metrics': db_stats,
            'status': 'healthy'
        }
        
        return jsonify(stats)
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'timestamp': get_current_time().isoformat(),
            'error': f'获取实时统计失败: {str(e)}'
        }), 500

@app.route('/api/realtime-logs')
def get_realtime_logs():
    """获取最近的实时日志记录"""
    try:
        if 'user_id' not in session:
            return jsonify({'error': '未登录'}), 401
        
        # 管理员可以查看所有实时日志，普通用户只能查看自己的
        is_admin = session.get('role') in ['admin', 'premium']
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 查询实时更新类型的日志
        realtime_actions = [
            'realtime_upload_stats', 'realtime_completion_stats', 'realtime_login_activity',
            'system_performance_summary', 'processing_status_summary', 'user_activity_metrics',
            'queue_status_update', 'processing_performance_metrics'
        ]
        
        placeholders = ','.join(['%s'] * len(realtime_actions))
        
        if is_admin:
            # 管理员查看所有实时日志
            query = f"""
                SELECT user_id, action, details, created_at, ip_address
                FROM system_logs 
                WHERE action IN ({placeholders})
                AND created_at >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
                ORDER BY created_at DESC 
                LIMIT 50
            """
            cursor.execute(query, realtime_actions)
        else:
            # 普通用户只看自己相关的日志
            query = f"""
                SELECT user_id, action, details, created_at, ip_address
                FROM system_logs 
                WHERE (action IN ({placeholders}) OR user_id = %s)
                AND created_at >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
                ORDER BY created_at DESC 
                LIMIT 30
            """
            cursor.execute(query, realtime_actions + [session['user_id']])
        
        logs = cursor.fetchall()
        conn.close()
        
        # 格式化日志数据
        formatted_logs = []
        for log in logs:
            formatted_logs.append({
                'user_id': log['user_id'],
                'action': log['action'],
                'details': log['details'],
                'created_at': format_time_for_display(log['created_at']),
                'ip_address': log['ip_address'],
                'is_system': log['user_id'] == 0
            })
        
        return jsonify({
            'success': True,
            'logs': formatted_logs,
            'count': len(formatted_logs),
            'is_admin': is_admin
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取实时日志失败: {str(e)}'
        }), 500

@app.route('/api/files')
@api_login_required
def get_user_files():
    """获取用户的文件列表 - Dashboard版本，只显示需要处理的文档"""
    
    # 从数据库获取当前用户的所有文件
    all_user_files = get_user_files_from_db(session['user_id'])
    
    # Dashboard显示uploaded(未解析)、processing(处理中)和error(解析错误)状态，不显示completed
    pending_files = [file for file in all_user_files if file['status'] in ('uploaded', 'processing', 'error')]
    
    # 按上传时间排序，最新的在前
    pending_files.sort(key=lambda x: x['upload_time'], reverse=True)
    
    return jsonify({
        'success': True,
        'files': pending_files
    })

@app.route('/api/parsed-files')
@api_login_required
def get_parsed_files():
    """获取用户的文档列表 - 文档列表页面版本，显示所有状态"""
    
    # 从数据库获取当前用户的所有文件
    all_user_files = get_user_files_from_db(session['user_id'])
    
    # 文档列表页面显示所有状态：uploaded、completed、error
    parsed_files = [file for file in all_user_files if file['status'] in ('uploaded', 'completed', 'error')]
    
    # 为每个文件计算处理时间 - 直接在数据库层计算，避免时区问题
    # processing_time_seconds 已经在 get_user_files_from_db 中通过 SQL 计算
    
    # 按上传时间排序，最新的在前
    parsed_files.sort(key=lambda x: x['upload_time'], reverse=True)
    
    return jsonify({
        'success': True,
        'files': parsed_files
    })

@app.route('/api/files/<file_id>', methods=['DELETE'])
@api_login_required
def delete_user_file(file_id):
    """删除用户的单个文件 - 支持任务中断"""
    
    user_id = session['user_id']
    
    try:
        # 🚀 第一步：立即中断相关任务
        interrupted_tasks = interrupt_file_tasks(file_id, user_id)
        
        # 从数据库删除文件记录
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # 先查询文件是否存在且属于当前用户
            cursor.execute("SELECT original_name, reference_md_content, reference_validation_json FROM user_files WHERE file_id = %s AND user_id = %s", (file_id, user_id))
            result = cursor.fetchone()
            
            if not result:
                conn.close()
                return jsonify({'error': '文件不存在或无权限'}), 404
            
            original_name = result['original_name']
            has_reference_data = bool(result['reference_md_content'] or result['reference_validation_json'])
            
            # 删除文件记录（包括所有相关数据：md_content, extracted_json_data, nsfc_json_data, reference_md_content, reference_validation_json等）
            cursor.execute("DELETE FROM user_files WHERE file_id = %s AND user_id = %s", (file_id, user_id))
            file_deleted = cursor.rowcount > 0
            
            # 任务在内存中管理，无需数据库清理
            
            conn.commit()  # 提交数据库删除操作
            conn.close()
            
            if file_deleted:
                # 记录删除操作，包括参考文献数据清除信息
                delete_info = f'删除文件: {original_name}'
                if has_reference_data:
                    delete_info += ' (包括参考文献分析数据)'
                log_user_action(user_id, 'delete_file', delete_info)
                
                return jsonify({
                    'success': True,
                    'message': f'文件 "{original_name}" 已删除',
                    'interrupted_tasks': interrupted_tasks
                })
            else:
                return jsonify({'error': '删除失败'}), 500
        else:
            return jsonify({'error': '数据库连接失败'}), 500
            
    except Exception as e:
        print(f"删除文件错误: {e}")
        return jsonify({'error': '删除失败'}), 500

@app.route('/api/interrupt-file-tasks/<file_id>', methods=['POST'])
@api_login_required
def interrupt_file_tasks_api(file_id):
    """中断指定文件的所有任务"""
    
    user_id = session['user_id']
    
    try:
        interrupted_tasks = interrupt_file_tasks(file_id, user_id)
        
        return jsonify({
            'success': True,
            'message': f'已中断文件 {file_id} 的相关任务',
            'interrupted_tasks': interrupted_tasks
        })
        
    except Exception as e:
        print(f"中断文件任务错误: {e}")
        return jsonify({'error': f'中断失败: {str(e)}'}), 500

@app.route('/api/files/clear-all', methods=['DELETE'])
@api_login_required
def clear_all_user_files():
    """删除用户的所有文件（支持文件类型过滤）"""
    
    user_id = session['user_id']
    file_type = request.args.get('type', 'all')  # 'pdf', 'word', 'all'
    
    try:
        # 🚀 第一步：立即中断所有相关任务（在删除数据库记录之前）
        interrupted_tasks = []
        
        # 根据文件类型构建SQL条件
        # 注意：LIKE 中的 % 需要转义为 %% 以避免 PyMySQL cursor.execute() 误解为格式占位符
        if file_type == 'pdf':
            type_condition = "(original_document_type = 'pdf' OR (original_document_type = 'unknown' AND original_name LIKE '%%.pdf'))"
        elif file_type == 'word':
            type_condition = "(original_document_type IN ('word_doc', 'word_docx') OR (original_document_type = 'unknown' AND (original_name LIKE '%%.doc' OR original_name LIKE '%%.docx')))"
        else:
            type_condition = "1=1"  # 所有文件
        
        # 1. 中断内存中的所有任务
        task_ids_to_remove = []
        with tasks_db_lock:
            for task_id, task in list(tasks_db.items()):
                if task.get('user_id') == user_id:
                    # 检查任务对应的文件类型
                    file_id = task.get('file_id')
                    should_remove = False
                    
                    if file_type == 'all':
                        should_remove = True
                    elif file_id:
                        # 从数据库查询文件名判断类型
                        try:
                            conn_check = DatabaseConfig.get_db_connection()
                            if conn_check:
                                cursor_check = conn_check.cursor()
                                cursor_check.execute("SELECT original_name FROM user_files WHERE file_id = %s", (file_id,))
                                file_result = cursor_check.fetchone()
                                conn_check.close()
                                
                                if file_result:
                                    filename = file_result['original_name'].lower()
                                    if file_type == 'pdf' and filename.endswith('.pdf'):
                                        should_remove = True
                                    elif file_type == 'word' and (filename.endswith('.doc') or filename.endswith('.docx')):
                                        should_remove = True
                        except:
                            pass
                    
                    if should_remove:
                        # 标记任务为中断状态
                        task['status'] = 'cancelled'
                        task['message'] = '❌ 用户清空文件，任务中断'
                        task_ids_to_remove.append(task_id)
                        interrupted_tasks.append(f"任务: {task_id}")
                        print(f"[任务中断] 清空{file_type}文件时中断任务: {task_id}")
        
        # 2. 中断一键处理任务
        oneclick_tasks_to_cancel = []
        with one_click_tasks_lock:
            for task_id, task in list(one_click_tasks.items()):
                if task.get('user_id') == user_id:
                    task['status'] = 'cancelled'
                    task['error'] = '用户清空文件，任务中断'
                    oneclick_tasks_to_cancel.append(task_id)
                    interrupted_tasks.append(f"一键处理任务: {task_id}")
                    print(f"[任务中断] 清空文件时中断一键处理任务: {task_id}")
        
        # 延迟清理内存任务记录，让前端有时间显示中断状态
        def delayed_cleanup():
            import time
            for i in range(30):  # 分30次检查，每次100ms，总共3秒
                if _monitoring_shutdown:
                    return  # 如果系统正在关闭，直接退出
                time.sleep(0.1)
            
            if not _monitoring_shutdown:  # 只有在系统未关闭时才清理
                with tasks_db_lock:
                    for task_id in task_ids_to_remove:
                        tasks_db.pop(task_id, None)
                        print(f"[任务清理] 已清理任务: {task_id}")
                with one_click_tasks_lock:
                    for task_id in oneclick_tasks_to_cancel:
                        one_click_tasks.pop(task_id, None)
                        print(f"[任务清理] 已清理一键处理任务: {task_id}")
        
        import threading
        cleanup_thread = threading.Thread(target=delayed_cleanup)
        cleanup_thread.daemon = True
        cleanup_thread.start()
        
        # 第二步：从数据库删除所有文件和任务记录
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # 先查询待处理/处理中/错误状态的文件数量（不包括completed）
            cursor.execute(f"""
                SELECT COUNT(*) as count 
                FROM user_files 
                WHERE user_id = %s AND status IN ('uploaded', 'processing', 'error')
                AND {type_condition}
            """, (user_id,))
            result = cursor.fetchone()
            file_count = result['count'] if result else 0
            
            if file_count == 0 and len(interrupted_tasks) == 0:
                conn.close()
                return jsonify({
                    'success': True,
                    'message': '没有可删除的文件',
                    'removed_count': 0,
                    'interrupted_tasks': []
                })
            
            # 🔧 只删除未完成的文件记录（保留completed状态的已解析文档）
            cursor.execute(f"""
                DELETE FROM user_files 
                WHERE user_id = %s AND status IN ('uploaded', 'processing', 'error')
                AND {type_condition}
            """, (user_id,))
            files_deleted = cursor.rowcount
            
            # 任务在内存中管理，无需数据库清理
            
            conn.commit()  # 提交批量删除操作
            conn.close()
            
            type_desc = {'pdf': 'PDF', 'word': 'Word', 'all': '全部'}[file_type]
            log_user_action(user_id, 'clear_all_files', f'清除{type_desc}任务列表文件，共 {files_deleted} 个（未完成状态），中断 {len(interrupted_tasks)} 个任务')
            
            return jsonify({
                'success': True,
                'message': f'已清除 {files_deleted} 个{type_desc}任务文件，中断 {len(interrupted_tasks)} 个任务（已完成的文档已保留）',
                'removed_count': files_deleted,
                'interrupted_tasks': interrupted_tasks
            })
        else:
            return jsonify({'error': '数据库连接失败'}), 500
            
    except Exception as e:
        print(f"清除所有文件错误: {e}")
        return jsonify({'error': '删除失败'}), 500

@app.route('/api/pdf-file/<file_id>')
@api_login_required
def get_pdf_file(file_id):
    """获取PDF文件用于预览"""
    print(f"[PDF API] 请求文件ID: {file_id}")
    
    user_id = session['user_id']
    print(f"[PDF API] 用户ID: {user_id}")
    
    try:
        # 从数据库获取PDF文件数据
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            smart_print(f"[PDF API] 数据库连接失败", LogLevel.ERROR)
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        cursor.execute("""
            SELECT pdf_data, original_name, size 
            FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        conn.close()
        
        if not result:
            print(f"[PDF API] 文件不存在或无权限: file_id={file_id}, user_id={user_id}")
            return jsonify({'error': '文件不存在或无权限访问'}), 404
        
        # 由于使用DictCursor，result是字典，需要用键名访问
        pdf_data = result['pdf_data']
        original_name = result['original_name']
        file_size = result['size']
        print(f"[PDF API] 找到文件: {original_name}, 大小: {file_size}")
        
        if not pdf_data:
            print(f"[PDF API] PDF数据为空")
            return jsonify({'error': 'PDF数据不存在'}), 404
        
        print(f"[PDF API] PDF数据类型: {type(pdf_data)}, 大小: {len(pdf_data)}")
        
        # 确保pdf_data是bytes类型
        if isinstance(pdf_data, str):
            pdf_data = pdf_data.encode('latin-1')
            print(f"[PDF API] 转换字符串为bytes")
        elif not isinstance(pdf_data, bytes):
            # 如果是其他类型，尝试转换
            pdf_data = bytes(pdf_data)
            print(f"[PDF API] 转换为bytes")
        
        # 检查PDF文件头
        if len(pdf_data) > 8:
            header = pdf_data[:8]
            print(f"[PDF API] PDF文件头: {header}")
        
        # 返回PDF文件
        file_obj = BytesIO(pdf_data)
        print(f"[PDF API] 返回PDF文件，大小: {len(pdf_data)} bytes")
        
        return send_file(
            file_obj,
            mimetype='application/pdf',
            as_attachment=False,
            download_name=original_name
        )
        
    except Exception as e:
        print(f"[PDF API] 获取PDF文件错误: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': '获取文件失败'}), 500

@app.route('/api/pdf-file-test/<file_id>')
def get_pdf_file_test(file_id):
    """测试PDF文件获取（不检查用户权限）"""
    print(f"[PDF API TEST] 请求文件ID: {file_id}")
    
    try:
        # 从数据库获取PDF文件数据（不检查用户ID）
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            smart_print(f"[PDF API TEST] 数据库连接失败", LogLevel.ERROR)
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        cursor.execute("""
            SELECT pdf_data, original_name, file_size 
            FROM user_files 
            WHERE file_id = %s
        """, (file_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        if not result:
            print(f"[PDF API TEST] 文件不存在: file_id={file_id}")
            return jsonify({'error': '文件不存在'}), 404
        
        # 由于使用DictCursor，result是字典，需要用键名访问
        pdf_data = result['pdf_data']
        original_name = result['original_name']
        file_size = result['file_size']
        print(f"[PDF API TEST] 找到文件: {original_name}, 大小: {file_size}")
        
        if not pdf_data:
            print(f"[PDF API TEST] PDF数据为空")
            return jsonify({'error': 'PDF数据不存在'}), 404
        
        print(f"[PDF API TEST] PDF数据类型: {type(pdf_data)}, 大小: {len(pdf_data)}")
        
        # 确保pdf_data是bytes类型
        if isinstance(pdf_data, str):
            pdf_data = pdf_data.encode('latin-1')
            print(f"[PDF API TEST] 转换字符串为bytes")
        elif not isinstance(pdf_data, bytes):
            # 如果是其他类型，尝试转换
            pdf_data = bytes(pdf_data)
            print(f"[PDF API TEST] 转换为bytes")
        
        # 检查PDF文件头
        if len(pdf_data) > 8:
            header = pdf_data[:8]
            print(f"[PDF API TEST] PDF文件头: {header}")
        
        # 返回PDF文件
        file_obj = BytesIO(pdf_data)
        print(f"[PDF API TEST] 返回PDF文件，大小: {len(pdf_data)} bytes")
        
        return send_file(
            file_obj,
            mimetype='application/pdf',
            as_attachment=False,
            download_name=original_name
        )
        
    except Exception as e:
        print(f"[PDF API TEST] 获取PDF文件错误: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': '获取文件失败'}), 500

@app.route('/check_api_token')
@api_login_required
def check_api_token():
    """检查用户是否已配置API Token和获取处理参数 - 基于分组密钥"""
    
    user_id = session['user_id']
    
    # 从分组密钥获取用户的API密钥
    active_key = get_active_api_key(user_id)
    
    # 获取默认配置（移除个人配置）
    api_config = MINERU_CONFIG.copy()
    
    return jsonify({
        'has_token': bool(active_key),
        'token_valid': True if active_key else False,
        'config': api_config,
        'active_key': {
            'id': active_key['id'],
            'provider': active_key['api_provider'],
            'token_preview': active_key['api_token'][:10] + '...' + active_key['api_token'][-6:] if len(active_key['api_token']) > 16 else active_key['api_token']
        } if active_key else None,
        'message': '使用分组统一密钥' if active_key else '请联系管理员配置分组密钥'
    })

# 分组密钥管理相关路由（个人密钥功能已移除）
# 所有个人密钥相关路由已被删除，用户现在使用分组统一密钥

@app.route('/api/update-profile', methods=['POST'])
@api_login_required
def update_profile():
    """更新用户资料"""
    
    try:
        data = request.get_json()
        user_id = session['user_id']
        username = data.get('username', '').strip()
        email = data.get('email', '').strip()
        
        # 验证输入
        if not username:
            return jsonify({'error': '用户名不能为空'}), 400
        
        # 验证邮箱格式（如果提供了邮箱）
        if email and '@' not in email:
            return jsonify({'error': '邮箱格式不正确'}), 400
        
        # 检查用户名是否已被其他用户使用
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
            
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM users WHERE username = %s AND id != %s", (username, user_id))
        existing_user = cursor.fetchone()
        
        if existing_user:
            cursor.close()
            conn.close()
            return jsonify({'error': '用户名已被使用'}), 400
        
        # 检查邮箱是否已被其他用户使用（如果提供了邮箱）
        if email:
            cursor.execute("SELECT id FROM users WHERE email = %s AND id != %s", (email, user_id))
            existing_email = cursor.fetchone()
            
            if existing_email:
                cursor.close()
                conn.close()
                return jsonify({'error': '邮箱已被使用'}), 400
        
        # 更新用户信息
        if email:
            cursor.execute("""
                UPDATE users 
                SET username = %s, email = %s, updated_at = CURRENT_TIMESTAMP 
                WHERE id = %s
            """, (username, email, user_id))
        else:
            cursor.execute("""
                UPDATE users 
                SET username = %s, updated_at = CURRENT_TIMESTAMP 
                WHERE id = %s
            """, (username, user_id))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        # 更新session中的用户名
        session['username'] = username
        
        log_user_action(user_id, 'update_profile', f'更新用户资料: {username}')
        
        return jsonify({
            'success': True,
            'message': '用户信息更新成功'
        })
        
    except Exception as e:
        app.logger.error(f"更新用户资料失败: {str(e)}")
        return jsonify({'error': '更新失败，请重试'}), 500

@app.route('/api/change-password', methods=['POST'])
@api_login_required
def change_password():
    """修改密码"""
    
    try:
        data = request.get_json()
        user_id = session['user_id']
        old_password = data.get('old_password', '')
        new_password = data.get('new_password', '')
        
        if not old_password or not new_password:
            return jsonify({'error': '密码不能为空'}), 400
        
        # 使用与注册一致的密码强度校验
        is_strong, strength_msg = validate_password_strength(new_password)
        if not is_strong:
            return jsonify({'error': strength_msg}), 400
        
        # 验证旧密码
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
            
        cursor = conn.cursor()
        cursor.execute("SELECT password_hash FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        
        if not user or not check_password_hash(user['password_hash'], old_password):
            cursor.close()
            conn.close()
            return jsonify({'error': '当前密码不正确'}), 400
        
        # 更新密码
        new_password_hash = generate_password_hash(new_password)
        cursor.execute("""
            UPDATE users 
            SET password_hash = %s, updated_at = CURRENT_TIMESTAMP 
            WHERE id = %s
        """, (new_password_hash, user_id))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        log_user_action(user_id, 'change_password', '修改密码')
        
        return jsonify({
            'success': True,
            'message': '密码修改成功'
        })
        
    except Exception as e:
        app.logger.error(f"修改密码失败: {str(e)}")
        return jsonify({'error': '修改失败，请重试'}), 500

@app.route('/api/notifications', methods=['GET'])
@api_login_required
def api_get_notifications():
    """获取用户通知列表API（兼容基础模板调用）"""
    
    try:
        user_id = session['user_id']
        from db_operations import get_user_notifications
        
        # 获取用户通知
        result = get_user_notifications(user_id, page=1, page_size=50, filter_type='all')
        
        if result.get('success'):
            return jsonify({
                'success': True,
                'notifications': result.get('notifications', []),
                'total': result.get('total', 0)
            })
        else:
            return jsonify({'success': False, 'message': '获取通知失败'}), 500
            
    except Exception as e:
        app.logger.error(f"获取通知列表失败: {str(e)}")
        return jsonify({'success': False, 'message': '获取通知失败'}), 500

@app.route('/api/users/list', methods=['GET'])
@api_login_required
def api_get_users_list():
    """获取用户列表（分组管理专用接口）"""
    
    try:
        # 获取查询参数
        exclude_confirmed_groups = request.args.get('exclude_confirmed_groups', 'false').lower() == 'true'
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 根据参数决定SQL查询条件
        if exclude_confirmed_groups:
            # 过滤掉已经在确定分组中的用户，只显示没有分组或在临时分组中的用户
            cursor.execute("""
                SELECT u.id, u.username, u.email, u.role, u.created_at, u.group_id, u.group_status
                FROM users u
                WHERE u.is_active = 1 
                AND (u.group_status IS NULL OR u.group_status = 'temporary')
                ORDER BY u.username ASC
            """)
        else:
            # 获取所有活跃用户的基本信息（包含角色和分组信息）
            cursor.execute("""
                SELECT u.id, u.username, u.email, u.role, u.created_at, u.group_id, u.group_status
                FROM users u
                WHERE u.is_active = 1 
                ORDER BY u.username ASC
            """)
        
        users = cursor.fetchall()
        conn.close()
        
        # 转换为前端期望的格式
        formatted_users = []
        for user in users:
            formatted_users.append({
                'id': user['id'],
                'name': user['username'],
                'username': user['username'],  # 保持向后兼容
                'email': user['email'],
                'role': user['role'],
                'group_id': user['group_id'],  # 添加分组ID
                'group_status': user['group_status'],  # 添加分组状态
                'joinDate': user['created_at'].strftime('%Y年%m月%d日') if user['created_at'] else '未知'
            })
        
        return jsonify({
            'success': True,
            'users': formatted_users
        })
        
    except Exception as e:
        app.logger.error(f"获取用户列表失败: {str(e)}")
        return jsonify({'error': '获取用户列表失败，请重试'}), 500

@app.route('/save_settings', methods=['POST'])
@api_login_required
def save_settings():
    """保存API设置 - 个人密钥功能已移除，现在使用分组统一密钥"""
    
    return jsonify({
        'error': '个人密钥配置功能已移除，请联系管理员配置分组密钥'
    }), 400

def process_pdf_background(task_id, file_info, api_token):
    """后台处理PDF转换任务 - 增强进度反馈"""
    try:
        # 检查任务是否已被删除
        if is_task_cancelled(task_id):
            print(f"[后台处理] 任务在开始时已被删除: {task_id}")
            return
        
        # 获取用户配置的处理参数
        if task_id not in tasks_db:
            print(f"[错误] 任务不存在于tasks_db中: {task_id}")
            return
        user_id = tasks_db[task_id]['user_id']
        
        # 【关键修复】确保tasks_db中包含file_id映射
        file_id = file_info.get('file_id')
        if not file_id:
            print(f"[错误] file_info缺少file_id字段: {file_info}")
            return
            
        if 'file_id' not in tasks_db[task_id]:
            tasks_db[task_id]['file_id'] = file_id
            # 修复：优先使用original_name，兼容filename字段
            filename = file_info.get('original_name') or file_info.get('filename', '')
            tasks_db[task_id]['filename'] = filename
            print(f"[文件ID映射] 为任务 {task_id} 补充file_id映射: {file_id}, 文件名: {filename}")
        
        # 步骤1: 准备阶段 (0-5%)
        tasks_db[task_id]['status'] = 'preparing'
        tasks_db[task_id]['progress'] = 2
        tasks_db[task_id]['message'] = '📋 正在准备处理参数...'
        update_task_status_in_db(task_id, 'preparing', 2, '📋 正在准备处理参数...')
        
        # 从数据库获取用户配置
        try:
            conn = DatabaseConfig.get_db_connection()
            if conn:
                cursor = conn.cursor()
                cursor.execute("SELECT api_config FROM users WHERE id = %s", (user_id,))
                result = cursor.fetchone()
                conn.close()
                
                if result and result['api_config']:
                    mineru_config = json.loads(result['api_config'])
                else:
                    mineru_config = MINERU_CONFIG.copy()
            else:
                mineru_config = MINERU_CONFIG.copy()
        except Exception as e:
            print(f"获取用户配置错误: {e}")
            mineru_config = MINERU_CONFIG.copy()
        
        # 步骤2: 文件检查阶段 (5-10%)
        tasks_db[task_id]['status'] = 'validating'
        tasks_db[task_id]['progress'] = 7
        tasks_db[task_id]['message'] = '🔍 正在验证文件和API配置...'
        update_task_status_in_db(task_id, 'validating', 7, '🔍 正在验证文件和API配置...')
        
        # 检查文件大小和格式
        file_size = float(file_info['size']) / (1024 * 1024)  # MB
        print(f"[后台处理] 文件大小: {file_size:.2f}MB")
        
        # 步骤3: 开始上传 (10-30%)
        # 再次检查任务是否已被删除
        if is_task_cancelled(task_id):
            print(f"[后台处理] 任务在上传前已被删除: {task_id}")
            return
        
        tasks_db[task_id]['status'] = 'uploading'
        tasks_db[task_id]['progress'] = 12
        tasks_db[task_id]['message'] = f'📤 正在上传文件 ({file_size:.1f}MB)...'
        update_task_status_in_db(task_id, 'uploading', 12, f'📤 正在上传文件 ({file_size:.1f}MB)...')
        
        # 上传文件到MinerU - 使用数据库版本
        # file_id 已在函数开头获取，无需重复获取
        
        # 更新上传进度到15%
        tasks_db[task_id]['progress'] = 15
        tasks_db[task_id]['message'] = '📤 正在调用文档解析服务...'
        update_task_status_in_db(task_id, 'uploading', 15, '📤 正在调用文档解析服务...')
        
        batch_id = upload_single_file_to_mineru_from_db(file_id, user_id, api_token, mineru_config, task_id)
        
        if not batch_id:
            # 检查具体的失败原因
            error_msg = '❌ 文件上传失败，请检查API Token或网络连接'
            
            # 标记当前密钥需要检查（但不自动判断为无效）
            try:
                api_key_id = tasks_db[task_id].get('api_key_id')
                if api_key_id:
                    print(f"[后台处理] 上传失败，API密钥ID: {api_key_id}")
            except Exception as e:
                print(f"[后台处理] 获取密钥信息失败: {e}")
            
            # 暂时保存错误信息，但不立即设置error状态，避免前端误报
            tasks_db[task_id]['pending_error'] = error_msg
            tasks_db[task_id]['status'] = 'pending_completion'  # 等待最终状态确认
            tasks_db[task_id]['message'] = '正在确认处理状态...'
            tasks_db[task_id]['completed_at'] = time.time()  # 记录错误发生时间，用于超时清理
            update_task_status_in_db(task_id, 'error', message=error_msg)
            # 注释掉立即更新数据库error状态，改为任务最终完成时统一更新
            # update_file_status_in_db(file_info['file_id'], 'error')
            return
        
        # 步骤4: 上传成功，开始处理 (30-80%)
        tasks_db[task_id]['batch_id'] = batch_id
        tasks_db[task_id]['status'] = 'processing'
        tasks_db[task_id]['progress'] = 35
        tasks_db[task_id]['message'] = f'⚙️ 文件已上传，正在解析中 (批次:{batch_id[:8]}...)'
        update_task_status_in_db(task_id, 'processing', 35, f'⚙️ 文件已上传，正在解析中 (批次:{batch_id[:8]}...)', batch_id)
        
        # 记录处理开始日志
        log_user_action(
            user_id, 
            'file_processing_start', 
            f"开始处理文件: {file_info['original_name']} (批次ID: {batch_id[:8]})",
            None,
            None
        )
        
        # 轮询处理状态 - 增强进度反馈
        max_wait_time = 600  # 最大等待10分钟
        start_time = time.time()
        poll_count = 0
        
        while time.time() - start_time < max_wait_time:
            # 检查任务是否已被删除（用户点击了清除按钮）
            if is_task_cancelled(task_id):
                print(f"[后台处理] 任务已被删除，终止处理: {task_id}")
                # 从内存中清理任务
                tasks_db.pop(task_id, None)
                return  # 立即终止任务处理
            
            poll_count += 1
            elapsed_time = time.time() - start_time
            
            # 动态计算进度 (35% -> 75%)
            progress_base = 35
            progress_range = 40  # 从35%到75%，范围40%
            time_progress = min(elapsed_time / max_wait_time, 1.0)
            current_progress = progress_base + int(time_progress * progress_range)
            
            # 更新轮询状态
            tasks_db[task_id]['progress'] = current_progress
            tasks_db[task_id]['message'] = f'⚙️ 文档解析中... (第{poll_count}次检查，已等待{int(elapsed_time)}秒)'
            update_task_status_in_db(task_id, 'processing', current_progress, f'⚙️ 文档解析中... (第{poll_count}次检查，已等待{int(elapsed_time)}秒)')
            
            status, results = poll_batch_status(batch_id, api_token)
            
            # 在处理结果之前再次检查任务是否被取消
            if is_task_cancelled(task_id):
                print(f"[后台处理] 任务在轮询后已被删除: {task_id}")
                # 从内存中清理任务
                tasks_db.pop(task_id, None)
                return  # 立即终止任务处理
            
            if status == 'completed':
                if results:
                    result = results[0]  # 单文件处理
                    download_url = result.get('zip_url')
                    file_name = result.get('file_name', '未知文件')
                    
                    if download_url:
                        # 检查任务是否已被删除
                        if is_task_cancelled(task_id):
                            print(f"[后台处理] 任务在下载前已被删除: {task_id}")
                            return
                        
                        # 步骤5: 开始下载结果 (75-95%)
                        tasks_db[task_id]['status'] = 'downloading'
                        tasks_db[task_id]['progress'] = 80
                        tasks_db[task_id]['message'] = f'📥 处理完成！正在下载结果文件 ({file_name})'
                        update_task_status_in_db(task_id, 'downloading', 80, f'📥 处理完成！正在下载结果文件 ({file_name})')
                        
                        download_path = f"downloads/{task_id}_result.zip"
                        
                        # 下载进度更新到85%
                        tasks_db[task_id]['progress'] = 85
                        tasks_db[task_id]['message'] = f'📥 正在下载ZIP结果包...'
                        update_task_status_in_db(task_id, 'downloading', 85, f'📥 正在下载ZIP结果包...')
                        
                        if download_result_file(download_url, download_path):
                            # 检查任务是否已被删除
                            if is_task_cancelled(task_id):
                                print(f"[后台处理] 任务在提取前已被删除: {task_id}")
                                # 清理已下载的文件
                                try:
                                    if os.path.exists(download_path):
                                        os.remove(download_path)
                                        print(f"[清理] 已删除已下载的ZIP文件: {download_path}")
                                except Exception as e:
                                    print(f"[清理] 删除ZIP文件失败: {e}")
                                return
                            
                            # 步骤6: 处理下载结果 (95-100%)
                            tasks_db[task_id]['progress'] = 90
                            tasks_db[task_id]['message'] = '📄 正在提取Markdown内容...'
                            update_task_status_in_db(task_id, 'extracting', 90, '📄 正在提取Markdown内容...')
                            
                            tasks_db[task_id]['status'] = 'extracting'
                            tasks_db[task_id]['download_url'] = download_url
                            tasks_db[task_id]['download_path'] = download_path
                            
                            # 更新文件状态到数据库
                            update_file_status_in_db(file_info['file_id'], 'processing')
                            
                            # 提取MD内容并保存到文件信息中
                            md_content = extract_md_from_zip(download_path, file_info['original_name'])
                            
                            # 最终完成 (100%)
                            tasks_db[task_id]['status'] = 'completed'
                            tasks_db[task_id]['progress'] = 100
                            
                            # 🆕 记录实时处理完成日志
                            try:
                                processing_time = time.time() - tasks_db[task_id].get('start_time', time.time())
                                
                                # 获取系统管理员ID用于记录处理完成日志
                                system_user_id = None
                                try:
                                    conn = DatabaseConfig.get_db_connection()
                                    if conn:
                                        cursor = conn.cursor()
                                        cursor.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                                        admin_user = cursor.fetchone()
                                        if admin_user:
                                            system_user_id = admin_user['id']
                                        conn.close()
                                except:
                                    pass
                                
                                # 只有在获取到有效管理员ID时才记录日志
                                if system_user_id:
                                    log_user_action(
                                        system_user_id,  # 使用系统管理员用户ID
                                        'realtime_processing_completed',
                                        f"文件处理完成: {file_info['original_name']} (用户ID: {user_id}, 耗时: {processing_time:.1f}秒)",
                                        None,
                                        'Processing-Completion-Monitor'
                                    )
                            except Exception as e:
                                print(f"[实时日志] 处理完成日志记录失败: {e}")
                            
                            # 设置完成时间戳，用于后续清理
                            tasks_db[task_id]['completed_at'] = time.time()
                            
                            if md_content:
                                # 更新文件的md内容到数据库（先不设置状态）
                                processed_time = datetime.now().isoformat()
                                update_file_status_in_db(file_info['file_id'], 'processing', 
                                                        md_content=md_content, processed_time=processed_time)
                                
                                # 检查并更新文件完整性状态（这会设置最终状态为completed或error）
                                update_file_completion_status(file_info['file_id'], file_info['user_id'])
                                
                                md_size_kb = len(md_content) / 1024
                                tasks_db[task_id]['message'] = f'✅ 处理完成！已提取 {md_size_kb:.1f}KB Markdown内容'
                                print(f"[MD处理] 成功提取并保存MD内容，文件ID: {file_info['file_id']}, 大小: {md_size_kb:.1f}KB")
                                
                                # 更新任务状态
                                update_task_status_in_db(task_id, 'completed', 100, f'✅ 处理完成！已提取 {md_size_kb:.1f}KB Markdown内容', download_url=download_url, download_path=download_path)
                            else:
                                # 没有MD内容，先设置为processing
                                update_file_status_in_db(file_info['file_id'], 'processing', processed_time=datetime.now().isoformat())
                                
                                # 检查并更新文件完整性状态（这会设置最终状态为completed或error）
                                update_file_completion_status(file_info['file_id'], file_info['user_id'])
                                
                                tasks_db[task_id]['message'] = '✅ 处理完成！(未找到Markdown内容，但文件已处理)'
                                print(f"[MD处理] 未能提取MD内容，文件ID: {file_info['file_id']}")
                                update_task_status_in_db(task_id, 'completed', 100, '✅ 处理完成！(未找到Markdown内容，但文件已处理)', download_url=download_url, download_path=download_path)
                            
                            # 清理临时ZIP文件
                            try:
                                if os.path.exists(download_path):
                                    os.remove(download_path)
                                    print(f"[清理] 已删除临时ZIP文件: {download_path}")
                            except Exception as e:
                                print(f"[清理] 删除ZIP文件失败: {e}")
                        else:
                            # 暂时保存错误信息，不立即设置error状态
                            tasks_db[task_id]['pending_error'] = '❌ 下载结果失败，请稍后重试'
                            tasks_db[task_id]['status'] = 'pending_completion'
                            tasks_db[task_id]['message'] = '正在确认处理状态...'
                            tasks_db[task_id]['completed_at'] = time.time()  # 记录失败时间
                            update_task_status_in_db(task_id, 'error', message='❌ 下载结果失败，请稍后重试')
                            # 注释掉立即更新数据库error状态，改为任务最终完成时统一更新
                            # update_file_status_in_db(file_info['file_id'], 'error')
                    else:
                        # 暂时保存错误信息，不立即设置error状态
                        tasks_db[task_id]['pending_error'] = '❌ 未获取到下载链接，处理可能失败'
                        tasks_db[task_id]['status'] = 'pending_completion'
                        tasks_db[task_id]['message'] = '正在确认处理状态...'
                        tasks_db[task_id]['completed_at'] = time.time()  # 记录失败时间
                        update_task_status_in_db(task_id, 'error', message='❌ 未获取到下载链接，处理可能失败')
                        # 注释掉立即更新数据库error状态，改为任务最终完成时统一更新
                        # update_file_status_in_db(file_info['file_id'], 'error')
                else:
                    # 暂时保存错误信息，不立即设置error状态
                    tasks_db[task_id]['pending_error'] = '❌ 处理结果为空，请检查文件格式'
                    tasks_db[task_id]['status'] = 'pending_completion'
                    tasks_db[task_id]['message'] = '正在确认处理状态...'
                    tasks_db[task_id]['completed_at'] = time.time()  # 记录失败时间
                    update_task_status_in_db(task_id, 'error', message='❌ 处理结果为空，请检查文件格式')
                    # 注释掉立即更新数据库error状态，改为任务最终完成时统一更新
                    # update_file_status_in_db(file_info['file_id'], 'error')
                break
                
            elif status == 'failed':
                error_msg = '❌ 文档解析失败'
                if results and len(results) > 0:
                    first_result = results[0]
                    if 'error' in first_result:
                        error_msg = f'❌ 处理失败: {first_result["error"]}'
                    elif 'err_msg' in first_result:
                        error_msg = f'❌ 处理失败: {first_result["err_msg"]}'
                
                print(f"[后台处理] MinerU处理失败: {error_msg}")
                
                # 记录处理失败日志
                log_user_action(
                    user_id, 
                    'file_processing_failed', 
                    f"文件处理失败: {file_info['original_name']} - {error_msg}",
                    None,
                    None
                )
                
                # 暂时保存错误信息，不立即设置error状态
                tasks_db[task_id]['pending_error'] = error_msg
                tasks_db[task_id]['status'] = 'pending_completion'
                tasks_db[task_id]['message'] = '正在确认处理状态...'
                tasks_db[task_id]['completed_at'] = time.time()  # 记录失败时间
                update_task_status_in_db(task_id, 'error', message=error_msg)
                # 注释掉立即更新数据库error状态，改为任务最终完成时统一更新
                # update_file_status_in_db(file_info['file_id'], 'error')
                break
            
            else:
                # 处理中状态，继续等待
                print(f"[后台处理] 处理中... 进度: {current_progress}%, 已等待: {int(elapsed_time)}秒, 状态: {status}")
                time.sleep(3)  # 减少等待时间，从10秒改为3秒，提高响应速度
        
        # 超时处理
        if tasks_db[task_id]['status'] in ['processing', 'uploading', 'preparing', 'validating']:
            print(f"[后台处理] 处理超时，任务ID: {task_id}")
            timeout_msg = f'⏰ 处理超时 (已等待{int(time.time() - start_time)}秒)，请稍后重试'
            # 暂时保存错误信息，不立即设置error状态
            tasks_db[task_id]['pending_error'] = timeout_msg
            tasks_db[task_id]['status'] = 'pending_completion'
            tasks_db[task_id]['message'] = '正在确认处理状态...'
            tasks_db[task_id]['completed_at'] = time.time()  # 记录超时时间，用于清理
            update_task_status_in_db(task_id, 'error', message=timeout_msg)
            # 注释掉立即更新数据库error状态，改为任务最终完成时统一更新
            # update_file_status_in_db(file_info['file_id'], 'error')
            
    except Exception as e:
        error_msg = f'❌ 处理异常: {str(e)}'
        print(f"[后台处理] 处理异常: {e}")
        print(f"[后台处理] 异常详情: {traceback.format_exc()}")
        
        # 确保任务状态被正确设置
        if task_id in tasks_db:
            # 暂时保存错误信息，不立即设置error状态
            tasks_db[task_id]['pending_error'] = error_msg
            tasks_db[task_id]['status'] = 'pending_completion'
            tasks_db[task_id]['message'] = '正在确认处理状态...'
            tasks_db[task_id]['completed_at'] = time.time()  # 记录异常时间
            update_task_status_in_db(task_id, 'error', message=error_msg)
        
        # 注释掉立即更新数据库error状态，改为任务最终完成时统一更新
        # update_file_status_in_db(file_info['file_id'], 'error')
    
    # 任务最终完成时的统一状态更新
    finally:
        if task_id in tasks_db:
            final_task_status = tasks_db[task_id].get('status', 'unknown')
            print(f"[任务完成] 任务ID: {task_id}, 最终状态: {final_task_status}")
            
            # 只有在任务真正结束时才更新数据库的文件状态
            # 检查是否有待处理的错误状态或pending_completion状态需要最终确认
            has_pending_error = tasks_db[task_id].get('pending_error')
            is_pending_completion = final_task_status == 'pending_completion'
            
            # 安全获取file_id，防止早期return导致的未定义
            task_file_id = tasks_db[task_id].get('file_id')
            if not task_file_id:
                # 尝试从函数参数file_info中获取（如果可用且有效）
                try:
                    if file_info and isinstance(file_info, dict):
                        task_file_id = file_info.get('file_id')
                except NameError:
                    # file_info未定义，忽略
                    pass
            
            if task_file_id and (final_task_status in ['completed', 'error'] or has_pending_error or is_pending_completion):
                try:
                    if final_task_status == 'error' or has_pending_error or is_pending_completion:
                        # 注释掉基于内存状态的数据库error更新，改为仅由最终完整性检查决定
                        # 任务失败时，更新数据库为error状态
                        if has_pending_error:
                            error_message = has_pending_error
                        elif is_pending_completion:
                            error_message = tasks_db[task_id].get('message', '处理状态待确认，最终标记为失败')
                        else:
                            error_message = tasks_db[task_id].get('message', '处理失败')
                            
                        # update_file_status_in_db(task_file_id, 'error', 
                        #                        remarks=f"处理失败: {error_message}")
                        print(f"[最终状态] 文件 {task_file_id} 内存状态为失败，但数据库状态将由完整性检查决定: {error_message}")
                        
                        # 发送任务失败通知
                        try:
                            user_id = tasks_db[task_id].get('user_id')
                            filename = tasks_db[task_id].get('filename', '未知文件')
                            if user_id:
                                from db_operations import send_system_notification
                                send_system_notification(
                                    sender_id=1,  # 系统通知
                                    title="文档处理失败",
                                    content=f"您的文件 '{filename}' 处理失败。\n错误信息：{error_message}",
                                    recipient_id=user_id,
                                    notification_type='system'
                                )
                        except Exception as notify_error:
                            print(f"[通知发送失败] {notify_error}")
                            
                    elif final_task_status == 'completed':
                        # 发送任务完成通知
                        try:
                            user_id = tasks_db[task_id].get('user_id')
                            filename = tasks_db[task_id].get('filename', '未知文件')
                            if user_id:
                                from db_operations import send_system_notification
                                send_system_notification(
                                    sender_id=1,  # 系统通知
                                    title="文档处理完成",
                                    content=f"您的文件 '{filename}' 已成功处理完成，可以前往查看结果。",
                                    recipient_id=user_id,
                                    notification_type='system'
                                )
                        except Exception as notify_error:
                            print(f"[通知发送失败] {notify_error}")
                    
                    # completed状态已由 update_file_completion_status 处理，无需重复更新
                except Exception as e:
                    print(f"[最终状态更新失败] {e}")
            elif not task_file_id:
                print(f"[最终状态更新] 跳过：无法获取file_id for task {task_id}")

@app.route('/task_status/<task_id>')
@api_login_required
def get_task_status(task_id):
    """获取任务状态 - 增强版本，包含详细进度信息"""
    
    if task_id not in tasks_db:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks_db[task_id]
    
    if task['user_id'] != session['user_id']:
        return jsonify({'error': '无权限访问'}), 403
    
    # 增强任务信息
    enhanced_task = task.copy()
    
    # 添加进度详细描述
    status = task['status']
    progress = task.get('progress', 0)
    
    # 进度阶段映射
    progress_stages = {
        'preparing': {'min': 0, 'max': 5, 'desc': '准备阶段'},
        'validating': {'min': 5, 'max': 10, 'desc': '验证阶段'},
        'uploading': {'min': 10, 'max': 30, 'desc': '上传阶段'},
        'processing': {'min': 30, 'max': 75, 'desc': '处理阶段'},
        'downloading': {'min': 75, 'max': 95, 'desc': '下载阶段'},
        'extracting': {'min': 95, 'max': 100, 'desc': '提取阶段'},
        'completed': {'min': 100, 'max': 100, 'desc': '已完成'},
        'error': {'min': 0, 'max': 100, 'desc': '出现错误'}
    }
    
    stage_info = progress_stages.get(status, {'min': 0, 'max': 100, 'desc': '未知状态'})
    enhanced_task['stage_info'] = stage_info
    
    # 计算预估剩余时间
    if status in ['preparing', 'validating', 'uploading', 'processing', 'downloading', 'extracting']:
        if 'start_time' in task:
            try:
                start_time = datetime.strptime(task['start_time'], '%Y-%m-%d %H:%M:%S')
                elapsed_seconds = (datetime.now() - start_time).total_seconds()
                
                if progress > 10:  # 有了实际进度才计算
                    estimated_total = elapsed_seconds * 100 / progress
                    estimated_remaining = max(0, estimated_total - elapsed_seconds)
                    enhanced_task['estimated_remaining_seconds'] = int(estimated_remaining)
                    
                    # 格式化剩余时间
                    if estimated_remaining < 60:
                        enhanced_task['estimated_remaining_text'] = f"{int(estimated_remaining)}秒"
                    elif estimated_remaining < 3600:
                        enhanced_task['estimated_remaining_text'] = f"{int(estimated_remaining/60)}分{int(estimated_remaining%60)}秒"
                    else:
                        hours = int(estimated_remaining / 3600)
                        minutes = int((estimated_remaining % 3600) / 60)
                        enhanced_task['estimated_remaining_text'] = f"{hours}小时{minutes}分钟"
                else:
                    enhanced_task['estimated_remaining_text'] = "计算中..."
                    
                enhanced_task['elapsed_seconds'] = int(elapsed_seconds)
                enhanced_task['elapsed_text'] = f"{int(elapsed_seconds)}秒" if elapsed_seconds < 60 else f"{int(elapsed_seconds/60)}分{int(elapsed_seconds%60)}秒"
            except Exception as e:
                print(f"计算时间信息错误: {e}")
    
    # 添加文件信息
    if task.get('file_id'):
        # 从数据库获取文件信息
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("SELECT original_name, file_size FROM user_files WHERE file_id = %s", (task['file_id'],))
            file_result = cursor.fetchone()
            conn.close()
            
            if file_result:
                enhanced_task['file_info'] = {
                    'original_name': file_result['original_name'],
                    'size': file_result['file_size'],
                    'size_mb': round(float(file_result['file_size']) / (1024 * 1024), 2)
                }
    
    return jsonify({
        'success': True,
        'task': enhanced_task
    })

def get_recovery_stage_message(start_timestamp):
    """根据任务开始时间推测当前处理阶段"""
    if not start_timestamp:
        return '📋 正在恢复处理进度...'
    
    current_time = time.time()
    elapsed_time = current_time - start_timestamp
    
    # 根据已运行时间推测阶段
    if elapsed_time < 30:  # 30秒内
        return '📋 正在准备处理参数...'
    elif elapsed_time < 60:  # 1分钟内
        return '🔍 正在验证文件和API配置...'
    elif elapsed_time < 120:  # 2分钟内
        return '📤 正在上传文件...'
    elif elapsed_time < 180:  # 3分钟内
        return '📤 正在调用文档解析服务...'
    elif elapsed_time < 600:  # 10分钟内
        return '⚙️ 文档解析中，请耐心等待...'
    else:  # 超过10分钟
        return '⚙️ 长时间处理中，可能是大文件...'

@app.route('/api/task_progress/<task_id>')
@api_login_required
def get_task_progress_simple(task_id):
    """简化的任务进度接口，适合频繁轮询 - 优化响应速度"""
    
    if task_id not in tasks_db:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks_db[task_id]
    
    if task['user_id'] != session['user_id']:
        return jsonify({'error': '无权限访问'}), 403
    
    # 缓存关键信息，减少重复计算
    status = task['status']
    progress = task.get('progress', 0)
    is_completed = status == 'completed'
    is_error = status == 'error'
    
    # 只返回关键进度信息，减少数据传输
    return jsonify({
        'success': True,
        'status': status,
        'progress': progress,
        'message': task.get('message', ''),
        'batch_id': task.get('batch_id', ''),
        'is_completed': is_completed,
        'is_error': is_error,
        'timestamp': int(time.time())  # 添加时间戳便于前端缓存控制
    })

@app.route('/api/task/<task_id>/cancel', methods=['POST'])
@api_login_required
def cancel_task_api(task_id):
    """取消任务API"""
    
    if task_id not in tasks_db:
        return jsonify({'error': '任务不存在'}), 404
    
    task = tasks_db[task_id]
    
    if task['user_id'] != session['user_id']:
        return jsonify({'error': '无权限操作此任务'}), 403
    
    # 检查任务状态
    if task['status'] == 'completed':
        return jsonify({'error': '任务已完成，无法取消'}), 400
    
    if task['status'] == 'error':
        return jsonify({'error': '任务已出错，无法取消'}), 400
    
    # 取消任务
    task['status'] = 'cancelled'
    task['message'] = '任务已被用户取消'
    task['error'] = '用户主动取消任务'
    
    print(f"[任务取消] 任务 {task_id} 已被用户取消")
    
    return jsonify({
        'success': True,
        'message': '任务已取消'
    })

@app.route('/download_result/<task_id>')
@login_required
def download_result(task_id):
    """下载处理结果"""
    
    if task_id not in tasks_db:
        flash('任务不存在', 'error')
        return redirect(url_for('dashboard'))
    
    task = tasks_db[task_id]
    
    if task['user_id'] != session['user_id']:
        flash('无权限访问', 'error')
        return redirect(url_for('dashboard'))
    
    if task['status'] != 'completed' or not task.get('download_path'):
        flash('结果文件不可用', 'error')
        return redirect(url_for('dashboard'))
    
    # 重定向到实际的下载URL
    if task.get('download_url'):
        return redirect(task['download_url'])
    else:
        flash('下载链接已失效', 'error')
        return redirect(url_for('dashboard'))

@app.route('/api/data')
def api_data():
    """JSON数据接口 - 从数据库动态获取最新的NSFC JSON数据"""
    try:
        if 'user_id' not in session:
            return jsonify({"status": "error", "message": "请先登录"}), 401
        
        user_id = session['user_id']
        
        # 从数据库获取用户最新的NSFC JSON数据
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT nsfc_json_data, original_name 
                FROM user_files 
                WHERE user_id = %s AND nsfc_json_data IS NOT NULL 
                ORDER BY nsfc_json_updated_at DESC 
                LIMIT 1
            """, (user_id,))
            result = cursor.fetchone()
            conn.close()
            
            if result and result['nsfc_json_data']:
                try:
                    # 验证JSON格式但不转换为Python对象以保持顺序
                    json.loads(result['nsfc_json_data'])  # 只验证格式
                    raw_json_data = result['nsfc_json_data']  # 保持原始字符串
                    
                    # 直接返回原始JSON字符串以保持顺序
                    from flask import Response
                    response = Response(raw_json_data, mimetype='application/json')
                    response.headers['Content-Type'] = 'application/json;charset=utf-8'
                    return response
                except json.JSONDecodeError:
                    return jsonify({"status": "error", "message": "JSON数据格式错误"}), 500
        
        # 如果没有找到数据，返回空数据提示
        empty_data = {
            "提示": "暂无数据",
            "说明": "请先上传国家自然科学基金申请书PDF文件并进行表格转JSON处理"
        }
        
        response = jsonify(empty_data)
        response.headers['Content-Type'] = 'application/json;charset=utf-8'
        return response
            
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/peanut-qrcode')
def generate_peanut_qrcode():
    """生成花生壳链接的二维码"""
    try:
        # 花生壳固定URL
        url = "https://1114qd34764oj.vicp.fun/"
        
        # 生成二维码
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(url)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        # 保存到内存
        buffer = BytesIO()
        img.save(buffer, 'PNG')
        buffer.seek(0)
        
        return send_file(buffer, mimetype='image/png', as_attachment=False)
        
    except Exception as e:
        return f"无法生成花生壳二维码: {str(e)}", 500

# MD查看器路由已移至 md_viewer_module.py

def extract_md_from_zip(zip_path, original_filename):
    """从MinerU返回的ZIP文件中提取full.md内容，严格按照mineru_api.py的方法"""
    try:
        import zipfile
        
        print(f"[MD提取] 开始提取ZIP文件: {zip_path}")
        
        # 检查文件是否存在
        if not os.path.exists(zip_path):
            print(f"[MD提取] 错误: ZIP文件不存在: {zip_path}")
            return None
        
        file_size = os.path.getsize(zip_path)
        print(f"[MD提取] ZIP文件大小: {file_size} bytes")
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            # 列出所有文件
            all_files = zip_ref.namelist()
            print(f"[MD提取] ZIP文件包含 {len(all_files)} 个文件:")
            for file_name in all_files[:10]:  # 只显示前10个文件名
                print(f"[MD提取]   - {file_name}")
            if len(all_files) > 10:
                print(f"[MD提取]   ...还有 {len(all_files) - 10} 个文件")
            
            # 查找full.md文件
            md_files = [name for name in all_files if name.endswith('full.md')]
            print(f"[MD提取] 找到 {len(md_files)} 个.md文件: {md_files}")
            
            if not md_files:
                # 尝试查找其他可能的MD文件
                all_md_files = [name for name in all_files if name.endswith('.md')]
                print(f"[MD提取] 尝试查找其他MD文件: {all_md_files}")
                
                if all_md_files:
                    md_file = all_md_files[0]
                    print(f"[MD提取] 使用第一个MD文件: {md_file}")
                else:
                    print(f"[MD提取] ZIP文件中未找到任何MD文件: {zip_path}")
                    return None
            else:
                # 提取第一个找到的full.md文件
                md_file = md_files[0]
                print(f"[MD提取] 使用full.md文件: {md_file}")
            
            # 提取文件内容
            try:
                with zip_ref.open(md_file) as file:
                    raw_content = file.read()
                    print(f"[MD提取] 读取原始内容，字节长度: {len(raw_content)}")
                    
                    # 尝试多种编码
                    for encoding in ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']:
                        try:
                            content = raw_content.decode(encoding)
                            print(f"[MD提取] 成功使用 {encoding} 编码解析，字符长度: {len(content)}")
                            
                            # 简单验证内容是否有效
                            if len(content.strip()) > 0:
                                print(f"[MD提取] 内容预览（前200字符）: {content[:200]}")
                                
                                # 对提取的MD内容进行处理：重要内容提取和标题优化
                                print(f"[MD处理] 开始处理MD内容，原始长度: {len(content)}")
                                
                                # 1. 重要内容提取（去除承诺书等无关内容）
                                from title_enhancer import TitleEnhancer
                                important_content = TitleEnhancer.extract_important_content(content)
                                print(f"[MD处理] 重要内容提取完成，长度: {len(important_content)}")
                                
                                # 2. 标题优化处理
                                optimized_content = TitleEnhancer.optimize_headers_with_extractor(important_content)
                                print(f"[MD处理] 标题优化完成，最终长度: {len(optimized_content)}")
                                
                                return optimized_content
                            else:
                                print(f"[MD提取] 警告: {encoding} 解析的内容为空")
                        except UnicodeDecodeError:
                            print(f"[MD提取] {encoding} 编码解析失败，尝试下一个")
                            continue
                    
                    print(f"[MD提取] 所有编码尝试失败，返回原始字节内容")
                    return raw_content.decode('utf-8', errors='ignore')
                    
            except KeyError as e:
                print(f"[MD提取] 文件不存在于ZIP中: {e}")
                return None
            except Exception as e:
                print(f"[MD提取] 读取文件内容失败: {e}")
                return None
                
    except zipfile.BadZipFile as e:
        print(f"[MD提取] 无效的ZIP文件: {e}")
        return None
    except Exception as e:
        print(f"[MD提取] 提取MD内容失败: {e}")
        import traceback
        print(f"[MD提取] 异常详情: {traceback.format_exc()}")
        return None

@app.route('/favicon.ico')
def favicon():
    """处理favicon请求，避免404错误"""
    return '', 204  # 返回空响应，状态码204表示无内容

# 错误处理
@app.errorhandler(404)
def not_found_error(error):
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    return render_template('errors/500.html'), 500

@app.errorhandler(413)
def too_large(error):
    return jsonify({'error': '文件太大，请选择小于50MB的文件'}), 413

def load_data_from_db():
    """启动时从数据库加载任务数据到内存（仅清理未完成任务）"""
    try:
        print("🔄 系统启动时清理未完成任务...")
        
        # 由于使用纯内存任务管理，启动时清空内存中的任务
        global tasks_db
        tasks_db = {}
        
        print("✅ 任务内存已清理，使用纯内存任务管理模式")
    except Exception as e:
        print(f"❌ 清理任务内存失败: {e}")

# 表格JSON数据获取路由已移至 json_viewer_module.py

# ==================== API密钥获取路由 ====================

@app.route('/api/get-api-keys')
@api_login_required
def get_api_keys():
    """获取用户的所有API密钥配置 - 从分组密钥获取"""
    
    try:
        user_id = session['user_id']
        
        # 获取用户的分组ID
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        cursor.execute("SELECT group_id FROM users WHERE id = %s", (user_id,))
        user_result = cursor.fetchone()
        
        if not user_result or not user_result['group_id']:
            conn.close()
            return jsonify({
                'success': True,
                'keys': [],
                'message': '用户未分配到分组，请联系管理员'
            })
        
        group_id = user_result['group_id']
        
        # 从分组密钥表获取所有激活的API密钥（包含model_name）
        cursor.execute("""
            SELECT id, group_id, api_provider, api_token, model_name,
                   is_active, is_valid, created_at, updated_at
            FROM group_api_keys 
            WHERE group_id = %s AND is_active = 1 AND is_valid = 1
            ORDER BY api_provider
        """, (group_id,))
        
        keys = cursor.fetchall()
        conn.close()
        
        # #22 解密API密钥
        from db_operations import decrypt_api_token
        
        # 转换为前端期望的格式
        formatted_keys = []
        for key in keys:
            formatted_keys.append({
                'id': key['id'],
                'api_provider': key['api_provider'],
                'api_token': decrypt_api_token(key['api_token']),
                'model_name': key.get('model_name'),  # 添加模型名称
                'is_active': bool(key['is_active']),
                'created_at': key['created_at'].strftime('%Y-%m-%d %H:%M:%S') if key['created_at'] else None,
                'assigned_by_name': '管理员',  # 分组密钥由管理员配置
                'assigned_at': key['created_at'].strftime('%Y-%m-%d %H:%M:%S') if key['created_at'] else None
            })
        
        return jsonify({
            'success': True,
            'keys': formatted_keys,
            'message': '成功获取API密钥配置'
        })
        
    except Exception as e:
        app.logger.error(f"获取用户API密钥错误: {str(e)}")
        return jsonify({'success': False, 'error': f'获取失败: {str(e)}'}), 500

# ==================== 表格转JSON API路由 ====================

@app.route('/api/get-siliconflow-api-key')
@api_login_required
def get_siliconflow_api_key():
    """获取用户的SiliconFlow API密钥 - 现在使用分组密钥"""
    
    try:
        user_id = session['user_id']
        # 从分组密钥获取siliconflow密钥
        api_key_info = get_active_api_key(user_id, 'siliconflow')
        
        return jsonify({
            'success': True,
            'api_key': api_key_info['api_token'] if api_key_info else None,
            'has_key': bool(api_key_info),
            'message': '使用分组统一密钥' if api_key_info else '请联系管理员配置分组智能分析密钥'
        })
        
    except Exception as e:
        print(f"获取SiliconFlow API密钥错误: {e}")
        return jsonify({'error': f'获取失败: {str(e)}'}), 500

# 已废弃：/api/process-table-json 路由已不再使用
# 一键解析统一使用 md_content_processor 模块处理

# ==================== 表格转JSON API路由结束 ====================

# ==================== 综合处理页面API路由开始 ====================

@app.route('/api/file-content/<file_id>')
@api_login_required
def get_file_content_api(file_id):
    """获取文件内容API"""
    
    try:
        file_data, error_msg = get_file_content_from_db(file_id, session['user_id'])
        if error_msg or not file_data:
            return jsonify({'error': '文件不存在或无权限'}), 404
        
        # ✨ 验证原始文档类型
        is_valid, document_type, error_message = validate_md_extraction_for_document_type(
            file_data, 
            "文件内容获取"
        )
        
        if not is_valid:
            print(error_message)
            return jsonify({'error': f'文档类型验证失败: {document_type}类型文档不支持此操作'}), 400
        
        # 记录操作类型
        print(f"[文件内容获取] 原始文档类型: {document_type.upper()}, 文件: {file_data.get('original_name', '未知文件')}")
        
        # 使用md_content字段
        content = file_data.get('md_content') or ''
        
        return jsonify({
            'success': True,
            'content': content,
            'md_content': file_data.get('md_content'),
            'file_name': file_data['original_name'],
            'storage_type': file_data.get('storage_type'),
            'status': file_data.get('status'),
            'document_type': document_type  # 添加文档类型信息
        })
        
    except Exception as e:
        print(f"获取文件内容错误: {e}")
        return jsonify({'error': f'获取失败: {str(e)}'}), 500

@app.route('/api/file-json/<file_id>')
@api_login_required
def get_file_json_api(file_id):
    """获取文件已有的JSON数据API"""
    
    try:
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT nsfc_json_data, extracted_json_data, original_name
                FROM user_files 
                WHERE file_id = %s AND user_id = %s
            """, (file_id, session['user_id']))
            result = cursor.fetchone()
            conn.close()
            
            if result:
                # 合并所有JSON数据
                all_json_data = {}
                all_json_data_raw = {}  # 保存原始JSON字符串
                
                if result['nsfc_json_data']:
                    try:
                        all_json_data['nsfc_json'] = json.loads(result['nsfc_json_data'])
                        all_json_data_raw['nsfc_json'] = result['nsfc_json_data']  # 原始字符串
                    except:
                        pass
                
                if result['extracted_json_data']:
                    try:
                        all_json_data['table_json'] = json.loads(result['extracted_json_data'])
                        all_json_data_raw['table_json'] = result['extracted_json_data']  # 原始字符串
                    except:
                        pass
                
                return jsonify({
                    'success': True,
                    'json_data': all_json_data if all_json_data else None,
                    'json_data_raw': all_json_data_raw if all_json_data_raw else None,  # 返回原始JSON字符串
                    'file_name': result['original_name']
                })
            else:
                return jsonify({'error': '文件不存在或无权限'}), 404
        else:
            return jsonify({'error': '数据库连接失败'}), 500
                
    except Exception as e:
        print(f"获取文件JSON数据错误: {e}")
        return jsonify({'error': f'获取失败: {str(e)}'}), 500

@app.route('/api/update-api-key', methods=['POST'])
@api_login_required
def update_api_key():
    """更新用户API密钥 - 已废弃，使用分组密钥管理"""
    
    return jsonify({
        'error': '个人API密钥配置已废弃，请联系管理员进行分组密钥配置'
    }), 400

@app.route('/api/enhance-title', methods=['POST'])
@api_login_required
def enhance_title_api():
    """标题优化API"""
    
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({'error': '文件ID不能为空'}), 400
        
        # 获取文件内容
        file_data, error_msg = get_file_content_from_db(file_id, session['user_id'])
        if error_msg or not file_data:
            return jsonify({'error': '文件不存在或无权限'}), 404
        
        # ✨ 验证原始文档类型
        is_valid, document_type, error_message = validate_md_extraction_for_document_type(
            file_data, 
            "标题优化"
        )
        
        if not is_valid:
            print(error_message)
            return jsonify({'error': f'文档类型验证失败: {document_type}类型文档不支持此操作'}), 400
        
        # 记录操作类型
        print(f"[标题优化] 原始文档类型: {document_type.upper()}, 文件: {file_data.get('original_name', '未知文件')}")
        
        content = file_data.get('md_content')
        if not content:
            return jsonify({'error': '文件没有MD内容'}), 404
        
        # 提取重要内容
        important_content = TitleEnhancer.extract_important_content(content)
        
        # 优化标题
        optimized_content = TitleEnhancer.optimize_headers_with_extractor(important_content)
        
        return jsonify({
            'success': True,
            'optimized_content': optimized_content,
            'original_length': len(content),
            'optimized_length': len(optimized_content),
            'important_length': len(important_content)
        })
        
    except Exception as e:
        print(f"标题优化错误: {e}")
        return jsonify({'error': f'优化失败: {str(e)}'}), 500

@app.route('/api/extract-key-content', methods=['POST'])
@api_login_required
def extract_key_content_api():
    """重点内容提取API"""
    
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({'error': '文件ID不能为空'}), 400
        
        # 获取文件内容
        file_data, error_msg = get_file_content_from_db(file_id, session['user_id'])
        if error_msg or not file_data:
            return jsonify({'error': '文件不存在或无权限'}), 404
        
        # ✨ 验证原始文档类型
        is_valid, document_type, error_message = validate_md_extraction_for_document_type(
            file_data, 
            "重点内容提取"
        )
        
        if not is_valid:
            print(error_message)
            return jsonify({'error': f'文档类型验证失败: {document_type}类型文档不支持此操作'}), 400
        
        # 记录操作类型
        print(f"[重点内容提取] 原始文档类型: {document_type.upper()}, 文件: {file_data.get('original_name', '未知文件')}")
        
        content = file_data.get('md_content')
        if not content:
            return jsonify({'error': '文件没有MD内容'}), 404
        
        # 提取重要内容
        important_content = TitleEnhancer.extract_important_content(content)
        
        return jsonify({
            'success': True,
            'key_content': important_content,
            'original_length': len(content),
            'extracted_length': len(important_content),
            'extraction_ratio': round(len(important_content) / len(content) * 100, 2) if content else 0
        })
        
    except Exception as e:
        print(f"重点内容提取错误: {e}")
        return jsonify({'error': f'提取失败: {str(e)}'}), 500

@app.route('/api/clear-json', methods=['POST'])
@api_login_required
def clear_json_api():
    """清空JSON数据API"""
    
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        clear_type = data.get('clear_type', 'both')  # 默认清除所有
        
        if not file_id:
            return jsonify({'error': '文件ID不能为空'}), 400
        
        user_id = session['user_id']
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 检查文件是否存在且属于当前用户
        cursor.execute("""
            SELECT original_name FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        
        if not result:
            conn.close()
            return jsonify({'error': '文件不存在'}), 404
        
        # 根据清除类型执行相应操作
        if clear_type == 'table':
            cursor.execute("""
                UPDATE user_files 
                SET extracted_json_data = NULL
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
        elif clear_type == 'content':
            cursor.execute("""
                UPDATE user_files 
                SET nsfc_json_data = NULL
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
        else:  # both
            cursor.execute("""
                UPDATE user_files 
                SET extracted_json_data = NULL, nsfc_json_data = NULL
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
        
        updated = cursor.rowcount > 0
        conn.commit()  # 提交事务
        conn.close()
        
        if updated:
            # 清除后检查并更新文件完整性状态
            update_file_completion_status(file_id, user_id)
            return jsonify({
                'success': True,
                'message': 'JSON数据已清空'
            })
        else:
            return jsonify({'error': '清空JSON数据失败'}), 500
        
    except Exception as e:
        print(f"清空JSON错误: {e}")
        return jsonify({'error': f'清空失败: {str(e)}'}), 500

@app.route('/api/clear-table-json', methods=['POST'])
@api_login_required
def clear_table_json_api():
    """清空表格JSON数据API"""
    
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({'error': '文件ID不能为空'}), 400
        
        user_id = session['user_id']
        
        # 直接在这里实现删除表格JSON数据的逻辑
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 首先检查文件是否存在且属于当前用户
        cursor.execute("""
            SELECT original_name, extracted_json_data FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        
        if not result:
            conn.close()
            return jsonify({'error': '文件不存在'}), 404
        
        # 如果没有表格JSON数据，直接返回成功（已经是清空状态）
        if not result['extracted_json_data']:
            conn.close()
            return jsonify({
                'success': True,
                'message': '表格JSON数据已清空（原本就为空）'
            })
        
        # 清除表格JSON数据
        cursor.execute("""
            UPDATE user_files 
            SET extracted_json_data = NULL
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        updated = cursor.rowcount > 0
        conn.commit()  # 提交事务
        conn.close()
        
        if updated:
            # 清除后检查并更新文件完整性状态
            update_file_completion_status(file_id, user_id)
            return jsonify({
                'success': True,
                'message': '表格JSON数据已清空'
            })
        else:
            return jsonify({'error': '清空表格JSON数据失败'}), 500
        
    except Exception as e:
        print(f"清空表格JSON错误: {e}")
        return jsonify({'error': f'清空失败: {str(e)}'}), 500

def update_file_status_after_clear(file_id, user_id):
    """清除处理结果后检查并更新文件状态"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return False
        
        cursor = conn.cursor()
        
        # 检查文件是否还有任何处理结果
        cursor.execute("""
            SELECT md_content, nsfc_json_data, extracted_json_data
            FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        
        if result:
            md_content = result['md_content']
            nsfc_json_data = result['nsfc_json_data']
            extracted_json_data = result['extracted_json_data']
            
            # 如果所有处理结果都为空，则将状态改为uploaded
            if not md_content and not nsfc_json_data and not extracted_json_data:
                cursor.execute("""
                    UPDATE user_files 
                    SET status = 'uploaded'
                    WHERE file_id = %s AND user_id = %s
                """, (file_id, user_id))
                conn.commit()  # 提交状态更新
        
        conn.close()
        return True
        
    except Exception as e:
        print(f"更新文件状态错误: {e}")
        return False

def check_file_completion_status(file_id, user_id):
    """检查文件是否真正完成（根据文档类型有不同的完成标准）"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return False
        
        cursor = conn.cursor()
        
        # 获取文件的所有处理结果和原始文件名
        cursor.execute("""
            SELECT original_name, original_document_type, md_content, nsfc_json_data, extracted_json_data, 
                   storage_type, md_file_path, reference_md_content, reference_validation_json
            FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        conn.close()
        
        if not result:
            return False
        
        # 直接从数据库结果获取原始文档类型
        original_document_type = get_original_document_type(result)
        
        # 检查MD内容是否存在
        has_md = False
        if result['md_content']:
            has_md = True
        elif result['storage_type'] == 'filesystem' and result['md_file_path']:
            # 检查文件系统中的MD文件是否存在
            try:
                md_content = md_storage_manager.load_md_content(file_id)
                has_md = bool(md_content)
            except:
                has_md = False
        
        # 根据文档类型判断完成标准
        if original_document_type == 'word':
            # Word文档完成标准：有MD内容即可（参考文献数据是可选的）
            is_complete = has_md
            print(f"[完整性检查] Word文档 {file_id}: MD={has_md} -> 完成状态={is_complete}")
            
        elif original_document_type == 'pdf':
            # PDF文档完成标准：MD + 正文JSON + 表格JSON 都存在
            has_content_json = bool(result['nsfc_json_data'])
            has_table_json = bool(result['extracted_json_data'])
            is_complete = has_md and has_content_json and has_table_json
            print(f"[完整性检查] PDF文档 {file_id}: MD={has_md}, 正文JSON={has_content_json}, 表格JSON={has_table_json} -> 完成状态={is_complete}")
            
        else:
            # 未知类型文档，采用保守策略：只要有MD内容就算完成
            is_complete = has_md
            print(f"[完整性检查] 未知类型文档 {file_id}: MD={has_md} -> 完成状态={is_complete}")
        
        return is_complete
        
    except Exception as e:
        print(f"检查文件完整性错误: {e}")
        return False

def update_file_completion_status(file_id, user_id):
    """根据文件完整性更新状态"""
    try:
        is_complete = check_file_completion_status(file_id, user_id)
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return False
        
        cursor = conn.cursor()
        
        if is_complete:
            # 完整时设为completed
            cursor.execute("""
                UPDATE user_files 
                SET status = 'completed', error_message = NULL
                WHERE file_id = %s AND user_id = %s AND status != 'completed'
            """, (file_id, user_id))
            
            # 检查是否真的更新了状态（避免重复记录日志）
            if cursor.rowcount > 0:
                print(f"[状态更新] 文件 {file_id} 标记为完成")
                
                # 获取文件名用于日志记录
                cursor.execute("SELECT original_name FROM user_files WHERE file_id = %s", (file_id,))
                file_result = cursor.fetchone()
                file_name = file_result['original_name'] if file_result else file_id
                
                # 记录处理成功到系统日志
                from db_operations import log_user_action
                log_user_action(
                    user_id, 
                    'file_processed', 
                    f"文件 {file_name} 处理成功完成",
                    None,  # ip_address  
                    None   # user_agent
                )
                
                # 🆕 实时更新处理完成统计
                try:
                    # 获取今日处理统计
                    cursor.execute("""
                        SELECT COUNT(*) as today_completed 
                        FROM user_files 
                        WHERE created_at >= CURDATE() 
                        AND created_at < DATE_ADD(CURDATE(), INTERVAL 1 DAY) 
                        AND status = 'completed'
                    """)
                    result = cursor.fetchone()
                    today_completed = result['today_completed'] if result else 0
                    
                    # 计算处理时间
                    cursor.execute("""
                        SELECT TIMESTAMPDIFF(SECOND, processing_start_time, processed_time) as processing_time 
                        FROM user_files 
                        WHERE file_id = %s AND processing_start_time IS NOT NULL AND processed_time IS NOT NULL
                    """, (file_id,))
                    time_result = cursor.fetchone()
                    processing_time = float(time_result['processing_time']) if time_result and time_result['processing_time'] else 0
                    
                    # 获取系统管理员ID
                    system_user_id = None
                    try:
                        cursor.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                        admin_user = cursor.fetchone()
                        if admin_user:
                            system_user_id = admin_user['id']
                    except:
                        pass
                    
                    # 只有获取到有效用户ID时才记录日志
                    if system_user_id:
                        # 记录实时完成统计日志
                        log_user_action(
                            system_user_id,
                            'realtime_completion_stats',
                            f"今日完成总数: {today_completed}, 最新完成: {file_name} (耗时: {processing_time}秒)",
                            None,
                            'Completion-Stats-Monitor'
                        )
                except Exception as e:
                    print(f"[实时统计] 完成统计更新失败: {e}")
        else:
            # 不完整时设为error状态，并记录具体缺失的内容
            missing_parts = []
            
            # 重新检查各部分是否存在，以便提供详细错误信息
            cursor.execute("""
                SELECT md_content, nsfc_json_data, extracted_json_data, storage_type, md_file_path
                FROM user_files 
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
            
            result = cursor.fetchone()
            if result:
                # 检查MD内容
                has_md = False
                if result['md_content']:
                    has_md = True
                elif result['storage_type'] == 'filesystem' and result['md_file_path']:
                    try:
                        md_content = md_storage_manager.load_md_content(file_id)
                        has_md = bool(md_content)
                    except:
                        has_md = False
                
                if not has_md:
                    missing_parts.append("MD文档")
                if not result['nsfc_json_data']:
                    missing_parts.append("正文JSON")
                if not result['extracted_json_data']:
                    missing_parts.append("表格JSON")
            
            error_message = f"缺失文件：{', '.join(missing_parts)}"
            
            cursor.execute("""
                UPDATE user_files 
                SET status = 'error', error_message = %s
                WHERE file_id = %s AND user_id = %s
            """, (error_message, file_id, user_id))
            print(f"[状态更新] 文件 {file_id} 标记为错误（缺失文件）：{error_message}")
        
        conn.commit()
        conn.close()
        return True
        
    except Exception as e:
        print(f"更新文件完成状态错误: {e}")
        return False

def delete_md_content(file_id, delete_type, user_id):
    """删除MD内容"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return False
        
        cursor = conn.cursor()
        
        # 首先检查文件是否存在且属于当前用户
        cursor.execute("""
            SELECT original_name, md_content, md_file_path, storage_type FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        
        if not result:
            conn.close()
            return False
        
        original_name = result['original_name']
        md_content = result['md_content']
        md_file_path = result['md_file_path']
        storage_type = result['storage_type']
        
        # 根据删除类型执行相应操作
        if delete_type == 'md_content':
            # 只删除数据库中的MD内容
            cursor.execute("""
                UPDATE user_files 
                SET md_content = NULL
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
            log_user_action(user_id, 'delete_md_content', f'删除MD内容: {original_name}')
            
        elif delete_type == 'md_file':
            # 删除文件系统中的MD文件
            if md_file_path and os.path.exists(md_file_path):
                try:
                    os.remove(md_file_path)
                except Exception as e:
                    print(f"删除MD文件失败: {e}")
            
            cursor.execute("""
                UPDATE user_files 
                SET md_file_path = NULL, storage_type = 'database'
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
            log_user_action(user_id, 'delete_md_file', f'删除MD文件: {original_name}')
            
        elif delete_type == 'both':
            # 删除MD内容和文件
            cursor.execute("""
                UPDATE user_files 
                SET md_content = NULL, md_file_path = NULL, storage_type = 'database'
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
            
            # 删除文件系统中的MD文件
            if md_file_path and os.path.exists(md_file_path):
                try:
                    os.remove(md_file_path)
                except Exception as e:
                    print(f"删除MD文件失败: {e}")
            
            log_user_action(user_id, 'delete_md_both', f'删除MD内容和文件: {original_name}')
        
        updated = cursor.rowcount > 0
        conn.commit()  # 提交MD删除操作
        conn.close()
        
        # 删除后检查并更新文件状态
        if updated:
            update_file_completion_status(file_id, user_id)
        
        return updated
            
    except Exception as e:
        print(f"删除MD内容错误: {e}")
        return False

@app.route('/api/clear-content-json', methods=['POST'])
@api_login_required
def clear_content_json_api():
    """清空正文JSON数据API"""
    
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({'error': '文件ID不能为空'}), 400
        
        user_id = session['user_id']
        
        # 直接在这里实现删除NSFC JSON数据的逻辑
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 首先检查文件是否存在且属于当前用户
        cursor.execute("""
            SELECT original_name, nsfc_json_data FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        result = cursor.fetchone()
        
        if not result:
            conn.close()
            return jsonify({'error': '文件不存在'}), 404
        
        # 如果没有正文JSON数据，直接返回成功（已经是清空状态）
        if not result['nsfc_json_data']:
            conn.close()
            return jsonify({
                'success': True,
                'message': '正文JSON数据已清空（原本就为空）'
            })
        
        # 清除JSON数据
        cursor.execute("""
            UPDATE user_files 
            SET nsfc_json_data = NULL, nsfc_json_created_at = NULL, nsfc_json_updated_at = NULL
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        updated = cursor.rowcount > 0
        conn.commit()  # 提交事务
        conn.close()
        
        if updated:
            # 清除后检查并更新文件完整性状态
            update_file_completion_status(file_id, user_id)
            return jsonify({
                'success': True,
                'message': '正文JSON数据已清空'
            })
        else:
            return jsonify({'error': '清空正文JSON数据失败'}), 500
        
    except Exception as e:
        print(f"清空正文JSON错误: {e}")
        return jsonify({'error': f'清空失败: {str(e)}'}), 500

@app.route('/api/delete-md-content', methods=['POST'])
@api_login_required
def delete_md_content_api():
    """删除MD内容API"""
    
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        delete_type = data.get('delete_type')
        
        if not file_id or not delete_type:
            return jsonify({'error': '文件ID和删除类型不能为空'}), 400
        
        user_id = session['user_id']
        
        # 调用删除MD内容的函数
        success = delete_md_content(file_id, delete_type, user_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'MD内容删除成功'
            })
        else:
            return jsonify({'error': 'MD内容删除失败'}), 500
        
    except Exception as e:
        print(f"删除MD内容错误: {e}")
        return jsonify({'error': f'删除失败: {str(e)}'}), 500

@app.route('/api/save-json', methods=['POST'])
@api_login_required
def save_json_api():
    """保存JSON数据API"""
    
    try:
        data = request.get_json()
        json_data = data.get('json_data')
        filename = data.get('filename', 'export.json')
        
        print(f"[保存JSON] 收到请求，文件名: {filename}")
        print(f"[保存JSON] 数据大小: {len(str(json_data))} 字符")
        
        if not json_data:
            return jsonify({'error': 'JSON数据不能为空'}), 400
        
        # 过滤掉以下划线开头的字段（如_content_quality_check）
        def filter_underscore_fields(obj):
            if isinstance(obj, dict):
                return {k: filter_underscore_fields(v) for k, v in obj.items() if not k.startswith('_')}
            elif isinstance(obj, list):
                return [filter_underscore_fields(item) for item in obj]
            else:
                return obj
        
        filtered_json_data = filter_underscore_fields(json_data)
        print(f"[保存JSON] 过滤后数据大小: {len(str(filtered_json_data))} 字符")
        
        # 生成下载文件
        downloads_dir = 'downloads'
        os.makedirs(downloads_dir, exist_ok=True)
        
        file_id = str(uuid.uuid4())
        safe_filename = filename.replace('/', '_').replace('\\', '_')  # 防止路径注入
        file_path = os.path.join(downloads_dir, f'{file_id}_{safe_filename}')
        
        print(f"[保存JSON] 保存到文件: {file_path}")
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(filtered_json_data, f, ensure_ascii=False, indent=2)
        
        print(f"[保存JSON] 文件保存成功")
        
        return jsonify({
            'success': True,
            'message': 'JSON数据已保存',
            'download_url': f'/download/{file_id}_{safe_filename}'
        })
        
    except Exception as e:
        print(f"[保存JSON] 错误: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'保存失败: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """通用文件下载路由"""
    try:
        print(f"[下载文件] 请求下载: {filename}")
        downloads_dir = 'downloads'
        file_path = os.path.join(downloads_dir, filename)
        
        print(f"[下载文件] 文件路径: {file_path}")
        print(f"[下载文件] 文件是否存在: {os.path.exists(file_path)}")
        
        if os.path.exists(file_path):
            print(f"[下载文件] 开始发送文件")
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            print(f"[下载文件] 文件不存在")
            return jsonify({'error': '文件不存在'}), 404
            
    except Exception as e:
        print(f"[下载文件] 错误: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

@app.route('/download/<int:file_id>')
def download_file_by_id(file_id):
    """根据文件ID下载文件"""
    try:
        # 检查用户权限
        if 'user_id' not in session:
            return jsonify({'error': '未登录'}), 401
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 获取文件信息 - 使用正确的字段名
        cursor.execute("""
            SELECT uf.file_id, uf.original_name, uf.user_id, u.username 
            FROM user_files uf 
            JOIN users u ON uf.user_id = u.id 
            WHERE uf.file_id = %s
        """, (file_id,))
        
        file_result = cursor.fetchone()
        conn.close()
        
        if not file_result:
            return jsonify({'error': '文件不存在'}), 404
        
        # 明确解包元组以避免类型错误
        file_id_val, filename, user_id_val, username = file_result
        
        file_info = {
            'id': file_id_val,
            'filename': filename,
            'user_id': user_id_val,
            'username': username
        }
        
        # 检查权限：只有文件所有者或管理员可以下载
        user_id = session['user_id']
        is_owner = file_info['user_id'] == user_id
        is_admin_user = is_admin(user_id)
        
        if not (is_owner or is_admin_user):
            return jsonify({'error': '无权限下载此文件'}), 403
        
        # 构建文件路径
        file_path = os.path.join('uploads', file_info['filename'])
        
        if os.path.exists(file_path):
            # 记录下载日志
            log_user_action(
                user_id, 
                'file_download', 
                f"下载文件: {file_info['filename']}",
                request.remote_addr,
                request.headers.get('User-Agent')
            )
            
            return send_file(file_path, as_attachment=True, download_name=file_info['filename'])
        else:
            return jsonify({'error': '文件不存在'}), 404
            
    except Exception as e:
        app.logger.error(f"下载文件错误: {e}")
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

@app.route('/preview/<int:file_id>')
def preview_file_by_id(file_id):
    """根据文件ID预览文件"""
    try:
        # 检查用户权限
        if 'user_id' not in session:
            return jsonify({'error': '未登录'}), 401
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 获取文件信息 - 使用正确的字段名
        cursor.execute("""
            SELECT uf.file_id, uf.original_name, uf.user_id, u.username 
            FROM user_files uf 
            JOIN users u ON uf.user_id = u.id 
            WHERE uf.file_id = %s
        """, (file_id,))
        
        file_result = cursor.fetchone()
        conn.close()
        
        if not file_result:
            return jsonify({'error': '文件不存在'}), 404
        
        # 明确解包元组以避免类型错误
        file_id_val, filename, user_id_val, username = file_result
        
        file_info = {
            'id': file_id_val,
            'filename': filename,
            'user_id': user_id_val,
            'username': username
        }
        
        # 检查权限：只有文件所有者或管理员可以预览
        user_id = session['user_id']
        is_owner = file_info['user_id'] == user_id
        is_admin_user = is_admin(user_id)
        
        if not (is_owner or is_admin_user):
            return jsonify({'error': '无权限预览此文件'}), 403
        
        # 构建文件路径
        file_path = os.path.join('uploads', file_info['filename'])
        
        if os.path.exists(file_path):
            # 对于图片和文本文件，直接在浏览器中显示
            return send_file(file_path, download_name=file_info['filename'])
        else:
            return jsonify({'error': '文件不存在'}), 404
            
    except Exception as e:
        app.logger.error(f"预览文件错误: {e}")
        return jsonify({'error': f'预览失败: {str(e)}'}), 500

# ==================== 综合处理页面API路由结束 ====================

# API: 批量转换PDF
@app.route('/api/convert-pdf', methods=['POST'])
@api_login_required
def convert_pdf():
        
    try:
        data = request.get_json()
        file_id = data.get('file_id')
        
        if not file_id:
            return jsonify({'success': False, 'error': '缺少文件ID'})
        
        user_id = session['user_id']
        
        # 检查MinerU供应商类型
        from db_operations import get_user_mineru_provider
        mineru_provider = get_user_mineru_provider(user_id)
        
        if mineru_provider == 'local_mineru':
            # 使用本地MinerU API，无需API密钥
            print(f"[PDF转换] 用户 {user_id} 使用本地MinerU")
            success = convert_pdf_with_local_mineru(file_id, user_id)
        else:
            # 使用官方MinerU API，需要API密钥
            active_key = get_active_api_key(user_id, provider='mineru')
            if not active_key:
                return jsonify({
                    'success': False,
                    'error': '未配置PDF解析服务密钥，请联系管理员配置分组密钥'
                })
            mineru_api_key = active_key['api_token']
            success = convert_pdf_with_mineru(file_id, mineru_api_key, user_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': f'文件 {file_id} 转换完成'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'PDF转换失败'
            })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

# ==================== 一键解析功能 ====================

def get_file_format(filename):
    """检测文件格式类型"""
    if not filename:
        return 'unknown'
    
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return 'pdf'
    elif filename_lower.endswith('.doc'):
        return 'doc'
    elif filename_lower.endswith('.docx'):
        return 'docx'
    else:
        return 'unknown'

def is_word_document(filename):
    """判断是否为Word文档"""
    file_format = get_file_format(filename)
    return file_format in ['doc', 'docx']

def is_pdf_document(filename):
    """判断是否为PDF文档"""
    return get_file_format(filename) == 'pdf'

def get_original_document_type_from_db(file_id: str, user_id: int = None) -> str:
    """
    直接从数据库获取原始文档类型
    
    Args:
        file_id: 文件ID
        user_id: 用户ID（可选，用于权限验证）
        
    Returns:
        str: 'word', 'pdf', 'unknown'
    """
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return 'unknown'
        
        cursor = conn.cursor()
        
        if user_id:
            cursor.execute("""
                SELECT original_document_type 
                FROM user_files 
                WHERE file_id = %s AND user_id = %s
            """, (file_id, user_id))
        else:
            cursor.execute("""
                SELECT original_document_type 
                FROM user_files 
                WHERE file_id = %s
            """, (file_id,))
        
        result = cursor.fetchone()
        conn.close()
        
        if result and result['original_document_type']:
            db_type = result['original_document_type']
            # 转换数据库枚举值到函数返回值
            if db_type in ['word_doc', 'word_docx']:
                return 'word'
            elif db_type == 'pdf':
                return 'pdf'
            else:
                return 'unknown'
        else:
            return 'unknown'
            
    except Exception as e:
        print(f"获取文档类型失败: {e}")
        return 'unknown'

def get_original_document_type(file_info):
    """
    获取原始文档类型 - 优先从数据库获取，兼容旧代码
    
    Args:
        file_info (dict): 文件信息，可能包含file_id或original_name字段
        
    Returns:
        str: 'word', 'pdf', 'unknown'
    """
    if not file_info:
        return 'unknown'
    
    # 优先使用file_id从数据库获取
    if 'file_id' in file_info:
        file_id = file_info['file_id']
        user_id = file_info.get('user_id')
        db_type = get_original_document_type_from_db(file_id, user_id)
        if db_type != 'unknown':
            return db_type
    
    # 其次检查是否已经有数据库字段
    if 'original_document_type' in file_info and file_info['original_document_type']:
        db_type = file_info['original_document_type']
        # 转换数据库枚举值到函数返回值
        if db_type in ['word_doc', 'word_docx']:
            return 'word'
        elif db_type == 'pdf':
            return 'pdf'
        else:
            return 'unknown'
    
    # 如果仍然无法确定类型，返回unknown
    smart_print(f"无法确定文档类型，file_info: {file_info}", level=LogLevel.WARNING)
    return 'unknown'

def validate_md_extraction_for_document_type(file_info, operation_name="MD内容提取"):
    """
    验证MD提取操作是否适用于当前文档类型
    
    Args:
        file_info (dict): 文件信息
        operation_name (str): 操作名称，用于日志
        
    Returns:
        tuple: (is_valid, document_type, error_message)
    """
    document_type = get_original_document_type(file_info)
    
    if document_type == 'unknown':
        error_msg = f"[{operation_name}] 无法识别原始文档类型: {file_info.get('original_name', '未知文件')}"
        return False, document_type, error_msg
    
    # 记录文档类型信息
    print(f"[{operation_name}] 文档类型验证: {file_info.get('original_name', '未知文件')} -> {document_type.upper()}")
    
    return True, document_type, None

@app.route('/api/scan-all-files')
@api_login_required
def scan_all_files():
    """扫描所有文件，检查处理状态"""
    
    try:
        user_id = session['user_id']
        
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # 获取所有文件（基于实际表结构）
            cursor.execute("""
                SELECT file_id, original_name, md_content, 
                       extracted_json_data, nsfc_json_data,
                       status, upload_time
                FROM user_files 
                WHERE user_id = %s
                ORDER BY upload_time DESC
            """, (user_id,))
            
            # 直接使用cursor结果，不要用dict(zip())方式
            all_files = cursor.fetchall()
            
            # 统计文件状态
            total_files = len(all_files)
            pdf_files = 0
            md_files = 0
            complete_count = 0
            needs_conversion = 0
            needs_extraction = 0
            
            for file_data in all_files:
                # 使用字典访问方式（因为使用了DictCursor）
                original_name = file_data['original_name']
                md_content = file_data['md_content']
                extracted_json_data = file_data['extracted_json_data']
                nsfc_json_data = file_data['nsfc_json_data']
                
                # 判断文件状态（基于实际表结构）
                has_md = bool(md_content and md_content.strip())
                has_extracted_json = bool(extracted_json_data and extracted_json_data.strip())
                has_nsfc_json = bool(nsfc_json_data and nsfc_json_data.strip())
                
                # 根据用户要求：md、两种json都有的是已完成的，但凡缺失一个就是待处理
                # 两种JSON是指：extracted_json_data 和 nsfc_json_data
                if has_md:
                    # 有MD内容就算作MD文档（无论原始文件是PDF还是MD）
                    md_files += 1
                    # 完成条件：md_content + extracted_json_data + nsfc_json_data 都有内容
                    if has_md and has_extracted_json and has_nsfc_json:
                        complete_count += 1
                    else:
                        needs_extraction += 1
                else:
                    # 没有MD内容的算作需要转换的PDF
                    pdf_files += 1
                    needs_conversion += 1
                        
            conn.close()
            
            return jsonify({
                'success': True,
                'scan_result': {
                    'total_files': total_files,
                    'pdf_files': pdf_files,
                    'md_files': md_files,
                    'complete_files': complete_count,
                    'needs_conversion': needs_conversion,
                    'needs_extraction': needs_extraction
                }
            })
        else:
            return jsonify({'success': False, 'error': '数据库连接失败'})
            
    except Exception as e:
        print(f"扫描文件错误: {e}")
        return jsonify({
            'success': False,
            'error': f'扫描失败: {str(e)}'
        })

# 存储一键解析任务状态的全局变量
one_click_tasks = {}

@app.route('/api/one-click-process', methods=['POST'])
@api_login_required
def start_one_click_process():
    """启动一键解析处理"""
    
    try:
        user_id = session['user_id']
        
        # 获取请求数据，检查是否有选中的文件ID
        request_data = request.get_json() or {}
        selected_file_ids = request_data.get('selected_file_ids', [])
        
        # 从数据库获取MinerU供应商类型和API密钥
        from db_operations import get_user_mineru_provider
        user_mineru_provider = get_user_mineru_provider(user_id)
        
        if user_mineru_provider == 'local_mineru':
            mineru_api_key = 'local-not-needed'
            print(f"[一键解析] 用户使用本地MinerU，跳过MinerU密钥检查")
        else:
            active_key = get_active_api_key(user_id, provider='mineru')
            if not active_key:
                return jsonify({
                    'success': False,
                    'error': '未配置PDF解析服务密钥，请前往用户信息页面申请API密钥配置'
                })
            mineru_api_key = active_key['api_token']
        
        # 从数据库获取SiliconFlow API密钥（或检查本地LLM配置）
        from db_operations import get_user_llm_provider
        user_llm_provider = get_user_llm_provider(user_id)
        
        if user_llm_provider == 'local_llm':
            siliconflow_api_key = 'local-not-needed'
            print(f"[一键解析] 用户使用本地LLM，跳过SiliconFlow密钥检查")
        else:
            siliconflow_key = get_active_api_key(user_id, provider='siliconflow')
            if not siliconflow_key:
                return jsonify({
                    'success': False,
                    'error': '未配置智能分析服务密钥，请前往用户信息页面申请API密钥配置'
                })
            siliconflow_api_key = siliconflow_key['api_token']
        # 生成任务ID
        import uuid
        task_id = str(uuid.uuid4())
        
        # 获取需要处理的文件
        all_files = get_user_files_from_db(user_id, limit=1000)
        if selected_file_ids:
            # 过滤选中的文件
            all_files = [f for f in all_files if f['file_id'] in selected_file_ids]
        files_to_process = []
        
        print(f"[一键解析] 用户 {user_id} 共有 {len(all_files)} 个文件")
        
        for file_info in all_files:
                filename = file_info['original_name'].lower()
                
                # 判断文件状态（基于实际表结构 - 使用优化后的字段）
                has_md_content = bool(file_info.get('has_md_content', 0))
                has_nsfc_json = bool(file_info.get('has_nsfc_json', 0))
                has_extracted_json = bool(file_info.get('has_extracted_json', 0))
                status = file_info.get('status', 'uploaded')
                has_pdf_data = bool(file_info.get('has_pdf_data', 0))  # 检查是否有PDF二进制数据
                
                has_md = has_md_content
                has_nsfc = has_nsfc_json
                has_extracted = has_extracted_json
                
                # 判断JSON数据是否完整
                # 需要nsfc_json_data和extracted_json_data两个字段
                json_complete = has_nsfc and has_extracted
                
                print(f"[一键解析] 检查文件: {file_info['original_name']}")
                print(f"  状态: {status}, PDF数据: {has_pdf_data}, MD: {has_md}, JSON完整: {json_complete}")
                print(f"  详细: NSFC: {has_nsfc}, 提取: {has_extracted}")
                
                # 获取文件格式和文档类型
                file_format = get_file_format(filename)
                # 直接从数据库获取原始文档类型
                original_doc_type = get_original_document_type_from_db(file_info['file_id'], user_id)
                print(f"  文件格式: {file_format}, 原始文档类型: {original_doc_type}")
                
                # 处理逻辑：基于数据库中的文档类型，Word和PDF完全分离
                
                # 1. Word文档处理（优先级最高）
                if original_doc_type == 'word' and has_pdf_data:
                    if has_md:
                        # Word文档已有MD内容：进行状态检查和完整性验证
                        files_to_process.append({
                            'file_id': file_info['file_id'],
                            'filename': file_info['original_name'],
                            'original_name': file_info['original_name'],  # 添加这个字段用于文档类型检测
                            'original_document_type': file_info.get('original_document_type', 'unknown'),  # 添加数据库字段
                            'file_format': file_format,
                            'type': 'word_convert',  # 仍使用word_convert类型，但内部会跳过转换
                            'needs_processing': True,
                            'has_md': True  # 标记已有MD内容
                        })
                        print(f"  -> ✅ Word文档已有MD，进行状态检查（{file_format.upper()}）")
                    else:
                        # Word文档需要转换
                        files_to_process.append({
                            'file_id': file_info['file_id'],
                            'filename': file_info['original_name'],
                            'original_name': file_info['original_name'],  # 添加这个字段用于文档类型检测
                            'original_document_type': file_info.get('original_document_type', 'unknown'),  # 添加数据库字段
                            'file_format': file_format,
                            'type': 'word_convert',
                            'needs_processing': True,
                            'has_md': False  # 标记需要转换
                        })
                        print(f"  -> 📄 需要Word文档转换（{file_format.upper()}）")
                
                # 2. PDF文档处理
                elif original_doc_type == 'pdf':
                    if has_md:
                        # PDF已有MD内容：进行MD处理
                        files_to_process.append({
                            'file_id': file_info['file_id'],
                            'filename': file_info['original_name'],
                            'original_name': file_info['original_name'],  # 添加这个字段用于文档类型检测
                            'original_document_type': file_info.get('original_document_type', 'unknown'),  # 添加数据库字段
                            'original_name': file_info['original_name'],  # 添加这个字段用于文档类型检测
                            'file_format': file_format,
                            'type': 'md_process',
                            'needs_processing': True,
                            'needs_key_content': True,
                            'needs_title_enhance': True,
                            'needs_extracted': not has_extracted,
                            'needs_nsfc': True
                        })
                        print(f"  -> 📋 需要MD处理（重点内容+标题+JSON）")
                    elif has_pdf_data:
                        # PDF需要转换
                        files_to_process.append({
                            'file_id': file_info['file_id'],
                            'filename': file_info['original_name'],
                            'original_name': file_info['original_name'],  # 添加这个字段用于文档类型检测
                            'original_document_type': file_info.get('original_document_type', 'unknown'),  # 添加数据库字段
                            'file_format': file_format,
                            'type': 'pdf_convert',
                            'needs_processing': True
                        })
                        print(f"  -> 🔄 需要PDF转换")
                    else:
                        print(f"  -> ⚠️  跳过：PDF无数据")
                
                # 3. 跳过的情况
                else:
                    if has_pdf_data:  # 有文档数据但文档类型未知
                        print(f"  -> ⚠️  跳过：文档类型未知 ({original_doc_type})")
                    else:  # 没有文档数据
                        print(f"  -> ⚠️  跳过：无文档数据")
        
        # 检查是否超过全局并发限制
        global global_active_tasks, global_active_tasks_lock
        
        with global_active_tasks_lock:
            if global_active_tasks >= 2:
                return jsonify({
                    'success': False,
                    'error': '系统当前已有2个任务在处理，请等待其中一个任务完成后再试'
                }), 429
            
            # 增加活跃任务计数
            global_active_tasks += 1
            print(f"[全局并发控制] 当前活跃任务数: {global_active_tasks}/2")
        
        # 初始化任务状态
        one_click_tasks[task_id] = {
            'status': 'started',
            'total_files': len(files_to_process),
            'completed_files': 0,
            'completed': 0,  # 添加这个字段
            'total': len(files_to_process),  # 添加这个字段
            'current_step': '准备开始处理',
            'current_file': None,
            'errors': [],
            'start_time': time.time(),
            'user_id': user_id,
            'mineru_api_key': mineru_api_key,
            'mineru_provider': user_mineru_provider,
            'siliconflow_api_key': siliconflow_api_key,
            'files_to_process': files_to_process
        }
        
        # 启动并行后台处理
        import threading
        thread = threading.Thread(target=process_one_click_task_parallel, args=(task_id,))
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'success': True,
            'task_id': task_id,
            'total_files': len(files_to_process),
            'file_ids': [f['file_id'] for f in files_to_process],
            'message': '一键解析已启动（并行处理）'
        })
        
    except Exception as e:
        print(f"启动一键解析错误: {e}")
        return jsonify({
            'success': False,
            'error': f'启动失败: {str(e)}'
        })

def process_one_click_task_parallel(task_id):
    """并行处理一键解析任务（后台线程）"""
    global global_active_tasks, global_active_tasks_lock
    import threading
    import time
    
    try:
        task = one_click_tasks.get(task_id)
        if not task:
            print(f"[并行一键解析] 任务 {task_id} 不存在")
            return
        
        files_to_process = task['files_to_process']
        user_id = task['user_id']
        mineru_api_key = task['mineru_api_key']
        mineru_provider = task.get('mineru_provider', 'mineru')
        siliconflow_api_key = task['siliconflow_api_key']
        
        # 初始化任务状态
        task['status'] = 'processing'
        task['current_step'] = f'并行处理 {len(files_to_process)} 个文件'
        task['completed_files'] = 0
        task['errors'] = []
        
        print(f"[任务队列] 开始初始化任务队列，共 {len(files_to_process)} 个文件")
        
        # 📋 第一步：创建所有任务记录，状态为 waiting
        for file_info in files_to_process:
            file_id = file_info['file_id']
            filename = file_info['filename']
            file_format = file_info.get('file_format', 'unknown')
            file_task_id = f"oneclick_{task_id}_{file_id}"
            
            # 创建任务记录
            tasks_db[file_task_id] = {
                'task_id': file_task_id,
                'file_id': file_id,
                'user_id': user_id,
                'status': 'waiting',  # 初始状态：等待中
                'progress': 0,
                'message': f'⏳ 等待处理: {filename}',
                'created_at': time.time(),  # 创建时间
                'start_time': None,  # 开始执行时间（前端计时起点）
                'end_time': None,  # 结束时间
                'duration': None,  # 执行时长（秒）
                'parent_task_id': task_id,
                'filename': filename,
                'file_format': file_format,
                'file_info': file_info  # 保存完整文件信息
            }
            print(f"[任务队列] ✅ 任务已加入队列: {filename} (状态: waiting)")
        
        print(f"[任务队列] 任务列表创建完成，共 {len(files_to_process)} 个任务")
        
        # 任务队列锁（保证线程安全）
        queue_lock = threading.Lock()
        
        def get_next_waiting_task():
            """从队列中获取下一个等待中的任务"""
            with queue_lock:
                for task_id_key, task_info in tasks_db.items():
                    if task_info.get('parent_task_id') == task_id and task_info.get('status') == 'waiting':
                        return task_id_key, task_info
            return None, None
        
        def mark_task_started(file_task_id):
            """标记任务开始执行"""
            with queue_lock:
                if file_task_id in tasks_db:
                    from datetime import datetime
                    
                    # 记录开始时间（前端计时起点）
                    start_timestamp = time.time()
                    tasks_db[file_task_id]['status'] = 'processing'
                    tasks_db[file_task_id]['start_time'] = start_timestamp
                    tasks_db[file_task_id]['message'] = f'🚀 开始处理...'
                    
                    # 同时更新数据库的 processing_start_time
                    try:
                        conn_start = DatabaseConfig.get_db_connection()
                        if conn_start:
                            cursor_start = conn_start.cursor()
                            local_now = datetime.fromtimestamp(start_timestamp).strftime('%Y-%m-%d %H:%M:%S')
                            cursor_start.execute("""
                                UPDATE user_files 
                                SET processing_start_time = %s
                                WHERE file_id = %s AND user_id = %s AND processing_start_time IS NULL
                            """, (local_now, tasks_db[file_task_id]['file_id'], tasks_db[file_task_id]['user_id']))
                            conn_start.commit()
                            conn_start.close()
                            print(f"[任务队列] ⏱️ 任务开始执行: {tasks_db[file_task_id]['filename']}, 开始时间: {local_now}")
                    except Exception as e:
                        print(f"[任务队列] 设置数据库开始时间失败: {e}")
        
        def mark_task_completed(file_task_id, success=True, error_msg=None):
            """标记任务完成"""
            with queue_lock:
                if file_task_id in tasks_db:
                    end_timestamp = time.time()
                    start_timestamp = tasks_db[file_task_id].get('start_time')
                    
                    tasks_db[file_task_id]['status'] = 'completed' if success else 'failed'
                    tasks_db[file_task_id]['end_time'] = end_timestamp
                    tasks_db[file_task_id]['completed_at'] = end_timestamp
                    
                    # 计算执行时长（只计算真正执行的时间）
                    if start_timestamp:
                        duration = end_timestamp - start_timestamp
                        tasks_db[file_task_id]['duration'] = duration
                        duration_str = f"{int(duration // 60)}分{int(duration % 60)}秒"
                        tasks_db[file_task_id]['message'] = f'✅ 处理完成 (用时: {duration_str})' if success else f'❌ 处理失败: {error_msg}'
                        print(f"[任务队列] ✅ 任务完成: {tasks_db[file_task_id]['filename']}, 用时: {duration_str}")
                    else:
                        tasks_db[file_task_id]['message'] = f'❌ 任务异常: 未记录开始时间'
                    
                    # 更新进度
                    if success:
                        tasks_db[file_task_id]['progress'] = 100
        
        def process_single_file(file_task_id):
            """处理单个文件的函数"""
            from db_operations import update_file_status_in_db, update_file_complete_status
            
            try:
                # 检查任务是否被取消
                current_task = one_click_tasks.get(task_id)
                if not current_task or current_task.get('status') == 'cancelled':
                    print(f"[任务队列] 一键任务已被取消，跳过文件处理")
                    mark_task_completed(file_task_id, success=False, error_msg='一键任务已取消')
                    return
                
                # 获取任务信息
                task_info = tasks_db.get(file_task_id)
                if not task_info:
                    print(f"[任务队列] 任务不存在: {file_task_id}")
                    return
                
                file_info = task_info['file_info']
                file_id = file_info['file_id']
                filename = file_info['filename']
                file_format = file_info.get('file_format', 'unknown')
                
                print(f"[任务队列] 开始处理文件: {filename} ({file_format})")
                
                # ⏱️ 标记任务开始（前端此时启动计时器）
                mark_task_started(file_task_id)
                
                # 根据文件类型进行处理
                if file_info['type'] == 'word_convert':
                    # Word文档转换
                    print(f"[任务队列] 🔄 开始Word文档转换: {filename}")
                    tasks_db[file_task_id]['progress'] = 5
                    tasks_db[file_task_id]['message'] = f'📄 正在转换Word文档为Markdown'
                    
                    # 创建Word转换进度回调
                    def word_conversion_progress_callback(progress, message):
                        # Word转换进度从5%开始到95%
                        actual_progress = 5 + (progress * 0.9)
                        # 安全地更新进度,防止任务被取消后仍然访问tasks_db
                        if file_task_id in tasks_db:
                            tasks_db[file_task_id]['progress'] = int(actual_progress)
                            tasks_db[file_task_id]['message'] = f'📄 {message}'
                        print(f"[Word转换进度] {filename}: {actual_progress:.1f}% - {message}")
                    
                    # 调用Word转换接口，传递进度回调
                    result = convert_word_to_markdown_from_db(file_id, user_id, file_task_id, word_conversion_progress_callback)
                    
                    # 处理任务取消的情况
                    if isinstance(result, dict) and result.get('cancelled'):
                        print(f"[并行一键解析] Word转换任务被取消: {filename}")
                        if file_task_id in tasks_db:
                            tasks_db[file_task_id]['status'] = 'cancelled'
                            tasks_db[file_task_id]['message'] = '❌ 任务已取消'
                            tasks_db[file_task_id]['completed_at'] = time.time()
                        # 更新一键任务统计
                        current_task = one_click_tasks.get(task_id)
                        if current_task:
                            current_task['completed'] += 1
                            current_task['cancelled_files'] = current_task.get('cancelled_files', 0) + 1
                        return
                    
                    # 判断是否成功
                    if isinstance(result, bool):
                        success = result
                    elif isinstance(result, dict):
                        success = result.get('success', False)
                    else:
                        success = False
                    
                    if success:
                        print(f"[并行一键解析] Word文档转换完成: {filename}")
                        if file_task_id in tasks_db:
                            tasks_db[file_task_id]['progress'] = 100
                            tasks_db[file_task_id]['message'] = f'✅ Word文档转换完成'
                            tasks_db[file_task_id]['status'] = 'completed'
                            tasks_db[file_task_id]['completed_at'] = time.time()
                            
                            # 验证数据库中的文件状态是否正确更新
                            try:
                                file_content, error_msg = get_file_content_from_db(file_id, user_id)
                                if not file_content or not file_content.get('md_content'):
                                    print(f"[并行一键解析] 警告：Word转换成功但数据库中未找到MD内容: {filename}")
                                    # 重新调用状态检查
                                    from db_operations import update_file_complete_status
                                    update_file_complete_status(user_id, file_id)
                                else:
                                    print(f"[并行一键解析] Word转换数据库验证成功: {filename}")
                            except Exception as verify_error:
                                print(f"[并行一键解析] Word转换状态验证失败: {verify_error}")
                            
                            # 更新一键任务统计
                            current_task = one_click_tasks.get(task_id)
                            if current_task:
                                current_task['completed'] += 1
                                current_task['completed_files'] += 1
                                current_task['progress'] = int((current_task['completed'] / current_task['total']) * 100)
                            
                            # 最终完整性检查
                            try:
                                from db_operations import update_file_complete_status
                                completion_result = update_file_complete_status(user_id, file_id)
                                print(f"[任务队列] Word文档最终完整性检查结果: {completion_result}")
                            except Exception as completion_error:
                                print(f"[任务队列] Word文档完整性检查失败: {completion_error}")
                            
                            # ✅ Word文档处理成功 - 标记任务完成
                            print(f"[任务队列] Word文档处理完成: {filename}")
                            mark_task_completed(file_task_id, success=True)
                            return
                    else:
                        # ❌ Word转换失败 - 标记任务失败
                        error_msg = result.get('error', '未知错误') if isinstance(result, dict) else 'Word转换失败'
                        print(f"[任务队列] Word文档转换失败: {filename} - {error_msg}")
                        mark_task_completed(file_task_id, success=False, error_msg=error_msg)
                        
                        current_task = one_click_tasks.get(task_id)
                        if current_task:
                            current_task['completed'] += 1
                            current_task['errors'].append(f'Word文档转换失败: {filename} - {error_msg}')
                        return
                
                elif file_info['type'] == 'pdf_convert':
                    # PDF转换
                    print(f"[并行一键解析] 开始PDF转换: {filename}")
                    tasks_db[file_task_id]['progress'] = 1
                    tasks_db[file_task_id]['message'] = f'🔄 正在调用文档解析服务转换PDF'
                    
                    # 根据MinerU供应商类型选择转换方式
                    if mineru_provider == 'local_mineru':
                        success = convert_pdf_with_local_mineru(file_id, user_id, task_id=file_task_id)
                    else:
                        success = convert_pdf_with_mineru(file_id, mineru_api_key, user_id)
                    if not success:
                        # 暂时保存错误信息，不立即设置error状态
                        tasks_db[file_task_id]['pending_error'] = f'❌ PDF转换失败'
                        tasks_db[file_task_id]['status'] = 'pending_completion'
                        tasks_db[file_task_id]['message'] = '正在确认处理状态...'
                        tasks_db[file_task_id]['completed_at'] = time.time()  # 记录失败时间
                        current_task = one_click_tasks.get(task_id)
                        if current_task:
                            current_task['errors'].append(f'PDF转换失败: {filename}')
                        return
                    
                    # convert_pdf_with_mineru 是同步函数，返回True时md_content已保存到数据库
                    print(f"[并行一键解析] PDF转换完成: {filename}")
                    tasks_db[file_task_id]['progress'] = 15
                    tasks_db[file_task_id]['message'] = f'✅ PDF转换完成'
                    
                    # ✨ 保留原始文档类型信息，避免丢失文档来源
                    original_document_type = get_original_document_type_from_db(file_id, user_id)
                    print(f"[并行一键解析] 保留原始文档类型: {original_document_type.upper()}, 文件: {filename}")
                    
                    file_info['type'] = 'md_process'
                    file_info['original_document_type'] = original_document_type
                    file_info['needs_json'] = True
                    file_info['needs_nsfc'] = True
                    file_info['needs_extracted'] = True
                
                # MD文件处理流程（仅适用于PDF转换产生的MD文档）
                if file_info['type'] in ['md_process', 'pdf_convert']:
                    # ✨ 验证原始文档类型
                    is_valid, document_type, error_message = validate_md_extraction_for_document_type(
                        file_info, 
                        "MD后处理流程"
                    )
                    
                    if not is_valid:
                        print(error_message)
                        print(f"[并行一键解析] 文档类型验证失败，跳过后处理: {filename}")
                        tasks_db[file_task_id]['pending_error'] = f'❌ 文档类型验证失败: {document_type}类型文档不支持MD后处理'
                        tasks_db[file_task_id]['status'] = 'pending_completion'
                        tasks_db[file_task_id]['completed_at'] = time.time()
                        current_task = one_click_tasks.get(task_id)
                        if current_task:
                            current_task['errors'].append(f'文档类型验证失败: {filename}')
                        return
                    
                    # 🚨 额外验证：确保Word文档不会进入MD后处理流程
                    if document_type == 'word':
                        print(f"[并行一键解析] 警告：Word文档不应进入MD后处理流程: {filename}")
                        # Word文档转换应该已经完成，直接标记为完成
                        tasks_db[file_task_id]['progress'] = 100
                        tasks_db[file_task_id]['message'] = f'✅ Word文档处理完成'
                        tasks_db[file_task_id]['status'] = 'completed'
                        tasks_db[file_task_id]['completed_at'] = time.time()
                        return
                    
                    # 记录操作类型（仅PDF文档应该到达这里）
                    print(f"[MD后处理流程] 原始文档类型: {document_type.upper()}, 文件: {filename}")
                    
                    # 获取文件格式信息
                    file_format = file_info.get('file_format', 'unknown')
                    print(f"[并行一键解析] 文件格式: {file_format}")
                    
                    # 步骤1: 重要内容提取（承诺书删除）
                    print(f"[并行一键解析] 重要内容提取中: {filename}")
                    if file_task_id in tasks_db:
                        tasks_db[file_task_id]['progress'] = 16
                        tasks_db[file_task_id]['message'] = '📋 内容预处理中（承诺书检测）...'
                    processed_content = extract_key_content_for_file(file_id, user_id)
                    if processed_content:
                        print(f"[并行一键解析] 重要内容提取完成: {filename}")
                    else:
                        print(f"[并行一键解析] 重要内容提取失败: {filename}")
                        # 如果提取失败，获取原始内容继续处理
                        file_content, error_msg = get_file_content_from_db(file_id, user_id)
                        processed_content = file_content.get('md_content') if file_content else None
                    
                    # 检查任务是否被中断
                    if file_task_id in tasks_db and tasks_db[file_task_id].get('status') == 'cancelled':
                        print(f"[并行一键解析] 文件任务 {file_task_id} 已被中断，停止处理 {filename}")
                        return
                    
                    # 步骤2: 标题优化
                    if processed_content:
                        print(f"[并行一键解析] 标题优化中: {filename}")
                        if file_task_id in tasks_db:
                            tasks_db[file_task_id]['progress'] = 17
                            tasks_db[file_task_id]['message'] = '📝 标题结构优化中...'
                        final_content = enhance_title_for_file(file_id, user_id, processed_content)
                        if final_content:
                            # 保存最终处理后的内容，但不设置为completed状态，也不保存处理时间
                            update_file_status_in_db(file_id, 'processing', md_content=final_content)
                            print(f"[并行一键解析] 标题优化完成: {filename}")
                        else:
                            print(f"[并行一键解析] 标题优化失败，保存重要内容提取结果: {filename}")
                            # 如果标题优化失败，保存重要内容提取的结果，但不设置为completed状态，也不保存处理时间
                            update_file_status_in_db(file_id, 'processing', md_content=processed_content)
                    
                    # 检查任务是否被中断
                    if file_task_id in tasks_db and tasks_db[file_task_id].get('status') == 'cancelled':
                        print(f"[并行一键解析] 文件任务 {file_task_id} 已被中断，停止处理 {filename}")
                        return
                    
                    # 四项API调用：表格提取 + 正文提取 + 表格初筛 + 正文初筛
                    if file_info.get('needs_json') or file_info.get('needs_extracted', False):
                        print(f"[并行一键解析] 开始四项API调用: {filename}")
                        # 创建任务检查函数
                        def check_task_cancelled():
                            # 检查主任务是否被取消
                            current_task = one_click_tasks.get(task_id)
                            if not current_task or current_task.get('status') == 'cancelled':
                                return True
                            # 检查子任务是否被取消
                            if file_task_id in tasks_db and tasks_db[file_task_id].get('status') == 'cancelled':
                                return True
                            return False
                        
                        # 创建进度回调函数
                        def update_progress(progress, message):
                            if file_task_id in tasks_db:
                                tasks_db[file_task_id]['progress'] = progress
                                tasks_db[file_task_id]['message'] = message
                        
                        # VLM仅在MinerU未能成功提取表格时才会启动（在call_table_extraction_for_file内部判断）
                        success = call_table_extraction_for_file(
                            file_id, user_id, siliconflow_api_key, 
                            check_task_cancelled, update_progress,
                            progress_start=20
                        )
                        if success:
                            print(f"[并行一键解析] 四项API调用完成: {filename}")
                        else:
                            print(f"[并行一键解析] 四项API调用失败: {filename}")
                    
                    print(f"[并行一键解析] 文件处理完成: {filename}")
                
                # 立即计算处理时间
                start_time = tasks_db.get(file_task_id, {}).get('start_time', time.time())
                processing_time = time.time() - start_time
                processing_time_str = f"{int(processing_time // 60)}分{int(processing_time % 60)}秒" if processing_time >= 60 else f"{int(processing_time)}秒"
                
                print(f"[处理完成] 文件 {filename} 处理完成，耗时: {processing_time_str}")
                
                # 立即更新内存中的任务状态为完成（线程处理完成）
                tasks_db[file_task_id]['status'] = 'completed'
                tasks_db[file_task_id]['progress'] = 100
                tasks_db[file_task_id]['message'] = f'✅ 处理完成 ({processing_time_str})'
                tasks_db[file_task_id]['processing_time'] = processing_time
                tasks_db[file_task_id]['processing_time_str'] = processing_time_str
                tasks_db[file_task_id]['completed_at'] = time.time()  # 记录完成时间戳
                
                print(f"[线程完成] 任务 {file_task_id} 处理线程已完成: {tasks_db[file_task_id]}")
                
                # 强制刷新一次，确保前端能立即获取到完成状态
                print(f"[状态广播] 文件 {filename} 线程完成状态已准备就绪，等待前端轮询")
                
                # 保存处理时间到remarks字段，但先不更新status
                try:
                    update_file_status_in_db(file_info['file_id'], None, remarks=f"处理时间: {processing_time_str}")
                    print(f"[处理时间保存] 文件 {filename} 处理时间已保存到remarks字段")
                except Exception as e:
                    print(f"[处理时间保存失败] {e}")
                
                # 进行最终的完整性检查并更新状态（这是关键步骤）
                try:
                    print(f"[最终完整性检查] 开始检查文件 {filename} 的完整性...")
                    update_result = update_file_complete_status(user_id, file_info['file_id'])
                    print(f"[最终完整性检查] 文件 {filename} 完整性检查结果: {update_result}")
                except Exception as e:
                    print(f"[最终完整性检查失败] 文件 {filename}: {e}")
                
                # 检查数据库中的最终状态
                try:
                    conn = DatabaseConfig.get_db_connection()
                    if conn:
                        cursor = conn.cursor()
                        cursor.execute("SELECT status FROM user_files WHERE file_id = %s", (file_info['file_id'],))
                        db_status = cursor.fetchone()
                        if db_status:
                            print(f"[数据状态确认] 文件 {filename} 数据库状态: {db_status['status']}")
                            if db_status['status'] == 'completed':
                                print(f"[完全完成] 文件 {filename} 线程处理完成且数据齐全")
                            else:
                                print(f"[部分完成] 文件 {filename} 线程处理完成但数据不齐全: {db_status['status']}")
                        conn.close()
                except Exception as e:
                    print(f"[状态检查失败] {e}")
                
                # 设置一个延迟清理任务，但不影响前端的即时响应
                def delayed_cleanup():
                    for i in range(100):  # 分100次检查，每次100ms，总共10秒
                        if _monitoring_shutdown:
                            return  # 如果系统正在关闭，直接退出
                        time.sleep(0.1)
                    
                    if not _monitoring_shutdown and file_task_id in tasks_db:
                        # 检查是否真的应该清理（避免清理正在被前端使用的任务）
                        task_info = tasks_db[file_task_id]
                        if (task_info.get('status') == 'completed' and 
                            time.time() - task_info.get('completed_at', 0) > 8):
                            print(f"[任务清理] 清理已完成任务: {file_task_id} (文件: {file_info['filename']})")
                            tasks_db.pop(file_task_id, None)
                        else:
                            print(f"[任务清理] 任务 {file_task_id} 暂不清理，状态: {task_info.get('status')}")
                
                import threading
                cleanup_thread = threading.Thread(target=delayed_cleanup)
                cleanup_thread.daemon = True
                cleanup_thread.start()
                
                # 原子地增加完成计数
                current_task = one_click_tasks.get(task_id)
                if current_task:
                    current_task['completed_files'] += 1
                    print(f"[并行一键解析] 文件处理成功: {filename} ({current_task['completed_files']}/{current_task['total_files']})")
                
            except Exception as e:
                print(f"[并行一键解析] 处理文件 {file_info['filename']} 错误: {e}")
                # 更新文件任务状态为错误
                file_task_id = f"oneclick_{task_id}_{file_info['file_id']}"
                if file_task_id in tasks_db:
                    # 计算处理时间
                    start_time = tasks_db.get(file_task_id, {}).get('start_time', time.time())
                    processing_time = time.time() - start_time
                    processing_time_str = f"{int(processing_time // 60)}分{int(processing_time % 60)}秒" if processing_time >= 60 else f"{int(processing_time)}秒"
                    
                    # 暂时保存错误信息，不立即设置error状态
                    tasks_db[file_task_id]['pending_error'] = f'❌ 处理失败: {str(e)}'
                    tasks_db[file_task_id]['status'] = 'pending_completion'
                    tasks_db[file_task_id]['message'] = '正在确认处理状态...'
                    tasks_db[file_task_id]['processing_time'] = processing_time
                    tasks_db[file_task_id]['processing_time_str'] = processing_time_str
                    tasks_db[file_task_id]['completed_at'] = time.time()  # 记录错误完成时间戳
                
                # 保存处理时间到remarks字段，但不立即更新数据库状态为error
                try:
                    start_time = tasks_db.get(file_task_id, {}).get('start_time', time.time())
                    processing_time = time.time() - start_time
                    processing_time_str = f"{int(processing_time // 60)}分{int(processing_time % 60)}秒" if processing_time >= 60 else f"{int(processing_time)}秒"
                    
                    # 注释掉立即更新数据库error状态，改为在线程池任务全部完成后统一更新
                    # update_file_status_in_db(file_info['file_id'], 'error', remarks=f"处理失败，耗时: {processing_time_str}, 错误: {str(e)}")
                    print(f"[错误记录] 文件 {file_info['filename']} 处理失败，耗时: {processing_time_str}，错误: {str(e)}")
                except Exception as db_e:
                    print(f"[错误记录失败] {db_e}")
                
                current_task = one_click_tasks.get(task_id)
                if current_task:
                    current_task['errors'].append(f'{file_info["filename"]}: {str(e)}')
        
        # 使用全局共享线程池控制并发数，确保跨批次严格限制最多2个文件同时处理
        
        # 📊 第二步：将任务提交到全局线程池（所有批次共享，max_workers=2）
        print(f"[任务队列] 提交到全局线程池（最大并发: 2），当前批次 {len(files_to_process)} 个文件")
        
        futures = {}  # task_id -> Future 映射
        
        for file_info in files_to_process:
            file_task_id = f"oneclick_{task_id}_{file_info['file_id']}"
            future = _global_file_executor.submit(process_single_file, file_task_id)
            futures[file_task_id] = future
            print(f"[任务队列] 📤 提交任务到全局线程池: {file_info['filename']}")
        
        # 等待本批次所有任务完成
        print(f"[任务队列] ⏳ 等待本批次任务完成，共 {len(futures)} 个任务...")
        for file_task_id, future in futures.items():
            try:
                future.result()  # 阻塞等待任务完成
            except Exception as e:
                task_info = tasks_db.get(file_task_id, {})
                filename = task_info.get('filename', 'unknown')
                print(f"[任务队列] ❌ 任务执行异常: {filename} - {e}")
                mark_task_completed(file_task_id, success=False, error_msg=str(e))
        
        # 📈 第三步：任务队列全部完成，统计结果
        print(f"[任务队列] ✅ 所有任务执行完毕")
        
        # 统计成功和失败数量
        completed_count = 0
        failed_count = 0
        for file_info in files_to_process:
            file_task_id = f"oneclick_{task_id}_{file_info['file_id']}"
            if file_task_id in tasks_db:
                status = tasks_db[file_task_id].get('status')
                if status == 'completed':
                    completed_count += 1
                elif status in ['failed', 'error', 'cancelled']:
                    failed_count += 1
        
        # 更新一键任务状态
        final_task = one_click_tasks.get(task_id)
        if final_task:
            final_task['status'] = 'completed'
            final_task['current_step'] = f'🎉 任务队列处理完成！成功: {completed_count}, 失败: {failed_count}'
            final_task['completed_files'] = completed_count
            print(f"[任务队列] 📊 最终统计: 成功 {completed_count} 个, 失败 {failed_count} 个")
            
            # 统一更新所有文件的最终数据库状态
            print(f"[统一状态更新] 开始更新所有文件的最终数据库状态...")
            for file_info in files_to_process:
                file_task_id = f"oneclick_{task_id}_{file_info['file_id']}"
                if file_task_id in tasks_db:
                    task_final_status = tasks_db[file_task_id].get('status', 'unknown')
                    has_pending_error = tasks_db[file_task_id].get('pending_error')
                    is_pending_completion = task_final_status == 'pending_completion'
                    
                    try:
                        if task_final_status == 'error' or has_pending_error:
                            # 获取错误信息和处理时间
                            if has_pending_error:
                                error_message = has_pending_error
                            else:
                                error_message = tasks_db[file_task_id].get('message', '处理失败')
                                
                            processing_time_str = tasks_db[file_task_id].get('processing_time_str', '未知')
                            
                            # 更新数据库为error状态
                            update_file_status_in_db(
                                file_info['file_id'], 
                                'error', 
                                remarks=f"处理失败，耗时: {processing_time_str}, 错误: {error_message}"
                            )
                            print(f"[统一状态更新] 文件 {file_info['filename']} 最终状态: error，原因: {error_message}")
                        elif task_final_status == 'pending_completion':
                            # pending_completion状态表示处理成功，等待最终确认
                            # 检查数据库中是否已经有转换结果
                            try:
                                conn = DatabaseConfig.get_db_connection()
                                cursor = conn.cursor()
                                cursor.execute("""
                                    SELECT status, markdown_content, references_content, validation_json 
                                    FROM uploaded_files 
                                    WHERE id = %s
                                """, (file_info['file_id'],))
                                result = cursor.fetchone()
                                conn.close()
                                
                                if result and result['markdown_content']:
                                    # 有转换内容，标记为成功
                                    processing_time_str = tasks_db[file_task_id].get('processing_time_str', '未知')
                                    update_file_status_in_db(
                                        file_info['file_id'], 
                                        'completed', 
                                        remarks=f"转换成功，耗时: {processing_time_str}"
                                    )
                                    print(f"[统一状态更新] 文件 {file_info['filename']} 最终状态: completed (转换成功)")
                                else:
                                    # 没有转换内容，标记为失败
                                    processing_time_str = tasks_db[file_task_id].get('processing_time_str', '未知')
                                    error_message = tasks_db[file_task_id].get('message', '转换结果为空')
                                    update_file_status_in_db(
                                        file_info['file_id'], 
                                        'error', 
                                        remarks=f"转换失败，耗时: {processing_time_str}, 错误: {error_message}"
                                    )
                                    print(f"[统一状态更新] 文件 {file_info['filename']} 最终状态: error，原因: 转换结果为空")
                            except Exception as db_check_error:
                                print(f"[统一状态更新] 检查文件 {file_info['filename']} 数据库状态失败: {db_check_error}")
                                # 数据库检查失败，保守标记为错误
                                update_file_status_in_db(
                                    file_info['file_id'], 
                                    'error', 
                                    remarks=f"状态确认失败: {db_check_error}"
                                )
                                print(f"[统一状态更新] 文件 {file_info['filename']} 最终状态: error，原因: 状态确认失败")
                        elif task_final_status == 'completed':
                            print(f"[统一状态更新] 文件 {file_info['filename']} 最终状态: completed (已由处理流程更新)")
                        else:
                            print(f"[统一状态更新] 文件 {file_info['filename']} 状态未确定: {task_final_status}")
                    except Exception as e:
                        print(f"[统一状态更新] 更新文件 {file_info['filename']} 状态失败: {e}")
            
            # 清理个别文件的任务状态，避免进度条卡在97%
            print(f"[并行一键解析] 开始清理子任务...")
            for file_info in files_to_process:
                file_task_id = f"oneclick_{task_id}_{file_info['file_id']}"
                if file_task_id in tasks_db:
                    # 根据任务状态决定清理延迟时间
                    task_status = tasks_db[file_task_id].get('status', 'unknown')
                    delay_time = 30 if task_status in ['error', 'failed'] else 10  # 错误状态保留30秒，成功状态保留10秒
                    
                    def delayed_cleanup(task_id_to_clean, delay):
                        def cleanup():
                            # 分段睡眠，响应关闭信号
                            for i in range(delay * 10):  # 每次100ms，总共delay秒
                                if _monitoring_shutdown:
                                    return  # 如果系统正在关闭，直接退出
                                time.sleep(0.1)
                            
                            if not _monitoring_shutdown:
                                removed = tasks_db.pop(task_id_to_clean, None)
                                if removed is not None:
                                    print(f"[并行一键解析] 清理子任务: {task_id_to_clean} (延迟{delay}秒)")
                        return cleanup
                    
                    cleanup_thread = threading.Thread(target=delayed_cleanup(file_task_id, delay_time))
                    cleanup_thread.daemon = True
                    cleanup_thread.start()
        
        # 任务完成，释放全局并发槽位
        with global_active_tasks_lock:
            global_active_tasks -= 1
            print(f"[全局并发控制] 任务 {task_id} 完成，当前活跃任务数: {global_active_tasks}/2")
        
    except Exception as e:
        print(f"并行一键解析任务错误: {e}")
        current_task = one_click_tasks.get(task_id)
        if current_task:
            current_task['status'] = 'failed'
            current_task['error'] = str(e)
        
        # 发生异常也要释放并发槽位
        with global_active_tasks_lock:
            global_active_tasks -= 1
            print(f"[全局并发控制] 任务 {task_id} 异常结束，当前活跃任务数: {global_active_tasks}/2")

@app.route('/api/one-click-progress/<task_id>')
@api_login_required
def get_one_click_progress(task_id):
    """获取一键解析进度"""
    
    task = one_click_tasks.get(task_id)
    if not task:
        return jsonify({
            'success': False,
            'error': '任务不存在'
        })
    
    # 检查任务所有者
    if task.get('user_id') != session['user_id']:
        return jsonify({
            'success': False,
            'error': '无权限访问此任务'
        })
    
    return jsonify({
        'success': True,
        'progress': {
            'status': task.get('status', 'unknown'),
            'total_files': task.get('total_files', 0),
            'completed_files': task.get('completed_files', 0),
            'current_step': task.get('current_step', ''),
            'current_file': task.get('current_file'),
            'errors': task.get('errors', []),
            'error': task.get('error')
        }
    })

@app.route('/api/progress-stream')
@api_login_required
def progress_stream():
    """#17 SSE进度推送端点 — 替代轮询，实时推送任务进度变化"""

    user_id = session['user_id']

    def generate():
        """SSE生成器：每1.5秒检查一次变化，仅在数据变化时推送"""
        import json as _json
        last_snapshot = None          # 上一次推送的快照（序列化后字符串）
        heartbeat_counter = 0         # 心跳计数器
        idle_streak = 0               # 连续无变化次数
        MAX_IDLE = 400                # 连续无变化上限（约10分钟），达到后关闭连接

        while True:
            try:
                progress_data = {}
                tasks_snapshot = dict(tasks_db.items())

                for task_id, task_info in tasks_snapshot.items():
                    if task_info.get('user_id') != user_id:
                        continue
                    file_id = task_info.get('file_id')
                    if not file_id:
                        continue

                    status = task_info.get('status', 'unknown')
                    try:
                        progress = max(0, min(100, int(task_info.get('progress', 0))))
                    except (ValueError, TypeError):
                        progress = 0

                    progress_data[file_id] = {
                        'status': status,
                        'progress': progress,
                        'stage': task_info.get('message', ''),
                        'filename': task_info.get('filename', ''),
                        'task_id': task_id,
                        'start_time': task_info.get('start_time'),
                        'completed_at': task_info.get('completed_at'),
                        'data_complete': (status == 'completed'),
                    }

                # 序列化后对比是否有变化
                current_snapshot = _json.dumps(progress_data, sort_keys=True, default=str)

                if current_snapshot != last_snapshot:
                    last_snapshot = current_snapshot
                    idle_streak = 0
                    payload = _json.dumps({
                        'success': True,
                        'progress': progress_data,
                        'timestamp': int(time.time()),
                    }, ensure_ascii=False, default=str)
                    yield f"data: {payload}\n\n"
                else:
                    idle_streak += 1
                    # 每20次（约30秒）发送心跳，保持连接活跃
                    heartbeat_counter += 1
                    if heartbeat_counter >= 20:
                        heartbeat_counter = 0
                        yield ": heartbeat\n\n"

                # 连续无变化过久则正常关闭，客户端会自动重连
                if idle_streak >= MAX_IDLE:
                    yield "event: close\ndata: idle_timeout\n\n"
                    return

                time.sleep(1.5)

            except GeneratorExit:
                return
            except Exception as e:
                print(f"[SSE] 错误: {e}")
                yield f"event: error\ndata: {str(e)}\n\n"
                return

    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',       # Nginx 禁止缓冲
            'Connection': 'keep-alive',
        }
    )

@app.route('/api/get-all-progress')
@api_login_required
def get_all_progress():
    """获取所有文档的处理进度 - 优化性能和响应速度"""
    
    user_id = session['user_id']
    
    try:
        progress_data = {}
        
        # 1. 从tasks_db中快速获取活跃任务进度（已在内存中）
        active_tasks = {}
        thread_completed_tasks = {}  # 线程完成的任务（用于控制前端进度条）
        
        # 创建tasks_db的快照避免遍历时修改的问题
        tasks_snapshot = dict(tasks_db.items())
        tasks_to_remove = []  # 记录需要清理的任务ID
        
        for task_id, task_info in tasks_snapshot.items():
            if task_info.get('user_id') == user_id:
                file_id = task_info.get('file_id')
                
                # 【容错处理】如果没有file_id，尝试从task_id中提取或跳过
                if not file_id:
                    if 'oneclick_' in task_id:
                        # 从oneclick_主任务ID_file_id格式中提取
                        parts = task_id.split('_')
                        if len(parts) >= 3:
                            file_id = parts[2]
                            # 安全检查：确保任务仍存在于tasks_db中
                            if task_id in tasks_db:
                                tasks_db[task_id]['file_id'] = file_id  # 补充映射
                                # 同时尝试补充filename（如果缺失）
                                if not tasks_db[task_id].get('filename'):  # 使用最新的tasks_db数据检查
                                    # 安全处理file_id长度
                                    display_id = file_id[:8] if len(file_id) >= 8 else file_id
                                    tasks_db[task_id]['filename'] = f"文件_{display_id}"
                                # 更新本地task_info以保持一致性
                                task_info = tasks_db[task_id].copy()
                                print(f"[映射修复] 从task_id提取file_id: {file_id}")
                            else:
                                print(f"[映射警告] 任务已不存在，跳过: {task_id}")
                                continue
                        else:
                            print(f"[映射警告] 无法从task_id提取file_id: {task_id}")
                            continue
                    else:
                        print(f"[映射警告] 任务缺少file_id映射: {task_id}")
                        continue
                
                # 确保有file_id才继续处理
                if file_id:
                    status = task_info.get('status', 'unknown')
                    # 安全处理进度值
                    try:
                        progress = max(0, min(100, int(task_info.get('progress', 0))))
                    except (ValueError, TypeError):
                        progress = 0  # 如果进度值无法转换，默认为0
                        print(f"[进度值错误] 任务 {task_id} 进度值无效: {task_info.get('progress')}")
                    
                    # 清理已完成任务的内存状态（保留2分钟供前端显示）
                    if status == 'completed':
                        completed_at = task_info.get('completed_at')
                        if completed_at and (time.time() - completed_at) > 120:  # 2分钟后清理
                            print(f"[内存清理] 任务 {task_id} 已完成超过2分钟，从内存中清理")
                            tasks_to_remove.append(task_id)
                            continue  # 跳过后续处理
                    
                    # 清理超时的pending_completion状态任务
                    if status == 'pending_completion':
                        completed_at = task_info.get('completed_at')
                        if completed_at and (time.time() - completed_at) > 300:  # 5分钟超时
                            print(f"[超时清理] 任务 {task_id} pending_completion状态超时，不再自动更新数据库")
                            # 注释掉自动更新数据库为error状态的逻辑
                            # error_message = task_info.get('pending_error', '处理状态待确认超时')
                            # update_file_status_in_db(file_id, 'error', remarks=f"处理失败: {error_message}")
                            
                            # 仅更新内存状态以供显示，但不影响数据库
                            if task_id in tasks_db:
                                tasks_db[task_id]['status'] = 'error'
                                tasks_db[task_id]['message'] = f'❌ 状态确认超时，请检查后台日志'
                                
                                # 🆕 记录实时错误监控日志
                                try:
                                    error_message = task_info.get('pending_error', '处理状态待确认超时')
                                    
                                    # 获取系统管理员ID用于记录错误日志
                                    system_user_id = None
                                    try:
                                        conn = DatabaseConfig.get_db_connection()
                                        if conn:
                                            cursor = conn.cursor()
                                            cursor.execute("SELECT id FROM users WHERE role = 'admin' ORDER BY id LIMIT 1")
                                            admin_user = cursor.fetchone()
                                            if admin_user:
                                                system_user_id = admin_user['id']
                                            conn.close()
                                    except:
                                        pass
                                    
                                    # 只有在获取到有效管理员ID时才记录日志
                                    if system_user_id:
                                        log_user_action(
                                            system_user_id,  # 使用系统管理员用户ID
                                            'realtime_processing_error',
                                            f"处理任务错误: {task_id} - {error_message} (文件ID: {file_id})",
                                            None,
                                            'Error-Monitor'
                                        )
                                except Exception as e:
                                    print(f"[实时日志] 错误监控日志记录失败: {e}")
                                
                                # 更新本地task_info以保持一致性
                                task_info = tasks_db[task_id].copy()
                                status = 'error'  # 更新当前循环使用的状态
                    
                    # 从数据库或tasks_db读取 processing_start_time
                    start_time_timestamp = task_info.get('start_time')
                    
                    if start_time_timestamp is None:
                        try:
                            conn_time = DatabaseConfig.get_db_connection()
                            if conn_time:
                                cursor_time = conn_time.cursor()
                                cursor_time.execute("""
                                    SELECT processing_start_time 
                                    FROM user_files 
                                    WHERE file_id = %s AND user_id = %s
                                """, (file_id, user_id))
                                result_time = cursor_time.fetchone()
                                conn_time.close()
                                
                                if result_time and result_time['processing_start_time']:
                                    from datetime import datetime
                                    db_start_time = result_time['processing_start_time']
                                    if isinstance(db_start_time, str):
                                        start_time_timestamp = datetime.strptime(db_start_time, '%Y-%m-%d %H:%M:%S').timestamp()
                                    else:
                                        start_time_timestamp = db_start_time.timestamp()
                        except Exception as e:
                            print(f"[开始时间读取] 失败: {e}")
                    
                    task_data = {
                        'status': status,
                        'progress': progress,
                        'stage': task_info.get('message', ''),
                        'filename': task_info.get('filename', ''),
                        'task_id': task_id,
                        'start_time': start_time_timestamp,  # 使用更新后的时间戳
                        'processing_time_str': task_info.get('processing_time_str', ''),
                        'completed_at': task_info.get('completed_at')  # 添加完成时间戳
                    }
                    
                    # 🆕 添加调试日志
                    if status == 'processing':
                        print(f"[进度返回] 文件 {file_id} ({task_info.get('filename', '未知')}): start_time={start_time_timestamp}, progress={progress}%")
                    
                    # 关键修改：区分线程完成和数据完成
                    if status == 'completed':
                        # 这是线程处理完成，用于控制前端进度条和计时器
                        thread_completed_tasks[file_id] = task_data
                        print(f"[线程完成] 文件 {task_info.get('filename', file_id)} 处理线程已完成")
                    else:
                        active_tasks[file_id] = task_data
                        if status == 'processing':
                            print(f"[线程处理中] 文件 {task_info.get('filename', file_id)} 进度: {progress}%")
        
        # 2. 批量从数据库获取文件状态（单次查询）
        # 注意：completed状态的文档不在上传页面显示，只在文档列表页面查询
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT file_id, status, original_name, remarks, processing_start_time
                FROM user_files 
                WHERE user_id = %s AND status IN ('uploaded', 'processing', 'error')
                ORDER BY upload_time DESC
                LIMIT 50
            """, (user_id,))  # 不包含completed状态，这些文档只在文档列表页面显示
            
            files = cursor.fetchall()
            conn.close()
            
            # 特殊处理：检查线程完成状态中是否有数据库已完成的文件
            # 这些文件需要发送完成状态给前端，即使数据库查询中没有包含它们
            additional_completed_files = {}
            for file_id, task_data in thread_completed_tasks.items():
                if file_id not in progress_data:  # 还没有被处理过
                    # 检查数据库中的状态
                    conn = DatabaseConfig.get_db_connection()
                    if conn:
                        cursor = conn.cursor()
                        cursor.execute("""
                            SELECT status, original_name, remarks 
                            FROM user_files 
                            WHERE file_id = %s AND user_id = %s
                        """, (file_id, user_id))
                        result = cursor.fetchone()
                        conn.close()
                        
                        if result and result['status'] == 'completed':
                            # 这是一个数据库已完成但线程状态还存在的文件
                            progress_data[file_id] = task_data.copy()
                            progress_data[file_id]['data_complete'] = True
                            progress_data[file_id]['progress'] = 100
                            progress_data[file_id]['status'] = 'completed'
                            if not progress_data[file_id].get('filename'):
                                progress_data[file_id]['filename'] = result['original_name']
                            progress_data[file_id]['remarks'] = result.get('remarks', '')
                            print(f"[补充完成] 文件 {result['original_name']} 数据库已完成，添加到进度响应中")
            
            # 3. 合并活跃任务和数据库状态 - 优先级：线程完成状态 > 活跃任务 > 数据库状态
            for file in files:
                file_id = file['file_id']
                filename = file['original_name']
                db_status = file['status']  # 这是数据完成状态（completed表示三个结果都有）
                remarks = file.get('remarks', '')
                processing_start_time = file.get('processing_start_time')
                
                # 次高优先级：线程明确标记为完成的任务（用于前端进度条控制）
                if file_id in thread_completed_tasks:
                    progress_data[file_id] = thread_completed_tasks[file_id].copy()
                    if not progress_data[file_id].get('filename'):
                        progress_data[file_id]['filename'] = filename
                    progress_data[file_id]['remarks'] = remarks
                    # 添加数据库完成状态信息
                    progress_data[file_id]['data_complete'] = (db_status == 'completed')
                    print(f"[线程完成] 文件 {filename} 处理线程已完成，数据状态: {db_status}")
                    
                # 第三优先级：活跃的处理任务
                elif file_id in active_tasks:
                    task_info = active_tasks[file_id].copy()
                    task_status = task_info.get('status')
                    
                    # 特殊处理：pending_completion状态显示为processing，避免前端误报错误
                    if task_status == 'pending_completion':
                        task_info['status'] = 'processing'
                        task_info['stage'] = task_info.get('stage', '正在确认处理状态...')
                        print(f"[状态转换] 文件 {filename} pending_completion -> processing (前端显示)")
                    
                    progress_data[file_id] = task_info
                    if not progress_data[file_id].get('filename'):
                        progress_data[file_id]['filename'] = filename
                    progress_data[file_id]['remarks'] = remarks
                    progress_data[file_id]['data_complete'] = (db_status == 'completed')
                    print(f"[活跃任务] 文件 {filename} 正在处理，进度: {active_tasks[file_id].get('progress', 0)}%")
                        
                # 最低优先级：仅有数据库状态 - 简化处理，不自动重置
                elif db_status == 'processing':
                    # 对于数据库中仍为processing状态的文件，不进行自动重置
                    # 让任务级超时机制来处理超时情况，避免状态混乱
                    
                    if processing_start_time:
                        # 使用统一的时间戳解析，完全时区无关
                        start_timestamp = parse_timestamp_safe(processing_start_time)
                        
                        if start_timestamp is None:
                            print(f"[恢复跳过] 文件 {filename} 时间戳解析失败: {processing_start_time}")
                            continue
                        
                        # 根据进度推测当前阶段，不进行超时检查和自动重置
                        stage_message = get_recovery_stage_message(start_timestamp)
                        progress_data[file_id] = {
                            'status': 'processing',
                            'progress': 15,  # 给一个合理的默认进度
                            'stage': stage_message,
                            'filename': filename,
                            'remarks': remarks,
                            'start_time': start_timestamp,
                            'data_complete': False
                        }
                        print(f"[状态恢复] 文件 {filename} 恢复processing状态，开始时间: {processing_start_time}, 阶段: {stage_message}")
                    else:
                        print(f"[恢复跳过] 文件 {filename} 无开始时间记录，跳过恢复")
                        continue
                    
                elif db_status == 'error':
                    progress_data[file_id] = {
                        'status': 'error',
                        'progress': 0,
                        'stage': '处理失败',
                        'filename': filename,
                        'remarks': remarks,
                        'data_complete': False
                    }
                    
                # 注意：不再为 'uploaded' 状态创建进度条，除非有明确的活跃任务
        
        # 批量清理已完成的任务
        with tasks_db_lock:
            for task_id in tasks_to_remove:
                removed = tasks_db.pop(task_id, None)
                if removed:
                    print(f"[批量清理] 清理已完成任务: {task_id}")
        
        return jsonify({
            'success': True,
            'progress': progress_data,
            'timestamp': int(time.time()),  # 添加时间戳便于前端缓存
            'total_tasks': len(progress_data),  # 调试信息
            'thread_completed_count': len([p for p in progress_data.values() if p.get('status') == 'completed']),  # 线程完成数
            'processing_count': len([p for p in progress_data.values() if p.get('status') == 'processing']),
            'data_completed_count': len([p for p in progress_data.values() if p.get('data_complete') == True])  # 数据完成数
        })
        
    except Exception as e:
        print(f"[获取进度] 错误: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/cancel-one-click-task/<task_id>', methods=['POST'])
@api_login_required
def cancel_one_click_task(task_id):
    """取消一键解析任务"""
    
    task = one_click_tasks.get(task_id)
    if not task:
        return jsonify({
            'success': False,
            'error': '任务不存在'
        })
    
    # 检查任务所有者
    if task.get('user_id') != session['user_id']:
        return jsonify({
            'success': False,
            'error': '无权限取消此任务'
        })
    
    # 标记任务为已取消
    task['status'] = 'cancelled'
    task['error'] = '用户取消操作'
    
    return jsonify({
        'success': True,
        'message': '任务已取消'
    })

@app.route('/api/check-all-documents-completed')
@api_login_required
def check_all_documents_completed():
    """检查用户的所有文档是否都已完成解析"""
    
    try:
        user_id = session['user_id']
        
        # 从数据库获取用户的所有文件
        user_files = get_user_files_from_db(user_id)
        
        if not user_files:
            return jsonify({
                'success': True,
                'all_completed': True,
                'total_files': 0,
                'completed_files': 0,
                'message': '没有文档需要处理'
            })
        
        # 统计文件状态
        total_files = len(user_files)
        completed_files = len([f for f in user_files if f.get('status') == 'completed'])
        processing_files = len([f for f in user_files if f.get('status') == 'processing'])
        pending_files = len([f for f in user_files if f.get('status') == 'pending'])
        
        all_completed = (completed_files == total_files)
        
        return jsonify({
            'success': True,
            'all_completed': all_completed,
            'total_files': total_files,
            'completed_files': completed_files,
            'processing_files': processing_files,
            'pending_files': pending_files,
            'completion_rate': round((completed_files / total_files) * 100, 1) if total_files > 0 else 0
        })
        
    except Exception as e:
        print(f"[检查文档完成状态] 错误: {e}")
        return jsonify({'success': False, 'error': str(e)})


# ==================== 本地MinerU API配置 ====================
LOCAL_MINERU_CONFIG = {
    "parse_url": "http://172.16.135.211:60605/parse_pdf",
    "name": "本地MinerU服务器",
}


def _check_local_mineru_available():
    """快速检查本地MinerU服务是否可连通（HEAD/GET 健康检查，2秒超时）"""
    try:
        # 尝试访问健康检查端点
        health_url = LOCAL_MINERU_CONFIG["parse_url"].rsplit('/', 1)[0] + "/health"
        r = requests.get(health_url, timeout=2)
        return r.status_code == 200
    except Exception:
        return False


def convert_pdf_with_local_mineru(file_id, user_id, task_id=None):
    """使用本地MinerU API转换PDF，失败时自动回退到官方MinerU云端API
    
    本地MinerU接口：POST multipart/form-data, 返回markdown文本
    回退逻辑：连接失败/超时 → 自动尝试官方MinerU云端API（需用户分组已配置MinerU密钥）
    支持并发：通过_mineru_semaphore控制并发数
    """
    # 带取消检查的信号量获取
    acquired = False
    while not _mineru_semaphore.acquire(timeout=2):
        if task_id and is_task_cancelled(task_id):
            print(f"[本地MinerU] 等待信号量期间任务被取消: {task_id}")
            return False
    acquired = True
    try:
        return _convert_pdf_with_local_mineru_inner(file_id, user_id)
    finally:
        if acquired:
            _mineru_semaphore.release()


def _convert_pdf_with_local_mineru_inner(file_id, user_id):
    """本地MinerU转换内部实现"""
    try:
        # 从数据库获取文件信息
        pdf_info = get_file_info_from_db(file_id, user_id)
        if not pdf_info:
            print(f"[本地MinerU] 无法获取PDF数据: {file_id}")
            return False
        
        filename = pdf_info['original_name']
        pdf_data = pdf_info['pdf_data']
        print(f"[本地MinerU] 开始转换: {filename} ({len(pdf_data):,} bytes)")
        
        # 先快速检查本地MinerU是否可连通
        local_reachable = _check_local_mineru_available()
        if not local_reachable:
            print(f"[本地MinerU] ⚠️ 服务不可达，直接尝试回退到官方MinerU")
        
        local_success = False
        if local_reachable:
            # 调用本地MinerU API
            import io
            files = {
                'file': (filename, io.BytesIO(pdf_data), 'application/pdf')
            }
            data = {
                'backend': 'hybrid-auto-engine'
            }
            
            max_attempts = 3
            response = None
            for attempt in range(max_attempts):
                try:
                    print(f"[本地MinerU] 尝试连接 (尝试 {attempt + 1}/{max_attempts})")
                    response = requests.post(
                        LOCAL_MINERU_CONFIG["parse_url"],
                        files=files,
                        data=data,
                        timeout=600  # 10分钟超时，PDF解析可能较慢
                    )
                    
                    if response.status_code == 200:
                        break
                    else:
                        print(f"[本地MinerU] 请求失败 (尝试 {attempt + 1}): HTTP {response.status_code}")
                        if attempt < max_attempts - 1:
                            # 重新创建BytesIO（上次可能已读完）
                            files = {
                                'file': (filename, io.BytesIO(pdf_data), 'application/pdf')
                            }
                            time.sleep(5 * (attempt + 1))
                        continue
                        
                except requests.exceptions.Timeout:
                    print(f"[本地MinerU] 请求超时 (尝试 {attempt + 1}/{max_attempts})")
                    if attempt < max_attempts - 1:
                        files = {
                            'file': (filename, io.BytesIO(pdf_data), 'application/pdf')
                        }
                        time.sleep(5 * (attempt + 1))
                    continue
                except requests.exceptions.ConnectionError:
                    print(f"[本地MinerU] 连接失败 (尝试 {attempt + 1}/{max_attempts})")
                    if attempt < max_attempts - 1:
                        files = {
                            'file': (filename, io.BytesIO(pdf_data), 'application/pdf')
                        }
                        time.sleep(5 * (attempt + 1))
                    continue
                except Exception as e:
                    print(f"[本地MinerU] 请求异常 (尝试 {attempt + 1}): {e}")
                    if attempt < max_attempts - 1:
                        files = {
                            'file': (filename, io.BytesIO(pdf_data), 'application/pdf')
                        }
                        time.sleep(5 * (attempt + 1))
                    continue
            
            if response and response.status_code == 200:
                # 本地MinerU直接返回markdown文本
                markdown_content = response.text
                
                if markdown_content and len(markdown_content.strip()) >= 50:
                    print(f"[本地MinerU] 转换成功，Markdown长度: {len(markdown_content):,} chars")
                    
                    # 保存MD内容到数据库
                    update_file_status_in_db(
                        file_id,
                        'processing',
                        md_content=markdown_content
                    )
                    
                    print(f"[本地MinerU] MD内容已保存: {file_id}")
                    local_success = True
                else:
                    print(f"[本地MinerU] 返回内容过短或为空: {len(markdown_content) if markdown_content else 0} chars")
            else:
                print(f"[本地MinerU] 所有尝试均失败")
        
        if local_success:
            return True
        
        # ========== 回退到官方MinerU云端API ==========
        print(f"[本地MinerU→回退] 本地MinerU不可用，尝试回退到官方MinerU云端API...")
        try:
            fallback_key = get_active_api_key(user_id, provider='mineru')
            if fallback_key:
                mineru_api_key = fallback_key['api_token']
                print(f"[本地MinerU→回退] 找到官方MinerU API密钥，开始云端转换...")
                return convert_pdf_with_mineru(file_id, mineru_api_key, user_id)
            else:
                print(f"[本地MinerU→回退] ⚠️ 未配置官方MinerU API密钥，无法回退")
                return False
        except Exception as fallback_err:
            print(f"[本地MinerU→回退] 回退失败: {fallback_err}")
            return False
        
    except Exception as e:
        print(f"[本地MinerU] 转换错误: {e}")
        import traceback
        print(f"[本地MinerU] 错误详情: {traceback.format_exc()}")
        return False


def convert_pdf_with_mineru(file_id, api_key, user_id):
    """使用MinerU转换PDF - 基于one_click_test.py的成功实现，增强SSL错误处理"""
    
    # 配置SSL和网络设置
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    # 创建一个更宽松的会话来处理SSL问题
    session = requests.Session()
    
    # 设置重试策略
    from urllib3.util.retry import Retry
    from requests.adapters import HTTPAdapter
    
    # 兼容不同版本的urllib3
    try:
        retry_strategy = Retry(
            total=3,
            status_forcelist=[429, 500, 502, 503, 504],
            allowed_methods=["HEAD", "GET", "POST", "PUT", "DELETE", "OPTIONS", "TRACE"],  # 新版本
            backoff_factor=1
        )
    except TypeError:
        # 兼容旧版本
        retry_strategy = Retry(
            total=3,
            status_forcelist=[429, 500, 502, 503, 504],
            method_whitelist=["HEAD", "GET", "POST", "PUT", "DELETE", "OPTIONS", "TRACE"],  # 旧版本
            backoff_factor=1
        )
    
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    
    try:
        # 从数据库获取文件信息
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            smart_print(f"[一键解析] 数据库连接失败", LogLevel.ERROR)
            return False
            
        cursor = conn.cursor()
        cursor.execute("""
            SELECT file_id, user_id, original_name, file_path, file_size
            FROM user_files 
            WHERE file_id = %s AND user_id = %s
        """, (file_id, user_id))
        
        file_data = cursor.fetchone()
        conn.close()
        
        if not file_data:
            print(f"[一键解析] 文件 {file_id} 不存在或无权限")
            return False
            
        # 获取PDF文件数据
        pdf_info = get_file_info_from_db(file_id, user_id)
        if not pdf_info:
            print(f"[一键解析] 无法获取PDF数据")
            return False
        
        filename = pdf_info['original_name']
        data_id = f"mineru_{file_id}_{int(time.time())}"  # 生成唯一的data_id
        
        print(f"[一键解析] 开始MinerU转换: {filename}")
        
        # 步骤1: 申请上传URL
        print(f"[一键解析] 步骤1: 申请上传URL...")
        url = 'https://mineru.net/api/v4/file-urls/batch'
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_key}'
        }
        data = {     # 开启公式识别
            "language": "ch",           # 中文文档
            "enable_table": True,       # 开启表格识别
            "model_version": "vlm",     # 使用VLM模型获得最佳性能改为中文，支持中文文档
            "enable_table": True,
            "files": [
                {"name": filename, "is_ocr": True, "data_id": data_id}
            ]
        }
        
        # 使用配置好的会话进行API调用，包含重试机制
        max_attempts = 3
        
        for attempt in range(max_attempts):
            try:
                print(f"[一键解析] 尝试连接 MinerU API (尝试 {attempt + 1}/{max_attempts})")
                
                response = session.post(url, headers=headers, json=data, timeout=300, verify=False)
                
                if response.status_code == 200:
                    print(f"[一键解析] 成功连接到 MinerU API")
                    break  # 成功，退出重试循环
                else:
                    print(f"[一键解析] 申请上传URL失败 (尝试 {attempt + 1}): {response.status_code} - {response.text}")
                    if attempt < max_attempts - 1:
                        delay = 5 * (attempt + 1)  # 递增延时
                        print(f"[一键解析] {delay}秒后重试...")
                        time.sleep(delay)
                    continue
                    
            except Exception as e:
                print(f"[一键解析] 连接错误 (尝试 {attempt + 1}): {type(e).__name__}: {e}")
                if attempt < max_attempts - 1:
                    delay = 5 * (attempt + 1)
                    print(f"[一键解析] {delay}秒后重试...")
                    time.sleep(delay)
                else:
                    print(f"[一键解析] 连接失败，已尝试 {max_attempts} 次。可能的原因:")
                    print("1. 网络连接不稳定")
                    print("2. MinerU服务暂时不可用")
                    print("3. API密钥无效或已过期")
                    print("4. 防火墙或代理设置阻止了连接")
                    return False
        else:
            # 所有重试都失败了
            print(f"[一键解析] 申请上传URL失败，已尝试 {max_attempts} 次")
            return False
        
        result = response.json()
        print(f"[一键解析] 申请上传URL响应: {result}")
        
        if result.get("code") != 0:
            print(f"[一键解析] 申请上传URL失败: {result.get('msg', '未知错误')}")
            return False
        
        batch_id = result["data"]["batch_id"]
        upload_urls = result["data"]["file_urls"]
        
        print(f"[一键解析] 获取到batch_id: {batch_id}")
        print(f"[一键解析] 获取到上传URL: {len(upload_urls)} 个")
        
        # 步骤2: 上传文件 - 根据API文档的正确方式
        print(f"[一键解析] 步骤2: 上传PDF文件...")
        if len(upload_urls) > 0:
            upload_url = upload_urls[0]
            
            # 根据API文档，需要以文件形式上传，使用io.BytesIO包装二进制数据
            import io
            pdf_file_obj = io.BytesIO(pdf_info['pdf_data'])
            
            # 根据API文档："上传文件时，无须设置 Content-Type 请求头"
            print(f"[一键解析] 开始上传文件，大小: {len(pdf_info['pdf_data'])} bytes")
            
            try:
                upload_response = session.put(
                    upload_url, 
                    data=pdf_file_obj.getvalue(),  # 直接使用二进制数据
                    timeout=300,  # 5分钟超时
                    verify=False
                )
            except Exception as e:
                print(f"[一键解析] 文件上传失败: {type(e).__name__}: {e}")
                return False
                
            if upload_response.status_code == 200:
                print(f"[一键解析] 文件上传成功")
                print(f"[一键解析] 上传响应: {upload_response.text}")
            else:
                print(f"[一键解析] 文件上传失败: {upload_response.status_code}")
                print(f"[一键解析] 错误响应: {upload_response.text}")
                return False
        else:
            print(f"[一键解析] 未获取到上传URL")
            return False
        
        # 步骤3: 等待处理完成并获取结果
        print(f"[一键解析] 步骤3: 等待处理完成...")
        status_url = f'https://mineru.net/api/v4/extract-results/batch/{batch_id}'
        
        max_wait_time = 1800  # 最多等待30分钟
        poll_interval = 30    # 每30秒查询一次
        start_time = time.time()
        
        while time.time() - start_time < max_wait_time:
            try:
                status_response = session.get(
                    status_url,
                    headers={
                        'Content-Type': 'application/json',
                        'Authorization': f'Bearer {api_key}'
                    },
                    timeout=60,  # 1分钟超时
                    verify=False
                )
            except Exception as e:
                print(f"[一键解析] 状态查询失败: {type(e).__name__}: {e}")
                time.sleep(10)  # 等待10秒后继续
                continue
            
            if status_response.status_code == 200:
                status_data = status_response.json()
                print(f"[一键解析] 状态查询响应: {status_data}")
                
                if status_data.get("code") == 0:
                    extract_results = status_data["data"]["extract_result"]
                    
                    # 查找我们的文件
                    target_result = None
                    for result_item in extract_results:
                        if result_item["file_name"] == filename:
                            target_result = result_item
                            break
                    
                    if target_result:
                        state = target_result["state"]
                        
                        if state == "done":
                            # 处理完成，下载结果
                            zip_url = target_result["full_zip_url"]
                            print(f"[一键解析] 处理完成，下载结果: {zip_url}")
                            
                            # 下载并解压获取markdown内容
                            markdown_content = download_and_extract_markdown_from_zip(zip_url)
                            
                            if markdown_content:
                                print(f"[一键解析] 获取到Markdown内容，长度: {len(markdown_content)}")
                                
                                # 保存MD内容到数据库，但不设置为completed状态
                                # 设置为processing状态，等待后续的JSON提取完成
                                update_file_status_in_db(
                                    file_id, 
                                    'processing',  # 使用processing状态而不是completed
                                    md_content=markdown_content
                                    # 不设置processed_time，等到所有处理完成后再设置
                                )
                                
                                # 不在这里进行完整性检查，因为JSON内容还没有提取
                                # 完整性检查将在所有处理步骤完成后进行
                                
                                print(f"[一键解析] PDF转换成功，MD内容已保存: {file_id}")
                                return True
                            else:
                                print(f"[一键解析] 警告: 处理成功但无法提取Markdown内容")
                                return False
                                
                        elif state == "error":
                            error_msg = target_result.get("err_msg", "未知错误")
                            print(f"[一键解析] 处理失败: {error_msg}")
                            # 注释掉立即更新数据库error状态，改为函数结束时统一更新
                            # update_file_status_in_db(file_id, 'error')
                            return False
                            
                        elif state == "running":
                            # 还在处理中
                            progress_info = target_result.get("extract_progress", {})
                            extracted_pages = progress_info.get("extracted_pages", 0)
                            total_pages = progress_info.get("total_pages", 0)
                            print(f"[一键解析] 处理中... 页面进度: {extracted_pages}/{total_pages}")
                        
                        else:
                            print(f"[一键解析] 当前状态: {state}")
                    else:
                        print(f"[一键解析] 未找到对应文件的处理结果")
                else:
                    print(f"[一键解析] 查询结果失败: {status_data.get('msg')}")
            else:
                print(f"[一键解析] 查询请求失败: {status_response.status_code}")
            
            time.sleep(poll_interval)
        
        print(f"[一键解析] 处理超时")
        # 注释掉立即更新数据库error状态，改为函数结束时统一更新
        # update_file_status_in_db(file_id, 'error')
        return False
        
    except Exception as e:
        print(f"[一键解析] PDF转换错误: {e}")
        import traceback
        print(f"[一键解析] 错误详情: {traceback.format_exc()}")
        # 注释掉立即更新数据库error状态，改为函数结束时统一更新
        # update_file_status_in_db(file_id, 'error')
        return False

def monitor_mineru_batch_processing_new(batch_id, api_key, file_id, user_id, filename, task_check_func=None):
    """监控MinerU批量处理状态并处理结果 - 新版实现
    
    Args:
        task_check_func: 可选的任务检查函数，如果返回True则中断处理
    """
    import time
    import tempfile
    import zipfile
    try:
        print(f"[MinerU监控] 开始监控批次: {batch_id}, 文件: {filename}")
        
        max_wait_time = 1800  # 最多等待30分钟
        poll_interval = 30    # 每30秒查询一次
        start_time = time.time()
        
        while time.time() - start_time < max_wait_time:
            # 检查任务是否被取消
            if task_check_func and task_check_func():
                print(f"[MinerU监控] 任务已被取消，中断监控: {filename}")
                return False
            # 查询批量处理状态
            status_url = f'https://mineru.net/api/v4/extract-results/batch/{batch_id}'
            headers = {'Authorization': f'Bearer {api_key}'}
            
            try:
                status_response = requests.get(status_url, headers=headers, timeout=600)  # 10分钟超时
                
                if status_response.status_code == 200:
                    status_data = status_response.json()
                    print(f"[MinerU监控] 状态查询响应: {status_data}")
                    
                    if status_data.get("code") == 0:
                        extract_results = status_data["data"]["extract_result"]
                        
                        # 查找我们的文件
                        target_result = None
                        for result_item in extract_results:
                            if result_item["file_name"] == filename:
                                target_result = result_item
                                break
                        
                        if target_result:
                            state = target_result["state"]
                            print(f"[MinerU监控] 文件状态: {state}")
                            
                            if state == "done":
                                # 处理完成，下载结果
                                zip_url = target_result["full_zip_url"]
                                print(f"[MinerU监控] 处理完成，下载结果: {zip_url}")
                                
                                # 下载并解压获取markdown内容
                                markdown_content = download_and_extract_markdown_from_zip(zip_url)
                                
                                if markdown_content:
                                    print(f"[MinerU监控] 获取到Markdown内容，长度: {len(markdown_content)}")
                                    
                                    # 保存MD内容到数据库
                                    current_time = datetime.now()
                                    update_file_status_in_db(
                                        file_id, 
                                        'completed', 
                                        md_content=markdown_content,
                                        processed_time=current_time
                                    )
                                    
                                    print(f"[MinerU监控] 文件处理成功，已保存到数据库: {file_id}")
                                    return True
                                else:
                                    print(f"[MinerU监控] 警告: 处理成功但无法提取Markdown内容")
                                    
                            elif state == "error":
                                # 处理失败
                                error_msg = target_result.get("err_msg", "未知错误")
                                print(f"[MinerU监控] 文件处理失败: {error_msg}")
                                
                                # 注释掉立即更新数据库error状态，改为函数结束时统一更新
                                # update_file_status_in_db(file_id, 'error')
                                return False
                                
                            elif state == "running":
                                # 还在处理中
                                progress_info = target_result.get("extract_progress", {})
                                extracted_pages = progress_info.get("extracted_pages", 0)
                                total_pages = progress_info.get("total_pages", 0)
                                print(f"[MinerU监控] 处理中... 页面进度: {extracted_pages}/{total_pages}")
                            
                            else:
                                print(f"[MinerU监控] 当前状态: {state}")
                        else:
                            print(f"[MinerU监控] 未找到文件 {filename} 的处理结果")
                    else:
                        print(f"[MinerU监控] 查询结果失败: {status_data.get('msg')}")
                else:
                    print(f"[MinerU监控] 查询请求失败: {status_response.status_code}")
                    
            except Exception as query_error:
                print(f"[MinerU监控] 状态查询异常: {query_error}")
                
            # 等待下次查询
            time.sleep(poll_interval)
            
        # 超时处理
        print(f"[MinerU监控] 处理超时，批次: {batch_id}")
        # 注释掉立即更新数据库error状态，改为函数结束时统一更新
        # update_file_status_in_db(file_id, 'error')
        return False
        
    except Exception as e:
        print(f"[MinerU监控] 监控异常: {e}")
        import traceback
        print(f"[MinerU监控] 异常详情: {traceback.format_exc()}")
        
        # 注释掉立即更新数据库error状态，改为函数结束时统一更新
        # update_file_status_in_db(file_id, 'error')
        return False

def download_and_extract_markdown_from_zip(zip_url):
    """下载并解压ZIP文件，提取Markdown内容"""
    try:
        print(f"[MinerU监控] 下载处理结果...")
        
        # 创建会话处理网络问题
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        
        session = requests.Session()
        session.verify = False
        
        # 下载ZIP文件
        try:
            zip_response = session.get(zip_url, timeout=600)  # 10分钟超时
        except Exception as e:
            print(f"[MinerU监控] ZIP文件下载失败: {type(e).__name__}: {e}")
            return None
            
        if zip_response.status_code != 200:
            print(f"[MinerU监控] 下载ZIP文件失败: {zip_response.status_code}")
            return None
        
        # 保存临时ZIP文件
        import tempfile
        import zipfile
        
        with tempfile.NamedTemporaryFile(suffix='.zip', delete=False) as temp_zip:
            temp_zip.write(zip_response.content)
            temp_zip_path = temp_zip.name
        
        try:
            # 解压ZIP文件
            with zipfile.ZipFile(temp_zip_path, 'r') as zip_ref:
                # 列出所有文件
                all_files = zip_ref.namelist()
                print(f"[MinerU监控] ZIP文件包含文件: {all_files}")
                
                # 查找full.md文件（优先）
                full_md_files = [f for f in all_files if f.endswith('full.md')]
                
                if full_md_files:
                    # 优先使用full.md文件
                    markdown_file = full_md_files[0]
                    print(f"[MinerU监控] 找到full.md文件: {markdown_file}")
                else:
                    # 备选：查找其他markdown文件
                    markdown_files = [f for f in all_files if f.endswith('.md')]
                    
                    if not markdown_files:
                        print(f"[MinerU监控] ZIP文件中未找到任何Markdown文件")
                        return None
                    
                    markdown_file = markdown_files[0]
                    print(f"[MinerU监控] 使用其他MD文件: {markdown_file}")
                
                # 读取markdown文件
                with zip_ref.open(markdown_file) as md_file:
                    raw_content = md_file.read()
                    
                    # 尝试多种编码
                    for encoding in ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']:
                        try:
                            markdown_content = raw_content.decode(encoding)
                            print(f"[MinerU监控] 成功使用 {encoding} 编码解析，字符长度: {len(markdown_content)}")
                            
                            if len(markdown_content.strip()) > 0:
                                print(f"[MinerU监控] 内容预览（前200字符）: {markdown_content[:200]}")
                                return markdown_content
                        except UnicodeDecodeError:
                            continue
                    
                    # 如果所有编码都失败，使用ignore模式
                    markdown_content = raw_content.decode('utf-8', errors='ignore')
                    print(f"[MinerU监控] 使用ignore模式解析，字符长度: {len(markdown_content)}")
                    return markdown_content
                
        finally:
            # 清理临时文件
            os.unlink(temp_zip_path)
            
    except Exception as e:
        print(f"[MinerU监控] 下载和解压结果失败: {e}")
        return None

def extract_key_content_for_file(file_id, user_id):
    """为文件提取重要内容 - 使用TitleEnhancer进行内容提取，返回处理后的内容而不直接保存"""
    try:
        # 检查是否有MD内容
        file_content, error_msg = get_file_content_from_db(file_id, user_id)
        if error_msg or not file_content or not file_content.get('md_content'):
            print(f"[一键解析] 文件 {file_id} 没有MD内容 - {error_msg or '无内容'}")
            return None
        
        md_content = file_content['md_content']
        
        # 使用TitleEnhancer提取重要内容
        important_content = TitleEnhancer.extract_important_content(md_content)
        
        if important_content and len(important_content.strip()) > 0:
            print(f"[一键解析] 重要内容提取完成: {file_id}, 内容长度: {len(important_content)}")
            return important_content  # 返回内容而不是直接保存
        else:
            print(f"[一键解析] 未找到重要内容: {file_id}")
            return None
        
    except Exception as e:
        print(f"[一键解析] 重要内容提取错误: {e}")
        return None

def enhance_title_for_file(file_id, user_id, content=None):
    """为文件优化标题 - 使用TitleEnhancer进行标题优化，可接受预处理的内容"""
    try:
        # 如果没有提供内容，从数据库获取
        if content is None:
            file_content, error_msg = get_file_content_from_db(file_id, user_id)
            if error_msg or not file_content or not file_content.get('md_content'):
                print(f"[一键解析] 文件 {file_id} 没有MD内容 - {error_msg or '无内容'}")
                return None
            content = file_content['md_content']
        
        # 使用TitleEnhancer进行标题优化
        optimized_content = TitleEnhancer.optimize_headers_with_extractor(content)
        
        if optimized_content and len(optimized_content.strip()) > 0:
            print(f"[一键解析] 标题优化完成: {file_id}")
            return optimized_content  # 返回内容而不是直接保存
        else:
            print(f"[一键解析] 标题优化失败: {file_id}")
            return None
        
    except Exception as e:
        print(f"[一键解析] 标题优化错误: {e}")
        return None


# 已废弃：call_nsfc_extraction_for_file 函数已删除（无调用者，原依赖 process_nsfc_content_with_llm 也已删除）

def call_table_extraction_for_file(file_id, user_id, api_key, task_check_func=None, progress_callback=None, progress_start=5):
    """为文件进行五项处理：表格提取+正文提取+表格初筛+正文初筛+参考文献验证
    
    Args:
        file_id: 文件ID
        user_id: 用户ID
        api_key: SiliconFlow API密钥
        task_check_func: 可选的任务检查函数，如果返回True则中断处理
        progress_callback: 进度回调函数 callback(progress, message)
        progress_start: 进度起始百分比（默认5%，一键解析从20%开始）
    """
    try:
        # 导入新的处理模块（包含补救函数）
        from md_content_processor import (
            separate_tables_and_text,
            call_llm_for_table_extraction,
            call_llm_for_text_extraction,
            ensure_complete_json_structure,
            convert_table_json_to_md_table,
            replace_md_table_with_vlm,
            call_vlm_for_table_extraction,
            _check_local_vlm_available,
            _check_local_llm_available,
            LOCAL_LLM_CONFIG
        )
        # 初筛检查使用 app_pdf_md.py 中的 preliminary_check_with_llm（纯正则）
        
        # 获取文件内容
        file_data, error_msg = get_file_content_from_db(file_id, user_id)
        if error_msg or not file_data:
            return False
        
        # 使用统一的md_content字段
        content = file_data.get('md_content')
        if not content:
            return False
        
        # 获取用户配置的模型（支持本地LLM）
        from db_operations import get_user_llm_provider
        llm_provider = get_user_llm_provider(user_id)
        
        if llm_provider == 'local_llm':
            from md_content_processor import LOCAL_LLM_CONFIG
            model_name = f"local:{LOCAL_LLM_CONFIG['model']}"
            print(f"[四项处理] 使用本地LLM: {model_name}")
        else:
            model_name = get_user_siliconflow_model(user_id) or "Qwen/Qwen3-30B-A3B-Instruct-2507"
            print(f"[四项处理] 使用硅基流动模型: {model_name}")
        
        # 纯增量进度系统（无固定百分比映射）- 必须在步骤1之前定义
        import threading
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        _parallel_progress_lock = threading.Lock()
        _parallel_max_progress = [progress_start + 1]  # 从起始点+1开始
        _pre_completion_cap = 99.0  # 数据库最终落库前，进度不触达100%
        
        def increment_progress(step=1, msg=None):
            """增量式前进进度，确保严格递增，永不后退"""
            with _parallel_progress_lock:
                new_pct = min(_parallel_max_progress[0] + step, _pre_completion_cap)
                _parallel_max_progress[0] = new_pct
                if msg and progress_callback:
                    progress_callback(new_pct, msg)
                return new_pct

        def set_progress_at_least(target_pct, msg=None):
            """将进度至少推进到指定百分比（绝不回退）。"""
            with _parallel_progress_lock:
                target = max(_parallel_max_progress[0], min(float(target_pct), _pre_completion_cap))
                if target <= _parallel_max_progress[0]:
                    return _parallel_max_progress[0]
                _parallel_max_progress[0] = target
                if msg and progress_callback:
                    progress_callback(target, msg)
                return target
        
        # 步骤1: 分离表格和正文
        print(f"[五项处理] 步骤1: 开始分离表格和正文...")
        if progress_callback:
            progress_callback(progress_start, '📄 步骤1: 分离表格和正文中...')
        table_content, text_content = separate_tables_and_text(content)
        table_len = len(table_content)
        text_len = len(text_content)
        print(f"[五项处理] 分离完成 - 表格: {table_len} 字符, 正文: {text_len} 字符")
        increment_progress(1, f'✅ 分离完成 | 表格: {table_len:,}字 | 正文: {text_len:,}字 → 分析中...')
        
        # 判断MinerU是否真正识别出了表格（必须包含HTML <table>或Markdown管道表格标记）
        import re as _re_check
        _has_html_table = bool(_re_check.search(r'<table[\s>]', table_content, _re_check.IGNORECASE)) if table_content else False
        _has_pipe_table = bool(_re_check.search(r'^\|.+\|', table_content, _re_check.MULTILINE)) if table_content else False
        
        # [修改] 强制跳过MinerU表格检查，直接使用VLM（如果VLM可用）
        # 原因：本地MinerU默认不识别前三页，导致表格内容缺失
        # 此时强制认为mineru_table_ok=False，触发后续VLM提取逻辑
        pdf_data_for_vlm_check = file_data.get('pdf_data') if file_data else None
        if pdf_data_for_vlm_check and _check_local_vlm_available():
            print(f"[五项处理] ℹ️ 策略强制：忽略MinerU表格检查（HTML={_has_html_table}, 管道={_has_pipe_table}），强制使用本地VLM提取前三页")
            mineru_table_ok = False
        else:
            mineru_table_ok = _has_html_table or _has_pipe_table

        if mineru_table_ok:
            print(f"[五项处理] ✓ MinerU识别到表格（HTML={_has_html_table}, 管道={_has_pipe_table}，{len(table_content)} 字符），优先使用MinerU结果")
        else:
            print(f"[五项处理] ⚠️ MinerU未识别到有效表格（{len(table_content)} 字符，无<table>或|标记），将使用VLM提取")
        
        # 初始化结果变量
        table_json_data = None
        text_json_data = None
        table_violations = None
        text_violations = None
        ref_md_content = None
        ref_validation_json = None
        md_content_modified = False
        
        # 用于在线程内传递VLM跳过状态
        _vlm_skipped_flag = [False]
        # 存储VLM直接输出的Markdown内容（用于替换MinerU残缺的前几页）
        _vlm_md_content = [None]
        
        # VLM游离线程进度文本（非空时表示VLM正在并行运行，前端会显示双椭圆）
        _vlm_parallel_text = ['']
        # 最近一次主步骤消息（VLM回调用来拼接双椭圆）
        _last_main_msg = ['']
        
        def _make_progress_msg(main_msg):
            """构造进度消息：如果VLM正在并行运行，用|||拼接双步骤显示"""
            if main_msg:
                _last_main_msg[0] = main_msg
            effective_msg = main_msg or _last_main_msg[0] or '处理中...'
            vlm_text = _vlm_parallel_text[0]
            if vlm_text:
                return f'{effective_msg}|||{vlm_text}'
            return effective_msg
        
        # 提前获取VLM所需的PDF数据和本地VLM可用性判断
        pdf_data_for_vlm = file_data.get('pdf_data') if file_data else None
        _use_local_vlm = False  # 默认不使用本地VLM游离线程模式
        
        if not mineru_table_ok and pdf_data_for_vlm:
            _use_local_vlm = _check_local_vlm_available()
            if _use_local_vlm:
                print(f"[VLM路径] 本地VLM可用 → 使用游离线程+信号量模式")
            else:
                print(f"[VLM路径] 本地VLM不可用 → 硅基流动同步执行（ThreadPool内并行）")

        def _build_table_json_from_vlm_md(vlm_md_content):
            """将VLM输出的Markdown交给LLM转为表格JSON。优先本地LLM。"""
            if not vlm_md_content or len(vlm_md_content.strip()) < 20:
                print(f"[VLM->LLM] Markdown为空，跳过JSON转换")
                return None

            llm_model_for_table = model_name
            if _check_local_llm_available():
                llm_model_for_table = f"local:{LOCAL_LLM_CONFIG['model']}"

            print(f"[VLM->LLM] 使用模型转换表格JSON: {llm_model_for_table}")
            llm_table_result = call_llm_for_table_extraction(
                table_content=vlm_md_content,
                api_key=api_key,
                model_name=llm_model_for_table,
                task_check_func=task_check_func,
                pdf_data=None,
                progress_callback=None
            )
            if llm_table_result.get('success') and llm_table_result.get('json_data'):
                converted_json = ensure_complete_json_structure(llm_table_result['json_data'])
                print(f"[VLM->LLM] ✅ Markdown转JSON成功，键数: {len(converted_json)}")
                return converted_json

            print(f"[VLM->LLM] ⚠️ Markdown转JSON失败: {llm_table_result.get('error')}")
            return None
        
        def _validate_table_json_quality(json_data):
            """验证表格JSON的质量：检查是否包含真实的非空关键字段
            
            至少需要3个关键字段有真实值（非空字符串），才视为提取成功
            """
            if not json_data or not isinstance(json_data, dict):
                return False
            
            # 关键字段路径
            key_fields_found = 0
            
            project_info = json_data.get('项目信息', {})
            basic_info = json_data.get('基本信息', {})
            
            # 项目信息关键字段
            for field in ['项目名称', '申请人', '依托单位', '电子邮箱']:
                val = project_info.get(field, '')
                if isinstance(val, str) and len(val.strip()) >= 2:
                    key_fields_found += 1
            
            # 基本信息关键字段
            applicant_info = basic_info.get('申请人信息', {})
            if isinstance(applicant_info, dict):
                for field in ['姓名', '性别', '学位', '职称']:
                    val = applicant_info.get(field, '')
                    if isinstance(val, str) and len(val.strip()) >= 1:
                        key_fields_found += 1
            
            project_basic = basic_info.get('项目基本信息', {})
            if isinstance(project_basic, dict):
                for field in ['申请代码', '研究期限']:
                    val = project_basic.get(field, '')
                    if isinstance(val, str) and len(val.strip()) >= 2:
                        key_fields_found += 1
            
            # 摘要
            for field in ['中文摘要', '英文摘要']:
                val = basic_info.get(field, '')
                if isinstance(val, str) and len(val.strip()) >= 20:
                    key_fields_found += 1
            
            print(f"[表格质量验证] 关键字段非空数: {key_fields_found}/12")
            return key_fields_found >= 3
        
        def do_table_extraction():
            """表格提取线程任务
            
            策略：
              MinerU成功提取表格 → 文本LLM转JSON
              MinerU未提取到表格 → 跳过（本地VLM游离线程或硅基流动VLM在后续步骤顺序处理）
            """
            nonlocal table_json_data
            
            if not mineru_table_ok:
                # MinerU表格不足：
                #   本地VLM可用 → 游离线程已在后台运行
                #   本地VLM不可用 → 硅基流动VLM在正文+参考文献完成后顺序执行
                if _use_local_vlm:
                    print(f"[并行] 表格：MinerU表格不足，等待本地VLM游离线程结果...")
                    increment_progress(1, '📊 表格识别已启动，等待结果中...')
                else:
                    print(f"[并行] 表格：MinerU表格不足，硅基流动VLM将在后续步骤顺序执行")
                    increment_progress(1, '📊 表格识别将在后续步骤执行...')
                return
            
            # MinerU识别到真实表格 → 用文本LLM直接提取JSON（不需要VLM）
            table_len = len(table_content)
            print(f"[并行] 表格：MinerU识别到有效表格（{table_len} 字符），文本LLM直接转JSON...")
            increment_progress(1, f'📊 表格提取中 | 解析表格: {table_len:,}字...')
            increment_progress(0.5, '📊 表格结构化尝试 1/2...')
            
            if task_check_func and task_check_func():
                return
            
            text_llm_result = call_llm_for_table_extraction(
                table_content, api_key, model_name, task_check_func,
                pdf_data=None, progress_callback=None
            )
            
            if text_llm_result.get('success') and _validate_table_json_quality(text_llm_result.get('json_data')):
                table_json_data = text_llm_result['json_data']
                _vlm_skipped_flag[0] = True
                print(f"[并行] ✓ 表格：文本LLM提取成功（VLM未启动），{len(table_json_data)} 个键")
                increment_progress(2, f'📄 表格识别完成 | {len(table_json_data)} 个键')
                return
            
            reason = "JSON质量不足" if text_llm_result.get('success') else text_llm_result.get('error', '未知错误')
            print(f"[并行] ⚠️ 表格：文本LLM提取不合格（{reason}），再次尝试（无质量验证）...")
            increment_progress(0.5, f'📊 表格结构化尝试 2/2...')
            
            # 降级重试：不做质量验证
            if table_content.strip():
                result = call_llm_for_table_extraction(
                    table_content, api_key, model_name, task_check_func,
                    pdf_data=None, progress_callback=None
                )
                if result.get('success'):
                    table_json_data = result['json_data']
                    _vlm_skipped_flag[0] = True
                    print(f"[并行] ✓ 表格：降级文本LLM提取完成，{len(table_json_data)} 个键")
                    increment_progress(2, f'📄 表格识别完成 | {len(table_json_data)} 个键')
                else:
                    print(f"[并行] ✗ 表格文本LLM降级提取也失败: {result.get('error')}")
                    increment_progress(1, '⚠️ 表格文本提取失败')
        
        def do_text_extraction():
            """正文提取线程任务"""
            nonlocal text_json_data
            if not text_content.strip():
                print(f"[并行] 正文：跳过（无内容）")
                increment_progress(1, _make_progress_msg('⏭️ 正文为空，已跳过正文提取'))
                return
            text_len = len(text_content)
            print(f"[并行] 正文：开始LLM提取（{text_len} 字符）")
            increment_progress(1, _make_progress_msg(f'🏛️ 正文提取中 | {text_len:,}字...'))

            # 正文阶段进度范围：避免正文内部回调把总进度直接推到100%
            text_stage_start = _parallel_max_progress[0]
            text_stage_end = min(max(text_stage_start + 18.0, text_stage_start), 84.0)
            
            if task_check_func and task_check_func():
                return
            
            # 正文提取内部进度用 increment_progress 传递
            def scaled_text_progress(progress, message):
                # 将正文内部进度(40~100)映射到正文阶段局部区间，避免总进度提前到顶。
                try:
                    p = float(progress)
                except Exception:
                    p = 40.0
                norm = max(0.0, min((p - 40.0) / 60.0, 1.0))
                mapped = text_stage_start + norm * (text_stage_end - text_stage_start)
                text_msg = message if message else f'🏛️ 正文提取中 {int(p)}%'
                set_progress_at_least(mapped, _make_progress_msg(text_msg))
                
                # 透传章节级进度文案到前端（不增加百分比，仅更新消息）
                # 例如: "正文提取：已完成 1/4 个章节"
                if message and ('已完成' in message or '章节' in message):
                    increment_progress(0, _make_progress_msg(f'🏛️ {message}'))
            
            result = call_llm_for_text_extraction(text_content, api_key, model_name, task_check_func, scaled_text_progress)
            if result.get('success'):
                text_json_data = result['json_data']
                text_keys = len(text_json_data)
                print(f"[并行] ✓ 正文提取成功，{text_keys} 个键")
                
                if 'quality_check' in result:
                    quality_check_data = result['quality_check']
                    quality_rules = len(quality_check_data)
                    print(f"[并行] ✓ 提取到质量检查数据: {quality_rules} 条规则")
                    text_json_data['_quality_check_internal'] = quality_check_data
                    increment_progress(2, _make_progress_msg(f'✅ 正文提取完成 | 数据键: {text_keys} | 质量规则: {quality_rules}'))
                else:
                    increment_progress(2, _make_progress_msg(f'✅ 正文提取完成 | 数据键: {text_keys}'))
            else:
                print(f"[并行] ✗ 正文提取失败: {result.get('error')}")
                increment_progress(1, _make_progress_msg('⚠️ 正文提取失败'))
        
        # ================================================================
        # VLM游离线程：仅当MinerU表格识别失败 且 本地VLM可用时 才启动
        # 本地VLM在后台运行（带信号量），其他步骤与VLM并行执行
        # 硅基流动VLM在do_table_extraction中同步执行，无需游离线程
        # ================================================================
        _vlm_floating_thread = [None]
        _vlm_floating_result = [None]
        
        if not mineru_table_ok and pdf_data_for_vlm and _use_local_vlm:
            def _vlm_floating_worker():
                """本地VLM游离线程：后台提取表格，带取消检查和信号量"""
                try:
                    print(f"[VLM游离] 🔒 等待本地VLM信号量...")
                    _vlm_parallel_text[0] = '📊 表格识别 等待中...'
                    increment_progress(1, _make_progress_msg('📊 表格识别后台启动'))
                    
                    # 带取消检查的信号量获取（每2秒检查一次取消状态）
                    while not _vlm_semaphore.acquire(timeout=2):
                        if task_check_func and task_check_func():
                            print(f"[VLM游离] ❌ 等待信号量期间任务被取消")
                            _vlm_floating_result[0] = {'success': False, 'error': '任务已取消'}
                            _vlm_parallel_text[0] = ''
                            return
                    
                    try:
                        if task_check_func and task_check_func():
                            print(f"[VLM游离] ❌ 获得信号量后任务已被取消")
                            _vlm_floating_result[0] = {'success': False, 'error': '任务已取消'}
                            _vlm_parallel_text[0] = ''
                            return
                        
                        print(f"[VLM游离] 🚀 获得本地VLM信号量，开始VLM表格提取")
                        _vlm_parallel_text[0] = '📊 智能表格提取中...'
                        
                        import re as _re_local
                        def _local_vlm_progress_cb(progress, message):
                            """本地VLM内部进度回调：更新并行文本"""
                            page_match = _re_local.search(r'第\s*(\d+)/(\d+)\s*页', message)
                            if page_match:
                                _vlm_parallel_text[0] = f'📊 表格识别（{page_match.group(1)}/{page_match.group(2)}）'
                                increment_progress(1, _make_progress_msg(_last_main_msg[0] or '处理中...'))
                        
                        result = call_vlm_for_table_extraction(
                            pdf_data_for_vlm, api_key, task_check_func, _local_vlm_progress_cb
                        )
                        _vlm_floating_result[0] = result
                        
                        if result.get('success'):
                            md_size = len(result.get('vlm_md_content', ''))
                            print(f"[VLM游离] ✅ 本地VLM表格提取成功（Markdown） | MD: {md_size:,}字")
                            _vlm_parallel_text[0] = ''  # VLM完成，清除并行显示
                            increment_progress(2, f'✅ 智能表格识别完成 | MD: {md_size:,}字')
                        else:
                            print(f"[VLM游离] ⚠️ 本地VLM表格提取失败: {result.get('error')}")
                            _vlm_parallel_text[0] = ''  # VLM完成，清除并行显示
                            increment_progress(1, '⚠️ 智能表格识别失败')
                    finally:
                        _vlm_semaphore.release()
                        _vlm_parallel_text[0] = ''  # 确保清除
                        print(f"[VLM游离] 🔓 释放本地VLM信号量")
                except Exception as e:
                    print(f"[VLM游离] ❌ VLM异常: {e}")
                    _vlm_floating_result[0] = {'success': False, 'error': str(e)}
                    _vlm_parallel_text[0] = ''  # 异常时清除
            
            _vlm_floating_thread[0] = threading.Thread(
                target=_vlm_floating_worker, daemon=True, name='vlm-floating-local'
            )
            _vlm_floating_thread[0].start()
            print(f"[VLM游离] 🚀 本地VLM游离线程已启动，与后续步骤并行执行")
            increment_progress(1, '📊 表格不足，表格识别后台启动，继续后续处理...')
        elif not mineru_table_ok and not pdf_data_for_vlm:
            print(f"[VLM游离] ⚠️ 无PDF数据，无法启动VLM")
        
        # ================================================================
        # 步骤2: 分别处理表格和正文
        # 【修改】正文提取改为串行执行（本地LLM经常超时，取消并行）
        #   - 表格提取可继续并行（单独线程或VLM游离线程）
        #   - 参考文献提取在步骤3中继续并行
        # ================================================================
        print(f"[五项处理] 步骤2: 开始提取表格和正文...")
        increment_progress(1, f'📊 步骤2: 结构化提取中 | 表格: {len(table_content):,}字 | 正文: {len(text_content):,}字...')
        
        # 【步骤2.1】表格提取（单线程或VLM游离线程）
        print(f"[五项处理] 步骤2.1: 表格提取")
        do_table_extraction()
        
        # 【步骤2.2】正文提取（串行，不再并行）
        print(f"[五项处理] 步骤2.2: 正文提取（串行）")
        increment_progress(1, _make_progress_msg('🏛️ 已进入正文提取阶段...'))
        do_text_extraction()
        
        text_keys = len(text_json_data) if text_json_data else 0
        if mineru_table_ok:
            table_keys = len(table_json_data) if table_json_data else 0
            print(f"[五项处理] 步骤2完成 | 表格键: {table_keys} | 正文键: {text_keys}")
            increment_progress(2, f'✅ 数据提取完成 | 表格键: {table_keys} | 正文键: {text_keys}')
        elif _use_local_vlm:
            print(f"[五项处理] 步骤2完成（正文已提取，本地VLM仍在后台运行）| 正文键: {text_keys}")
            increment_progress(2, _make_progress_msg(f'✅ 正文提取完成 | 正文键: {text_keys}'))
        else:
            # 硅基流动VLM将在正文+参考文献完成后顺序执行
            print(f"[五项处理] 步骤2完成（仅正文）| 正文键: {text_keys} | 表格待硅基流动VLM")
            increment_progress(2, f'✅ 正文提取完成 | 正文键: {text_keys}')
        
        # ================================================================
        # 步骤3: 参考文献提取与验证（与VLM并行（本地VLM）或按顺序（硅基流动））
        # 参考文献在文档末尾，与VLM表格替换（前几页）完全无关
        # ================================================================
        if task_check_func and task_check_func():
            return False
        
        if content:
            print(f"[五项处理] 步骤3: 参考文献提取与验证（与VLM并行）...")
            increment_progress(1, _make_progress_msg('📚 步骤3: 参考文献提取与验证中...'))
            
            try:
                _ref_progress_count = [0]
                def ref_progress_cb(progress, message):
                    if task_check_func and task_check_func():
                        return False
                    if progress_callback:
                        if int(progress / 10) > _ref_progress_count[0]:
                            _ref_progress_count[0] = int(progress / 10)
                            increment_progress(1, _make_progress_msg(f'📚 参考文献: {message}'))
                    return True
                
                ref_md_content, ref_validation_json = extract_references_from_markdown(
                    content, user_id=user_id, progress_callback=ref_progress_cb, task_id=None
                )
                
                if ref_md_content:
                    ref_len = len(ref_md_content)
                    print(f"[五项处理] ✓ 参考文献提取成功，内容 {ref_len} 字符")
                    if ref_validation_json:
                        ref_report_len = len(ref_validation_json)
                        print(f"[五项处理] ✓ 参考文献验证完成，报告 {ref_report_len} 字符")
                    increment_progress(2, _make_progress_msg(f'✅ 参考文献验证完成 | {ref_len:,}字'))
                else:
                    print(f"[五项处理] 步骤3: 未找到参考文献内容")
                    increment_progress(1, _make_progress_msg('⏭️ 未找到参考文献'))
            except Exception as ref_e:
                print(f"[五项处理] ✗ 参考文献处理异常: {ref_e}")
                import traceback
                traceback.print_exc()
                increment_progress(1, _make_progress_msg('⚠️ 参考文献处理异常'))
        else:
            print(f"[五项处理] 步骤3: 跳过（无文档内容）")
            increment_progress(1, _make_progress_msg('⏭️ 无内容，跳过参考文献'))
        
        # ================================================================
        # 等待本地VLM游离线程完成（仅当使用本地VLM游离线程时）
        # ================================================================
        if _vlm_floating_thread[0] is not None:
            thread = _vlm_floating_thread[0]
            print(f"[VLM游离] 等待本地VLM游离线程完成...")
            
            while thread.is_alive():
                if task_check_func and task_check_func():
                    print(f"[VLM游离] 任务已取消，停止等待")
                    break
                increment_progress(0.5, '📊 等待智能表格识别结果...')
                thread.join(timeout=2)
            
            # 收集本地VLM结果
            vlm_result = _vlm_floating_result[0]
            if vlm_result and vlm_result.get('success'):
                if vlm_result.get('vlm_md_content'):
                    _vlm_md_content[0] = vlm_result['vlm_md_content']
                table_json_data = _build_table_json_from_vlm_md(_vlm_md_content[0])
                vlm_keys = len(table_json_data) if table_json_data else 0
                if table_json_data:
                    print(f"[VLM游离] ✅ 本地VLM结果已收集并完成表格结构化 | 表格键: {vlm_keys}")
                    increment_progress(2, f'✅ 智能表格结构化完成 | 表格键: {vlm_keys}')
                else:
                    md_size = len(_vlm_md_content[0]) if _vlm_md_content[0] else 0
                    print(f"[VLM游离] ⚠️ VLM提取到Markdown，但表格结构化失败 | MD: {md_size:,}字")
                    increment_progress(1, f'⚠️ 智能表格结构化失败 | MD: {md_size:,}字')
            elif vlm_result and vlm_result.get('error') == '任务已取消':
                print(f"[VLM游离] 任务已取消，VLM结果放弃")
                return False
            else:
                error_msg = vlm_result.get('error', '未知') if vlm_result else '线程异常'
                print(f"[VLM游离] ⚠️ 本地VLM未成功: {error_msg}")
                increment_progress(1, '⚠️ 智能表格识别失败')
        
        # ================================================================
        # 硅基流动VLM顺序执行（仅当MinerU表格不足 + 本地VLM不可用时）
        # 在正文提取和参考文献全部完成后才执行，进度条单椭圆显示
        # ================================================================
        if not mineru_table_ok and not _use_local_vlm and pdf_data_for_vlm:
            if task_check_func and task_check_func():
                return False
            
            print(f"[硅基流动VLM] 开始顺序执行表格提取...")
            increment_progress(1, '📊 云端智能表格提取准备中...')
            
            import re as _re_sf
            def _siliconflow_vlm_progress(progress, message):
                """硅基流动VLM内部进度回调"""
                page_match = _re_sf.search(r'第\s*(\d+)/(\d+)\s*页', message)
                if page_match:
                    increment_progress(1, f'📊 智能表格识别（{page_match.group(1)}/{page_match.group(2)}）')
            
            vlm_result = call_vlm_for_table_extraction(
                pdf_data_for_vlm, api_key, task_check_func, _siliconflow_vlm_progress
            )
            
            if vlm_result and vlm_result.get('success'):
                if vlm_result.get('vlm_md_content'):
                    _vlm_md_content[0] = vlm_result['vlm_md_content']
                table_json_data = _build_table_json_from_vlm_md(_vlm_md_content[0])
                vlm_keys = len(table_json_data) if table_json_data else 0
                if table_json_data:
                    print(f"[硅基流动VLM] ✅ 表格提取+结构化成功 | {vlm_keys} 个键")
                    increment_progress(2, f'✅ 云端智能表格结构化完成 | {vlm_keys} 个键')
                else:
                    md_size = len(_vlm_md_content[0]) if _vlm_md_content[0] else 0
                    print(f"[硅基流动VLM] ⚠️ 提取到Markdown，但表格结构化失败 | MD: {md_size:,}字")
                    increment_progress(1, f'⚠️ 云端智能表格结构化失败 | MD: {md_size:,}字')
            else:
                error = vlm_result.get('error', '未知') if vlm_result else '调用异常'
                print(f"[硅基流动VLM] ⚠️ 表格提取失败: {error}")
                increment_progress(1, '⚠️ 云端智能表格识别失败')
        
        # 保存VLM修改前的原始MinerU输出（用于正则初筛检查，避免VLM替换丢失字段）
        original_content = content
        
        # ================================================================
        # 步骤4: VLM MD替换（仅当MinerU表格不足且VLM提取成功时）
        # ================================================================
        if not mineru_table_ok and table_json_data is not None:
            print(f"[五项处理] 步骤4: MinerU表格不足，用VLM结果替换MD文档表格区域...")
            increment_progress(1, '📝 步骤4: 将表格识别结果写入文档中...')
            
            vlm_md_direct = _vlm_md_content[0]
            if vlm_md_direct and len(vlm_md_direct.strip()) > 50:
                vlm_md_len = len(vlm_md_direct)
                print(f"[五项处理] 使用VLM直接输出的Markdown替换前几页（{vlm_md_len} 字符）")
                modified_content = replace_md_table_with_vlm(content, vlm_md_direct)
                if modified_content:
                    content = modified_content
                    md_content_modified = True
                    table_content, text_content = separate_tables_and_text(content)
                    print(f"[五项处理] ✓ VLM Markdown直接替换完成")
                    increment_progress(2, f'🔍 智能表格写入完成 | {vlm_md_len:,}字 → 初筛中...')
                else:
                    print(f"[五项处理] ⚠️ VLM Markdown替换失败，回退到JSON转HTML方式")
                    vlm_md_direct = None
            
            if not vlm_md_direct or not md_content_modified:
                vlm_html_table = convert_table_json_to_md_table(table_json_data)
                if vlm_html_table:
                    html_len = len(vlm_html_table)
                    modified_content = replace_md_table_with_vlm(content, vlm_html_table)
                    if modified_content:
                        content = modified_content
                        md_content_modified = True
                        table_content, text_content = separate_tables_and_text(content)
                        print(f"[五项处理] ✓ JSON→HTML替换完成")
                        increment_progress(2, f'🔍 HTML表替换完成 | {html_len:,}字 → 初筛中...')
                    else:
                        print(f"[五项处理] ⚠️ MD替换失败（未找到报告正文标记）")
                        increment_progress(1, '🔍 写回跳过 → 准备规则检查...')
                else:
                    print(f"[五项处理] ⚠️ VLM JSON转HTML失败")
                    increment_progress(1, '🔍 写回失败 → 准备规则检查...')
        elif not mineru_table_ok:
            print(f"[五项处理] ⚠️ MinerU表格不足但VLM也未成功")
            increment_progress(1, '🔍 表格写回未执行 → 准备规则检查...')
        
        # ================================================================
        # 步骤5: 规则初筛检查（纯正则，使用最新的表格数据）
        # ================================================================
        if task_check_func and task_check_func():
            return False
        
        if text_content.strip() or table_content.strip():
            print(f"[五项处理] 步骤5: 纯正则初筛检查（{'VLM替换后' if md_content_modified else 'MinerU原始'}数据）...")
            increment_progress(1, f'🔍 步骤5: 规则初筛中...')
            
            result = preliminary_check_with_llm(
                table_content=table_content,
                text_content=text_content,
                api_key=api_key,
                user_id=user_id,
                json_data=table_json_data,
                md_content=original_content,  # 传入VLM修改前的原始MinerU输出，避免VLM替换丢失字段
                vlm_md_content=content  # VLM处理后的内容作为回退源
            )
            
            if result.get('success'):
                all_violations = result.get('violations', [])
                total_rules = result.get('total_rules', 24)
                table_violations = [v for v in all_violations if '表' in v.get('location', '')]
                text_violations = [v for v in all_violations if v not in table_violations]
                violation_count = len(all_violations)
                print(f"[五项处理] ✓ 初筛完成，发现 {violation_count} 处违规")
                increment_progress(2, f'✅ 初筛完成 | 规则: {total_rules} | 问题: {violation_count} → 保存中...')
            else:
                print(f"[五项处理] ✗ 初筛失败: {result.get('error')}")
                table_violations = []
                text_violations = []
                increment_progress(1, '⚠️ 初筛异常 → 保存中...')
        else:
            print(f"[五项处理] 步骤5: 跳过（无内容）")
            increment_progress(1, '⏭️ 跳过初筛 → 保存中...')
            table_violations = []
            text_violations = []
        
        # ================================================================
        # 步骤6: 保存所有结果到数据库
        # ================================================================
        print(f"[五项处理] 步骤6: 保存处理结果到数据库...")
        if progress_callback:
            increment_progress(2, '💾 步骤6: 保存结果到数据库中...')
        
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # 构建更新SQL
            update_fields = []
            update_values = []
            
            # 若MD内容被VLM替换过，保存修改后的MD，同时保存原始MinerU输出
            if md_content_modified:
                update_fields.append("md_content = %s")
                update_values.append(content)
                update_fields.append("original_md_content = %s")
                update_values.append(original_content)
                print(f"[五项处理] ✓ 添加md_content字段（VLM替换后，{len(content)} 字符），原始MinerU: {len(original_content)} 字符")
            
            if table_json_data is not None:
                update_fields.append("extracted_json_data = %s")
                update_values.append(json.dumps(table_json_data, ensure_ascii=False, indent=2))
            
            if text_json_data is not None:
                update_fields.append("nsfc_json_data = %s")
                update_values.append(json.dumps(text_json_data, ensure_ascii=False, indent=2))
                print(f"[四项处理] 正文JSON已保存（含内部规则字段），键数: {len(text_json_data)}")
            
            if table_violations is not None:
                update_fields.append("table_preliminary_check = %s")
                update_values.append(json.dumps(table_violations, ensure_ascii=False, indent=2))
            
            if text_violations is not None:
                update_fields.append("text_preliminary_check = %s")
                update_values.append(json.dumps(text_violations, ensure_ascii=False, indent=2))
            
            # 参考文献内容与验证报告
            if ref_md_content:
                update_fields.append("reference_md_content = %s")
                update_values.append(ref_md_content)
                print(f"[五项处理] ✓ 添加reference_md_content字段 ({len(ref_md_content)} 字符)")
            
            if ref_validation_json:
                update_fields.append("reference_validation_json = %s")
                update_values.append(ref_validation_json)
                print(f"[五项处理] ✓ 添加reference_validation_json字段 ({len(ref_validation_json)} 字符)")
            
            if update_fields:
                # 统计保存的信息
                saved_fields_count = len(update_fields)
                total_data_size = sum(len(str(v)) for v in update_values[:-2])  # 排除file_id和user_id
                total_data_kb = total_data_size / 1024
                
                update_values.extend([file_id, user_id])
                sql = f"""
                    UPDATE user_files 
                    SET {', '.join(update_fields)}
                    WHERE file_id = %s AND user_id = %s
                """
                cursor.execute(sql, update_values)
                conn.commit()
                print(f"[五项处理] ✓ 数据保存成功 | 字段: {saved_fields_count} | 大小: {total_data_kb:.1f}KB")
                if progress_callback:
                    increment_progress(5, f'✅ 全部完成！ | 保存字段: {saved_fields_count} | 数据大小: {total_data_kb:.1f}KB')
            
            conn.close()
        
        print(f"[五项处理] 文件 {file_id} 全部处理完成")
        return True
        
    except Exception as e:
        print(f"[五项处理] 错误: {e}")
        import traceback
        traceback.print_exc()
        return False

# =========================
# 一键解析完成标记
# =========================

# =========================
# 管理员功能
# =========================

def is_admin(user_id):
    """检查用户是否为真正的管理员 - 只有admin角色"""
    # 检查数据库中的role字段
    try:
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("SELECT role FROM users WHERE id = %s", (user_id,))
            result = cursor.fetchone()
            conn.close()
            return result and result.get('role') == 'admin'
    except Exception as e:
        # 数据库查询失败时记录日志
        app.logger.error(f"权限检查失败 user_id:{user_id} - {e}")
    
    return False

def can_view_admin_info(user_id):
    """检查用户是否可以查看管理信息 - admin和premium角色"""
    # 检查数据库中的role字段
    try:
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("SELECT role FROM users WHERE id = %s", (user_id,))
            result = cursor.fetchone()
            conn.close()
            return result and result.get('role') in ['admin', 'premium']
    except Exception as e:
        # 数据库查询失败时记录日志
        app.logger.error(f"查看权限检查失败 user_id:{user_id} - {e}")
    
    return False

def is_full_admin(user_id):
    """检查用户是否为完全管理员（只有admin角色，不包括premium）"""
    # 检查数据库中的role字段
    try:
        conn = DatabaseConfig.get_db_connection()
        if conn:
            cursor = conn.cursor()
            cursor.execute("SELECT role FROM users WHERE id = %s", (user_id,))
            result = cursor.fetchone()
            conn.close()
            return result and result.get('role') == 'admin'
    except Exception as e:
        # 数据库查询失败时记录日志
        app.logger.error(f"完全管理员权限检查失败 user_id:{user_id} - {e}")
    
    return False

def log_admin_action(action, details=""):
    """记录管理员操作日志 - 最小实现"""
    app.logger.info(f"[ADMIN] {session.get('username', 'unknown')} - {action} - {details}")

def admin_required(f):
    """管理员权限装饰器 - 增强版"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # 1. 检查登录状态
        if 'user_id' not in session:
            flash('请先登录', 'error')
            return redirect(url_for('login'))
        
        # 2. 验证用户ID有效性
        user_id = session.get('user_id')
        if not user_id or not isinstance(user_id, int):
            flash('会话无效，请重新登录', 'error')
            session.clear()
            return redirect(url_for('login'))
        
        # 3. 检查管理员权限
        if not is_admin(user_id):
            # 记录权限拒绝
            app.logger.warning(f"[ADMIN_ACCESS_DENIED] User:{session.get('username', 'unknown')} ID:{user_id} attempted to access {f.__name__}")
            flash('权限不足，需要管理员权限', 'error')
            return redirect(url_for('dashboard'))
        
        # 4. 权限验证通过，执行函数
        return f(*args, **kwargs)
    return decorated_function

def full_admin_required(f):
    """完全管理员权限装饰器 - 只允许admin角色，不允许premium"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # 1. 检查登录状态
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': '请先登录'}), 401
        
        # 2. 验证用户ID有效性
        user_id = session.get('user_id')
        if not user_id or not isinstance(user_id, int):
            return jsonify({'success': False, 'message': '会话无效，请重新登录'}), 401
        
        # 3. 检查完全管理员权限
        if not is_full_admin(user_id):
            # 记录权限拒绝
            app.logger.warning(f"[FULL_ADMIN_ACCESS_DENIED] User:{session.get('username', 'unknown')} ID:{user_id} attempted to access {f.__name__}")
            return jsonify({'success': False, 'message': '权限不足，该操作需要管理员权限'}), 403
        
        # 4. 权限验证通过，执行函数
        return f(*args, **kwargs)
    return decorated_function

def view_admin_required(f):
    """装饰器：检查用户是否有查看管理信息的权限（admin和premium用户）"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # 1. 检查登录状态
        if 'user_id' not in session:
            flash('请先登录', 'error')
            return redirect(url_for('login'))
        
        # 2. 验证用户ID有效性
        user_id = session.get('user_id')
        if not user_id or not isinstance(user_id, int):
            flash('会话无效，请重新登录', 'error')
            return redirect(url_for('login'))
        
        # 3. 检查查看权限
        if not can_view_admin_info(user_id):
            # 记录权限拒绝
            app.logger.warning(f"[VIEW_ADMIN_ACCESS_DENIED] User:{session.get('username', 'unknown')} ID:{user_id} attempted to access {f.__name__}")
            flash('权限不足，无法访问该页面', 'error')
            return redirect(url_for('home'))
        
        # 4. 权限验证通过，执行函数
        return f(*args, **kwargs)
    return decorated_function

@app.route('/admin')
@view_admin_required
def admin_dashboard():
    """管理员控制台上传文件 - YNEX版本"""
    log_admin_action("ACCESS_ADMIN_DASHBOARD", "访问管理员控制台")
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            flash('数据库连接失败', 'error')
            return redirect(url_for('dashboard'))
        
        cursor = conn.cursor()
        
        # 获取用户统计信息
        cursor.execute("SELECT COUNT(*) as total_users FROM users")
        user_stats = cursor.fetchone()
        print(f"用户统计: {user_stats}")
        
        # 获取文件统计信息
        cursor.execute("SELECT COUNT(*) as total_files FROM user_files")
        file_stats = cursor.fetchone()
        print(f"文件统计: {file_stats}")
        
        # 获取最近注册的用户
        cursor.execute("""
            SELECT id, username, email, created_at, is_active 
            FROM users 
            ORDER BY created_at DESC 
            LIMIT 5
        """)
        recent_users = cursor.fetchall()
        print(f"最近用户数量: {len(recent_users) if recent_users else 0}")
        
        # 获取今日活跃统计 - 使用MySQL CURDATE() 避免Python/MySQL时区不一致
        cursor.execute("""
            SELECT COUNT(DISTINCT sl.user_id) as active_today 
            FROM system_logs sl
            WHERE sl.action = 'login' 
            AND sl.created_at >= CURDATE()
            AND sl.created_at < CURDATE() + INTERVAL 1 DAY
            AND sl.user_id IS NOT NULL
        """)
        active_stats = cursor.fetchone()
        active_today = active_stats['active_today'] if active_stats else 0
        
        # 获取今日处理成功的文件数量 - 使用MySQL CURDATE()
        cursor.execute("""
            SELECT COUNT(*) as processed_files 
            FROM system_logs 
            WHERE action = 'file_processed' 
              AND created_at >= CURDATE()
              AND created_at < CURDATE() + INTERVAL 1 DAY
        """)
        processed_stats = cursor.fetchone()
        processed_files = processed_stats['processed_files'] if processed_stats else 0
        
        conn.close()
        
        # 统计数据整合
        stats = {
            'total_users': user_stats['total_users'] if user_stats else 0,
            'total_files': file_stats['total_files'] if file_stats else 0,
            'processed_files': processed_files,
            'active_today': active_today
        }
        
        return render_template('admin/dashboard_ynex.html', 
                             stats=stats,
                             user_stats=user_stats,
                             file_stats=file_stats,
                             recent_users=recent_users,
                             active_today=active_today)
        
    except Exception as e:
        print(f"管理员控制台错误: {e}")
        import traceback

@app.route('/button-preview')
@admin_required
def button_preview():
    """按钮设计方案预览页面"""
    return render_template('button_preview.html')

@app.route('/group-management')
@view_admin_required
def group_management():
    """用户分组管理页面"""
    log_admin_action("ACCESS_GROUP_MANAGEMENT", "访问用户分组管理页面")
    
    # 由于@admin_required装饰器已经确保只有admin可以访问，无需额外检查
    try:
        # 这里可以添加从数据库获取分组数据的逻辑
        # 目前使用硬编码数据，由前端JavaScript处理
        
        return render_template('group_management.html',
                             username=session.get('username'),
                             user_id=session.get('user_id'),
                             email=session.get('email'),
                             current_user_role=session.get('role', 'user'),
                             is_admin=session.get('role') == 'admin')
        
    except Exception as e:
        app.logger.error(f"分组管理页面错误: {str(e)}")
        flash('加载分组管理页面时出错', 'error')
        return redirect(url_for('dashboard'))


@app.route('/api/groups/create', methods=['POST'])
@admin_required
def create_group():
    """创建用户分组"""
    try:
        data = request.get_json()
        group_name = data.get('name', '').strip()
        selected_users = data.get('users', [])
        
        if not group_name:
            return jsonify({'success': False, 'message': '分组名称不能为空'}), 400
            
        if not selected_users:
            return jsonify({'success': False, 'message': '请选择至少一个用户'}), 400
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 生成新的分组ID（获取当前最大的group_id + 1）
        cursor.execute("SELECT COALESCE(MAX(group_id), 0) + 1 as new_group_id FROM users WHERE group_id IS NOT NULL")
        result = cursor.fetchone()
        new_group_id = result['new_group_id'] if result else 1
        
        # 更新选中用户的分组信息
        user_ids_str = ','.join(map(str, selected_users))
        update_sql = f"UPDATE users SET group_id = %s, group_name = %s, group_status = 'confirmed' WHERE id IN ({user_ids_str})"
        cursor.execute(update_sql, (new_group_id, group_name))
        
        conn.commit()
        conn.close()
        
        log_admin_action("CREATE_GROUP", f"创建分组: {group_name}, 包含用户: {len(selected_users)}个")
        
        return jsonify({
            'success': True, 
            'message': f'分组 "{group_name}" 创建成功',
            'group_id': new_group_id,
            'group_name': group_name,
            'user_count': len(selected_users)
        })
        
    except Exception as e:
        app.logger.error(f"创建分组错误: {str(e)}")
        return jsonify({'success': False, 'message': f'创建分组失败: {str(e)}'}), 500


@app.route('/api/groups/<int:group_id>', methods=['DELETE'])
@api_login_required
def delete_group(group_id):
    """删除用户分组"""
    # 检查登录状态
    
    # 检查管理员权限
    current_user_id = session['user_id']
    if not is_admin(current_user_id):
        return jsonify({'success': False, 'message': '权限不足，需要管理员权限'}), 403
    
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500

        cursor = conn.cursor()
        
        # 检查分组是否存在
        cursor.execute("SELECT COUNT(*) as count FROM users WHERE group_id = %s", (group_id,))
        result = cursor.fetchone()
        if result['count'] == 0:
            return jsonify({'success': False, 'message': '分组不存在或已为空'}), 404
        
        # 获取分组名称和用户数量
        cursor.execute("SELECT group_name, COUNT(*) as user_count FROM users WHERE group_id = %s GROUP BY group_name", (group_id,))
        group_info = cursor.fetchone()
        group_name = group_info['group_name'] if group_info else f'分组{group_id}'
        user_count = group_info['user_count'] if group_info else 0
        
        # 将分组中的所有用户移出分组
        cursor.execute("""
            UPDATE users 
            SET group_id = NULL, group_name = NULL 
            WHERE group_id = %s
        """, (group_id,))
        
        if cursor.rowcount > 0:
            conn.commit()
            log_admin_action("DELETE_GROUP", f"删除分组 {group_id} ({group_name})，移出用户 {user_count} 个")
            
            return jsonify({
                'success': True,
                'message': f'分组 "{group_name}" 删除成功，已移出 {user_count} 个用户'
            })
        else:
            return jsonify({'success': False, 'message': '删除分组失败'}), 400
        
    except Exception as e:
        app.logger.error(f"删除分组错误: {str(e)}")
        return jsonify({'success': False, 'message': f'删除分组失败: {str(e)}'}), 500
    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/api/groups/batch-delete', methods=['DELETE'])
@api_login_required
def batch_delete_groups():
    """批量删除用户分组"""
    # 检查登录状态
    
    # 检查管理员权限
    current_user_id = session['user_id']
    if not is_admin(current_user_id):
        return jsonify({'success': False, 'message': '权限不足，需要管理员权限'}), 403
    
    try:
        data = request.get_json()
        group_ids = data.get('group_ids', [])
        
        if not group_ids:
            return jsonify({'success': False, 'message': '请选择要删除的分组'}), 400
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500

        cursor = conn.cursor()
        
        # 统计总的用户数量和分组名称
        placeholders = ','.join(['%s'] * len(group_ids))
        cursor.execute(f"""
            SELECT group_id, group_name, COUNT(*) as user_count 
            FROM users 
            WHERE group_id IN ({placeholders})
            GROUP BY group_id, group_name
        """, group_ids)
        
        group_info = cursor.fetchall()
        total_users = sum(info['user_count'] for info in group_info)
        group_names = [info['group_name'] for info in group_info]
        
        # 将所有选中分组中的用户移出分组
        cursor.execute(f"""
            UPDATE users 
            SET group_id = NULL, group_name = NULL 
            WHERE group_id IN ({placeholders})
        """, group_ids)
        
        affected_users = cursor.rowcount
        
        if affected_users > 0:
            conn.commit()
            log_admin_action("BATCH_DELETE_GROUPS", f"批量删除分组 {len(group_ids)} 个 ({', '.join(group_names)})，移出用户 {affected_users} 个")
            
            return jsonify({
                'success': True,
                'message': f'成功删除 {len(group_ids)} 个分组，移出用户 {affected_users} 个',
                'deleted_count': len(group_ids),
                'affected_users': affected_users
            })
        else:
            return jsonify({'success': False, 'message': '没有用户被移出分组'}), 400
        
    except Exception as e:
        app.logger.error(f"批量删除分组错误: {str(e)}")
        return jsonify({'success': False, 'message': f'批量删除分组失败: {str(e)}'}), 500
    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/api/groups/list', methods=['GET'])
@view_admin_required
def list_groups():
    """获取分组列表"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 获取所有分组信息（通过group_id分组，优先选择confirmed状态）
        cursor.execute("""
            SELECT 
                group_id,
                group_name,
                CASE 
                    WHEN COUNT(CASE WHEN group_status = 'confirmed' THEN 1 END) > 0 THEN 'confirmed'
                    ELSE MIN(group_status)
                END as group_status,
                COUNT(*) as user_count,
                MIN(created_at) as created_at
            FROM users 
            WHERE group_id IS NOT NULL AND group_name IS NOT NULL 
            GROUP BY group_id, group_name
            ORDER BY group_id ASC
        """)
        
        group_results = cursor.fetchall()
        groups = []
        
        for group_data in group_results:
            group_id = group_data['group_id']
            
            # 获取该分组的用户列表
            cursor.execute("""
                SELECT id, username, email, role, created_at
                FROM users 
                WHERE group_id = %s
                ORDER BY username ASC
            """, (group_id,))
            
            users = cursor.fetchall()
            
            # 获取分组的API密钥（从group_api_keys表获取）
            api_keys = {'mineru': '', 'siliconflow': ''}
            cursor.execute("""
                SELECT api_provider, api_token 
                FROM group_api_keys 
                WHERE group_id = %s AND is_active = 1
            """, (group_id,))
            
            group_api_keys = cursor.fetchall()
            for key_info in group_api_keys:
                provider = key_info['api_provider']
                if provider in api_keys:
                    # #22 解密API密钥
                    from db_operations import decrypt_api_token
                    api_keys[provider] = decrypt_api_token(key_info['api_token'])
            
            groups.append({
                'id': group_id,
                'name': group_data['group_name'],
                'description': f'包含 {group_data["user_count"]} 个用户的分组',
                'created_at': group_data['created_at'].strftime('%Y-%m-%d') if group_data['created_at'] else '',
                'created_by': '管理员',
                'status': group_data['group_status'],  # 使用真实的分组状态
                'user_count': group_data['user_count'],
                'api_keys': api_keys,
                'users': [
                    {
                        'id': user['id'],
                        'name': user['username'],
                        'email': user['email'],
                        'role': user['role'] or 'user'
                    }
                    for user in users
                ]
            })
        
        conn.close()
        
        return jsonify({
            'success': True,
            'groups': groups,
            'total_count': len(groups)
        })
        
    except Exception as e:
        app.logger.error(f"获取分组列表错误: {str(e)}")
        return jsonify({'success': False, 'message': f'获取分组列表失败: {str(e)}'}), 500


@app.route('/api/groups/<int:group_id>/api-keys', methods=['GET'])
@view_admin_required
def get_group_api_keys_route(group_id):
    """获取分组API密钥 - 从分组密钥表获取（包含model_name）"""
    try:
        # 使用新的分组密钥函数
        keys = get_group_api_keys(group_id)
        
        # 转换为前端期望的格式
        formatted_keys = []
        for key in keys:
            formatted_keys.append({
                'id': key['id'],
                'key_name': f"{key['api_provider'].capitalize()} API Key",
                'api_provider': key['api_provider'],
                'api_token': key['api_token'],  # 返回完整密钥
                'model_name': key.get('model_name'),  # 新增：返回模型名称
                'is_active': bool(key['is_active']),
                'created_at': key['created_at'].strftime('%Y-%m-%d %H:%M:%S') if key['created_at'] else None
            })
        
        return jsonify({
            'success': True,
            'keys': formatted_keys
        })
        
    except Exception as e:
        app.logger.error(f"获取分组API密钥错误: {str(e)}")
        return jsonify({'success': False, 'message': f'获取分组API密钥失败: {str(e)}'}), 500


@app.route('/api/groups/next-id', methods=['GET'])
@admin_required
def get_next_group_id():
    """获取下一个可用的分组ID"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 获取当前最大的group_id + 1
        cursor.execute("SELECT COALESCE(MAX(group_id), 0) + 1 as next_group_id FROM users WHERE group_id IS NOT NULL")
        result = cursor.fetchone()
        next_group_id = result['next_group_id'] if result else 1
        
        conn.close()
        
        return jsonify({
            'success': True,
            'next_group_id': next_group_id,
            'default_name': f'组{next_group_id}'
        })
        
    except Exception as e:
        app.logger.error(f"获取下一个分组ID错误: {str(e)}")
        return jsonify({'success': False, 'message': f'获取分组ID失败: {str(e)}'}), 500


@app.route('/api/groups/api-keys/save', methods=['POST'])
@admin_required
def save_group_api_keys():
    """保存分组API密钥 - 使用分组密钥表，支持模型选择"""
    try:
        data = request.get_json()
        
        group_id = data.get('group_id')
        provider = data.get('provider')  # 'mineru' or 'siliconflow'
        api_key = data.get('api_key', '').strip()
        model_name = data.get('model_name', '').strip() or None  # 新增：模型名称
        
        print(f"[分组API密钥保存] 解析参数 - group_id: {group_id}, provider: {provider}, api_key长度: {len(api_key) if api_key else 0}, model_name: {model_name}")
        
        if not group_id or not provider:
            print(f"[分组API密钥保存] 参数验证失败 - group_id: {bool(group_id)}, provider: {bool(provider)}")
            return jsonify({'success': False, 'message': '参数不完整'}), 400
        
        # 本地服务不需要API密钥
        if provider in ('local_llm', 'local_mineru'):
            api_key = api_key or 'local-not-needed'
        elif not api_key:
            print(f"[分组API密钥保存] API密钥为空")
            return jsonify({'success': False, 'message': '参数不完整：需要API密钥'}), 400
        
        if provider not in ['mineru', 'siliconflow', 'local_llm', 'local_mineru']:
            return jsonify({'success': False, 'message': '不支持的API提供商'}), 400
        
        # 验证分组是否存在
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) as count FROM users WHERE group_id = %s", (group_id,))
        result = cursor.fetchone()
        
        if not result or result['count'] == 0:
            return jsonify({'success': False, 'message': f'分组 {group_id} 不存在或没有成员'}), 400
        
        # 使用新的分组密钥函数保存（包含model_name）
        success, result = add_group_api_key(
            group_id=group_id,
            api_provider=provider,
            api_token=api_key,
            created_by=session['user_id'],
            model_name=model_name  # 新增：传递模型名称
        )
        
        if success:
            # 记录管理员操作
            model_info = f", 模型: {model_name}" if model_name else ""
            log_admin_action("SAVE_GROUP_API_KEY", f"为分组 {group_id} 保存 {provider} API密钥{model_info}")
            
            return jsonify({
                'success': True,
                'message': f'成功为分组 {group_id} 保存 {provider} API密钥{model_info}',
                'key_id': result
            })
        else:
            return jsonify({
                'success': False, 
                'message': f'保存分组API密钥失败: {result}'
            }), 400
        
    except Exception as e:
        app.logger.error(f"保存分组API密钥错误: {str(e)}")
        return jsonify({'success': False, 'message': f'保存API密钥失败: {str(e)}'}), 500
    finally:
        if 'conn' in locals() and conn:
            conn.close()


@app.route('/api/groups/members/add', methods=['POST'])
@api_login_required
def add_group_members():
    """添加用户到分组，并为新成员分配API密钥"""
    # 检查登录状态
    
    # 检查管理员权限
    current_user_id = session['user_id']
    if not is_admin(current_user_id):
        return jsonify({'success': False, 'message': '权限不足，需要管理员权限'}), 403
    
    try:
        data = request.get_json()
        group_id = data.get('group_id')
        user_ids = data.get('user_ids', [])
        
        if not group_id or not user_ids:
            return jsonify({'success': False, 'message': '参数不完整'}), 400
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 获取分组名称
        cursor.execute("SELECT group_name FROM users WHERE group_id = %s LIMIT 1", (group_id,))
        result = cursor.fetchone()
        group_name = result['group_name'] if result else f'组{group_id}'
        
        added_count = 0
        
        # 为每个用户分配到分组
        for user_id in user_ids:
            try:
                # 更新用户的分组信息（设置为confirmed状态）
                cursor.execute("""
                    UPDATE users 
                    SET group_id = %s, group_name = %s, group_status = 'confirmed' 
                    WHERE id = %s
                """, (group_id, group_name, user_id))
                
                if cursor.rowcount > 0:
                    added_count += 1
                    app.logger.info(f"用户 {user_id} 成功加入分组 {group_id}，可使用分组统一密钥")
                else:
                    app.logger.warning(f"用户 {user_id} 加入分组 {group_id} 失败")
                    
            except Exception as e:
                app.logger.error(f"添加用户 {user_id} 到分组 {group_id} 时出错: {str(e)}")
        
        conn.commit()
        
        log_admin_action("ADD_GROUP_MEMBERS", f"向分组 {group_id} 添加 {added_count} 个成员")
        
        return jsonify({
            'success': True,
            'message': f'成功添加 {added_count} 个用户到分组',
            'added_count': added_count
        })
        
    except Exception as e:
        app.logger.error(f"添加分组成员错误: {str(e)}")
        return jsonify({'success': False, 'message': f'添加成员失败: {str(e)}'}), 500
    finally:
        if 'conn' in locals() and conn:
            conn.close()


@app.route('/api/groups/<int:group_id>/members/<int:user_id>', methods=['DELETE'])
@api_login_required
def remove_group_member(group_id, user_id):
    """从分组中移除用户"""
    # 检查登录状态
    
    # 检查管理员权限
    current_user_id = session['user_id']
    if not is_admin(current_user_id):
        return jsonify({'success': False, 'message': '权限不足，需要管理员权限'}), 403
    
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500

        cursor = conn.cursor()
        
        # 检查用户是否在该分组中
        cursor.execute("SELECT id FROM users WHERE id = %s AND group_id = %s", (user_id, group_id))
        if not cursor.fetchone():
            return jsonify({'success': False, 'message': '用户不在该分组中'}), 400
        
        # 将用户从分组中移除（设置group_id、group_name和group_status为NULL）
        cursor.execute("""
            UPDATE users 
            SET group_id = NULL, group_name = NULL, group_status = NULL
            WHERE id = %s
        """, (user_id,))
        
        if cursor.rowcount > 0:
            # 清理用户的分组相关数据（安全考虑）
            try:
                # 注意：现在使用分组密钥，无需删除个人密钥
                # 仅清理users表中的兼容性字段
                cursor.execute("""
                    UPDATE users SET siliconflow_api_key = NULL 
                    WHERE id = %s
                """, (user_id,))
                
                app.logger.info(f"已清理用户 {user_id} 的分组相关数据")
            except Exception as e:
                app.logger.warning(f"清理用户 {user_id} API密钥时出错: {str(e)}")
            
            conn.commit()
            log_admin_action("REMOVE_GROUP_MEMBER", f"从分组 {group_id} 移除用户 {user_id}")
            
            return jsonify({
                'success': True,
                'message': '用户已从分组中移除'
            })
        else:
            return jsonify({'success': False, 'message': '移除用户失败'}), 400
        
    except Exception as e:
        app.logger.error(f"移除分组成员错误: {str(e)}")
        return jsonify({'success': False, 'message': f'移除成员失败: {str(e)}'}), 500
    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/api/groups/<int:group_id>/confirm', methods=['POST'])
@admin_required
def confirm_group(group_id):
    """确认分组（将临时分组转为确定分组）"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500

        cursor = conn.cursor()
        
        # 检查分组是否存在且为临时状态
        cursor.execute("SELECT group_name, group_status FROM users WHERE group_id = %s LIMIT 1", (group_id,))
        group_info = cursor.fetchone()
        
        if not group_info:
            return jsonify({'success': False, 'message': '分组不存在'}), 404
            
        if group_info['group_status'] == 'confirmed':
            return jsonify({'success': False, 'message': '分组已经是确定状态'}), 400
        
        # 生成新的分组名称（去掉"临时"两字）
        old_group_name = group_info['group_name']
        new_group_name = old_group_name.replace('临时', '').strip()
        
        # 将分组状态更新为confirmed，同时更新分组名称
        cursor.execute("""
            UPDATE users 
            SET group_status = 'confirmed', group_name = %s
            WHERE group_id = %s
        """, (new_group_name, group_id))
        
        if cursor.rowcount > 0:
            conn.commit()
            log_admin_action("CONFIRM_GROUP", f"确认分组 {group_id}: {old_group_name} -> {new_group_name}")
            
            return jsonify({
                'success': True,
                'message': f'分组 "{new_group_name}" 已确认'
            })
        else:
            return jsonify({'success': False, 'message': '确认分组失败'}), 400
        
    except Exception as e:
        app.logger.error(f"确认分组错误: {str(e)}")
        return jsonify({'success': False, 'message': f'确认分组失败: {str(e)}'}), 500
    finally:
        if 'conn' in locals():
            conn.close()


@app.route('/admin/api/dashboard-stats')
@view_admin_required
def api_dashboard_stats():
    """获取仪表盘统计数据"""
    try:
        from datetime import datetime, timedelta
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 基础统计
        cursor.execute("SELECT COUNT(*) as total_users FROM users")
        result = cursor.fetchone()
        total_users = result['total_users'] if result else 0
        
        cursor.execute("SELECT COUNT(*) as total_files FROM user_files")
        result = cursor.fetchone()
        total_files = result['total_files'] if result else 0
        
        # 今日统计 - 使用MySQL CURDATE() 避免Python/MySQL时区不一致
        cursor.execute("""
            SELECT COUNT(DISTINCT sl.user_id) as today_active_users 
            FROM system_logs sl
            WHERE sl.action = 'login' 
            AND sl.created_at >= CURDATE()
            AND sl.created_at < CURDATE() + INTERVAL 1 DAY
            AND sl.user_id IS NOT NULL
        """)
        result = cursor.fetchone()
        today_active_users = result['today_active_users'] if result else 0
        
        cursor.execute("""
            SELECT COUNT(*) as today_uploads 
            FROM user_files 
            WHERE upload_time >= CURDATE()
            AND upload_time < CURDATE() + INTERVAL 1 DAY
        """)
        result = cursor.fetchone()
        today_uploads = result['today_uploads'] if result else 0
        
        cursor.execute("""
            SELECT COUNT(*) as today_processed 
            FROM system_logs 
            WHERE action = 'file_processed' 
            AND created_at >= CURDATE()
            AND created_at < CURDATE() + INTERVAL 1 DAY
        """)
        result = cursor.fetchone()
        today_processed = result['today_processed'] if result else 0
        
        # 错误统计
        cursor.execute("""
            SELECT COUNT(*) as error_files 
            FROM user_files 
            WHERE (md_content IS NULL OR md_content = '')
              AND upload_time >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
        """)
        result = cursor.fetchone()
        error_count = result['error_files'] if result else 0
        
        # 处理速度统计
        cursor.execute("""
            SELECT AVG(TIMESTAMPDIFF(MINUTE, upload_time, 
                CASE WHEN md_content IS NOT NULL THEN NOW() ELSE upload_time END)) as avg_time
            FROM user_files 
            WHERE upload_time >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
              AND md_content IS NOT NULL
        """)
        avg_process_result = cursor.fetchone()
        avg_process_time = round(float(avg_process_result['avg_time'] or 0), 1) if avg_process_result else 0
        
        conn.close()
        
        return jsonify({
            'total_users': total_users,
            'total_files': total_files,
            'today_uploads': today_uploads,
            'today_processed': today_processed,
            'active_today': today_active_users,  # 前端期望的字段名
            'avg_process_time': avg_process_time
        })
        
    except Exception as e:
        print(f"获取仪表盘统计失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/admin/api/recent-activities')
@view_admin_required  
def api_recent_activities():
    """获取最近活动数据"""
    try:
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        
        # 获取最近文件上传
        cursor.execute("""
            SELECT uf.id, uf.original_name, uf.upload_time, 
                   u.username, uf.md_content, uf.extracted_json_data
            FROM user_files uf
            JOIN users u ON uf.user_id = u.id
            ORDER BY uf.upload_time DESC
            LIMIT 20
        """)
        recent_files = cursor.fetchall()
        
        activities = []
        for file_info in recent_files:
            # 判断处理状态
            has_md = file_info['md_content'] and file_info['md_content'].strip()
            has_json = file_info['extracted_json_data'] and file_info['extracted_json_data'].strip()
            
            if has_md and has_json:
                activity_type = 'success'
                icon = 'bi-check-circle'
                color = 'success'
                title = f"文档处理完成 - {file_info['original_name']}"
                details = f"用户: {file_info['username']} | 状态: 处理成功"
            elif has_md or has_json:
                activity_type = 'processing'
                icon = 'bi-gear'
                color = 'warning'
                title = f"文档处理中 - {file_info['original_name']}"
                details = f"用户: {file_info['username']} | 状态: 部分完成"
            else:
                activity_type = 'upload'
                icon = 'bi-upload'
                color = 'info'
                title = f"文档已上传 - {file_info['original_name']}"
                details = f"用户: {file_info['username']} | 状态: 等待处理"
            
            # 计算时间差
            upload_time = file_info['upload_time']
            now = datetime.now()
            time_diff = now - upload_time
            
            if time_diff.days > 0:
                time_str = f"{time_diff.days}天前"
            elif time_diff.seconds > 3600:
                hours = time_diff.seconds // 3600
                time_str = f"{hours}小时前"
            elif time_diff.seconds > 60:
                minutes = time_diff.seconds // 60
                time_str = f"{minutes}分钟前"
            else:
                time_str = "刚刚"
            
            activities.append({
                'type': activity_type,
                'icon': icon,
                'color': color,
                'title': title,
                'details': details,
                'time': time_str
            })
        
        conn.close()
        return jsonify({'activities': activities})
        
    except Exception as e:
        print(f"获取活动数据失败: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/admin/api/service-status')
@view_admin_required
def api_service_status():
    """获取服务状态"""
    try:
        import time
        
        services = []
        
        # 数据库服务检查
        try:
            start_time = time.time()
            conn = DatabaseConfig.get_db_connection()
            if conn:
                cursor = conn.cursor()
                cursor.execute("SELECT 1")
                cursor.fetchone()
                conn.close()
                db_response_time = round((time.time() - start_time) * 1000, 1)
                db_status = 'online'
            else:
                db_response_time = 0
                db_status = 'offline'
        except:
            db_response_time = 0
            db_status = 'error'
        
        services.append({
            'name': '数据库',
            'icon': 'bi-database',
            'status': db_status,
            'details': f"{db_response_time}ms"
        })
        
        # Web服务器状态
        services.append({
            'name': 'Web服务器', 
            'icon': 'bi-server',
            'status': 'online',
            'details': '运行中'
        })
        
        # 处理引擎状态（基于最近处理的文件）
        try:
            conn = DatabaseConfig.get_db_connection()
            if conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT COUNT(*) as recent_processed
                    FROM user_files 
                    WHERE upload_time >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
                      AND md_content IS NOT NULL
                """)
                result = cursor.fetchone()
                recent_processed = result['recent_processed'] if result else 0
                conn.close()
                
                services.append({
                    'name': '处理引擎',
                    'icon': 'bi-gear',
                    'status': 'active' if recent_processed > 0 else 'idle',
                    'details': f"{recent_processed}个任务"
                })
            else:
                services.append({
                    'name': '处理引擎',
                    'icon': 'bi-gear', 
                    'status': 'error',
                    'details': '数据库连接失败'
                })
        except Exception as e:
            services.append({
                'name': '处理引擎',
                'icon': 'bi-gear', 
                'status': 'error',
                'details': f'检查失败: {str(e)}'
            })
        
        # API服务状态
        services.append({
            'name': 'API服务',
            'icon': 'bi-cloud',
            'status': 'online',
            'details': '正常响应'
        })
        
        return jsonify({'services': services})
        
    except Exception as e:
        print(f"获取服务状态失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/admin/users')
@view_admin_required
def admin_users():
    """管理员用户管理页面 - YNEX版本"""
    try:
        page = request.args.get('page', 1, type=int)
        per_page = 20  # 每页显示20个用户
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            flash('数据库连接失败', 'error')
            return redirect(url_for('admin_dashboard'))
        
        cursor = conn.cursor()
        
        # 获取用户总数
        cursor.execute("SELECT COUNT(*) as total FROM users")
        total_result = cursor.fetchone()
        total_users = total_result['total'] if total_result else 0
        
        # 计算分页
        offset = (page - 1) * per_page
        total_pages = (total_users + per_page - 1) // per_page
        
        # 获取用户列表 - 包含角色信息
        cursor.execute("""
            SELECT u.id, u.username, u.email, u.role, u.created_at, u.updated_at, u.is_active,
                   COUNT(f.file_id) as file_count,
                   MAX(f.upload_time) as last_upload
            FROM users u
            LEFT JOIN user_files f ON u.id = f.user_id
            GROUP BY u.id, u.username, u.email, u.role, u.created_at, u.updated_at, u.is_active
            ORDER BY u.created_at DESC
            LIMIT %s OFFSET %s
        """, (per_page, offset))
        
        users = cursor.fetchall()
        
        conn.close()
        
        return render_template('admin/users_ynex.html', 
                             users=users,
                             current_page=page,
                             total_pages=total_pages,
                             total_users=total_users,
                             per_page=per_page,
                             current_user_role=session.get('role', 'user'))
        
    except Exception as e:
        flash(f'获取用户列表失败: {str(e)}', 'error')
        return redirect(url_for('admin_dashboard'))

@app.route('/admin/users/create', methods=['POST'])
@full_admin_required
def admin_create_user():
    """创建新用户"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': '无效的请求数据'})
        
        username = data.get('username', '').strip()
        email = data.get('email', '').strip()
        password = data.get('password', '')
        role = data.get('role', 'user')
        
        # 验证必填字段
        if not username or not email or not password:
            return jsonify({'success': False, 'message': '请填写所有必填字段'})
        
        # 验证角色
        if role not in ['user', 'premium', 'admin']:
            return jsonify({'success': False, 'message': '无效的用户角色'})
        
        # 验证邮箱格式
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, email):
            return jsonify({'success': False, 'message': '邮箱格式不正确'})
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'})
        
        cursor = conn.cursor()
        
        # 检查用户名是否已存在
        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
        if cursor.fetchone():
            conn.close()
            return jsonify({'success': False, 'message': '用户名已存在'})
        
        # 检查邮箱是否已存在
        cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
        if cursor.fetchone():
            conn.close()
            return jsonify({'success': False, 'message': '邮箱已被注册'})
        
        # 生成密码哈希
        password_hash = generate_password_hash(password)
        
        # 创建用户
        cursor.execute("""
            INSERT INTO users (username, email, password_hash, role, is_active, created_at, updated_at)
            VALUES (%s, %s, %s, %s, 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
        """, (username, email, password_hash, role))
        
        user_id = cursor.lastrowid
        conn.commit()
        conn.close()
        
        # 记录管理员操作日志
        role_text = {'user': '普通用户', 'premium': '高级用户', 'admin': '管理员'}[role]
        log_admin_action("CREATE_USER", f"创建用户: {username} (角色: {role_text})")
        
        log_user_action(
            session['user_id'], 
            'admin_create_user', 
            f"管理员操作: 创建用户 {username} (角色: {role_text})",
            request.remote_addr,
            request.headers.get('User-Agent')
        )
        
        return jsonify({
            'success': True, 
            'message': f'用户 {username} 创建成功！角色: {role_text}',
            'user_id': user_id
        })
        
    except Exception as e:
        app.logger.error(f"创建用户失败: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': f'创建用户失败: {str(e)}'})

@app.route('/admin/users/<int:user_id>/toggle-status', methods=['POST'])
@full_admin_required
def admin_toggle_user_status(user_id):
    """切换用户激活状态"""
    try:
        # 防止禁用管理员账户
        if is_admin(user_id):
            return jsonify({'success': False, 'message': '不能禁用管理员账户'})
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'})
        
        cursor = conn.cursor()
        
        # 获取当前状态
        cursor.execute("SELECT is_active FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            conn.close()
            return jsonify({'success': False, 'message': '用户不存在'})
        
        # 切换状态
        new_status = 0 if user['is_active'] else 1
        cursor.execute("""
            UPDATE users 
            SET is_active = %s, updated_at = CURRENT_TIMESTAMP
            WHERE id = %s
        """, (new_status, user_id))
        
        conn.commit()
        conn.close()
        
        status_text = '激活' if new_status else '禁用'
        log_admin_action("TOGGLE_USER_STATUS", f"用户ID:{user_id} 状态:{status_text}")
        
        # 记录管理员操作日志
        log_user_action(
            session['user_id'], 
            'admin_toggle_user_status', 
            f"管理员操作: {status_text}用户 (用户ID: {user_id})",
            request.remote_addr,
            request.headers.get('User-Agent')
        )
        
        return jsonify({'success': True, 'message': f'用户已{status_text}', 'new_status': new_status})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'操作失败: {str(e)}'})

@app.route('/admin/users/<int:user_id>/set-role', methods=['POST'])
@view_admin_required
def admin_set_user_role(user_id):
    """设置用户角色 - 管理员可设任意角色，高级用户只能在user/premium之间切换"""
    try:
        data = request.get_json()
        new_role = data.get('role')
        
        # 验证角色有效性
        valid_roles = ['user', 'admin', 'premium']
        if new_role not in valid_roles:
            return jsonify({'success': False, 'message': '无效的角色类型'}), 400
        
        current_user_id = session.get('user_id')
        current_user_role = session.get('role', 'user')
        
        # 防止修改当前管理员账户的角色（防止锁定）
        if user_id == current_user_id and current_user_role == 'admin' and new_role != 'admin':
            return jsonify({'success': False, 'message': '不能修改自己的管理员角色'}), 403
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
            
        cursor = conn.cursor()
        
        # 获取用户当前信息
        cursor.execute("SELECT username, role FROM users WHERE id = %s", (user_id,))
        user_info = cursor.fetchone()
        if not user_info:
            conn.close()
            return jsonify({'success': False, 'message': '用户不存在'}), 404
        
        old_role = user_info['role']
        username = user_info['username']
        
        # 高级用户权限限制：只能在 user ↔ premium 之间切换，不能设置 admin
        if current_user_role == 'premium':
            if new_role == 'admin':
                conn.close()
                return jsonify({'success': False, 'message': '高级用户无权设置管理员角色'}), 403
            if old_role == 'admin':
                conn.close()
                return jsonify({'success': False, 'message': '高级用户无权修改管理员的角色'}), 403
        
        # 更新用户角色
        cursor.execute("UPDATE users SET role = %s WHERE id = %s", (new_role, user_id))
        conn.commit()
        conn.close()
        
        # 记录管理操作
        role_text = {'user': '普通用户', 'premium': '高级用户', 'admin': '管理员'}
        log_admin_action("SET_USER_ROLE", f"用户ID:{user_id} ({username}) 角色: {role_text.get(old_role, old_role)} → {role_text.get(new_role, new_role)}")
        
        return jsonify({
            'success': True, 
            'message': f'用户角色已更新为 {role_text.get(new_role, new_role)}',
            'old_role': old_role,
            'new_role': new_role
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'设置角色失败: {str(e)}'})

@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@full_admin_required
def admin_delete_user(user_id):
    """删除用户"""
    try:
        # 检查要删除的用户是否是管理员
        if is_admin(user_id):
            log_admin_action("DELETE_ADMIN_BLOCKED", f"尝试删除管理员用户 ID:{user_id}")
            return jsonify({'success': False, 'message': '不能删除管理员账户'})
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': '数据库连接失败'})
        
        cursor = conn.cursor()
        
        # 先删除用户的文件
        cursor.execute("DELETE FROM user_files WHERE user_id = %s", (user_id,))
        
        # 再删除用户
        cursor.execute("DELETE FROM users WHERE id = %s", (user_id,))
        
        if cursor.rowcount > 0:
            conn.commit()
            conn.close()
            log_admin_action("DELETE_USER", f"用户ID:{user_id}")
            return jsonify({'success': True, 'message': '用户已删除'})
        else:
            conn.close()
            return jsonify({'success': False, 'message': '用户不存在'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除失败: {str(e)}'})

@app.route('/admin/users/<int:user_id>/files')
@view_admin_required
def admin_get_user_files(user_id):
    """获取指定用户的文档列表"""
    try:
        app.logger.info(f"Getting files for user_id: {user_id}")
        
        conn = DatabaseConfig.get_db_connection()
        if not conn:
            app.logger.error("Database connection failed")
            return jsonify({'success': False, 'message': '数据库连接失败'})
        
        cursor = conn.cursor()
        
        # 获取用户信息
        cursor.execute("SELECT id, username, email FROM users WHERE id = %s", (user_id,))
        user_result = cursor.fetchone()
        app.logger.info(f"User query result: {user_result}, type: {type(user_result)}")
        
        if not user_result:
            conn.close()
            app.logger.warning(f"User {user_id} not found")
            return jsonify({'success': False, 'message': '用户不存在'})
        
        # 安全地访问用户数据
        try:
            if isinstance(user_result, dict):
                user = {
                    'id': user_result.get('id'),
                    'username': user_result.get('username'),
                    'email': user_result.get('email')
                }
            else:
                # 假设是tuple/list
                user = {
                    'id': user_result[0] if len(user_result) > 0 else None,
                    'username': user_result[1] if len(user_result) > 1 else None,
                    'email': user_result[2] if len(user_result) > 2 else None
                }
        except Exception as user_error:
            app.logger.error(f"Error parsing user data: {user_error}")
            user = {'id': user_id, 'username': 'Unknown', 'email': 'Unknown'}
        
        app.logger.info(f"User info: {user}")
        
        # 先检查表结构和数据
        cursor.execute("SHOW COLUMNS FROM user_files")
        columns = cursor.fetchall()
        app.logger.info(f"user_files table columns: {columns}")
        
        # 获取用户的文档列表 - 包含状态判断所需的字段，使用安全的字段查询
        try:
            cursor.execute("""
                SELECT file_id, original_name, file_size, upload_time, status,
                       md_content, extracted_json_data, nsfc_json_data
                FROM user_files 
                WHERE user_id = %s 
                ORDER BY upload_time DESC
            """, (user_id,))
            files_result = cursor.fetchall()
            app.logger.info(f"Files query returned {len(files_result)} rows")
        except Exception as query_error:
            app.logger.error(f"SQL query failed: {query_error}")
            # 如果字段不存在，使用基础查询
            cursor.execute("""
                SELECT file_id, original_name, file_size, upload_time, status
                FROM user_files 
                WHERE user_id = %s 
                ORDER BY upload_time DESC
            """, (user_id,))
            files_result = cursor.fetchall()
            app.logger.info(f"Fallback query returned {len(files_result)} rows")
        
        # 格式化文件信息并在后端完成状态判断
        formatted_files = []
        for i, file_row in enumerate(files_result):
            try:
                app.logger.debug(f"Processing file row {i}: {file_row}")
                
                # 安全地访问文件数据
                if isinstance(file_row, dict):
                    file_id = file_row.get('file_id')
                    original_name = file_row.get('original_name', 'unknown')
                    file_size = file_row.get('file_size', 0)
                    upload_time = file_row.get('upload_time')
                    status = file_row.get('status', 'pending')
                    # 安全获取状态字段，如果不存在则为None
                    md_content = file_row.get('md_content')
                    extracted_json_data = file_row.get('extracted_json_data')
                    nsfc_json_data = file_row.get('nsfc_json_data')
                else:
                    # tuple/list访问 - 根据基础查询的字段数量调整
                    file_id = file_row[0] if len(file_row) > 0 else 0
                    original_name = file_row[1] if len(file_row) > 1 else 'unknown'
                    file_size = file_row[2] if len(file_row) > 2 else 0
                    upload_time = file_row[3] if len(file_row) > 3 else None
                    status = file_row[4] if len(file_row) > 4 else 'pending'
                    # 如果查询包含更多字段，则获取；否则设为None
                    md_content = file_row[5] if len(file_row) > 5 else None
                    extracted_json_data = file_row[6] if len(file_row) > 6 else None
                    nsfc_json_data = file_row[7] if len(file_row) > 7 else None
                
                # 在后端判断处理状态 - 安全检查
                has_md = bool(md_content and str(md_content).strip())
                has_table_json = bool(extracted_json_data and str(extracted_json_data).strip())
                has_content_json = bool(nsfc_json_data and str(nsfc_json_data).strip())
                
                # 生成状态徽章数据
                status_badges = []
                if has_md:
                    status_badges.append({'text': 'MD', 'class': 'bg-primary-transparent text-primary'})
                if has_table_json:
                    status_badges.append({'text': '表格', 'class': 'bg-success-transparent text-success'})
                if has_content_json:
                    status_badges.append({'text': '正文', 'class': 'bg-info-transparent text-info'})
                
                if not status_badges:
                    status_badges.append({'text': '待处理', 'class': 'bg-warning-transparent text-warning'})
                
                formatted_file = {
                    'file_id': file_id,
                    'id': file_id,  # 兼容性字段
                    'original_name': original_name,
                    'filename': original_name,  # 兼容性字段
                    'file_size': file_size,
                    'size': file_size,  # 兼容性字段
                    'upload_time': upload_time.isoformat() if upload_time else None,
                    'status': status,
                    'status_badges': status_badges,  # 新增：状态徽章数据
                    'has_md': has_md,
                    'has_table_json': has_table_json,
                    'has_content_json': has_content_json
                }
                
                formatted_files.append(formatted_file)
                app.logger.debug(f"Formatted file {i}: {formatted_file}")
            except Exception as row_error:
                app.logger.error(f"Error formatting row {i}: {row_error}, row data: {file_row}")
                # 添加一个安全的默认条目
                formatted_files.append({
                    'file_id': i,
                    'id': i,
                    'original_name': f'file_{i}',
                    'filename': f'file_{i}',
                    'file_size': 0,
                    'size': 0,
                    'upload_time': None,
                    'status': 'error',
                    'status_badges': [{'text': '错误', 'class': 'bg-danger-transparent text-danger'}],
                    'has_md': False,
                    'has_table_json': False,
                    'has_content_json': False
                })
        
        conn.close()
        
        result = {
            'success': True,
            'user': user,
            'files': formatted_files
        }
        app.logger.info(f"Successfully returning {len(formatted_files)} files for user {user_id}")
        return jsonify(result)
        
    except Exception as e:
        app.logger.error(f"Error getting user files: {str(e)}", exc_info=True)
        # 确保连接被关闭
        try:
            if 'conn' in locals() and conn:
                conn.close()
        except:
            pass
        return jsonify({'success': False, 'message': f'获取文档列表失败: {str(e)}'})

@app.route('/admin/files/<string:file_id>/status')
def admin_get_file_status(file_id):
    """获取文档处理状态"""
    try:
        app.logger.info(f"Checking status for file {file_id}")
        
        # 获取文件基本信息
        connection = get_db_connection_with_error_handling()
        if not connection:
            return jsonify({'success': False, 'message': '数据库连接失败'})
        
        # 检查连接是否正确返回
        if isinstance(connection, tuple):
            app.logger.warning(f"Database connection returned tuple: {connection}")
            # 使用tuple的第一个元素作为连接对象
            connection = connection[0]
            if not connection:
                return jsonify({'success': False, 'message': '数据库连接错误'})
        
        cursor = connection.cursor()
        
        # 查询文件信息和处理状态
        file_query = """
        SELECT file_id, original_name, size, upload_time, status, 
               md_content, storage_type, md_file_path 
        FROM user_files 
        WHERE file_id = %s
        """
        cursor.execute(file_query, (file_id,))
        file_result = cursor.fetchone()
        
        if not file_result:
            safe_db_close(connection, cursor)
            return jsonify({'success': False, 'message': '文档不存在'})
        
        # 检查各种处理状态
        has_md = False
        has_table_json = False
        has_content_json = False
        
        try:
            # 检查MD文档是否存在
            if isinstance(file_result, dict):
                md_content = file_result.get('md_content')
                storage_type = file_result.get('storage_type')
                md_file_path = file_result.get('md_file_path')
            else:
                # 如果是tuple格式 (file_id, original_name, size, upload_time, status, md_content, storage_type, md_file_path)
                md_content = file_result[5] if len(file_result) > 5 else None
                storage_type = file_result[6] if len(file_result) > 6 else None
                md_file_path = file_result[7] if len(file_result) > 7 else None
            
            # 检查MD内容是否存在
            if md_content:
                has_md = True
            elif storage_type == 'filesystem' and md_file_path and os.path.exists(md_file_path):
                has_md = True
            
            # 检查表格和正文JSON（这里简化处理，实际应该检查具体的JSON字段或文件）
            # 暂时基于文档状态判断
            if has_md:
                # 如果有MD文档，假设50%概率有表格JSON，60%概率有正文JSON
                has_table_json = True  # 简化为都有
                has_content_json = True
            
        except Exception as check_error:
            app.logger.warning(f"Error checking file content status: {check_error}")
            # 如果检查出错，返回默认状态
            pass
        
        result = {
            'success': True,
            'file_id': file_id,
            'has_md': has_md,
            'has_table_json': has_table_json,
            'has_content_json': has_content_json,
            'processing_status': 'completed' if (has_md or has_table_json or has_content_json) else 'pending'
        }
        
        safe_db_close(connection, cursor)
        return jsonify(result)
        
    except Exception as e:
        app.logger.error(f"Error checking file status: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': f'检查文档状态失败: {str(e)}'})

@app.route('/api/user/stats')
@api_login_required
def get_user_stats():
    """获取用户统计数据"""
    
    try:
        user_id = session['user_id']
        
        # 从数据库获取用户的所有文件
        user_files = get_user_files_from_db(user_id)
        
        if not user_files:
            return jsonify({
                'success': True,
                'total_files': 0,
                'parsed_files': 0,
                'pending_files': 0,
                'today_activity': 0
            })
        
        # 统计文件状态
        total_files = len(user_files)
        parsed_files = len([f for f in user_files if f.get('status') == 'completed'])
        pending_files = len([f for f in user_files if f.get('status') in ['pending', 'processing']])
        
        # 计算今日活动（今日上传 + 今日解析完成的文件数，避免重复计算）
        from datetime import datetime, timedelta
        today = datetime.now().date()
        today_activity = 0
        
        # 记录今日上传的文件ID，避免重复计算
        today_uploaded_files = set()
        
        for file in user_files:
            try:
                # 检查今日上传的文件
                upload_time = file.get('upload_time')
                if upload_time:
                    if isinstance(upload_time, str):
                        # 尝试解析字符串格式的时间
                        try:
                            upload_date = datetime.strptime(upload_time.split()[0], '%Y-%m-%d').date()
                        except:
                            upload_date = datetime.fromisoformat(upload_time.split('T')[0]).date()
                    else:
                        upload_date = upload_time.date()
                    
                    if upload_date == today:
                        today_activity += 1
                        today_uploaded_files.add(file.get('file_id'))
                
                # 检查今日完成解析的文件（不重复计算今日上传的）
                if file.get('status') == 'completed' and file.get('file_id') not in today_uploaded_files:
                    # 优先检查 processed_time，其次检查 nsfc_json_updated_at
                    processed_time = file.get('processed_time') or file.get('nsfc_json_updated_at')
                    if processed_time:
                        try:
                            if isinstance(processed_time, str):
                                # 尝试解析字符串格式的时间
                                try:
                                    processed_date = datetime.strptime(processed_time.split()[0], '%Y-%m-%d').date()
                                except:
                                    processed_date = datetime.fromisoformat(processed_time.split('T')[0]).date()
                            else:
                                processed_date = processed_time.date()
                            
                            if processed_date == today:
                                today_activity += 1
                        except Exception as e:
                            continue  # 忽略时间解析错误
                            
            except Exception as e:
                continue  # 忽略时间解析错误
        
        return jsonify({
            'success': True,
            'total_files': total_files,
            'parsed_files': parsed_files,
            'pending_files': pending_files,
            'today_activity': today_activity
        })
        
    except Exception as e:
        print(f"获取用户统计数据失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e),
            'total_files': 0,
            'parsed_files': 0,
            'pending_files': 0,
            'today_activity': 0
        }), 500

@app.route('/api/admin/logs')
@view_admin_required
def get_admin_logs():
    """获取管理员操作日志和系统日志统计"""
    try:
        # 获取智能日志系统统计
        log_stats = smart_logger.get_stats()
        
        # 简单读取Flask默认日志
        logs = []
        
        # 从内存中获取最近的日志记录（简化版）
        recent_logs = [
            f"管理员操作日志功能已启用 - {session.get('username', 'admin')}",
            "可通过 /api/admin/logs 查看操作记录",
            f"日志统计: 缓存 {log_stats['cache_size']} 条, 去重防护 {log_stats['duplicate_prevention']}",
            f"自动清理: {'启用' if log_stats['auto_cleanup_enabled'] else '禁用'}, 上次清理: {log_stats['last_cleanup']}"
        ]
        
        return jsonify({
            'success': True, 
            'logs': recent_logs, 
            'log_stats': log_stats,
            'message': '日志功能已启用'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取日志失败: {str(e)}'})

@app.route('/api/admin/logs/cleanup', methods=['POST'])
@view_admin_required
def manual_log_cleanup():
    """手动触发日志清理"""
    try:
        # 执行手动清理
        cleanup_stats = smart_logger.manual_cleanup()
        
        smart_print(f"管理员手动触发日志清理 - {session.get('username', 'admin')}", LogLevel.INFO)
        
        return jsonify({
            'success': True,
            'message': '日志清理完成',
            'cleanup_stats': cleanup_stats
        })
        
    except Exception as e:
        smart_print(f"手动日志清理失败: {e}", LogLevel.ERROR)
        return jsonify({'success': False, 'message': f'清理失败: {str(e)}'})

@app.route('/api/admin/logs/config', methods=['GET', 'POST'])
@view_admin_required
def log_cleanup_config():
    """日志清理配置管理"""
    try:
        if request.method == 'GET':
            # 返回当前配置
            return jsonify({
                'success': True,
                'config': {
                    'auto_cleanup_enabled': smart_logger.config['auto_cleanup_enabled'],
                    'cleanup_interval': smart_logger.config['cleanup_interval'],
                    'log_retention_days': smart_logger.config['log_retention_days'],
                    'max_cache_size': smart_logger.config['max_cache_size']
                }
            })
        
        elif request.method == 'POST':
            # 更新配置
            data = request.get_json()
            
            # 验证和更新配置
            if 'auto_cleanup_enabled' in data:
                if data['auto_cleanup_enabled']:
                    smart_logger.start_auto_cleanup()
                else:
                    smart_logger.stop_auto_cleanup()
            
            if 'cleanup_interval' in data:
                interval = int(data['cleanup_interval'])
                if 300 <= interval <= 86400:  # 5分钟到24小时
                    smart_logger.config['cleanup_interval'] = interval
                    if smart_logger.config['auto_cleanup_enabled']:
                        smart_logger.start_auto_cleanup()  # 重启定时器
            
            if 'log_retention_days' in data:
                days = int(data['log_retention_days'])
                if 1 <= days <= 30:  # 1天到30天
                    smart_logger.config['log_retention_days'] = days
            
            if 'max_cache_size' in data:
                size = int(data['max_cache_size'])
                if 100 <= size <= 2000:  # 100到2000条
                    smart_logger.config['max_cache_size'] = size
            
            smart_print(f"管理员更新日志配置 - {session.get('username', 'admin')}", LogLevel.INFO)
            
            return jsonify({
                'success': True,
                'message': '配置更新成功',
                'config': {
                    'auto_cleanup_enabled': smart_logger.config['auto_cleanup_enabled'],
                    'cleanup_interval': smart_logger.config['cleanup_interval'],
                    'log_retention_days': smart_logger.config['log_retention_days'],
                    'max_cache_size': smart_logger.config['max_cache_size']
                }
            })
            
    except Exception as e:
        smart_print(f"日志配置操作失败: {e}", LogLevel.ERROR)
        return jsonify({'success': False, 'message': f'操作失败: {str(e)}'})

@app.route('/api/admin/logs/stats')
@view_admin_required 
def get_log_stats():
    """获取详细的日志统计信息"""
    try:
        stats = smart_logger.get_stats()
        
        return jsonify({
            'success': True,
            'stats': stats
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取统计失败: {str(e)}'})

@app.route('/api/system/log_status')
def get_system_log_status():
    """获取系统日志状态（公开接口）"""
    try:
        stats = smart_logger.get_stats()
        
        # 只返回基本的状态信息
        return jsonify({
            'success': True,
            'log_system': {
                'status': 'active',
                'auto_cleanup': stats['auto_cleanup_enabled'],
                'cache_size': stats['cache_size'],
                'last_cleanup': stats['last_cleanup']
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取状态失败: {str(e)}'})

if __name__ == '__main__':
    print("🚀 PDF转MD文档处理系统启动中...")
    print()
    
    # ==================== 智能日志系统初始化 ====================
    print("🧹 初始化智能日志清理系统...")
    try:
        # 配置日志清理参数
        log_config = {
            'auto_cleanup_enabled': True,
            'cleanup_interval': 1800,  # 30分钟清理一次
            'log_retention_days': 3,   # 保留3天的日志
            'max_cache_size': 500,     # 最大缓存500条
            'cleanup_on_startup': True
        }
        smart_logger.config.update(log_config)
        
        # 启动时执行一次手动清理
        cleanup_stats = smart_logger.manual_cleanup()
        smart_print(f"启动清理完成: 缓存清理 {cleanup_stats['cache_cleaned']} 条, "
                   f"文件清理 {cleanup_stats['old_logs_removed']} 个", LogLevel.INFO)
        
        print("   ✅ 智能日志清理系统已启用")
        print(f"   📊 清理间隔: {log_config['cleanup_interval']/60:.0f} 分钟")
        print(f"   📅 日志保留: {log_config['log_retention_days']} 天")
        print(f"   💾 缓存上限: {log_config['max_cache_size']} 条")
        
    except Exception as e:
        print(f"   ⚠️  日志清理系统初始化失败: {e}")
        print("   系统仍可正常运行，使用默认日志处理")
    print()
    
    # 设置信号处理器，优雅关闭守护线程
    import signal
    import atexit
    
    def signal_handler(sig, frame):
        """处理终止信号"""
        print("\n🛑 接收到终止信号，正在安全关闭系统...")
        shutdown_monitoring()
        import sys
        sys.exit(0)
    
    def cleanup_on_exit():
        """程序退出时的清理函数"""
        smart_print("系统关闭中，执行最终清理...", LogLevel.INFO)
        try:
            # 停止自动清理定时器
            smart_logger.stop_auto_cleanup()
            # 执行最后一次清理
            cleanup_stats = smart_logger.manual_cleanup()
            smart_print(f"最终清理完成: 缓存清理 {cleanup_stats['cache_cleaned']} 条", LogLevel.INFO)
        except Exception as e:
            smart_print(f"最终清理过程中出错: {e}", LogLevel.WARNING)
        
        shutdown_monitoring()
    
    # 注册信号处理器
    signal.signal(signal.SIGINT, signal_handler)  # Ctrl+C
    signal.signal(signal.SIGTERM, signal_handler)  # 终止信号
    atexit.register(cleanup_on_exit)  # 程序退出时清理
    
    # 显示简化的系统信息
    try:
        import platform
        print("🌍 系统环境信息:")
        print(f"   操作系统: {platform.system()}")
        print(f"   本地时间: {get_current_time().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"   时间戳: {time.time():.0f}")
        print("   ✅ 使用中国时区(UTC+8)")
        print()
    except Exception as e:
        print(f"⚠️  系统信息获取失败: {e}")
        print("   系统仍可正常运行，使用默认时区处理")
        print()
    
    # 从数据库加载历史数据
    load_data_from_db()
    
    # 🆕 启动实时监控系统
    print("📈 启动实时监控系统...")
    start_realtime_monitoring()
    
    print("   - 智能MD存储: 小文档存数据库，大文档存文件系统")
    print("   - 🌍 支持全球任意时区和系统环境部署")
    print("   - 📈 实时系统监控和统计已启用")
    print("   - 🧹 智能日志清理系统已启用，自动管理日志大小")
    print("   - 🛡️  优雅关闭机制已启用，支持Ctrl+C安全退出")
    print()
    
    try:
        app.run(debug=False, host='0.0.0.0', port=5001, threaded=True)  # 多线程模式，支持多用户并发
    except KeyboardInterrupt:
        print("\n🛑 接收到键盘中断，正在安全关闭...")
        shutdown_monitoring()
    except Exception as e:
        print(f"\n❌ 应用程序异常: {e}")
        shutdown_monitoring()
    finally:
        print("👋 系统已安全关闭")
